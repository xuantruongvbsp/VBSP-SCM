# -*- coding: utf-8 -*-
"""Chạy một lần: ghi data/giao_ban.py — có thể xóa file này sau."""
from pathlib import Path

CONTENT = r'''"""
Tính số liệu giao ban xã và tạo 3 bảng động cho mẫu Word.
Dùng chung cho ws_operation (CBTD) và ws_management (PGD).
"""
import pandas as pd
from datetime import date
from dateutil.relativedelta import relativedelta
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from config import (
    COT_TEN_PGD, COT_MA_KH, COT_SO_KU, COT_TONG_DU_NO,
    COT_DU_NO_QH, COT_TEN_CT, COT_TEN_XA,
)

COT_DU_NO_KHOANH   = "Dư nợ khoanh"
COT_NGAY_DH_GH     = "Ngày ĐH theo Gia hạn"
COT_DVUT           = "Tên ĐVUT"
COT_TEN_TO         = "Tên tổ"
COT_DS_CV_THANG    = "Giải ngân trong tháng"
COT_TN_TH_THANG    = "Thu nợ TH tháng"
COT_TN_QH_THANG    = "Thu nợ QH tháng"
COT_TIEN_GUI       = "Số dư tiền gửi 105"
COT_MON_3M         = "is_3m_inactive"   # từ danh_dau_khong_hd()


# ── Lọc dữ liệu theo xã / điểm giao dịch ─────────────────────────────────────

def loc_theo_xa(df: pd.DataFrame, ten_xa: str) -> pd.DataFrame:
    return df[df[COT_TEN_XA] == ten_xa].copy()


# ── Tính số liệu văn xuôi ─────────────────────────────────────────────────────

def tinh_so_lieu_van_xuoi(
    df_xa: pd.DataFrame,
    df_baseline: pd.DataFrame | None,
    nam_moc: int,
) -> dict:
    """
    Trả về dict các tag văn xuôi cho mẫu Word.
    df_xa       : HSTD đã lọc theo xã hiện tại
    df_baseline : HSTD mốc 31/12 (toàn CN, chưa lọc) — None nếu chưa có
    nam_moc     : năm của baseline (VD: 2025)
    """
    dn  = pd.to_numeric(df_xa[COT_TONG_DU_NO], errors="coerce").fillna(0)
    qh  = pd.to_numeric(df_xa[COT_DU_NO_QH],   errors="coerce").fillna(0)
    kh  = pd.to_numeric(df_xa.get(COT_DU_NO_KHOANH, 0), errors="coerce").fillna(0)
    tg  = pd.to_numeric(df_xa.get(COT_TIEN_GUI, 0),     errors="coerce").fillna(0)
    gn  = pd.to_numeric(df_xa.get(COT_DS_CV_THANG, 0),  errors="coerce").fillna(0)
    tn_th = pd.to_numeric(df_xa.get(COT_TN_TH_THANG, 0), errors="coerce").fillna(0)
    tn_qh = pd.to_numeric(df_xa.get(COT_TN_QH_THANG, 0), errors="coerce").fillna(0)

    tong_dn     = dn.sum()
    tong_qh     = qh.sum()
    tong_kh     = kh.sum()
    tong_tg     = tg.sum()
    tl_qh       = tong_qh / tong_dn * 100 if tong_dn > 0 else 0.0
    tl_kh       = tong_kh / tong_dn * 100 if tong_dn > 0 else 0.0
    so_kh       = df_xa[df_xa[COT_TONG_DU_NO] > 0][COT_MA_KH].nunique()
    so_to       = df_xa[COT_TEN_TO].nunique() if COT_TEN_TO in df_xa.columns else 0

    # So với tháng trước (tính ngược)
    dn_thang_truoc  = tong_dn - gn.sum() + tn_th.sum() + tn_qh.sum()
    chenh_lech_thang = tong_dn - dn_thang_truoc
    tang_giam_thang  = "tăng" if chenh_lech_thang >= 0 else "giảm"

    # So với đầu năm (baseline)
    chenh_lech_dn = pct_dau_nam = 0.0
    tang_giam_dau_nam = "tăng"
    if df_baseline is not None and COT_TEN_XA in df_baseline.columns:
        df_bl_xa = df_baseline[df_baseline[COT_TEN_XA] == df_xa[COT_TEN_XA].iloc[0]]
        dn_bl = pd.to_numeric(df_bl_xa[COT_TONG_DU_NO], errors="coerce").fillna(0).sum()
        chenh_lech_dn    = tong_dn - dn_bl
        pct_dau_nam      = chenh_lech_dn / dn_bl * 100 if dn_bl > 0 else 0.0
        tang_giam_dau_nam = "tăng" if chenh_lech_dn >= 0 else "giảm"

    from utils import fmt, fmt_pct
    return {
        "{{tong_du_no}}":        fmt(tong_dn),
        "{{so_kh}}":             str(so_kh),
        "{{so_to}}":             str(so_to),
        "{{du_no_qh}}":          fmt(tong_qh),
        "{{ty_le_nqh}}":         f"{tl_qh:.2f}",
        "{{du_no_khoanh}}":      fmt(tong_kh),
        "{{ty_le_khoanh}}":      f"{tl_kh:.2f}",
        "{{tien_gui_105}}":      fmt(tong_tg),
        "{{tang_giam_thang}}":   tang_giam_thang,
        "{{chenh_lech_thang}}":  fmt(abs(chenh_lech_thang)),
        "{{tang_giam_dau_nam}}": tang_giam_dau_nam,
        "{{chenh_lech_dau_nam}}":fmt(abs(chenh_lech_dn)),
        "{{pct_dau_nam}}":       f"{abs(pct_dau_nam):.1f}",
        "{{nam_moc}}":           str(nam_moc),
    }


# ── Tạo bảng ĐVUT ─────────────────────────────────────────────────────────────

def tao_bang_dvut(doc: Document, df_xa: pd.DataFrame) -> None:
    """Chèn bảng kết quả theo ĐVUT vào doc tại vị trí placeholder."""
    from data import danh_dau_khong_hd
    df_m = danh_dau_khong_hd(df_xa)

    DVUT_ORDER = ["Hội nông dân", "Hội liên hiệp phụ nữ",
                  "Hội cựu chiến binh", "Đoàn thanh niên"]

    t = df_m.groupby(COT_DVUT).agg(
        so_to    =(COT_TEN_TO,      "nunique"),
        so_kh    =(COT_MA_KH,       "nunique"),
        ds_cv    =(COT_DS_CV_THANG, "sum"),
        ds_tn    =(COT_TN_TH_THANG, "sum"),
        tong_dn  =(COT_TONG_DU_NO,  "sum"),
        nqh      =(COT_DU_NO_QH,    "sum"),
        no_kh    =(COT_DU_NO_KHOANH,"sum"),
        mon_3m   =(COT_MON_3M,      "sum"),
    ).reset_index()

    # Sắp xếp theo thứ tự chuẩn, bỏ ĐVUT không có dư nợ
    t = t[t["tong_dn"] > 0]
    t["_ord"] = t[COT_DVUT].apply(
        lambda x: DVUT_ORDER.index(x) if x in DVUT_ORDER else 99)
    t = t.sort_values("_ord").drop(columns="_ord")

    tong_dn_all = t["tong_dn"].sum()

    HEADERS = ["Stt", "Đơn vị nhận ủy thác", "Số Tổ", "Số KH",
               "DS cho vay", "DS thu nợ", "Tổng dư nợ",
               "Tỷ trọng %", "NQH", "Nợ khoanh", "Món 3T KHĐ"]

    tbl = doc.add_table(rows=1, cols=len(HEADERS))
    tbl.style = "Table Grid"
    # Header
    for i, h in enumerate(HEADERS):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(9)

    from utils import fmt
    for idx, row in enumerate(t.itertuples(), 1):
        ty_trong = row.tong_dn / tong_dn_all * 100 if tong_dn_all > 0 else 0
        vals = [
            str(idx), row[1],
            str(int(row.so_to)), str(int(row.so_kh)),
            fmt(row.ds_cv), fmt(row.ds_tn),
            fmt(row.tong_dn), f"{ty_trong:.1f}%",
            fmt(row.nqh), fmt(row.no_kh), str(int(row.mon_3m)),
        ]
        tr = tbl.add_row()
        for i, v in enumerate(vals):
            tr.cells[i].text = v
            tr.cells[i].paragraphs[0].alignment = (
                WD_ALIGN_PARAGRAPH.CENTER if i != 1
                else WD_ALIGN_PARAGRAPH.LEFT)
            tr.cells[i].paragraphs[0].runs[0].font.size = Pt(9)

    # Dòng Cộng
    tr = tbl.add_row()
    cong_vals = [
        "", "Cộng",
        str(int(t["so_to"].sum())), str(int(t["so_kh"].sum())),
        fmt(t["ds_cv"].sum()), fmt(t["ds_tn"].sum()),
        fmt(tong_dn_all), "100%",
        fmt(t["nqh"].sum()), fmt(t["no_kh"].sum()),
        str(int(t["mon_3m"].sum())),
    ]
    for i, v in enumerate(cong_vals):
        tr.cells[i].text = v
        run = tr.cells[i].paragraphs[0].runs[0] if tr.cells[i].paragraphs[0].runs else tr.cells[i].paragraphs[0].add_run(v)
        run.bold = True
        run.font.size = Pt(9)


# ── Tạo bảng chương trình so sánh đầu năm ────────────────────────────────────

def tao_bang_chuong_trinh(
    doc: Document,
    df_xa: pd.DataFrame,
    df_baseline: pd.DataFrame | None,
    nam_moc: int,
) -> None:
    if df_baseline is not None and COT_TEN_XA in df_baseline.columns:
        ten_xa = df_xa[COT_TEN_XA].iloc[0]
        df_bl  = df_baseline[df_baseline[COT_TEN_XA] == ten_xa]
    else:
        df_bl  = None

    g_ht = df_xa.groupby(COT_TEN_CT)[COT_TONG_DU_NO].sum().reset_index()
    g_ht.columns = [COT_TEN_CT, "dn_ht"]

    if df_bl is not None and not df_bl.empty:
        g_bl = df_bl.groupby(COT_TEN_CT)[COT_TONG_DU_NO].sum().reset_index()
        g_bl.columns = [COT_TEN_CT, "dn_bl"]
        mg = g_ht.merge(g_bl, on=COT_TEN_CT, how="left").fillna(0)
    else:
        mg = g_ht.copy()
        mg["dn_bl"] = 0

    mg = mg[mg["dn_ht"] > 0].copy()
    mg["chenh_lech"] = mg["dn_ht"] - mg["dn_bl"]
    mg["pct"] = mg.apply(
        lambda r: r["chenh_lech"] / r["dn_bl"] * 100 if r["dn_bl"] > 0 else 0,
        axis=1)

    HEADERS = ["Stt", "Chương trình tín dụng",
               f"Dư nợ 31/12/{nam_moc}", "Dư nợ hiện tại",
               "Tăng/giảm", "%"]

    tbl = doc.add_table(rows=1, cols=len(HEADERS))
    tbl.style = "Table Grid"
    for i, h in enumerate(HEADERS):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(9)

    from utils import fmt
    for idx, row in enumerate(mg.itertuples(), 1):
        dau = "+" if row.chenh_lech >= 0 else ""
        vals = [
            str(idx), row[1],
            fmt(row.dn_bl), fmt(row.dn_ht),
            f"{dau}{fmt(row.chenh_lech)}", f"{dau}{row.pct:.1f}%",
        ]
        tr = tbl.add_row()
        for i, v in enumerate(vals):
            tr.cells[i].text = v
            tr.cells[i].paragraphs[0].alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if i == 1
                else WD_ALIGN_PARAGRAPH.CENTER)
            tr.cells[i].paragraphs[0].runs[0].font.size = Pt(9)

    # Dòng Cộng
    tr = tbl.add_row()
    tong_bl = mg["dn_bl"].sum()
    tong_ht = mg["dn_ht"].sum()
    tong_cl = tong_ht - tong_bl
    dau = "+" if tong_cl >= 0 else ""
    pct_tong = tong_cl / tong_bl * 100 if tong_bl > 0 else 0
    cong = ["", "Cộng", fmt(tong_bl), fmt(tong_ht),
            f"{dau}{fmt(tong_cl)}", f"{dau}{pct_tong:.1f}%"]
    for i, v in enumerate(cong):
        tr.cells[i].text = v
        run = tr.cells[i].paragraphs[0].runs[0] if tr.cells[i].paragraphs[0].runs else tr.cells[i].paragraphs[0].add_run(v)
        run.bold = True
        run.font.size = Pt(9)


# ── Tạo bảng kế hoạch thu nợ / giải ngân ─────────────────────────────────────

def tao_bang_ke_hoach(
    doc: Document,
    df_xa: pd.DataFrame,
    giai_ngan_input: dict | None = None,
) -> None:
    """
    df_xa            : HSTD lọc theo xã hiện tại
    giai_ngan_input  : dict {(ten_dvut, ten_to, ten_ct): so_tien} — CBTD nhập tay
                       None = để trống cột giải ngân
    """
    thang_toi = date.today() + relativedelta(months=1)
    thang_toi_dau = date(thang_toi.year, thang_toi.month, 1)

    # Lọc món đến hạn tháng tới
    ngay_dh = pd.to_datetime(df_xa[COT_NGAY_DH_GH], errors="coerce")
    mask = (ngay_dh.dt.year  == thang_toi_dau.year) & \
           (ngay_dh.dt.month == thang_toi_dau.month)
    df_dh = df_xa[mask].copy()

    g = df_dh.groupby([COT_DVUT, COT_TEN_TO, COT_TEN_CT])[COT_TONG_DU_NO] \
             .sum().reset_index()
    g = g[g[COT_TONG_DU_NO] > 0]

    HEADERS = ["Stt", "Đơn vị nhận ủy thác / Tổ TK&VV",
               "Chương trình cho vay", "Số tiền thu nợ", "Số tiền giải ngân"]

    tbl = doc.add_table(rows=1, cols=len(HEADERS))
    tbl.style = "Table Grid"
    for i, h in enumerate(HEADERS):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(9)

    from utils import fmt
    stt = 0
    for dvut, grp_dvut in g.groupby(COT_DVUT):
        # Dòng tên Hội (in đậm)
        tr = tbl.add_row()
        tr.cells[0].text = ""
        run = tr.cells[1].paragraphs[0].add_run(dvut.upper())
        run.bold = True
        run.font.size = Pt(9)
        for i in range(2, len(HEADERS)):
            tr.cells[i].text = ""

        for to, grp_to in grp_dvut.groupby(COT_TEN_TO):
            for _, row in grp_to.iterrows():
                stt += 1
                gn_val = ""
                if giai_ngan_input:
                    key = (dvut, to, row[COT_TEN_CT])
                    gn_val = fmt(giai_ngan_input.get(key, 0)) \
                             if giai_ngan_input.get(key, 0) > 0 else ""
                vals = [
                    str(stt), to, row[COT_TEN_CT],
                    fmt(row[COT_TONG_DU_NO]), gn_val,
                ]
                tr = tbl.add_row()
                for i, v in enumerate(vals):
                    tr.cells[i].text = v
                    tr.cells[i].paragraphs[0].alignment = (
                        WD_ALIGN_PARAGRAPH.LEFT if i in (1, 2)
                        else WD_ALIGN_PARAGRAPH.CENTER)
                    if tr.cells[i].paragraphs[0].runs:
                        tr.cells[i].paragraphs[0].runs[0].font.size = Pt(9)

    if len(tbl.rows) == 1:
        tr = tbl.add_row()
        tr.cells[1].text = "Không có món đến hạn tháng tới"


# ── Xuất file Word hoàn chỉnh ─────────────────────────────────────────────────

def xuat_bien_ban_giao_ban(
    df_xa: pd.DataFrame,
    df_baseline: pd.DataFrame | None,
    nam_moc: int,
    template_path: str,
    giai_ngan_input: dict | None = None,
) -> bytes:
    """
    Trả về bytes file .docx đã điền đầy đủ.
    Gọi từ ws_operation hoặc ws_management.
    """
    import io, copy
    from docx import Document as _Doc

    doc = _Doc(template_path)
    tag_values = tinh_so_lieu_van_xuoi(df_xa, df_baseline, nam_moc)

    # Điền tag văn xuôi trong toàn bộ paragraph
    for para in doc.paragraphs:
        for tag, val in tag_values.items():
            if tag in para.text:
                for run in para.runs:
                    if tag in run.text:
                        run.text = run.text.replace(tag, val)

    # Thay tag bảng động
    BANG_FUNCS = {
        "{{bang_dvut}}":         lambda p: tao_bang_dvut(doc, df_xa),
        "{{bang_chuong_trinh}}": lambda p: tao_bang_chuong_trinh(
                                     doc, df_xa, df_baseline, nam_moc),
        "{{bang_ke_hoach}}":     lambda p: tao_bang_ke_hoach(
                                     doc, df_xa, giai_ngan_input),
    }
    for para in doc.paragraphs:
        for tag, fn in BANG_FUNCS.items():
            if tag in para.text:
                # Chèn bảng ngay sau paragraph này
                fn(para)
                # Xóa paragraph chứa tag
                p = para._element
                p.getparent().remove(p)
                break

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
'''

def main() -> None:
    root = Path(__file__).resolve().parent
    out = root / "data" / "giao_ban.py"
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text(CONTENT, encoding="utf-8")
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
