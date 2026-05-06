"""Tab thống kê theo Đơn vị Ủy thác (Hội đoàn thể) và Tổ TK&VV.

Phân hệ: ws_management — Báo cáo tổng hợp theo Hội đoàn thể và Tổ TK&VV.
"""
from __future__ import annotations

import socket
from typing import TYPE_CHECKING, Any

import pandas as pd
import streamlit as st

from config import (
    COT_DVUT,
    COT_TEN_TO,
    COT_TEN_PGD,
    COT_TEN_XA,
    COT_TONG_DU_NO,
    COT_DU_NO_QH,
    COT_DU_NO_TH,
    COT_MA_KH,
    COT_SO_KU,
    DS_PGD,
    PGD_XA_MAP,
)

import db
from utils import fmt, fmt_so, fmt_ty

if TYPE_CHECKING:
    from streamlit.delta_generator import DeltaGenerator


import re as _re


def _norm_ten_xa(ten: str) -> str:
    """Bỏ prefix Phường/Xã/Thị trấn để so sánh."""
    return _re.sub(
        r"^(Phường|Xã|Thị trấn|TT|Thi tran)\s+",
        "",
        str(ten).strip(),
        flags=_re.IGNORECASE,
    ).strip().lower()


def _hostname() -> str:
    """Lấy hostname máy để log audit."""
    try:
        return socket.gethostname()
    except Exception:
        return "unknown"


def _get_df_hstd(kwargs: dict[str, Any]) -> pd.DataFrame:
    """Lấy DataFrame HSTD từ kwargs."""
    df = kwargs.get("df_full")
    if df is None or df.empty:
        df = kwargs.get("df")
    if df is None:
        return pd.DataFrame()
    return df if isinstance(df, pd.DataFrame) else pd.DataFrame()


def render(tab: DeltaGenerator, **kwargs: Any) -> None:
    """Render tab thống kê Hội đoàn thể và Tổ TK&VV."""
    role = kwargs.get("role", "user")
    username = kwargs.get("username", "unknown")
    df_h = _get_df_hstd(kwargs)
    hn = _hostname()

    with tab:
        st.subheader("📊 Thống kê theo Hội đoàn thể & Tổ TK&VV")
        st.caption("Báo cáo tổng hợp theo Đơn vị Ủy thác và Tổ Tiết kiệm & Vay vốn.")

        t_hdt, t_to, t_xuat = st.tabs(
            ["📊 Theo Hội đoàn thể", "🏘️ Theo Tổ TK&VV", "📄 Xuất Thông báo KL"]
        )

        with t_hdt:
            _render_theo_hoi_doan_the(df_h, username, hn)

        with t_to:
            _render_theo_to_tk_vv(df_h, username, hn)

        with t_xuat:
            _render_xuat_tb_klgb(df_h, username, hn)


def _render_theo_hoi_doan_the(df_h: pd.DataFrame, username: str, hn: str) -> None:
    """Sub-tab 1: Thống kê theo Hội đoàn thể (ĐVUT)."""
    st.markdown("### 📊 Thống kê theo Hội đoàn thể")

    # Bộ lọc
    col_pgd, col_xa, col_ky = st.columns([2, 2, 2])

    with col_pgd:
        ds_pgd = ["Tất cả"] + DS_PGD
        chon_pgd = st.selectbox("Chọn PGD", ds_pgd, key="hdt_loc_pgd")

    with col_xa:
        ds_xa = ["Tất cả"]
        if chon_pgd != "Tất cả":
            ds_xa = ds_xa + [str(x).strip() for x in (PGD_XA_MAP.get(chon_pgd, []) or [])]
        chon_xa = st.selectbox("Chọn Xã/Phường", ds_xa, key="hdt_loc_xa")

    with col_ky:
        chon_ky = st.selectbox("Chọn kỳ", ["Tháng", "Quý", "Năm"], key="hdt_loc_ky")

    st.divider()

    if df_h.empty:
        st.info("Không có dữ liệu HSTD để thống kê.")
        db.ghi_audit(username, "view_hoi_doan_the", f"[{hn}] Không có dữ liệu HSTD")
        return

    # Lọc dữ liệu vectorized
    df_filtered = df_h.copy()
    if chon_pgd != "Tất cả":
        mask_pgd = df_filtered[COT_TEN_PGD].astype(str).str.strip() == chon_pgd
        df_filtered = df_filtered[mask_pgd]
    if chon_xa != "Tất cả":
        mask_xa = df_filtered[COT_TEN_XA].astype(str).apply(_norm_ten_xa) == _norm_ten_xa(chon_xa)
        df_filtered = df_filtered[mask_xa]

    if df_filtered.empty:
        st.warning("Không có dữ liệu sau khi lọc.")
        return

    # Group by ĐVUT — lazy load với nút bấm
    @st.cache_data(ttl=300, show_spinner=False)
    def _tinh_tong_hop_hdt(df: pd.DataFrame, ky: str) -> pd.DataFrame:
        """Tính tổng hợp theo ĐVUT với vectorization."""
        col_map = {
            "Tháng": {
                "gn": "Giải ngân trong tháng",
                "thu_th": "Thu nợ TH tháng",
                "thu_qh": "Thu nợ QH tháng",
                "thu_khoanh": "Thu nợ khoanh tháng",
            },
            "Quý": {
                "gn": "Giải ngân trong Quý",
                "thu_th": "Thu nợ TH Quý",
                "thu_qh": "Thu nợ QH Quý",
                "thu_khoanh": "Thu nợ Khoanh Quý",
            },
            "Năm": {
                "gn": "Giải ngân Năm",
                "thu_th": "Thu nợ TH Năm",
                "thu_qh": "Thu nợ QH Năm",
                "thu_khoanh": "Thu nợ Khoanh Năm",
            },
        }
        cols = col_map.get(ky, col_map["Tháng"])

        # Đảm bảo các cột số tồn tại
        numeric_cols = [COT_TONG_DU_NO, COT_DU_NO_QH, COT_DU_NO_TH, "Dư nợ khoanh",
                       cols["gn"], cols["thu_th"], cols["thu_qh"], cols["thu_khoanh"]]
        for c in numeric_cols:
            if c not in df.columns:
                df[c] = 0.0
            else:
                df[c] = pd.to_numeric(df[c], errors="coerce").fillna(0)

        # Kiểm tra cột is_3m_inactive
        has_3m = "is_3m_inactive" in df.columns

        # Group by ĐVUT
        agg_dict = {
            COT_TEN_TO: "nunique",   # Số tổ
            COT_MA_KH: "nunique",    # Số KH
            COT_TONG_DU_NO: "sum",
            COT_DU_NO_QH: "sum",
            cols["gn"]: "sum",       # DS cho vay
            cols["thu_th"]: "sum",   # Thu nợ TH
            cols["thu_qh"]: "sum",    # Thu nợ QH
            cols["thu_khoanh"]: "sum", # Thu nợ khoanh
        }
        if has_3m:
            agg_dict["is_3m_inactive"] = lambda x: (x == True).sum()  # 3m KHĐ

        grouped = df.groupby(COT_DVUT, as_index=False).agg(agg_dict)

        # Đổi tên cột
        new_cols = ["Tên ĐVUT", "Số tổ", "Số KH", "Tổng dư nợ", "Dư nợ QH",
                    f"DS cho vay {ky}", "Thu nợ TH", "Thu nợ QH", "Thu nợ khoanh"]
        if has_3m:
            new_cols.append("3m KHĐ")
        grouped.columns = new_cols

        # Tính DS thu nợ = TH + QH + khoanh
        thu_cols = ["Thu nợ TH", "Thu nợ QH", "Thu nợ khoanh"]
        grouped[f"DS thu nợ {ky}"] = grouped[thu_cols].sum(axis=1)

        # Sắp xếp theo Tổng dư nợ giảm dần
        grouped = grouped.sort_values("Tổng dư nợ", ascending=False)
        return grouped

    # Lazy load — chỉ tính khi nhấn nút
    if not st.button("📊 Tải dữ liệu", type="primary", key="hdt_load"):
        st.info("👆 Nhấn nút **Tải dữ liệu** để xem thống kê.")
        return

    with st.spinner("Đang tính toán..."):
        df_summary = _tinh_tong_hop_hdt(df_filtered, chon_ky)

    if df_summary.empty:
        st.info("Không có dữ liệu ĐVUT để hiển thị.")
        return

    # Tính tổng cho metrics
    tong_dn = df_summary["Tổng dư nợ"].sum()
    tong_nqh = df_summary["Dư nợ QH"].sum()
    tong_dsv = df_summary[f"DS cho vay {chon_ky}"].sum()
    tl_nqh = (tong_nqh / tong_dn * 100) if tong_dn > 0 else 0

    # 3 metrics phía trên
    c1, c2, c3 = st.columns(3)
    c1.metric("Tổng dư nợ", fmt_ty(tong_dn))
    c2.metric("Tổng NQH", fmt_ty(tong_nqh), delta=f"{tl_nqh:.2f}%")
    c3.metric("DS cho vay kỳ này", fmt_ty(tong_dsv))

    st.divider()

    # Bảng chi tiết
    st.markdown("#### 📋 Chi tiết theo Đơn vị Ủy thác")

    # Thêm hàng Cộng
    df_with_total = df_summary.copy()
    total_row = {"Tên ĐVUT": "Cộng"}
    for c in df_with_total.columns:
        if c != "Tên ĐVUT":
            if c in ["Số tổ", "Số KH", "3m KHĐ"]:
                total_row[c] = df_with_total[c].sum()
            elif c in ["Tổng dư nợ", "Dư nợ QH", f"DS cho vay {chon_ky}", f"DS thu nợ {chon_ky}",
                       "Thu nợ TH", "Thu nợ QH", "Thu nợ khoanh"]:
                total_row[c] = df_with_total[c].sum()
            else:
                total_row[c] = ""
    df_with_total = pd.concat([df_with_total, pd.DataFrame([total_row])], ignore_index=True)

    # Format số tiền cho hiển thị
    df_display = df_with_total.copy()
    money_cols = [c for c in df_display.columns if any(x in c.lower() for x in ["dư nợ", "ds cho vay", "ds thu nợ", "thu nợ"])]
    for c in money_cols:
        df_display[c] = df_display[c].apply(lambda x: fmt_ty(x) if pd.notna(x) and isinstance(x, (int, float)) and x != "" else x)

    # Highlight hàng Cộng
    def _style_total(row):
        if row["Tên ĐVUT"] == "Cộng":
            return ["background-color: #f0f0f0; font-weight: bold"] * len(row)
        return [""] * len(row)

    st.dataframe(df_display.style.apply(_style_total, axis=1), use_container_width=True, height=500)

    # Export Excel
    if not df_summary.empty:
        buf = st.session_state.get("_hdt_export_buf")
        if buf is None:
            import io
            buf = io.BytesIO()
            with pd.ExcelWriter(buf, engine="openpyxl") as w:
                df_summary.to_excel(w, index=False, sheet_name="HoiDoanThe")
            st.session_state["_hdt_export_buf"] = buf

        st.download_button(
            "📥 Tải xuống Excel",
            data=st.session_state["_hdt_export_buf"].getvalue(),
            file_name=f"thong_ke_hdt_{chon_ky.lower()}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            key="hdt_export_btn",
        )

    db.ghi_audit(username, "view_hoi_doan_the", f"[{hn}] Xem TK HĐT: PGD={chon_pgd}, Xã={chon_xa}, Kỳ={chon_ky}")


def _render_theo_to_tk_vv(df_h: pd.DataFrame, username: str, hn: str) -> None:
    """Sub-tab 2: Thống kê theo Tổ TK&VV."""
    st.markdown("### �️ Thống kê theo Tổ TK&VV")

    # Bộ lọc cascade: PGD → Xã → Hội
    col_pgd, col_xa, col_hoi, col_ky = st.columns([2, 2, 2, 2])

    with col_pgd:
        ds_pgd = ["Tất cả"] + DS_PGD
        chon_pgd = st.selectbox("Chọn PGD", ds_pgd, key="to_loc_pgd")

    with col_xa:
        ds_xa = ["Tất cả"]
        if chon_pgd != "Tất cả":
            ds_xa = ds_xa + [str(x).strip() for x in (PGD_XA_MAP.get(chon_pgd, []) or [])]
        chon_xa = st.selectbox("Chọn Xã/Phường", ds_xa, key="to_loc_xa")

    with col_hoi:
        ds_hoi = ["Tất cả"]
        if not df_h.empty and COT_DVUT in df_h.columns:
            # Lọc theo PGD/Xã đã chọn để danh sách Hội phù hợp
            df_temp = df_h.copy()
            if chon_pgd != "Tất cả":
                df_temp = df_temp[df_temp[COT_TEN_PGD].astype(str).str.strip() == chon_pgd]
            if chon_xa != "Tất cả":
                df_temp = df_temp[df_temp[COT_TEN_XA].astype(str).apply(_norm_ten_xa) == _norm_ten_xa(chon_xa)]
            if not df_temp.empty:
                ds_hoi = ds_hoi + sorted(df_temp[COT_DVUT].dropna().astype(str).unique().tolist())
        chon_hoi = st.selectbox("Chọn Hội đoàn thể", ds_hoi, key="to_loc_hoi")

    with col_ky:
        chon_ky = st.selectbox("Chọn kỳ", ["Tháng", "Quý", "Năm"], key="to_loc_ky")

    st.divider()

    if df_h.empty:
        st.info("Không có dữ liệu HSTD để thống kê.")
        return

    # Lọc dữ liệu vectorized
    df_filtered = df_h.copy()
    if chon_pgd != "Tất cả":
        mask_pgd = df_filtered[COT_TEN_PGD].astype(str).str.strip() == chon_pgd
        df_filtered = df_filtered[mask_pgd]
    if chon_xa != "Tất cả":
        mask_xa = df_filtered[COT_TEN_XA].astype(str).apply(_norm_ten_xa) == _norm_ten_xa(chon_xa)
        df_filtered = df_filtered[mask_xa]
    if chon_hoi != "Tất cả":
        df_filtered = df_filtered[df_filtered[COT_DVUT].astype(str) == chon_hoi]

    if df_filtered.empty:
        st.warning("Không có dữ liệu sau khi lọc.")
        return

    @st.cache_data(ttl=300, show_spinner=False)
    def _tinh_tong_hop_to(df: pd.DataFrame, ky: str) -> pd.DataFrame:
        """Tính tổng hợp theo [COT_DVUT, COT_TEN_TO] với vectorization."""
        col_map = {
            "Tháng": {
                "gn": "Giải ngân trong tháng",
                "thu_th": "Thu nợ TH tháng",
                "thu_qh": "Thu nợ QH tháng",
                "thu_khoanh": "Thu nợ khoanh tháng",
            },
            "Quý": {
                "gn": "Giải ngân trong Quý",
                "thu_th": "Thu nợ TH Quý",
                "thu_qh": "Thu nợ QH Quý",
                "thu_khoanh": "Thu nợ Khoanh Quý",
            },
            "Năm": {
                "gn": "Giải ngân Năm",
                "thu_th": "Thu nợ TH Năm",
                "thu_qh": "Thu nợ QH Năm",
                "thu_khoanh": "Thu nợ Khoanh Năm",
            },
        }
        cols = col_map.get(ky, col_map["Tháng"])

        numeric_cols = [COT_TONG_DU_NO, COT_DU_NO_QH, COT_DU_NO_TH, "Dư nợ khoanh",
                       cols["gn"], cols["thu_th"], cols["thu_qh"], cols["thu_khoanh"]]
        for c in numeric_cols:
            if c not in df.columns:
                df[c] = 0.0
            else:
                df[c] = pd.to_numeric(df[c], errors="coerce").fillna(0)

        # Group theo [ĐVUT, Tên tổ]
        grouped = df.groupby([COT_DVUT, COT_TEN_TO], as_index=False).agg({
            COT_MA_KH: "nunique",
            COT_SO_KU: "nunique",
            COT_TONG_DU_NO: "sum",
            COT_DU_NO_QH: "sum",
            COT_DU_NO_TH: "sum",
            "Dư nợ khoanh": "sum",
            cols["gn"]: "sum",
            cols["thu_th"]: "sum",
            cols["thu_qh"]: "sum",
            cols["thu_khoanh"]: "sum",
        })

        grouped.columns = [
            "ĐVUT", "Tên tổ", "Số KH", "Số KU", "Tổng dư nợ", "Dư nợ QH",
            "Dư nợ TH", "Dư nợ khoanh", f"Giải ngân {ky}",
            f"Thu nợ TH {ky}", f"Thu nợ QH {ky}", f"Thu nợ khoanh {ky}"
        ]

        grouped = grouped.sort_values(["ĐVUT", "Tổng dư nợ"], ascending=[True, False])
        return grouped

    # Lazy load — chỉ tính khi nhấn nút
    if not st.button("📊 Tải dữ liệu", type="primary", key="to_load"):
        st.info("👆 Nhấn nút **Tải dữ liệu** để xem thống kê.")
        return

    with st.spinner("Đang tính toán..."):
        df_summary = _tinh_tong_hop_to(df_filtered, chon_ky)

    if df_summary.empty:
        st.info("Không có dữ liệu Tổ TK&VV để hiển thị.")
        return

    # Tính tổng cho metrics
    tong_dn = df_summary["Tổng dư nợ"].sum()
    tong_nqh = df_summary["Dư nợ QH"].sum()
    tong_dsv = df_summary[f"Giải ngân {chon_ky}"].sum()
    tl_nqh = (tong_nqh / tong_dn * 100) if tong_dn > 0 else 0

    # 3 metrics
    c1, c2, c3 = st.columns(3)
    c1.metric("Tổng dư nợ", fmt_ty(tong_dn))
    c2.metric("Tổng NQH", fmt_ty(tong_nqh), delta=f"{tl_nqh:.2f}%")
    c3.metric("DS cho vay kỳ này", fmt_ty(tong_dsv))

    st.divider()

    st.markdown("#### 📋 Chi tiết theo Tổ TK&VV")

    df_display = df_summary.copy()
    money_cols = [c for c in df_display.columns if any(x in c for x in ["dư nợ", "Giải ngân", "Thu nợ"])]
    for c in money_cols:
        df_display[c] = df_display[c].apply(lambda x: fmt(x) if pd.notna(x) else "")

    st.dataframe(df_display, use_container_width=True, height=500)

    if not df_summary.empty:
        import io
        buf = io.BytesIO()
        with pd.ExcelWriter(buf, engine="openpyxl") as w:
            df_summary.to_excel(w, index=False, sheet_name="ToTKVV")

        st.download_button(
            "📥 Tải xuống Excel",
            data=buf.getvalue(),
            file_name=f"thong_ke_to_tk_vv_{chon_ky.lower()}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            key="to_export_btn",
        )

    db.ghi_audit(username, "view_to_tk_vv", f"[{hn}] Xem TK Tổ: PGD={chon_pgd}, Xã={chon_xa}, Kỳ={chon_ky}")


def _render_chi_tiet_kh(df_h: pd.DataFrame, username: str, hn: str, role: str) -> None:
    """Sub-tab 3: Chi tiết khách hàng theo ĐVUT/Tổ."""
    st.markdown("### 📋 Chi tiết Khách hàng")

    col_pgd, col_xa, col_hdt, col_to = st.columns([2, 2, 2, 2])

    with col_pgd:
        ds_pgd = ["Tất cả"] + DS_PGD
        chon_pgd = st.selectbox("Chọn PGD", ds_pgd, key="ct_loc_pgd")

    with col_xa:
        ds_xa = ["Tất cả"]
        if chon_pgd != "Tất cả":
            ds_xa = ds_xa + [str(x).strip() for x in (PGD_XA_MAP.get(chon_pgd, []) or [])]
        chon_xa = st.selectbox("Chọn Xã/Phường", ds_xa, key="ct_loc_xa")

    with col_hdt:
        ds_hdt = ["Tất cả"]
        if not df_h.empty and COT_DVUT in df_h.columns:
            ds_hdt = ds_hdt + sorted(df_h[COT_DVUT].dropna().astype(str).unique().tolist())
        chon_hdt = st.selectbox("Chọn Hội đoàn thể", ds_hdt, key="ct_loc_hdt")

    with col_to:
        ds_to = ["Tất cả"]
        if not df_h.empty and COT_TEN_TO in df_h.columns:
            ds_to = ds_to + sorted(df_h[COT_TEN_TO].dropna().astype(str).unique().tolist())
        chon_to = st.selectbox("Chọn Tổ TK&VV", ds_to, key="ct_loc_to")

    st.divider()

    if df_h.empty:
        st.info("Không có dữ liệu HSTD.")
        return

    # Lọc vectorized
    df_filtered = df_h.copy()
    if chon_pgd != "Tất cả":
        df_filtered = df_filtered[df_filtered[COT_TEN_PGD].astype(str).str.strip() == chon_pgd]
    if chon_xa != "Tất cả":
        df_filtered = df_filtered[df_filtered[COT_TEN_XA].astype(str).apply(_norm_ten_xa) == _norm_ten_xa(chon_xa)]
    if chon_hdt != "Tất cả":
        df_filtered = df_filtered[df_filtered[COT_DVUT].astype(str) == chon_hdt]
    if chon_to != "Tất cả":
        df_filtered = df_filtered[df_filtered[COT_TEN_TO].astype(str) == chon_to]

    if df_filtered.empty:
        st.warning("Không có dữ liệu sau khi lọc.")
        return

    # Chọn cột hiển thị
    display_cols = [
        COT_MA_KH, "Tên khách hàng", COT_DVUT, COT_TEN_TO, COT_TEN_PGD, COT_TEN_XA,
        COT_SO_KU, COT_TONG_DU_NO, COT_DU_NO_QH, COT_DU_NO_TH, "Dư nợ khoanh"
    ]
    cols_available = [c for c in display_cols if c in df_filtered.columns]
    df_display = df_filtered[cols_available].copy()

    # Format số tiền
    money_cols = [c for c in df_display.columns if "dư nợ" in c.lower() or "nợ" in c.lower()]
    for c in money_cols:
        df_display[c] = pd.to_numeric(df_display[c], errors="coerce").apply(lambda x: fmt(x) if pd.notna(x) else "")

    # Masking thông tin nhạy cảm nếu không phải admin/manager
    if role not in ("admin", "manager"):
        if "Tên khách hàng" in df_display.columns:
            df_display["Tên khách hàng"] = df_display["Tên khách hàng"].astype(str).str[:-3] + "***"
        if COT_MA_KH in df_display.columns:
            df_display[COT_MA_KH] = df_display[COT_MA_KH].astype(str).str[:4] + "***" + df_display[COT_MA_KH].astype(str).str[-4:]

    st.markdown(f"**Tổng số KH:** {fmt_so(len(df_display))}")
    st.dataframe(df_display, use_container_width=True, height=550)

    db.ghi_audit(username, "view_chi_tiet_kh_hdt", f"[{hn}] Xem chi tiết: PGD={chon_pgd}, Xã={chon_xa}, HĐT={chon_hdt}, Tổ={chon_to}")


def _render_xuat_tb_klgb(df_h: pd.DataFrame, username: str, hn: str) -> None:
    """Sub-tab 3: Xuất Thông báo Kết luận Giao ban."""
    st.markdown("### 📄 Xuất Thông báo Kết luận Giao ban")

    # Form nhập thông tin
    col1, col2 = st.columns(2)

    with col1:
        ds_pgd = ["Chọn PGD..."] + DS_PGD
        chon_pgd = st.selectbox("Chọn PGD", ds_pgd, key="klgb_pgd")

        ds_xa_dgd = ["Chọn Xã/ĐGD..."]
        if chon_pgd != "Chọn PGD...":
            ds_xa = PGD_XA_MAP.get(chon_pgd, []) or []
            ds_xa_dgd = ds_xa_dgd + [str(x).strip() for x in ds_xa]
        chon_xa = st.selectbox("Chọn Xã/Điểm GD", ds_xa_dgd, key="klgb_xa")

    with col2:
        ngay_hop = st.text_input("Ngày họp (dd/mm/yyyy)", value="", key="klgb_ngay")
        thang_nam = st.text_input("Tháng/năm (VD: tháng 5/2026)", value="", key="klgb_thang")
        so_vb = st.text_input("Số văn bản", value="", key="klgb_sovb")

    col3, col4 = st.columns(2)

    with col3:
        nguoi_ky = st.text_input("Người ký", value="", key="klgb_nguoiky")
        chuc_danh = st.text_input("Chức danh", value="", key="klgb_chucdanh")

    with col4:
        chinhsach_moi = st.text_area("Chính sách mới trong tháng", height=80, key="klgb_cs")
        ton_tai = st.text_area("Tồn tại hạn chế", height=80, key="klgb_tt")
        nhiem_vu = st.text_area("Nhiệm vụ tháng tiếp theo", height=80, key="klgb_nv")

    st.divider()

    if df_h.empty:
        st.info("Không có dữ liệu HSTD để tính bảng II.")
        return

    # Preview bảng II (tự động tính từ df lọc theo xã)
    if chon_pgd != "Chọn PGD..." and chon_xa != "Chọn Xã/ĐGD...":
        st.markdown("#### 📊 Preview Bảng II — Tình hình hoạt động theo Hội đoàn thể")

        # Lọc theo PGD và xã
        df_filtered = df_h[
            (df_h[COT_TEN_PGD].astype(str).str.strip() == chon_pgd) &
            (df_h[COT_TEN_XA].astype(str).apply(_norm_ten_xa) == _norm_ten_xa(chon_xa))
        ].copy()

        if not df_filtered.empty:
            # Tính bảng II.1 theo Hội
            @st.cache_data(ttl=60, show_spinner=False)
            def _tinh_bang_ii(df: pd.DataFrame) -> pd.DataFrame:
                """Tính bảng II.1 theo Hội đoàn thể."""
                numeric_cols = [COT_TONG_DU_NO, COT_DU_NO_QH, "Giải ngân trong tháng", "Thu nợ TH tháng", "Thu nợ QH tháng", "Thu nợ khoanh tháng"]
                for c in numeric_cols:
                    if c not in df.columns:
                        df[c] = 0.0
                    else:
                        df[c] = pd.to_numeric(df[c], errors="coerce").fillna(0)

                grouped = df.groupby(COT_DVUT, as_index=False).agg({
                    COT_MA_KH: "nunique",
                    COT_TONG_DU_NO: "sum",
                    COT_DU_NO_QH: "sum",
                    "Giải ngân trong tháng": "sum",
                    "Thu nợ TH tháng": "sum",
                    "Thu nợ QH tháng": "sum",
                    "Thu nợ khoanh tháng": "sum",
                })

                grouped.columns = [
                    "Tên Hội đoàn thể", "Số KH", "Tổng dư nợ", "Dư nợ QH",
                    "Giải ngân", "Thu nợ TH", "Thu nợ QH", "Thu nợ khoanh"
                ]

                # Thêm hàng Cộng
                total_row = {"Tên Hội đoàn thể": "Cộng"}
                for c in grouped.columns[1:]:
                    total_row[c] = grouped[c].sum()
                grouped = pd.concat([grouped, pd.DataFrame([total_row])], ignore_index=True)

                return grouped.sort_values("Tổng dư nợ", ascending=False)

            with st.spinner("Đang tính bảng II..."):
                df_bang_ii = _tinh_bang_ii(df_filtered)

            # Format hiển thị
            df_display = df_bang_ii.copy()
            money_cols = [c for c in df_display.columns if any(x in c.lower() for x in ["dư nợ", "giải ngân", "thu nợ"])]
            for c in money_cols:
                df_display[c] = df_display[c].apply(lambda x: fmt(x) if pd.notna(x) else "")

            st.dataframe(df_display, use_container_width=True, height=300)
        else:
            st.warning("Không có dữ liệu cho PGD/Xã đã chọn.")

    st.divider()

    # Nút xuất
    col_btn1, col_btn2, _ = st.columns([1, 1, 3])

    with col_btn1:
        xuat_disabled = chon_pgd == "Chọn PGD..." or chon_xa == "Chọn Xã/ĐGD..."
        if st.button("📄 Xuất Word", type="primary", disabled=xuat_disabled, key="klgb_xuat_word"):
            try:
                # Lọc lại df
                df_filtered = df_h[
                    (df_h[COT_TEN_PGD].astype(str).str.strip() == chon_pgd) &
                    (df_h[COT_TEN_XA].astype(str).apply(_norm_ten_xa) == _norm_ten_xa(chon_xa))
                ].copy()

                # Tính bảng II
                df_bang_ii = _tinh_bang_ii(df_filtered) if not df_filtered.empty else pd.DataFrame()

                # Xuất Word
                buf = _xuat_word_klgb(
                    df_bang_ii=df_bang_ii,
                    so_vb=so_vb,
                    ngay_hop=ngay_hop,
                    thang_nam=thang_nam,
                    nguoi_ky=nguoi_ky,
                    chuc_danh=chuc_danh,
                    chinhsach_moi=chinhsach_moi,
                    ton_tai=ton_tai,
                    nhiem_vu=nhiem_vu,
                    ten_pgd=chon_pgd,
                    ten_xa=chon_xa,
                )

                st.download_button(
                    "📥 Tải file Word",
                    data=buf.getvalue(),
                    file_name=f"Thong_bao_KL_GB_{chon_pgd.replace(' ', '_')}_{thang_nam.replace('/', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="klgb_download_word",
                )

                db.ghi_audit(username, "xuat_tb_klgb_word", f"[{hn}] Xuất Word KLGB: PGD={chon_pgd}, Xã={chon_xa}")
                st.success("✅ Đã tạo file Word!")

            except Exception as e:
                st.error(f"❌ Lỗi xuất Word: {e}")
                db.ghi_audit(username, "xuat_tb_klgb_loi", f"[{hn}] Lỗi xuất Word: {e}")

    with col_btn2:
        st.caption("🖨️ Xuất PDF (sắp có)")


def _xuat_word_klgb(
    df_bang_ii: pd.DataFrame,
    so_vb: str,
    ngay_hop: str,
    thang_nam: str,
    nguoi_ky: str,
    chuc_danh: str,
    chinhsach_moi: str,
    ton_tai: str,
    nhiem_vu: str,
    ten_pgd: str,
    ten_xa: str,
) -> Any:
    """Sinh file Word Thông báo Kết luận Giao ban theo chuẩn NĐ 30.

    Returns:
        BytesIO: Buffer chứa file Word
    """
    from docx import Document
    from docx.shared import Cm, Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsmap
    from docx.oxml import OxmlElement
    import io

    doc = Document()

    # Cài đặt trang A4, lề
    sections = doc.sections
    for section in sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)

    def _set_font(run, font_name="Times New Roman", size=13, bold=False):
        run.font.name = font_name
        run.font.size = Pt(size)
        run.font.bold = bold
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    def _add_centered_para(text, bold=False, size=13, italic=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        _set_font(run, size=size, bold=bold)
        run.font.italic = italic
        return p

    def _add_justified_para(text, bold=False, size=13):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(text)
        _set_font(run, size=size, bold=bold)
        return p

    # ===== HEADER: Quốc hiệu + Tiêu ngữ (2 cột) =====
    table_header = doc.add_table(rows=1, cols=2)
    table_header.autofit = False
    table_header.columns[0].width = Inches(3.0)
    table_header.columns[1].width = Inches(3.0)

    cell_left = table_header.cell(0, 0)
    cell_left.text = "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM"
    cell_left.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(cell_left.paragraphs[0].runs[0], bold=True, size=13)

    p_left2 = cell_left.add_paragraph("Độc lập - Tự do - Hạnh phúc")
    p_left2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(p_left2.runs[0], bold=True, size=13)
    p_left2.paragraph_format.space_after = Pt(6)

    cell_right = table_header.cell(0, 1)
    cell_right.text = "NGÂN HÀNG CHÍNH SÁCH XÃ HỘI"
    cell_right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(cell_right.paragraphs[0].runs[0], bold=True, size=13)

    p_right2 = cell_right.add_paragraph(f"{ten_pgd}")
    p_right2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(p_right2.runs[0], bold=True, size=13)

    # Gạch dưới
    p_gach = cell_right.add_paragraph("—" * 15)
    p_gach.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_gach.paragraph_format.space_after = Pt(12)

    doc.add_paragraph()  # Khoảng cách

    # ===== Số văn bản + Địa danh ngày tháng =====
    p_so = doc.add_paragraph()
    p_so.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_so = p_so.add_run(f"Số: {so_vb}/TB-" if so_vb else "Số: …/TB-")
    _set_font(run_so, size=13)

    p_ngay = doc.add_paragraph()
    p_ngay.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_ngay = p_ngay.add_run(f"{ten_xa}, ngày {ngay_hop}" if ngay_hop else f"{ten_xa}, ngày … tháng … năm …")
    _set_font(run_ngay, size=13, italic=True)
    p_ngay.paragraph_format.space_after = Pt(12)

    # ===== TIÊU ĐỀ =====
    _add_centered_para("THÔNG BÁO", bold=True, size=14)
    _add_centered_para(f"Kết luận cuộc họp giao ban {thang_nam}" if thang_nam else "Kết luận cuộc họp giao ban", bold=True, size=14)
    p_tieude = doc.add_paragraph()
    p_tieude.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_tieude = p_tieude.add_run("—" * 20)
    _set_font(run_tieude, size=13)
    p_tieude.paragraph_format.space_after = Pt(12)

    # ===== NỘI DUNG =====
    _add_centered_para("I. CHÍNH SÁCH MỚI TRONG THÁNG", bold=True, size=13)
    if chinhsach_moi:
        _add_justified_para(chinhsach_moi)
    else:
        _add_justified_para("(Không)")

    doc.add_paragraph()

    # ===== II. TÌNH HÌNH HOẠT ĐỘNG =====
    _add_centered_para("II. TÌNH HÌNH HOẠT ĐỘNG", bold=True, size=13)
    _add_justified_para("1. Kết quả hoạt động theo Hội đoàn thể:", bold=True)

    # Bảng II.1
    if not df_bang_ii.empty:
        table = doc.add_table(rows=len(df_bang_ii) + 1, cols=8)
        table.style = 'Table Grid'

        # Header
        headers = ["TT", "Tên Hội đoàn thể", "Số KH", "Tổng dư nợ", "Dư nợ QH", "Giải ngân", "Thu nợ", "Tỷ lệ NQH"]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            cell.paragraphs[0].runs[0].font.bold = True
            _set_font(cell.paragraphs[0].runs[0], bold=True, size=12)
            # Nền xanh nhạt
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'D9E2F3')
            cell._tc.get_or_add_tcPr().append(shading)

        # Data
        for idx, row in enumerate(df_bang_ii.itertuples(), 1):
            is_total = row[1] == "Cộng"
            table_row = table.rows[idx]

            # TT
            table_row.cells[0].text = "" if is_total else str(idx)
            # Tên Hội
            table_row.cells[1].text = str(row[1])
            # Số KH
            table_row.cells[2].text = fmt_so(int(row[2])) if pd.notna(row[2]) else "0"
            # Tổng dư nợ
            table_row.cells[3].text = fmt(row[3]) if pd.notna(row[3]) else "0"
            # Dư nợ QH
            table_row.cells[4].text = fmt(row[4]) if pd.notna(row[4]) else "0"
            # Giải ngân
            table_row.cells[5].text = fmt(row[5]) if pd.notna(row[5]) else "0"
            # Thu nợ (tổng TH+QH+khoanh)
            thu_noi = (row[6] if pd.notna(row[6]) else 0) + (row[7] if pd.notna(row[7]) else 0) + (row[8] if pd.notna(row[8]) else 0)
            table_row.cells[6].text = fmt(thu_noi)
            # Tỷ lệ NQH
            ty_le = (row[4] / row[3] * 100) if pd.notna(row[3]) and row[3] > 0 else 0
            table_row.cells[7].text = f"{ty_le:.2f}%"

            # Format hàng Cộng
            if is_total:
                for cell in table_row.cells:
                    for run in cell.paragraphs[0].runs:
                        run.font.bold = True

        # Căn chỉnh bảng
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = 1  # Center
    else:
        _add_justified_para("Không có dữ liệu.")

    doc.add_paragraph()
    _add_justified_para("2. Tồn tại hạn chế:", bold=True)
    if ton_tai:
        _add_justified_para(ton_tai)
    else:
        _add_justified_para("(Không)")

    doc.add_paragraph()

    # ===== III. NHIỆM VỤ THÁNG TIẾP THEO =====
    _add_centered_para("III. NHIỆM VỤ THÁNG TIẾP THEO", bold=True, size=13)

    _add_justified_para("1. Kế hoạch thu nợ và giải ngân:", bold=True)

    # Bảng III.1 (để trống — user tự điền)
    table_iii = doc.add_table(rows=3, cols=4)
    table_iii.style = 'Table Grid'

    headers_iii = ["Chỉ tiêu", "Kế hoạch tháng", "Thực hiện", "Tỷ lệ (%)"]
    for i, h in enumerate(headers_iii):
        cell = table_iii.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        _set_font(cell.paragraphs[0].runs[0], bold=True, size=12)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D9E2F3')
        cell._tc.get_or_add_tcPr().append(shading)

    table_iii.rows[1].cells[0].text = "Thu nợ"
    table_iii.rows[2].cells[0].text = "Giải ngân"

    doc.add_paragraph()

    if nhiem_vu:
        _add_justified_para(nhiem_vu)
    else:
        _add_justified_para("(Không)")

    doc.add_paragraph()
    doc.add_paragraph()

    # ===== PHẦN KÝ TÊN =====
    table_ky = doc.add_table(rows=2, cols=2)
    table_ky.autofit = False

    # Nơi nhận
    cell_nhan = table_ky.cell(0, 0)
    cell_nhan.text = "Nơi nhận:\n- Như trên;\n- Lưu: VT."
    for run in cell_nhan.paragraphs[0].runs:
        _set_font(run, size=12, italic=True)

    # Người ký
    cell_ky = table_ky.cell(0, 1)
    cell_ky.text = f"{chuc_danh if chuc_danh else 'Giám đốc'}\n\n\n\n{nguoi_ky if nguoi_ky else '(Ký, đóng dấu)'}"
    cell_ky.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cell_ky.paragraphs[0].runs:
        _set_font(run, size=13, bold=True)
    cell_ky.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Xuất buffer
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
