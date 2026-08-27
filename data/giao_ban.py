"""
Tính số liệu giao ban xã và tạo 3 bảng động cho mẫu Word.
Dùng chung cho ws_operation (CBTD) và ws_management (PGD).
"""
from __future__ import annotations

import pandas as pd
from datetime import date
from io import BytesIO
from dateutil.relativedelta import relativedelta
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from config import (
    COT_MA_KH,
    COT_TONG_DU_NO,
    COT_DU_NO_QH,
    COT_TEN_CT,
    COT_TEN_PGD,
    COT_TEN_XA,
    COT_LAI_TON,
    COT_LAI_TON_QH,
)
from utils import fmt

COT_DU_NO_KHOANH = "Dư nợ khoanh"
COT_NGAY_DH_GH = "Ngày ĐH theo Gia hạn"
COT_DVUT = "Tên ĐVUT"
COT_TEN_TO = "Tên tổ"
COT_DS_CV_THANG = "Giải ngân trong tháng"
COT_TN_TH_THANG = "Thu nợ TH tháng"
COT_TN_QH_THANG = "Thu nợ QH tháng"
COT_TIEN_GUI = "Số dư tiền gửi 105"
COT_MON_3M = "is_3m_inactive"  # từ danh_dau_khong_hd()


def loc_theo_xa(df: pd.DataFrame, ten_xa: str) -> pd.DataFrame:
    return df[df[COT_TEN_XA] == ten_xa].copy()


def loc_baseline_cung_xa_pgd(
    df_baseline: pd.DataFrame | None,
    df_xa: pd.DataFrame,
    ten_xa: str | None = None,
) -> pd.DataFrame | None:
    """Lọc baseline cùng xã và, khi xác định được, cùng đúng một PGD với dữ liệu hiện tại."""
    if df_baseline is None or df_baseline.empty:
        return df_baseline
    df_bl = df_baseline.copy()

    xa = str(ten_xa or "").strip()
    if not xa and COT_TEN_XA in df_xa.columns and not df_xa.empty:
        xa = str(df_xa[COT_TEN_XA].iloc[0] or "").strip()
    if xa and COT_TEN_XA in df_bl.columns:
        df_bl = df_bl[df_bl[COT_TEN_XA] == xa].copy()

    if COT_TEN_PGD in df_xa.columns and COT_TEN_PGD in df_bl.columns:
        pgd_vals = (
            df_xa[COT_TEN_PGD]
            .dropna()
            .astype(str)
            .str.strip()
        )
        pgd_unique = [v for v in pgd_vals.unique().tolist() if v]
        if len(pgd_unique) == 1:
            pgd = pgd_unique[0]
            df_bl = df_bl[df_bl[COT_TEN_PGD].astype(str).str.strip() == pgd].copy()

    return df_bl


def tinh_so_lieu_van_xuoi(
    df_xa: pd.DataFrame,
    df_baseline: pd.DataFrame | None,
    nam_moc: int,
) -> dict:
    """
    Trả về dict các tag văn xuôi cho mẫu Word.
    df_xa       : HSTD đã lọc theo xã hiện tại
    df_baseline : HSTD mốc 31/12 (toàn CN, chưa lọc) — None nếu chưa có
    nam_moc     : năm của baseline (VD: 2025)
    """
    _z = pd.Series(0.0, index=df_xa.index)
    dn    = pd.to_numeric(df_xa[COT_TONG_DU_NO], errors="coerce").fillna(0)
    qh    = pd.to_numeric(df_xa[COT_DU_NO_QH], errors="coerce").fillna(0)
    kh    = pd.to_numeric(df_xa[COT_DU_NO_KHOANH], errors="coerce").fillna(0) if COT_DU_NO_KHOANH in df_xa.columns else _z
    tg    = pd.to_numeric(df_xa[COT_TIEN_GUI], errors="coerce").fillna(0) if COT_TIEN_GUI in df_xa.columns else _z
    gn    = pd.to_numeric(df_xa[COT_DS_CV_THANG], errors="coerce").fillna(0) if COT_DS_CV_THANG in df_xa.columns else _z
    tn_th = pd.to_numeric(df_xa[COT_TN_TH_THANG], errors="coerce").fillna(0) if COT_TN_TH_THANG in df_xa.columns else _z
    tn_qh = pd.to_numeric(df_xa[COT_TN_QH_THANG], errors="coerce").fillna(0) if COT_TN_QH_THANG in df_xa.columns else _z

    tong_dn = dn.sum()
    tong_qh = qh.sum()
    tong_kh = kh.sum()
    tong_tg = tg.sum()
    tl_qh = tong_qh / tong_dn * 100 if tong_dn > 0 else 0.0
    tl_kh = tong_kh / tong_dn * 100 if tong_dn > 0 else 0.0
    so_kh = df_xa[df_xa[COT_TONG_DU_NO] > 0][COT_MA_KH].nunique()
    so_to = df_xa[COT_TEN_TO].nunique() if COT_TEN_TO in df_xa.columns else 0

    dn_thang_truoc = tong_dn - gn.sum() + tn_th.sum() + tn_qh.sum()
    chenh_lech_thang = tong_dn - dn_thang_truoc
    tang_giam_thang = "tăng" if chenh_lech_thang >= 0 else "giảm"

    chenh_lech_dn = pct_dau_nam = 0.0
    tang_giam_dau_nam = "tăng"
    if df_baseline is not None and COT_TEN_XA in df_baseline.columns and not df_xa.empty and COT_TEN_XA in df_xa.columns:
        df_bl_xa = loc_baseline_cung_xa_pgd(df_baseline, df_xa)
        dn_bl = pd.to_numeric(df_bl_xa[COT_TONG_DU_NO], errors="coerce").fillna(0).sum()
        chenh_lech_dn = tong_dn - dn_bl
        pct_dau_nam = chenh_lech_dn / dn_bl * 100 if dn_bl > 0 else 0.0
        tang_giam_dau_nam = "tăng" if chenh_lech_dn >= 0 else "giảm"

    return {
        "{{tong_du_no}}": fmt(tong_dn),
        "{{so_kh}}": str(so_kh),
        "{{so_to}}": str(so_to),
        "{{du_no_qh}}": fmt(tong_qh),
        "{{ty_le_nqh}}": f"{tl_qh:.2f}",
        "{{du_no_khoanh}}": fmt(tong_kh),
        "{{ty_le_khoanh}}": f"{tl_kh:.2f}",
        "{{tien_gui_105}}": fmt(tong_tg),
        "{{tang_giam_thang}}": tang_giam_thang,
        "{{chenh_lech_thang}}": fmt(abs(chenh_lech_thang)),
        "{{tang_giam_dau_nam}}": tang_giam_dau_nam,
        "{{chenh_lech_dau_nam}}": fmt(abs(chenh_lech_dn)),
        "{{pct_dau_nam}}": f"{abs(pct_dau_nam):.1f}",
        "{{nam_moc}}": str(nam_moc),
    }


def tao_bang_dvut(doc: Document, df_xa: pd.DataFrame) -> None:
    """Chèn bảng kết quả theo ĐVUT vào doc tại vị trí placeholder."""
    from data import danh_dau_khong_hd

    df_m = danh_dau_khong_hd(df_xa)

    DVUT_ORDER = [
        "Hội nông dân",
        "Hội liên hiệp phụ nữ",
        "Hội cựu chiến binh",
        "Đoàn thanh niên",
    ]

    t = df_m.groupby(COT_DVUT).agg(
        so_to=(COT_TEN_TO, "nunique"),
        so_kh=(COT_MA_KH, "nunique"),
        ds_cv=(COT_DS_CV_THANG, "sum"),
        ds_tn=(COT_TN_TH_THANG, "sum"),
        tong_dn=(COT_TONG_DU_NO, "sum"),
        nqh=(COT_DU_NO_QH, "sum"),
        no_kh=(COT_DU_NO_KHOANH, "sum"),
        mon_3m=(COT_MON_3M, "sum"),
    ).reset_index()

    t = t[t["tong_dn"] > 0]
    t["_ord"] = t[COT_DVUT].apply(
        lambda x: DVUT_ORDER.index(x) if x in DVUT_ORDER else 99
    )
    t = t.sort_values("_ord").drop(columns="_ord")

    tong_dn_all = t["tong_dn"].sum()

    HEADERS = [
        "Stt",
        "Đơn vị nhận ủy thác",
        "Số Tổ",
        "Số KH",
        "DS cho vay",
        "DS thu nợ",
        "Tổng dư nợ",
        "Tỷ trọng %",
        "NQH",
        "Nợ khoanh",
        "Món 3T KHĐ",
    ]

    tbl = doc.add_table(rows=1, cols=len(HEADERS))
    tbl.style = "Table Grid"
    for i, h in enumerate(HEADERS):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(9)

    from utils import fmt

    for idx, row in enumerate(t.itertuples(), 1):
        ty_trong = row.tong_dn / tong_dn_all * 100 if tong_dn_all > 0 else 0
        vals = [
            str(idx),
            row[1],
            str(int(row.so_to)),
            str(int(row.so_kh)),
            fmt(row.ds_cv),
            fmt(row.ds_tn),
            fmt(row.tong_dn),
            f"{ty_trong:.1f}%",
            fmt(row.nqh),
            fmt(row.no_kh),
            str(int(row.mon_3m)),
        ]
        tr = tbl.add_row()
        for i, v in enumerate(vals):
            tr.cells[i].text = v
            tr.cells[i].paragraphs[0].alignment = (
                WD_ALIGN_PARAGRAPH.CENTER if i != 1 else WD_ALIGN_PARAGRAPH.LEFT
            )
            tr.cells[i].paragraphs[0].runs[0].font.size = Pt(9)

    tr = tbl.add_row()
    cong_vals = [
        "",
        "Cộng",
        str(int(t["so_to"].sum())),
        str(int(t["so_kh"].sum())),
        fmt(t["ds_cv"].sum()),
        fmt(t["ds_tn"].sum()),
        fmt(tong_dn_all),
        "100%",
        fmt(t["nqh"].sum()),
        fmt(t["no_kh"].sum()),
        str(int(t["mon_3m"].sum())),
    ]
    for i, v in enumerate(cong_vals):
        tr.cells[i].text = v
        para = tr.cells[i].paragraphs[0]
        run = para.runs[0] if para.runs else para.add_run(v)
        run.bold = True
        run.font.size = Pt(9)


def _tao_bang_chi_tiet_to(doc: Document, df_xa: pd.DataFrame) -> None:
    """Bảng 2 — Chi tiết thu nợ / giải ngân theo Tổ TK&VV (5 cột)."""
    DVUT_ORDER = [
        "Hội nông dân",
        "Hội liên hiệp phụ nữ",
        "Hội cựu chiến binh",
        "Đoàn thanh niên",
    ]

    HEADERS = [
        "Stt",
        "Đơn vị nhận ủy thác / Tổ TK&VV",
        "Chương trình cho vay",
        "Số tiền thu nợ",
        "Số tiền giải ngân",
    ]

    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Table Grid"
    for i, h in enumerate(HEADERS):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)

    co_ten_to = COT_TEN_TO in df_xa.columns
    dvut_list = [d for d in DVUT_ORDER if d in df_xa[COT_DVUT].dropna().unique()]
    stt = 0

    for dvut in dvut_list:
        df_dvut = df_xa[df_xa[COT_DVUT] == dvut]

        tr = tbl.add_row()
        merged = tr.cells[0].merge(tr.cells[4])
        merged.text = ""
        p = merged.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(dvut.upper())
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(11)

        tc_pr = merged._tc.get_or_add_tcPr()
        shading_elm = OxmlElement("w:shd")
        shading_elm.set(qn("w:fill"), "DEEAF1")
        shading_elm.set(qn("w:val"), "clear")
        tc_pr.append(shading_elm)

        if co_ten_to:
            ds_to = sorted(df_dvut[COT_TEN_TO].dropna().unique().tolist())
            for to in ds_to:
                stt += 1
                tr = tbl.add_row()
                vals = [str(stt), to, "", "", ""]
                for i, v in enumerate(vals):
                    tr.cells[i].text = v
                    tr.cells[i].paragraphs[0].alignment = (
                        WD_ALIGN_PARAGRAPH.LEFT if i == 1 else WD_ALIGN_PARAGRAPH.CENTER
                    )

    tr = tbl.add_row()
    for i in range(5):
        tr.cells[i].text = "Tổng cộng" if i == 1 else ""
        para = tr.cells[i].paragraphs[0]
        run = para.runs[0] if para.runs else para.add_run(tr.cells[i].text)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)

    _patch_font(tbl, size=11)


def tao_bang_chuong_trinh(
    doc: Document,
    df_xa: pd.DataFrame,
    df_baseline: pd.DataFrame | None,
    nam_moc: int,
) -> None:
    if df_baseline is not None and COT_TEN_XA in df_baseline.columns:
        ten_xa = df_xa[COT_TEN_XA].iloc[0]
        df_bl = loc_baseline_cung_xa_pgd(df_baseline, df_xa, ten_xa)
    else:
        df_bl = None

    g_ht = df_xa.groupby(COT_TEN_CT)[COT_TONG_DU_NO].sum().reset_index()
    g_ht.columns = [COT_TEN_CT, "dn_ht"]

    if df_bl is not None and not df_bl.empty:
        g_bl = df_bl.groupby(COT_TEN_CT)[COT_TONG_DU_NO].sum().reset_index()
        g_bl.columns = [COT_TEN_CT, "dn_bl"]
        mg = g_ht.merge(g_bl, on=COT_TEN_CT, how="left").fillna(0)
    else:
        mg = g_ht.copy()
        mg["dn_bl"] = 0

    mg = mg[mg["dn_ht"] > 0].copy()
    mg["chenh_lech"] = mg["dn_ht"] - mg["dn_bl"]
    mg["pct"] = mg.apply(
        lambda r: r["chenh_lech"] / r["dn_bl"] * 100 if r["dn_bl"] > 0 else 0,
        axis=1,
    )

    HEADERS = [
        "Stt",
        "Chương trình tín dụng",
        f"Dư nợ 31/12/{nam_moc}",
        "Dư nợ hiện tại",
        "Tăng/giảm",
        "%",
    ]

    tbl = doc.add_table(rows=1, cols=len(HEADERS))
    tbl.style = "Table Grid"
    for i, h in enumerate(HEADERS):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(9)

    from utils import fmt

    for idx, row in enumerate(mg.itertuples(), 1):
        dau = "+" if row.chenh_lech >= 0 else ""
        vals = [
            str(idx),
            row[1],
            fmt(row.dn_bl),
            fmt(row.dn_ht),
            f"{dau}{fmt(row.chenh_lech)}",
            f"{dau}{row.pct:.1f}%",
        ]
        tr = tbl.add_row()
        for i, v in enumerate(vals):
            tr.cells[i].text = v
            tr.cells[i].paragraphs[0].alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if i == 1 else WD_ALIGN_PARAGRAPH.CENTER
            )
            tr.cells[i].paragraphs[0].runs[0].font.size = Pt(9)

    tr = tbl.add_row()
    tong_bl = mg["dn_bl"].sum()
    tong_ht = mg["dn_ht"].sum()
    tong_cl = tong_ht - tong_bl
    dau = "+" if tong_cl >= 0 else ""
    pct_tong = tong_cl / tong_bl * 100 if tong_bl > 0 else 0
    cong = [
        "",
        "Cộng",
        fmt(tong_bl),
        fmt(tong_ht),
        f"{dau}{fmt(tong_cl)}",
        f"{dau}{pct_tong:.1f}%",
    ]
    for i, v in enumerate(cong):
        tr.cells[i].text = v
        para = tr.cells[i].paragraphs[0]
        run = para.runs[0] if para.runs else para.add_run(v)
        run.bold = True
        run.font.size = Pt(9)


def tao_bang_ke_hoach(
    doc: Document,
    df_xa: pd.DataFrame,
    giai_ngan_input: dict | None = None,
) -> None:
    """
    df_xa            : HSTD lọc theo xã hiện tại
    giai_ngan_input  : dict {(ten_dvut, ten_to, ten_ct): so_tien} — CBTD nhập tay
                       None = để trống cột giải ngân
    """
    thang_toi = date.today() + relativedelta(months=1)
    thang_toi_dau = date(thang_toi.year, thang_toi.month, 1)

    ngay_dh = pd.to_datetime(df_xa[COT_NGAY_DH_GH], errors="coerce")
    mask = (ngay_dh.dt.year == thang_toi_dau.year) & (
        ngay_dh.dt.month == thang_toi_dau.month
    )
    df_dh = df_xa[mask].copy()

    g = (
        df_dh.groupby([COT_DVUT, COT_TEN_TO, COT_TEN_CT])[COT_TONG_DU_NO]
        .sum()
        .reset_index()
    )
    g = g[g[COT_TONG_DU_NO] > 0]

    HEADERS = [
        "Stt",
        "Đơn vị nhận ủy thác / Tổ TK&VV",
        "Chương trình cho vay",
        "Số tiền thu nợ",
        "Số tiền giải ngân",
    ]

    tbl = doc.add_table(rows=1, cols=len(HEADERS))
    tbl.style = "Table Grid"
    for i, h in enumerate(HEADERS):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(9)

    from utils import fmt

    stt = 0
    for dvut, grp_dvut in g.groupby(COT_DVUT):
        tr = tbl.add_row()
        tr.cells[0].text = ""
        run = tr.cells[1].paragraphs[0].add_run(dvut.upper())
        run.bold = True
        run.font.size = Pt(9)
        for i in range(2, len(HEADERS)):
            tr.cells[i].text = ""

        for to, grp_to in grp_dvut.groupby(COT_TEN_TO):
            for _, row in grp_to.iterrows():
                stt += 1
                gn_val = ""
                if giai_ngan_input:
                    key = (dvut, to, row[COT_TEN_CT])
                    gn_val = (
                        fmt(giai_ngan_input.get(key, 0))
                        if giai_ngan_input.get(key, 0) > 0
                        else ""
                    )
                vals = [
                    str(stt),
                    to,
                    row[COT_TEN_CT],
                    fmt(row[COT_TONG_DU_NO]),
                    gn_val,
                ]
                tr = tbl.add_row()
                for i, v in enumerate(vals):
                    tr.cells[i].text = v
                    tr.cells[i].paragraphs[0].alignment = (
                        WD_ALIGN_PARAGRAPH.LEFT if i in (1, 2) else WD_ALIGN_PARAGRAPH.CENTER
                    )
                    if tr.cells[i].paragraphs[0].runs:
                        tr.cells[i].paragraphs[0].runs[0].font.size = Pt(9)

    if len(tbl.rows) == 1:
        tr = tbl.add_row()
        tr.cells[1].text = "Không có món đến hạn tháng tới"


def xuat_thong_bao_ket_luan_giao_ban(
    df_xa: pd.DataFrame,
    ten_pgd: str,
    ten_xa: str,
    ten_dgd: str,
    thang_bao_cao: int,
    nam_bao_cao: int,
    ngay_hop: str,
    chinh_sach_moi: str,
    ton_tai_han_che: str,
    nhiem_vu_tiep: str,
    so_van_ban: str = "",
    ten_nguoi_ky: str = "",
    giai_ngan_input: dict | None = None,
    df_baseline: pd.DataFrame | None = None,
    nam_moc: int = 2025,
) -> bytes:
    """
    Xuất Thông báo Kết luận Giao ban chuẩn NĐ30/2020 (python-docx, không template).
    Tái sử dụng: tinh_so_lieu_van_xuoi(), tao_bang_dvut(), tao_bang_ke_hoach().
    Trả về bytes .docx.
    """
    doc = Document()

    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.different_first_page_header_footer = True

    def _ten_pgd_hien_thi(s: str) -> str:
        t = (s or "").strip()
        if t.upper().startswith("PGD"):
            t = t[3:].strip().lstrip("-").strip()
        return t.upper() or "…………"

    ten_pgd_hd = _ten_pgd_hien_thi(ten_pgd)

    def _p(
        text: str = "",
        bold: bool = False,
        italic: bool = False,
        size: int = 14,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        space_before: float = 0,
        space_after: float = 6,
        indent_cm: float = 0,
        line_spacing: float = 1.25,
    ):
        para = doc.add_paragraph()
        para.alignment = align
        fmt_pf = para.paragraph_format
        fmt_pf.space_before = Pt(space_before)
        fmt_pf.space_after = Pt(space_after)
        if indent_cm and indent_cm > 0:
            fmt_pf.first_line_indent = Cm(indent_cm)
        fmt_pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        fmt_pf.line_spacing = line_spacing
        if text:
            r = para.add_run(text)
            r.bold = bold
            r.italic = italic
            r.font.name = "Times New Roman"
            r.font.size = Pt(size)
            r.font.color.rgb = RGBColor(0, 0, 0)
        return para

    def _heading(so: str, text: str) -> None:
        _p(
            f"{so}. {text}",
            bold=True,
            size=14,
            align=WD_ALIGN_PARAGRAPH.LEFT,
            space_before=6,
            space_after=3,
            indent_cm=0,
        )

    def _underline_para(para) -> None:
        p_pr = para._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"), "6")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), "000000")
        p_bdr.append(bot)
        p_pr.append(p_bdr)

    def _no_border(table) -> None:
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tc_pr = tc.get_or_add_tcPr()
                tc_borders = OxmlElement("w:tcBorders")
                for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                    b = OxmlElement(f"w:{side}")
                    b.set(qn("w:val"), "none")
                    tc_borders.append(b)
                tc_pr.append(tc_borders)

    def _cell_w(cell, cm: float) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_w = OxmlElement("w:tcW")
        tc_w.set(qn("w:w"), str(int(cm * 567)))
        tc_w.set(qn("w:type"), "dxa")
        tc_pr.append(tc_w)

    def _patch_font(table, size: int = 11) -> None:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(size)
                        run.font.color.rgb = RGBColor(0, 0, 0)

    def _fmt_ngay(s: str) -> tuple[str, str, str]:
        parts = (s or "").strip().split("/")
        if len(parts) != 3:
            raise ValueError("Ngày họp cần định dạng DD/MM/YYYY")
        d, m, y = parts[0].strip(), parts[1].strip(), parts[2].strip()
        return f"{int(d):02d}", f"{int(m):02d}", y

    ngay_s, thang_s, nam_s = _fmt_ngay(ngay_hop)
    thang_fmt = f"{thang_bao_cao:02d}"

    header = doc.sections[0].header
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_pg = hp.add_run()
    run_pg.font.size = Pt(13)
    run_pg.font.name = "Times New Roman"
    fc1 = OxmlElement("w:fldChar")
    fc1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = " PAGE "
    fc2 = OxmlElement("w:fldChar")
    fc2.set(qn("w:fldCharType"), "end")
    run_pg._r.extend([fc1, it, fc2])

    fh = doc.sections[0].first_page_header
    for fp in fh.paragraphs:
        fp.text = ""

    tbl_h = doc.add_table(rows=1, cols=2)
    _no_border(tbl_h)
    cl, cr = tbl_h.rows[0].cells
    _cell_w(cl, 8.5)
    _cell_w(cr, 8.5)

    p = cl.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("NGÂN HÀNG CHÍNH SÁCH XÃ HỘI")
    r.bold = False
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"
    r.font.color.rgb = RGBColor(0, 0, 0)

    p2 = cl.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"CHI NHÁNH ĐỒNG NAI - PGD {ten_pgd_hd}")
    r2.bold = True
    r2.font.size = Pt(12)
    r2.font.name = "Times New Roman"
    r2.font.color.rgb = RGBColor(0, 0, 0)
    _underline_para(p2)

    p3 = cl.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    so_vb_hien_thi = so_van_ban.strip() if so_van_ban.strip() else "      "
    r3 = p3.add_run(f"Số: {so_vb_hien_thi}/TB-KLGB")
    r3.font.size = Pt(12)
    r3.font.name = "Times New Roman"
    r3.font.color.rgb = RGBColor(0, 0, 0)

    p4 = cr.paragraphs[0]
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    r4.bold = True
    r4.font.size = Pt(13)
    r4.font.name = "Times New Roman"
    r4.font.color.rgb = RGBColor(0, 0, 0)

    p5 = cr.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r5 = p5.add_run("Độc lập - Tự do - Hạnh phúc")
    r5.bold = True
    r5.font.size = Pt(14)
    r5.font.name = "Times New Roman"
    r5.font.color.rgb = RGBColor(0, 0, 0)
    _underline_para(p5)

    p6 = cr.add_paragraph()
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r6 = p6.add_run(f"Đồng Nai, ngày {ngay_s} tháng {thang_s} năm {nam_s}")
    r6.italic = True
    r6.font.size = Pt(13)
    r6.font.name = "Times New Roman"
    r6.font.color.rgb = RGBColor(0, 0, 0)

    _p()

    _p(
        "THÔNG BÁO",
        bold=True,
        size=14,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=0,
        indent_cm=0,
    )
    p_ty = _p(
        f"Kết luận cuộc họp giao ban tháng {thang_fmt}/{nam_bao_cao}",
        bold=True,
        size=14,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=0,
        indent_cm=0,
    )
    _underline_para(p_ty)

    _p(
        f"Tại điểm giao dịch {ten_dgd} xã {ten_xa}",
        bold=True,
        size=13,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=6,
        indent_cm=0,
    )

    _p(
        f"Ngày {ngay_s}/{thang_s}/{nam_s}, NHCSXH và các tổ chức CT-XH nhận ủy thác "
        f"đã tổ chức cuộc họp giao ban tại điểm giao dịch {ten_dgd} xã {ten_xa}. "
        f"Các thành phần tham gia cuộc họp cùng thống nhất các nội dung sau:",
        size=14,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        indent_cm=1.0,
    )

    _heading("I", "CHÍNH SÁCH MỚI TRONG THÁNG")
    if chinh_sach_moi.strip():
        _p(
            chinh_sach_moi,
            size=14,
            align=WD_ALIGN_PARAGRAPH.JUSTIFY,
            indent_cm=1.0,
        )
    else:
        _p(
            "(Không có chính sách mới trong tháng)",
            italic=True,
            size=14,
            indent_cm=1.0,
        )

    _heading("II", "KẾT QUẢ THỰC HIỆN HOẠT ĐỘNG ỦY THÁC")
    _p(
        "1. Kết quả cho vay trong tháng",
        bold=True,
        size=14,
        indent_cm=0,
        space_after=3,
    )

    tong_dn = pd.to_numeric(df_xa[COT_TONG_DU_NO], errors="coerce").sum() / 1e6
    so_kh = df_xa[COT_MA_KH].nunique()
    so_to = df_xa[COT_TEN_TO].nunique() if COT_TEN_TO in df_xa.columns else 0
    nqh = pd.to_numeric(df_xa[COT_DU_NO_QH], errors="coerce").sum() / 1e6
    ty_le_nqh = nqh / tong_dn * 100 if tong_dn > 0 else 0

    van_xuoi = (
        f"Tổng dư nợ đạt {tong_dn:,.0f} triệu đồng, với {so_kh:,} khách hàng "
        f"còn dư nợ, thông qua {so_to} Tổ TK&VV. "
        f"Trong đó nợ quá hạn {nqh:,.0f} triệu đồng, tỷ lệ {ty_le_nqh:.2f}%."
    )
    lai_ton_th = pd.to_numeric(
        df_xa.get(COT_LAI_TON, pd.Series(dtype=float)), errors="coerce"
    ).fillna(0).sum() / 1e6
    lai_ton_qh = pd.to_numeric(
        df_xa.get(COT_LAI_TON_QH, pd.Series(dtype=float)), errors="coerce"
    ).fillna(0).sum() / 1e6
    tong_lai_ton = lai_ton_th + lai_ton_qh
    van_xuoi += f" Lãi tồn {fmt(tong_lai_ton * 1e6)} triệu đồng."
    if df_baseline is not None and COT_TEN_XA in df_baseline.columns:
        ten_xa_val = df_xa[COT_TEN_XA].iloc[0] if COT_TEN_XA in df_xa.columns else None
        if ten_xa_val:
            df_bl_xa = loc_baseline_cung_xa_pgd(df_baseline, df_xa, str(ten_xa_val))
            baseline = pd.to_numeric(df_bl_xa[COT_TONG_DU_NO], errors="coerce").sum()
            tang_giam = tong_dn - (baseline / 1e6)
            van_xuoi += f" (tăng/giảm {abs(tang_giam):,.0f} triệu so với cùng kỳ)"
    _p(van_xuoi, size=13, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent_cm=1.0)

    tien_gui = pd.to_numeric(
        df_xa.get(COT_TIEN_GUI, pd.Series(dtype=float)), errors="coerce"
    ).fillna(0).sum() / 1e6
    so_kh_gui = (
        df_xa[df_xa.get(COT_TIEN_GUI, pd.Series(0)) > 0][COT_MA_KH].nunique()
        if COT_TIEN_GUI in df_xa.columns else 0
    )
    ty_le_gui = so_kh_gui / so_kh * 100 if so_kh > 0 else 0.0
    van_xuoi_tg = (
        f"Tổng số dư tiền gửi tiết kiệm đạt {fmt(tien_gui * 1e6)} triệu đồng. "
        f"Tỷ lệ hộ vay có tham gia gửi tiết kiệm đạt {ty_le_gui:.0f}%."
    )
    _p(van_xuoi_tg, size=13, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent_cm=1.0)

    from data import danh_dau_khong_hd
    df_m = danh_dau_khong_hd(df_xa)
    so_mon_3m = int(df_m[COT_MON_3M].sum()) if COT_MON_3M in df_m.columns else 0
    if so_mon_3m > 0:
        van_xuoi_3m = (
            f"Toàn xã hiện còn {so_mon_3m} món vay từ 03 tháng trở lên "
            f"không hoạt động, tiềm ẩn nguy cơ phát sinh nợ quá hạn, "
            f"cần tiếp tục theo dõi và xử lý kịp thời."
        )
        _p(van_xuoi_3m, size=13, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent_cm=1.0)

    _p(
        "Đơn vị tính: tổ, khách hàng, triệu đồng, %",
        italic=True,
        size=11,
        indent_cm=0,
    )

    tao_bang_dvut(doc, df_xa)
    _patch_font(doc.tables[-1], size=11)

    _p(
        "Biểu chi tiết kết quả giao dịch kèm theo",
        italic=True,
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        indent_cm=0,
    )
    _tao_bang_chi_tiet_to(doc, df_xa)

    _p(
        "2. Tồn tại, hạn chế",
        bold=True,
        size=14,
        indent_cm=0,
        space_after=3,
    )
    if ton_tai_han_che.strip():
        _p(
            ton_tai_han_che,
            size=14,
            align=WD_ALIGN_PARAGRAPH.JUSTIFY,
            indent_cm=1.0,
        )
    else:
        _p("(Không có)", italic=True, size=14, indent_cm=1.0)

    _heading("III", "NHIỆM VỤ THÁNG TIẾP THEO")
    _p(
        "1. Kế hoạch thu nợ, giải ngân",
        bold=True,
        size=14,
        indent_cm=0,
        space_after=3,
    )

    tao_bang_ke_hoach(doc, df_xa, giai_ngan_input=giai_ngan_input)
    _patch_font(doc.tables[-1], size=11)

    _p(
        "(Danh sách nợ đến hạn kèm theo)",
        italic=True,
        size=12,
        indent_cm=0,
    )

    _p(
        "2. Kế hoạch kiểm tra, giám sát",
        bold=True,
        size=14,
        indent_cm=0,
        space_after=3,
    )
    _p(
        "3. Kế hoạch đôn đốc, xử lý nợ xấu",
        bold=True,
        size=14,
        indent_cm=0,
        space_after=3,
    )
    _p(
        "4. Kế hoạch triển khai các nội dung khác (nếu có)",
        bold=True,
        size=14,
        indent_cm=0,
        space_after=3,
    )
    if nhiem_vu_tiep.strip():
        _p(
            nhiem_vu_tiep,
            size=14,
            align=WD_ALIGN_PARAGRAPH.JUSTIFY,
            indent_cm=1.0,
        )

    _p()

    tbl_f = doc.add_table(rows=1, cols=2)
    _no_border(tbl_f)
    fl, fr = tbl_f.rows[0].cells
    _cell_w(fl, 6.4)
    _cell_w(fr, 9.6)

    p_nn = fl.paragraphs[0]
    r_nn = p_nn.add_run("Nơi nhận:")
    r_nn.bold = True
    r_nn.font.size = Pt(11)
    r_nn.font.name = "Times New Roman"
    r_nn.font.color.rgb = RGBColor(0, 0, 0)
    for dong in (
        "- Đảng ủy, UBND xã (để b/c);",
        "- Tổ chức CT-XH nhận ủy thác;",
        "- Lưu PGD.",
    ):
        p_ln = fl.add_paragraph()
        r_ln = p_ln.add_run(dong)
        r_ln.font.size = Pt(11)
        r_ln.font.name = "Times New Roman"
        r_ln.font.color.rgb = RGBColor(0, 0, 0)

    p_kt_gd = fr.paragraphs[0]
    p_kt_gd.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_kt_gd = p_kt_gd.add_run("KT.GIÁM ĐỐC")
    r_kt_gd.bold = True
    r_kt_gd.font.size = Pt(11)
    r_kt_gd.font.name = "Times New Roman"
    r_kt_gd.font.color.rgb = RGBColor(0, 0, 0)

    p_pgd = fr.add_paragraph()
    p_pgd.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_pgd = p_pgd.add_run("PHÓ GIÁM ĐỐC")
    r_pgd.bold = True
    r_pgd.font.size = Pt(11)
    r_pgd.font.name = "Times New Roman"
    r_pgd.font.color.rgb = RGBColor(0, 0, 0)

    p_kt = fr.add_paragraph()
    p_kt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_kt = p_kt.add_run("(Ký tên, đóng dấu)")
    r_kt.italic = True
    r_kt.font.size = Pt(11)
    r_kt.font.name = "Times New Roman"
    r_kt.font.color.rgb = RGBColor(0, 0, 0)

    for _ in range(3):
        p_sp = fr.add_paragraph()
        p_sp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_ht = fr.add_paragraph()
    p_ht.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ten_ky_hien_thi = ten_nguoi_ky.strip() if ten_nguoi_ky.strip() else "…………………………………"
    r_ht = p_ht.add_run(ten_ky_hien_thi)
    r_ht.font.size = Pt(11)
    r_ht.font.name = "Times New Roman"
    r_ht.font.color.rgb = RGBColor(0, 0, 0)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def xuat_bien_ban_giao_ban(
    df_xa: pd.DataFrame,
    df_baseline: pd.DataFrame | None,
    nam_moc: int,
    template_path: str,
    giai_ngan_input: dict | None = None,
) -> bytes:
    """
    Trả về bytes file .docx đã điền đầy đủ.
    Gọi từ ws_operation hoặc ws_management.
    """
    import io

    from docx import Document as _Doc

    doc = _Doc(template_path)
    tag_values = tinh_so_lieu_van_xuoi(df_xa, df_baseline, nam_moc)

    for para in doc.paragraphs:
        for tag, val in tag_values.items():
            if tag in para.text:
                for run in para.runs:
                    if tag in run.text:
                        run.text = run.text.replace(tag, val)

    BANG_FUNCS = {
        "{{bang_dvut}}": lambda p: tao_bang_dvut(doc, df_xa),
        "{{bang_chuong_trinh}}": lambda p: tao_bang_chuong_trinh(
            doc, df_xa, df_baseline, nam_moc
        ),
        "{{bang_ke_hoach}}": lambda p: tao_bang_ke_hoach(doc, df_xa, giai_ngan_input),
    }
    for para in doc.paragraphs:
        for tag, fn in BANG_FUNCS.items():
            if tag in para.text:
                fn(para)
                p = para._element
                p.getparent().remove(p)
                break

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
