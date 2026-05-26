"""Báo cáo Word định kỳ — Tổng quan + NQH + KHTD.

Tạo file .docx được định dạng chuyên nghiệp từ dữ liệu parquet.
Gọi từ tab Báo cáo định kỳ hoặc chạy độc lập.
"""
from __future__ import annotations

import os
from datetime import datetime, date
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from data.core import _duckdb_query
from config import (
    CACHE_HSTD,
    COT_TEN_PGD, COT_TEN_KH, COT_SO_KU, COT_MA_KH,
    COT_TONG_DU_NO, COT_DU_NO_QH, COT_DU_NO_TH, COT_DU_NO_KHOANH,
    COT_LAI_TON, COT_NGAY_DH, COT_TEN_CT, COT_NGUON_VON,
    DS_PGD, TEN_CHI_NHANH_HIEN_THI,
)
from utils import fmt_ty, fmt_so

GREEN = RGBColor(0x2E, 0x7D, 0x32)
DARK = RGBColor(0x33, 0x33, 0x33)
GRAY = RGBColor(0x66, 0x66, 0x66)


def _add_heading(doc, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = GREEN


def _add_table(doc, df: pd.DataFrame, col_widths: list[float] | None = None):
    """Thêm bảng vào document từ DataFrame."""
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, col in enumerate(df.columns):
        cell = table.rows[0].cells[i]
        cell.text = str(col)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell._tc.get_or_add_tcPr()
        shading = cell._tc.tcPr.makeelement(qn("w:shd"), {
            qn("w:fill"): "2E7D32",
            qn("w:val"): "clear",
        })
        cell._tc.tcPr.append(shading)

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            cells[i].text = str(row[col])
            for p in cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)


def _add_kpi_line(doc, label: str, value: str, bold: bool = False):
    p = doc.add_paragraph()
    run_label = p.add_run(f"{label}: ")
    run_label.font.size = Pt(11)
    run_label.font.color.rgb = GRAY
    run_val = p.add_run(value)
    run_val.font.size = Pt(11)
    run_val.font.color.rgb = DARK
    if bold:
        run_val.bold = True


def generate_word_report(parquet_path: str | None = None) -> bytes:
    """Tạo báo cáo Word tổng hợp. Trả về bytes .docx."""
    parquet_path = parquet_path or str(CACHE_HSTD)
    if not os.path.exists(parquet_path):
        raise FileNotFoundError(f"Không tìm thấy {parquet_path}")

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    now = datetime.now()

    doc.add_heading("BÁO CÁO TÓM TẮT ĐỊNH KỲ", level=0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{TEN_CHI_NHANH_HIEN_THI} — {now.strftime('%H:%M ngày %d/%m/%Y')}")
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    run.italic = True

    # ── 1. Tổng quan ──
    _add_heading(doc, "1. Tổng quan danh mục tín dụng", 1)

    sql_tq = f"""
        SELECT
            SUM("{COT_TONG_DU_NO}")   AS tong_du_no,
            COUNT(DISTINCT "{COT_MA_KH}") AS so_kh,
            COUNT("{COT_SO_KU}")       AS so_mon,
            SUM("{COT_DU_NO_QH}")      AS no_qh,
            SUM("{COT_DU_NO_TH}")      AS no_th,
            SUM("{COT_DU_NO_KHOANH}")  AS no_khoanh
        FROM read_parquet(?)
    """
    tq = _duckdb_query(sql_tq, [str(parquet_path)])
    if not tq.empty:
        r = tq.iloc[0]
        tdn = float(r.get("tong_du_no", 0))
        nqh = float(r.get("no_qh", 0))
        nkh = float(r.get("no_khoanh", 0))
        tlq = round(nqh / tdn * 100, 2) if tdn > 0 else 0
        tlk = round(nkh / tdn * 100, 2) if tdn > 0 else 0

        _add_kpi_line(doc, "Tổng dư nợ", f"{fmt_ty(tdn)} triệu đồng", bold=True)
        _add_kpi_line(doc, "Số khách hàng", fmt_so(int(r.get("so_kh", 0))))
        _add_kpi_line(doc, "Số món vay", fmt_so(int(r.get("so_mon", 0))))
        _add_kpi_line(doc, "Dư nợ quá hạn", f"{fmt_ty(nqh)} triệu đồng ({tlq}%)", bold=True)
        _add_kpi_line(doc, "Dư nợ khoanh", f"{fmt_ty(nkh)} triệu đồng ({tlk}%)")
        _add_kpi_line(doc, "Dư nợ trong hạn", f"{fmt_ty(float(r.get('no_th', 0)))} triệu đồng")
        _add_kpi_line(doc, "Tỷ lệ Nợ xấu (QH+Khoanh)", f"{round(tlq + tlk, 2)}%", bold=True)

    doc.add_page_break()

    # ── 2. Chi tiết theo PGD ──
    _add_heading(doc, "2. Dư nợ theo Phòng Giao dịch", 1)

    sql_pgd = f"""
        SELECT
            "{COT_TEN_PGD}"          AS "PGD",
            SUM("{COT_TONG_DU_NO}")  AS "Dư nợ",
            SUM("{COT_DU_NO_QH}")    AS "Nợ QH",
            SUM("{COT_DU_NO_TH}")    AS "Nợ TH",
            SUM("{COT_DU_NO_KHOANH}") AS "Khoanh"
        FROM read_parquet(?)
        WHERE "{COT_TONG_DU_NO}" IS NOT NULL
        GROUP BY "{COT_TEN_PGD}"
        ORDER BY "Dư nợ" DESC
    """
    df_pgd = _duckdb_query(sql_pgd, [str(parquet_path)])
    if not df_pgd.empty:
        df_pgd["Tỷ lệ NQH"] = (df_pgd["Nợ QH"] / df_pgd["Dư nợ"].replace(0, pd.NA) * 100).round(2)
        df_pgd["Dư nợ"] = df_pgd["Dư nợ"].apply(lambda x: f"{x:,.0f}")
        df_pgd["Nợ QH"] = df_pgd["Nợ QH"].apply(lambda x: f"{x:,.0f}")
        df_pgd["Nợ TH"] = df_pgd["Nợ TH"].apply(lambda x: f"{x:,.0f}")
        df_pgd["Khoanh"] = df_pgd["Khoanh"].apply(lambda x: f"{x:,.0f}")
        df_pgd["Tỷ lệ NQH"] = df_pgd["Tỷ lệ NQH"].apply(lambda x: f"{x:.2f}%")
        _add_table(doc, df_pgd, [5, 2.5, 2.5, 2.5, 2.5, 2])

    doc.add_paragraph()

    # ── 3. Top NQH ──
    _add_heading(doc, "3. Top 20 khoản vay Nợ quá hạn cao nhất", 1)

    sql_top = f"""
        SELECT
            "{COT_TEN_PGD}"   AS "PGD",
            "{COT_TEN_KH}"    AS "Tên KH",
            "{COT_SO_KU}"     AS "Số KU",
            "{COT_TONG_DU_NO}" AS "Dư nợ",
            "{COT_DU_NO_QH}"  AS "Nợ QH"
        FROM read_parquet(?)
        WHERE "{COT_DU_NO_QH}" > 0
        ORDER BY "{COT_DU_NO_QH}" DESC
        LIMIT 20
    """
    df_top = _duckdb_query(sql_top, [str(parquet_path)])
    if not df_top.empty:
        df_top["Dư nợ"] = df_top["Dư nợ"].apply(lambda x: f"{x:,.0f}")
        df_top["Nợ QH"] = df_top["Nợ QH"].apply(lambda x: f"{x:,.0f}")
        _add_table(doc, df_top, [3, 4, 3, 2.5, 2.5])

    doc.add_page_break()

    # ── 4. KHTD ──
    _add_heading(doc, "4. Kế hoạch Tín dụng", 1)

    import db
    khtd = db.doc_kv("khtd_cn")
    if isinstance(khtd, dict):
        _add_kpi_line(doc, "Số chương trình", str(len(khtd)))
        total_khtd = 0
        for ct, targets in khtd.items():
            if isinstance(targets, dict):
                for _, v in targets.items():
                    if isinstance(v, (int, float)):
                        total_khtd += v
        _add_kpi_line(doc, "Tổng KHTD toàn Chi nhánh", f"{fmt_ty(total_khtd)} triệu đồng", bold=True)
    else:
        p = doc.add_paragraph("⚠️ Chưa nhập Kế hoạch Tín dụng Chi nhánh.")
        p.runs[0].italic = True

    # ── Footer ──
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"VBSP-SCM · Tự động xuất lúc {now.strftime('%H:%M %d/%m/%Y')}")
    run.font.size = Pt(8)
    run.font.color.rgb = GRAY
    run.italic = True

    from io import BytesIO
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_word_report_file(output_path: str | None = None) -> str:
    """Tạo báo cáo Word và lưu ra file. Trả về đường dẫn file."""
    from config import CACHE_DIR
    out = output_path or str(Path(CACHE_DIR) / "reports" / f"BaoCao_Ngay_{datetime.now().strftime('%Y%m%d_%H%M')}.docx")
    os.makedirs(os.path.dirname(out), exist_ok=True)
    data = generate_word_report()
    with open(out, "wb") as f:
        f.write(data)
    return out
