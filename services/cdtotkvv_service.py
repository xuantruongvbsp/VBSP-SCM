"""
Các hàm xử lý dữ liệu thuần (không có st.*) cho tab Chấm điểm Tổ TK&VV.

Extract từ tabs/tab_cdtotkvv.py để tái sử dụng và kiểm thử độc lập.
"""
from __future__ import annotations

import os
import re
import unicodedata
from datetime import date as _date_today

import pandas as pd

from config import (
    DS_PGD,
    DON_VI_CHI_NHANH,
    CACHE_HSTD,
    COT_TEN_PGD as _COT_TEN_PGD_HSTD,
    COT_TEN_XA as _COT_TEN_XA_HSTD,
    COT_NGAY_SINH as _COT_NGAY_SINH_HSTD,
    COT_TEN_KH as _COT_TEN_KH_HSTD,
    COT_TEN_TO_TRUONG as _COT_TEN_TO_TRUONG_HSTD,
)
from data.pgd import duong_dan_pgd as _duong_dan_pgd
from data.core import ts_file
from utils import fmt_so
try:
    from logger import get_logger
    logger = get_logger(__name__)
except Exception:
    import logging
    logger = logging.getLogger(__name__)


def tong_hop_tu_pgd_data() -> "pd.DataFrame | None":
    """Đọc cdtotkvv_latest.xlsx từ pgd_data/{slug}/ của tất cả đơn vị, concat lại."""
    from data.cdtotkvv import doc_cdtotkvv_path

    tat_ca_dv = [DON_VI_CHI_NHANH] + DS_PGD
    frames = []
    for ten_dv in tat_ca_dv:
        try:
            dd = _duong_dan_pgd(ten_dv, "cdtotkvv")
            if not dd or not os.path.exists(dd):
                continue
            df = doc_cdtotkvv_path(dd, ts_file(dd))
            if df is None:
                continue
            if not df.empty:
                frames.append(df)
        except Exception as e:
            logger.warning("tong_hop_tu_pgd_data: bỏ qua đơn vị lỗi — %s", e, exc_info=True)
            continue
    return pd.concat(frames, ignore_index=True) if frames else None


def bang_trang_thai_cdtotkvv() -> pd.DataFrame:
    """
    Tạo bảng trạng thái CDTOTKVV cho 22 đơn vị (Chi nhánh + 21 PGD).
    Trả về DataFrame với cột: Đơn vị, Trạng thái, Cập nhật lần cuối.
    """
    from data.pgd import doc_trang_thai_file

    tat_ca_dv = [DON_VI_CHI_NHANH] + DS_PGD

    rows = []
    for ten_dv in tat_ca_dv:
        trang_thai_info = doc_trang_thai_file(ten_dv, "cdtotkvv")

        if not trang_thai_info["co_file"]:
            trang_thai = "❌ Chưa có"
            cap_nhat = "—"
        else:
            ngay_upload = trang_thai_info["ngay_upload"]
            so_ngay_cu = trang_thai_info["so_ngay_cu"]

            if trang_thai_info["canh_bao"] == "ok":
                trang_thai = f"✅ {ngay_upload.strftime('%d/%m')}"
            else:
                trang_thai = f"⚠️ {ngay_upload.strftime('%d/%m')} ({so_ngay_cu} ngày)"

            cap_nhat = ngay_upload.strftime('%d/%m/%Y %H:%M')

        rows.append({
            "Đơn vị": ten_dv,
            "Trạng thái": trang_thai,
            "Cập nhật lần cuối": cap_nhat,
        })

    return pd.DataFrame(rows)


def loc_df(df: pd.DataFrame, mode: str, pgd_user: str) -> pd.DataFrame:
    """
    Lọc dữ liệu theo chế độ hiển thị:
    - mode "cn": toàn Chi nhánh (dùng cho ws_management)
    - mode "pgd": chỉ PGD mình (dùng cho ws_operation)
    """
    if df is None or df.empty:
        return df
    if mode == "pgd" and pgd_user:
        if "ma_dv" in df.columns and "ten_dv" in df.columns:
            mask_ten = df["ten_dv"].astype(str).str.strip().str.lower() == pgd_user.strip().lower()
            if not mask_ten.any():
                mask_ma = df["ma_dv"].astype(str).str.strip().str.lower() == pgd_user.strip().lower()
                return df[mask_ma]
            return df[mask_ten]
        elif "ten_dv" in df.columns:
            return df[df["ten_dv"].astype(str).str.strip().str.lower() == pgd_user.strip().lower()]
        elif "ma_dv" in df.columns:
            return df[df["ma_dv"].astype(str).str.strip().str.lower() == pgd_user.strip().lower()]
    return df  # mode "cn" → trả về toàn bộ


def cdtotkvv_ten_sheet_excel(ten_hien_thi: str, da_dung: set[str]) -> str:
    """Tên sheet Excel ≤31 ký tự, không ký tự cấm, không trùng."""
    forbidden = set("[]:*?/\\")
    base = "".join("_" if c in forbidden else c for c in ten_hien_thi).strip() or "Tieu_chi"
    base = base[:28]
    ten = base[:31]
    n = 1
    while ten in da_dung:
        hau_to = f"_{n}"
        ten = (base[: 31 - len(hau_to)]).rstrip("_") + hau_to
        n += 1
    da_dung.add(ten)
    return ten


def fmt_xuat_to_khong_dat_vn(df: pd.DataFrame) -> pd.DataFrame:
    """Bản sao DataFrame đã format số/tiền chuẩn VN cho Excel/PDF (không dùng cho thống kê .mean())."""
    out = df.copy()
    if "Dư nợ" in out.columns:
        out["Dư nợ"] = pd.to_numeric(out["Dư nợ"], errors="coerce").map(
            lambda v: fmt_so(v) if pd.notna(v) else "—"
        )
    if "Số dư TK" in out.columns:
        out["Số dư TK"] = pd.to_numeric(out["Số dư TK"], errors="coerce").map(
            lambda v: fmt_so(v) if pd.notna(v) else "—"
        )
    for cot in ("Điểm đạt được", "Điểm tối đa", "Thiếu", "Tổng điểm"):
        if cot not in out.columns:
            continue
        out[cot] = out[cot].apply(
            lambda v: fmt_so(int(round(float(v)))) if pd.notna(v) and v != "" else "—"
        )
    return out


# ═══════════════════════════════════════════════════════════════════════════════
# Thống kê Tuổi Tổ trưởng theo PGD / theo Xã (bins nhóm tuổi)
# ═══════════════════════════════════════════════════════════════════════════════

_BINS_THONG_KE_TUOI: list[tuple[int, int, str]] = [
    (0,   29,  "Dưới 30 tuổi"),
    (30,  39,  "30 - 39 tuổi"),
    (40,  49,  "40 - 49 tuổi"),
    (50,  59,  "50 - 59 tuổi"),
    (60,  69,  "60 - 69 tuổi"),
    (70, 200,  "Từ 70 tuổi trở lên"),
]
_TEN_BINS = [b[2] for b in _BINS_THONG_KE_TUOI]


def _them_cot_chuan_to(df: "pd.DataFrame") -> "pd.DataFrame":
    """Thêm cột chuẩn PGD/Xã/Tổ để đếm unique Tổ nhất quán."""
    out = df.copy()
    original_cols = set(out.columns)
    for _std_col, _fallbacks in [
        ("ten_pgd_std", ["ten_pgd_std", "ten_dv", "Tên PGD", "PGD"]),
        ("ten_xa_std", ["ten_xa_std", "ten_xa", "Tên xã", "Xã/Phường", "Xã"]),
        ("ten_to_std", ["ten_to_std", "ma_to", "Tổ", "Tên tổ"]),
    ]:
        source_col = next((_name for _name in _fallbacks if _name in original_cols), None)
        if source_col is not None:
            out[_std_col] = out[source_col].astype("string").fillna("").str.strip()
        else:
            out[_std_col] = ""
    return out


def _df_unique_theo_to(df: "pd.DataFrame") -> "pd.DataFrame":
    """Dedupe theo PGD/Xã/Tổ; dòng thiếu mã/tên tổ vẫn giữ riêng từng dòng."""
    if df is None or df.empty or "ten_to_std" not in df.columns:
        return df.copy() if df is not None else pd.DataFrame()

    key_cols = [c for c in ["ten_pgd_std", "ten_xa_std", "ten_to_std"] if c in df.columns]
    key_df = df[key_cols].astype("string").fillna("").apply(lambda s: s.str.strip())
    has_to = key_df["ten_to_std"].ne("")
    if not bool(has_to.any()):
        return df.copy()

    with_key = df.loc[has_to].drop_duplicates(subset=key_cols, keep="first")
    without_key = df.loc[~has_to]
    return pd.concat([with_key, without_key], ignore_index=True)


def _dem_unique_to(df: "pd.DataFrame | None") -> int:
    if df is None or df.empty:
        return 0
    return int(len(_df_unique_theo_to(_them_cot_chuan_to(df))))


def _normalize_digits(s: "pd.Series", expected_len: int | None = None) -> "pd.Series":
    """Chuẩn hóa một Series về string chỉ chứa chữ số; zfill nếu có expected_len."""
    out = s.astype("string").fillna("").str.replace(r"\D", "", regex=True)
    if expected_len and expected_len > 0:
        out = out.str.zfill(expected_len)
    return out


def _strip_vn_accents(s: "pd.Series") -> "pd.Series":
    """Chuẩn hóa tiếng Việt: NFD bỏ dấu combining → uppercase → ký tự đặc biệt → strip prefix Ông/Bà/Thầy/Cô/Anh/Chị/Bác/Cũ."""
    def _norm_one(text: object) -> str:
        if text is None:
            return ""
        try:
            s_str = str(text)
        except Exception:
            return ""
        if not s_str:
            return ""
        nfd = unicodedata.normalize("NFD", s_str)
        no_marks = "".join(ch for ch in nfd if unicodedata.category(ch) != "Mn")
        up = no_marks.upper()
        cleaned = re.sub(r"[^A-Z0-9\s]", "", up)
        collapse = re.sub(r"\s+", " ", cleaned).strip()
        prefixes = ("ONG ", "BA ", "THAY ", "CO ", "ANH ", "CHI ", "BAC ", "CU ")
        for p in prefixes:
            if collapse.startswith(p):
                collapse = collapse[len(p):].strip()
                break
        return collapse
    return s.map(_norm_one).astype("string").fillna("")


def _tinh_tuoi_tu_ngay_sinh(ngay_sinh_series: "pd.Series") -> "pd.Series":
    """Tính tuổi từ một Series ngày sinh datetime; trả về float series NA nếu không parse được."""
    today = _date_today.today()
    dob = pd.to_datetime(ngay_sinh_series, errors="coerce", dayfirst=True)
    tuoi = today.year - dob.dt.year
    chua_sinh_nhat = (dob.dt.month > today.month) | (
        (dob.dt.month == today.month) & (dob.dt.day > today.day)
    )
    tuoi = tuoi - chua_sinh_nhat.astype(int)
    valid = tuoi.between(15, 110) & dob.notna()
    return tuoi.where(valid, pd.NA)


def enrich_tuoi_to_truong_fallback_tu_hstd(
    df_cdto: "pd.DataFrame | None",
) -> tuple["pd.DataFrame | None", str, int]:
    """
    Nếu df_cdto không có (hoặc quá ít) dữ liệu tuổi tổ trưởng từ file CDTOTKVV upload,
    fallback enrich từ HSTD: ước tính tuổi theo TB tuổi khách hàng trong Xã + độ lệch
    tổ trưởng thường lớn hơn KH trong tổ.
    Returns:
      (df_enriched, source_used: str, so_dong_duoc_fill: int)
      source_used ∈ {"CDTOTKVV upload (chính xác)",
                     "HSTD: ước tính từ Tuổi KH xã (Median + 8 năm)",
                     "Không có dữ liệu"}
    """
    if df_cdto is None or df_cdto.empty:
        return df_cdto, "Không có dữ liệu", 0

    df = df_cdto.copy()

    # 1) Kiểm tra: đã có dữ liệu tuổi từ file upload chưa?
    has_tuoi_col = any(c in df.columns for c in ("tuoi_to_truong", "tuoi", "Tuổi tổ trưởng", "Tuổi"))
    co_dl_tuoi = False
    # Khởi tạo cột phân loại nguồn + flag có vay vốn (sử dụng ở UI chia 2 danh sách)
    if "_nguon_tuoi_est" not in df.columns:
        df["_nguon_tuoi_est"] = pd.Series("", index=df.index, dtype="string")
    if "_nguon_chi_tiet" not in df.columns:
        df["_nguon_chi_tiet"] = pd.Series("", index=df.index, dtype="string")
    if "_co_vay_von" not in df.columns:
        df["_co_vay_von"] = pd.Series(pd.NA, index=df.index, dtype="Int64")  # 1 = Có vay vốn | 0 = Không xác định / Không vay vốn | NA = Upload thật (chưa kiểm tra)
    if has_tuoi_col:
        src_col = next(
            (c for c in ("tuoi_to_truong", "tuoi", "Tuổi tổ trưởng", "Tuổi") if c in df.columns),
            None,
        )
        if src_col:
            tuoi_num = pd.to_numeric(df[src_col], errors="coerce")
            valid = tuoi_num.between(18, 100)
            ty_le_valid = float(valid.mean()) if len(df) > 0 else 0.0
            if bool(valid.any()):
                df.loc[valid, "_nguon_chi_tiet"] = "Upload thật (CDTOTKVV chính xác)"
            if ty_le_valid >= 0.30:
                df["tuoi_to_truong"] = tuoi_num.where(valid, pd.NA)
                # Upload thật ≥30%: flag có vay vốn = NA (người dùng tự kiểm tra)
                df.loc[~valid, "_nguon_chi_tiet"] = "Upload thật (chưa có tuổi hợp lệ)"
                # Không fill _co_vay_von cho upload thật, giữ NA để UI hiển thị theo nguồn
                return (
                    df if "tuoi_to_truong" in df.columns else df.assign(tuoi_to_truong=tuoi_num.where(valid, pd.NA)),
                    "CDTOTKVV upload (chính xác — từ file chấm điểm Tổ)",
                    0,
                )
            co_dl_tuoi = ty_le_valid > 0

    # 2) Fallback: enrich từ HSTD (cache/hstd.parquet)
    try:
        import pyarrow.parquet as pq
        if not os.path.exists(CACHE_HSTD):
            return df, (
                "CDTOTKVV upload — thiếu tuổi tổ trưởng (HSTD chưa có dữ liệu)"
                if not co_dl_tuoi
                else "CDTOTKVV upload — ít dữ liệu tuổi (HSTD chưa có dữ liệu)"
            ), 0

        # Rule 6.16: DuckDB kiểm tra schema trước khi query
        schema = pq.read_schema(CACHE_HSTD)
        cols_hstd = [f.name for f in schema]
        needed_cols = [_COT_TEN_PGD_HSTD, _COT_TEN_XA_HSTD, _COT_NGAY_SINH_HSTD, _COT_TEN_KH_HSTD]
        missing = [c for c in needed_cols if c not in cols_hstd]
        if missing:
            logger.warning(
                "enrich_tuoi_to_truong_fallback: HSTD thiếu cột %s — không thể enrich",
                missing,
            )
            return df, (
                "CDTOTKVV upload — thiếu tuổi (HSTD thiếu cột ngày sinh)"
                if not co_dl_tuoi
                else "CDTOTKVV upload — ít tuổi (HSTD thiếu cột)"
            ), 0

        # Đọc chỉ các cột cần thiết từ parquet (tiết kiệm RAM)
        hstd = pq.read_table(
            CACHE_HSTD,
            columns=[_COT_TEN_PGD_HSTD, _COT_TEN_XA_HSTD, _COT_NGAY_SINH_HSTD, _COT_TEN_KH_HSTD],
        ).to_pandas()
        if hstd.empty:
            return df, (
                "CDTOTKVV upload — thiếu tuổi (HSTD rỗng)"
                if not co_dl_tuoi
                else "CDTOTKVV upload — ít tuổi (HSTD rỗng)"
            ), 0

        # Tính tuổi từng khách hàng HSTD
        hstd["_tuoi_kh"] = _tinh_tuoi_tu_ngay_sinh(hstd[_COT_NGAY_SINH_HSTD])
        hstd = hstd[hstd["_tuoi_kh"].notna()].copy()
        if hstd.empty:
            return df, (
                "CDTOTKVV upload — thiếu tuổi (HSTD không parse được ngày sinh KH)"
                if not co_dl_tuoi
                else "CDTOTKVV upload — ít tuổi (HSTD không parse được ngày sinh)"
            ), 0

        # Chuẩn hóa cột PGD/Xã trong df_cdto (cột ten_pgd_std/ten_xa_std đã có từ _them_cot_chuan_to)
        if "ten_pgd_std" not in df.columns or "ten_xa_std" not in df.columns:
            df = _them_cot_chuan_to(df)
        if "tuoi_to_truong" not in df.columns:
            df["tuoi_to_truong"] = pd.NA

        # Tạo key join chuẩn hóa PGD+Xã — dùng chung cho Bước 2 (link tên) và Bước 3 (median xã)
        df["_join_pgd"] = df["ten_pgd_std"].astype("string").fillna("").str.strip()
        df["_join_xa"]  = df["ten_xa_std"].astype("string").fillna("").str.strip()
        for c_std, c_goc in (("_join_pgd", "PGD"), ("_join_xa", "Xã/Phường")):
            if (df[c_std] == "").any():
                mask = df[c_std] == ""
                if c_goc in df.columns:
                    df.loc[mask, c_std] = (
                        df.loc[mask, c_goc].astype("string").fillna("").str.strip()
                    )

        # --- Bước 2 (TỐI ƯU MỚI): Link Tên tổ trưởng + Xã chuẩn hóa → HSTD KH cùng tên,
        #     chọn KH LỚN TUỔI NHẤT (drop_duplicates keep first sau sort desc tuổi)
        n_method_c = 0
        _mask_cur_valid = pd.to_numeric(df["tuoi_to_truong"], errors="coerce").between(18, 100)
        _need_fill = ~_mask_cur_valid
        if bool(_need_fill.any()):
            try:
                # Chuẩn hóa key HSTD: (PGD_std + Xã_std không dấu, Tên KH không dấu bỏ prefix Ông/Bà)
                hstd["_k_xa_std"] = _strip_vn_accents(
                    hstd[_COT_TEN_XA_HSTD].astype("string").fillna("")
                )
                hstd["_k_pgd_std"] = _strip_vn_accents(
                    hstd[_COT_TEN_PGD_HSTD].astype("string").fillna("")
                )
                hstd["_k_ten_std"] = _strip_vn_accents(hstd[_COT_TEN_KH_HSTD])
                # Loại KH key rỗng (không link được)
                _hstd_has_key = (hstd["_k_pgd_std"] != "") & (hstd["_k_xa_std"] != "") & (hstd["_k_ten_std"] != "")
                hstd_key = hstd.loc[_hstd_has_key, ["_k_pgd_std", "_k_xa_std", "_k_ten_std", "_tuoi_kh"]].copy()
                if not hstd_key.empty:
                    # Sort: cùng (PGD,Xã,Tên) → lớn tuổi nhất đứng trước
                    hstd_key = hstd_key.sort_values(
                        by=["_k_pgd_std", "_k_xa_std", "_k_ten_std", "_tuoi_kh"],
                        ascending=[True, True, True, False],
                        na_position="last",
                    )
                    hstd_key = hstd_key.drop_duplicates(
                        subset=["_k_pgd_std", "_k_xa_std", "_k_ten_std"], keep="first"
                    )
                    # Chuẩn hóa key CDTO
                    df["_k_xa_std"] = _strip_vn_accents(df["_join_xa"])
                    df["_k_pgd_std"] = _strip_vn_accents(df["_join_pgd"])
                    if "ten_to_truong" in df.columns:
                        df["_k_ten_std"] = _strip_vn_accents(df["ten_to_truong"])
                    else:
                        ten_col = next(
                            (c for c in ("Tên tổ trưởng", "Tổ trưởng", "Họ tên tổ trưởng") if c in df.columns),
                            None,
                        )
                        df["_k_ten_std"] = _strip_vn_accents(df[ten_col]) if ten_col else ""
                    _df_sub = df.loc[_need_fill, ["_k_pgd_std", "_k_xa_std", "_k_ten_std"]].copy()
                    _merged_c = _df_sub.merge(
                        hstd_key.rename(columns={"_tuoi_kh": "_tuoi_c"}),
                        how="left",
                        on=["_k_pgd_std", "_k_xa_std", "_k_ten_std"],
                    )
                    _tuoi_c_match = pd.to_numeric(_merged_c["_tuoi_c"], errors="coerce")
                    _tuoi_c_valid = _tuoi_c_match.between(18, 100)
                    # Fill ngược lại df theo boolean mask để không lệch khi df đã filter giữ index cũ.
                    _fill_c_mask_values = _need_fill.to_numpy(copy=True)
                    _fill_c_mask_values[_fill_c_mask_values] = _tuoi_c_valid.to_numpy()
                    _fill_c_mask = pd.Series(_fill_c_mask_values, index=df.index)
                    n_method_c = int(_fill_c_mask.sum())
                    if n_method_c > 0:
                        df.loc[_fill_c_mask, "tuoi_to_truong"] = (
                            _tuoi_c_match[_tuoi_c_valid].values.astype("int")
                        )
                        # Ghi nguồn
                        _src_c = df.get("_nguon_tuoi_est", pd.Series("", index=df.index, dtype="string"))
                        if not isinstance(_src_c, pd.Series) or len(_src_c) != len(df):
                            _src_c = pd.Series("", index=df.index, dtype="string")
                        _src_c = _src_c.astype("string").fillna("")
                        _src_c.loc[_fill_c_mask] = "HSTD: Tên tổ trùng KH (ngày sinh hồ sơ vay)"
                        df["_nguon_tuoi_est"] = _src_c
                        # Gắn _nguon_chi_tiet + _co_vay_von = 1 (match HSTD = CÓ hồ sơ vay vốn thực tế với VBSP)
                        _ct = df.get("_nguon_chi_tiet", pd.Series("", index=df.index, dtype="string"))
                        if not isinstance(_ct, pd.Series) or len(_ct) != len(df):
                            _ct = pd.Series("", index=df.index, dtype="string")
                        _ct = _ct.astype("string").fillna("")
                        _ct.loc[_fill_c_mask] = "Method C: Tên tổ trùng KH (Có hồ sơ vay vốn)"
                        df["_nguon_chi_tiet"] = _ct
                        _cv = df.get("_co_vay_von", pd.Series(pd.NA, index=df.index, dtype="Int64"))
                        if not isinstance(_cv, pd.Series) or len(_cv) != len(df):
                            _cv = pd.Series(pd.NA, index=df.index, dtype="Int64")
                        _cv.loc[_fill_c_mask] = 1
                        df["_co_vay_von"] = _cv.astype("Int64")
                # Cleanup helper key Bước 2
                for _c in ["_k_pgd_std", "_k_xa_std", "_k_ten_std"]:
                    if _c in df.columns:
                        df = df.drop(columns=[_c])
            except Exception as exc2:
                logger.warning(
                    "enrich_tuoi_to_truong_fallback: Bước 2 link tên tổ trưởng lỗi — %s",
                    exc2, exc_info=True,
                )
                n_method_c = 0

        # --- Bước 3 (FALLBACK CUỐI): Median tuổi KH xã + 8 năm — chỉ fill những dòng Bước 2 vẫn rỗng
        grp = (
            hstd.groupby([_COT_TEN_PGD_HSTD, _COT_TEN_XA_HSTD], dropna=False)["_tuoi_kh"]
            .agg(["median", "mean", "count"])
            .reset_index()
            .rename(columns={
                _COT_TEN_PGD_HSTD: "_h_pgd",
                _COT_TEN_XA_HSTD:  "_h_xa",
                "median": "_tuoi_median",
                "mean":   "_tuoi_mean",
                "count":  "_so_kh",
            })
        )
        for c in ("_h_pgd", "_h_xa"):
            grp[c] = grp[c].astype("string").fillna("").str.strip()

        merged = df.merge(
            grp,
            how="left",
            left_on=["_join_pgd", "_join_xa"],
            right_on=["_h_pgd", "_h_xa"],
            suffixes=("", "_h"),
        )

        tuoi_est = (merged["_tuoi_median"].fillna(merged["_tuoi_mean"]) + 8).round(0)
        tuoi_est = tuoi_est.clip(lower=18, upper=100)
        cur_valid = pd.to_numeric(merged["tuoi_to_truong"], errors="coerce").between(18, 100)
        fill_mask_3 = (~cur_valid) & tuoi_est.notna()
        so_fill_3 = int(fill_mask_3.sum())
        if so_fill_3 > 0:
            merged.loc[fill_mask_3, "tuoi_to_truong"] = tuoi_est[fill_mask_3].astype("Int64")
            _src = merged.get("_nguon_tuoi_est", pd.Series("", index=merged.index, dtype="string"))
            if not isinstance(_src, pd.Series) or len(_src) != len(merged):
                _src = pd.Series("", index=merged.index, dtype="string")
            _src = _src.astype("string").fillna("")
            _src.loc[fill_mask_3] = "HSTD ước tính (TB tuổi KH xã + 8y)"
            merged["_nguon_tuoi_est"] = _src
            # Gắn _nguon_chi_tiet = Method 3 (KHÔNG tìm được hồ sơ vay nào dưới tên tổ trưởng → KHÔNG VAY VỐN / tên không trùng)
            _ct3 = merged.get("_nguon_chi_tiet", pd.Series("", index=merged.index, dtype="string"))
            if not isinstance(_ct3, pd.Series) or len(_ct3) != len(merged):
                _ct3 = pd.Series("", index=merged.index, dtype="string")
            _ct3 = _ct3.astype("string").fillna("")
            _ct3.loc[fill_mask_3] = "Method 3: Ước tính TB xã (Chưa có hồ sơ vay / tên không trùng)"
            merged["_nguon_chi_tiet"] = _ct3
            _cv3 = merged.get("_co_vay_von", pd.Series(pd.NA, index=merged.index, dtype="Int64"))
            if not isinstance(_cv3, pd.Series) or len(_cv3) != len(merged):
                _cv3 = pd.Series(pd.NA, index=merged.index, dtype="Int64")
            _cv3.loc[fill_mask_3] = 0
            merged["_co_vay_von"] = _cv3.astype("Int64")

        # Drop helper cols
        drop_cols = [c for c in ["_join_pgd", "_join_xa", "_h_pgd", "_h_xa",
                                 "_tuoi_median", "_tuoi_mean", "_so_kh"] if c in merged.columns]
        merged = merged.drop(columns=drop_cols)

        so_fill_total = n_method_c + so_fill_3
        _n_total = len(merged) if len(merged) > 0 else 1
        _pct_c = (n_method_c / _n_total * 100.0) if n_method_c else 0.0
        _pct_3 = (so_fill_3 / _n_total * 100.0) if so_fill_3 else 0.0

        src_parts = []
        if co_dl_tuoi:
            src_parts.append("CDTOTKVV upload cho phần còn lại")
        if n_method_c:
            src_parts.append(
                f"Tên tổ trưởng trùng KH HSTD — ngày sinh hồ sơ vay chính: "
                f"{n_method_c:,} tổ ({_pct_c:.1f}%)"
            )
        if so_fill_3:
            src_parts.append(
                f"Ước tính từ TB tuổi KH xã + 8 năm: {so_fill_3:,} tổ ({_pct_3:.1f}%)"
            )
        if not src_parts:
            src_parts.append("CDTOTKVV upload — vẫn thiếu tuổi tổ trưởng" if not co_dl_tuoi else "CDTOTKVV upload — ít dữ liệu tuổi")
        source_msg = " · ".join(src_parts)
        return merged, source_msg, so_fill_total

    except Exception as exc:
        logger.error(
            "enrich_tuoi_to_truong_fallback_tu_hstd: Lỗi enrich fallback HSTD — %s",
            exc, exc_info=True,
        )
        return df, (
            f"CDTOTKVV upload — lỗi enrich HSTD ({exc})"
            if not co_dl_tuoi
            else f"CDTOTKVV upload — ít tuổi (lỗi enrich HSTD: {exc})"
        ), 0


def _loc_tuoi_nhan_bin(tuoi: int | float | None) -> str | None:
    if tuoi is None:
        return None
    try:
        t = int(tuoi)
    except Exception:
        return None
    if not 0 < t < 120:
        return None
    for lo, hi, label in _BINS_THONG_KE_TUOI:
        if lo <= t <= hi:
            return label
    return None


def _df_chi_tiet_so_huu_tuoi(df_raw: "pd.DataFrame | None") -> "pd.DataFrame":
    """Lọc chỉ những dòng có dữ liệu tuổi tổ trưởng hợp lệ. Trả về df enrich 'nhom_tuoi'.
    ⚠️ CHỈ LẤY các tổ trưởng thuộc 1 trong 2 loại:
      (a) Upload thật (từ file CDTOTKVV PGD nhập thủ công)
      (b) Method C: Tên tổ trùng KH (XÁC MINH CÓ HỒ SƠ VAY VỐN tại VBSP)
    Loại bỏ hoàn toàn Method 3 ước tính TB xã (không dự phóng cho thống kê PGD/Xã)."""
    if df_raw is None or df_raw.empty:
        return pd.DataFrame()
    df = df_raw.copy()
    # Filter: Chỉ giữ những dòng có flag = 1 (Có vay vốn Method C) HOẶC nguồn = Upload thật (nếu _co_vay_von NA mà _nguon_chi_tiet chứa Upload thật)
    _cv_raw = df.get("_co_vay_von")
    if isinstance(_cv_raw, pd.Series) and len(_cv_raw) == len(df):
        _cv = _cv_raw.copy()
    else:
        _cv = pd.Series(pd.NA, index=df.index, dtype="Int64")
    _nc_raw = df.get("_nguon_chi_tiet")
    if isinstance(_nc_raw, pd.Series) and len(_nc_raw) == len(df):
        _nc = _nc_raw.astype("string").fillna("")
    else:
        _nc = pd.Series("", index=df.index, dtype="string")
    # Trường hợp upload thật _co_vay_von NA → chấp nhận vì tuổi chính xác 100% từ PGD nhập thủ công
    _mask_upload = _nc.str.contains("Upload thật", na=False, case=False)
    _mask_covay = (_cv == 1)
    _mask_keep = _mask_covay | _mask_upload
    # Nếu _co_vay_von cột nào đó rỗng (hệ thống cũ chưa enrich) → giữ nguyên không filter (backward compat)
    if _cv.notna().any() or _mask_upload.any():
        df = df.loc[_mask_keep].copy()
    # Chuẩn hóa các cột tên đơn vị / xã: ưu tiên cột chuẩn, fallback alias
    if "tuoi_to_truong" not in df.columns:
        for _alias in ["tuoi", "tuoi_to", "Tuổi tổ trưởng", "Tuổi"]:
            if _alias in df.columns:
                df["tuoi_to_truong"] = pd.to_numeric(df[_alias], errors="coerce")
                break
    if "tuoi_to_truong" not in df.columns:
        return pd.DataFrame()
    df["_tuoi_num"] = pd.to_numeric(df["tuoi_to_truong"], errors="coerce")
    df = df[df["_tuoi_num"].notna()].copy()
    if df.empty:
        return pd.DataFrame()
    df["nhom_tuoi"] = df["_tuoi_num"].map(_loc_tuoi_nhan_bin)
    df = df[df["nhom_tuoi"].notna()].copy()
    df = _them_cot_chuan_to(df)
    return df.reset_index(drop=True)


def thong_ke_tuoi_theo_pgd(
    df_raw: "pd.DataFrame | None",
) -> tuple["pd.DataFrame", dict]:
    """
    Thống kê Tuổi Tổ trưởng THEO PGD.
    Trả về: (df_bins, summary)
      - df_bins: DataFrame với cột ['PGD'] + 6 bins nhóm tuổi + ['Tổng tổ trưởng có dữ liệu']
      - summary: dict tổng toàn CN: {tong_to, tb_tuoi, min_tuoi, max_tuoi, median_tuoi, khong_co_du_lieu, tong_to_tong_so}
    """
    df = _df_chi_tiet_so_huu_tuoi(df_raw)
    tong_to_tong_so = _dem_unique_to(df_raw)
    if df.empty:
        summary_empty = {
            "tong_to": 0,
            "tb_tuoi": None,
            "min_tuoi": None,
            "max_tuoi": None,
            "median_tuoi": None,
            "khong_co_du_lieu": tong_to_tong_so,
            "tong_to_tong_so": tong_to_tong_so,
        }
        cols = ["PGD"] + _TEN_BINS + ["Tổng tổ trưởng có dữ liệu"]
        return pd.DataFrame(columns=cols), summary_empty
    df_unique = _df_unique_theo_to(df)
    df_unique["_to_count"] = 1
    # Pivot theo PGD x nhóm tuổi
    df_pivot = df_unique.pivot_table(
        index="ten_pgd_std",
        columns="nhom_tuoi",
        values="_to_count",
        aggfunc="sum",
        fill_value=0,
    ).reset_index()
    # Đảm bảo đủ 6 bins (có thể thiếu nhãn nếu 0 tổ)
    for label in _TEN_BINS:
        if label not in df_pivot.columns:
            df_pivot[label] = 0
    df_pivot = df_pivot[["ten_pgd_std"] + _TEN_BINS].copy()
    df_pivot.columns = ["PGD"] + _TEN_BINS
    df_pivot["Tổng tổ trưởng có dữ liệu"] = df_pivot[_TEN_BINS].sum(axis=1)
    df_pivot = df_pivot.sort_values(by="Tổng tổ trưởng có dữ liệu", ascending=False).reset_index(drop=True)
    # Summary toàn tập hợp dữ liệu tuổi
    _ages = pd.to_numeric(df_unique["_tuoi_num"], errors="coerce").dropna()
    summary = {
        "tong_to": int(len(df_unique)),
        "tb_tuoi": round(float(_ages.mean()), 1) if not _ages.empty else None,
        "min_tuoi": int(_ages.min()) if not _ages.empty else None,
        "max_tuoi": int(_ages.max()) if not _ages.empty else None,
        "median_tuoi": round(float(_ages.median()), 1) if not _ages.empty else None,
        "khong_co_du_lieu": max(0, tong_to_tong_so - int(len(df_unique))),
        "tong_to_tong_so": tong_to_tong_so,
    }
    return df_pivot, summary


def thong_ke_tuoi_theo_xa(
    df_raw: "pd.DataFrame | None",
    pgd_loc: str | None = None,
) -> tuple["pd.DataFrame", dict]:
    """
    Thống kê Tuổi Tổ trưởng THEO Xã. Nếu pgd_loc (tên PGD) khác None → chỉ lọc PGD đó.
    Trả về: (df_bins, summary)
      - df_bins: ['PGD','Xã/Phường'] + 6 bins + ['Tổng tổ trưởng có dữ liệu']
      - summary: dict tương tự theo Xã (có thể thêm pgd_loc nếu đã lọc)
    """
    df = _df_chi_tiet_so_huu_tuoi(df_raw)
    df_raw_scope = df_raw
    if df_raw is not None and not df_raw.empty and pgd_loc:
        df_raw_std = _them_cot_chuan_to(df_raw)
        _pgd_norm_raw = str(pgd_loc).strip().lower()
        df_raw_scope = df_raw_std[df_raw_std["ten_pgd_std"].str.lower() == _pgd_norm_raw].copy()
    tong_to_tong_so = _dem_unique_to(df_raw_scope)
    if df.empty:
        summary_empty = {
            "tong_to": 0,
            "tb_tuoi": None,
            "min_tuoi": None,
            "max_tuoi": None,
            "median_tuoi": None,
            "khong_co_du_lieu": tong_to_tong_so,
            "tong_to_tong_so": tong_to_tong_so,
            "pgd_loc": pgd_loc,
        }
        cols = ["PGD", "Xã/Phường"] + _TEN_BINS + ["Tổng tổ trưởng có dữ liệu"]
        return pd.DataFrame(columns=cols), summary_empty
    # Lọc theo PGD nếu yêu cầu
    if pgd_loc:
        _pgd_norm = str(pgd_loc).strip().lower()
        df = df[df["ten_pgd_std"].str.lower() == _pgd_norm].copy()
        if df.empty:
            summary_empty = {
                "tong_to": 0,
                "tb_tuoi": None,
                "min_tuoi": None,
                "max_tuoi": None,
                "median_tuoi": None,
                "khong_co_du_lieu": tong_to_tong_so,
                "tong_to_tong_so": tong_to_tong_so,
                "pgd_loc": pgd_loc,
            }
            cols = ["PGD", "Xã/Phường"] + _TEN_BINS + ["Tổng tổ trưởng có dữ liệu"]
            return pd.DataFrame(columns=cols), summary_empty
    df_unique = _df_unique_theo_to(df)
    df_unique["_to_count"] = 1
    # Pivot theo (PGD, Xã) x nhóm tuổi: đếm unique tổ
    df_pivot = df_unique.pivot_table(
        index=["ten_pgd_std", "ten_xa_std"],
        columns="nhom_tuoi",
        values="_to_count",
        aggfunc="sum",
        fill_value=0,
    ).reset_index()
    for label in _TEN_BINS:
        if label not in df_pivot.columns:
            df_pivot[label] = 0
    df_pivot = df_pivot[["ten_pgd_std", "ten_xa_std"] + _TEN_BINS].copy()
    df_pivot.columns = ["PGD", "Xã/Phường"] + _TEN_BINS
    df_pivot["Tổng tổ trưởng có dữ liệu"] = df_pivot[_TEN_BINS].sum(axis=1)
    df_pivot = df_pivot.sort_values(
        by=["Tổng tổ trưởng có dữ liệu", "PGD", "Xã/Phường"],
        ascending=[False, True, True],
    ).reset_index(drop=True)
    _ages = pd.to_numeric(df_unique["_tuoi_num"], errors="coerce").dropna()
    summary = {
        "tong_to": int(len(df_unique)),
        "tb_tuoi": round(float(_ages.mean()), 1) if not _ages.empty else None,
        "min_tuoi": int(_ages.min()) if not _ages.empty else None,
        "max_tuoi": int(_ages.max()) if not _ages.empty else None,
        "median_tuoi": round(float(_ages.median()), 1) if not _ages.empty else None,
        "khong_co_du_lieu": max(0, tong_to_tong_so - int(len(df_unique))),
        "tong_to_tong_so": tong_to_tong_so,
        "pgd_loc": pgd_loc,
    }
    return df_pivot, summary


def _tao_word_thong_ke_tuoi_to_truong(
    df_bins: "pd.DataFrame",
    summary: dict,
    che_do_xem: str = "theo_pgd",
    tieu_de_pham_vi: str = "Toàn Chi nhánh",
    ten_pgd: str | None = None,
) -> bytes:
    """
    Tạo file Word (.docx) BÁO CÁO THỐNG KÊ TUỔI TỔ TRƯỞNG TK&VV.
    Layout chuẩn hành chính:
      - Header 2 cột (Tên đơn vị + Cộng hòa XHCN / Ngày tháng)
      - Tiêu đề báo cáo + Phạm vi
      - Phần TÓM TẮT: 4 thẻ (Tổng số tổ / Có dữ liệu / Chưa có / TB tuổi + Range)
      - Bảng dữ liệu CHI TIẾT: 6 bins độ tuổi + Tổng cộng
      - Footer: Người lập + Chức vụ ký tên
    Returns: bytes of .docx
    """
    import io
    from datetime import date as _date
    from docx import Document
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from config import TEN_CHI_NHANH_HIEN_THI

    BIN_LABELS = ["Dưới 30 tuổi", "30 - 39 tuổi", "40 - 49 tuổi",
                  "50 - 59 tuổi", "60 - 69 tuổi", "Từ 70 tuổi trở lên"]
    HDR_FILL = "BDD7EE"
    TONG_FILL = "FFF2CC"

    def _set_font(run, bold: bool = False, size: int = 13, italic: bool = False) -> None:
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        try:
            rPr = run._r.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn("w:ascii"), "Times New Roman")
            rFonts.set(qn("w:hAnsi"), "Times New Roman")
            rFonts.set(qn("w:eastAsia"), "Times New Roman")
        except Exception:
            pass

    def _shade_cell(cell, fill: str = HDR_FILL) -> None:
        try:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            for old in tcPr.findall(qn("w:shd")):
                tcPr.remove(old)
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill)
            tcPr.append(shd)
        except Exception:
            pass

    def _remove_borders(table) -> None:
        try:
            for row in table.rows:
                for cell in row.cells:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    for old in tcPr.findall(qn("w:tcBorders")):
                        tcPr.remove(old)
                    tcBorders = OxmlElement("w:tcBorders")
                    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                        el = OxmlElement(f"w:{side}")
                        el.set(qn("w:val"), "nil")
                        tcBorders.append(el)
                    tcPr.append(tcBorders)
        except Exception:
            pass

    def _set_fixed_table_widths(table, widths_cm: list[float]) -> None:
        """Khóa tblW/tblGrid/tcW để Word không tự autofit làm vỡ bảng báo cáo."""
        try:
            widths_dxa = [str(int(Cm(w).twips)) for w in widths_cm]
            table.autofit = False

            tbl = table._tbl
            tbl_pr = tbl.tblPr
            for old in tbl_pr.findall(qn("w:tblW")):
                tbl_pr.remove(old)
            tbl_w = OxmlElement("w:tblW")
            tbl_w.set(qn("w:type"), "dxa")
            tbl_w.set(qn("w:w"), str(sum(int(w) for w in widths_dxa)))
            tbl_pr.insert(0, tbl_w)

            for old in tbl_pr.findall(qn("w:tblLayout")):
                tbl_pr.remove(old)
            layout = OxmlElement("w:tblLayout")
            layout.set(qn("w:type"), "fixed")
            tbl_pr.append(layout)

            for old in tbl.findall(qn("w:tblGrid")):
                tbl.remove(old)
            grid = OxmlElement("w:tblGrid")
            for w in widths_dxa:
                grid_col = OxmlElement("w:gridCol")
                grid_col.set(qn("w:w"), w)
                grid.append(grid_col)
            tbl.insert(1, grid)

            for row in table.rows:
                for idx, cell in enumerate(row.cells[: len(widths_dxa)]):
                    cell.width = Cm(widths_cm[idx])
                    tc_pr = cell._tc.get_or_add_tcPr()
                    for old in tc_pr.findall(qn("w:tcW")):
                        tc_pr.remove(old)
                    tc_w = OxmlElement("w:tcW")
                    tc_w.set(qn("w:type"), "dxa")
                    tc_w.set(qn("w:w"), widths_dxa[idx])
                    tc_pr.insert(0, tc_w)
        except Exception:
            pass

    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(1.8)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(13)

    today = _date.today()
    tuoi_ngay = f"Đồng Nai, ngày {today.day} tháng {today.month} năm {today.year}"

    hdr = doc.add_table(rows=1, cols=2)
    _remove_borders(hdr)
    cell_l = hdr.rows[0].cells[0]
    for i, (txt, bold) in enumerate([
        ("NGÂN HÀNG CHÍNH SÁCH XÃ HỘI", True),
        (TEN_CHI_NHANH_HIEN_THI, False),
        (ten_pgd if ten_pgd else "Phòng KH-NV", False),
        ("────────────────────────", False),
        ("Số:      /BC-TKTuổi", False),
    ]):
        p = cell_l.paragraphs[0] if i == 0 else cell_l.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(p.add_run(txt), bold=bold, size=12 if i == 0 else 11)

    cell_r = hdr.rows[0].cells[1]
    for i, (txt, bold) in enumerate([
        ("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", True),
        ("Độc lập - Tự do - Hạnh phúc", True),
        ("──────────────────────────────", False),
        (tuoi_ngay, False),
    ]):
        p = cell_r.paragraphs[0] if i == 0 else cell_r.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(p.add_run(txt), bold=bold, size=12 if i == 0 else 11)
    _set_fixed_table_widths(hdr, [8.6, 8.6])

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(p.add_run("BÁO CÁO THỐNG KÊ TUỔI TỔ TRƯỞNG TIẾT KIỆM & VAY VỐN"), bold=True, size=16)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cd_txt = "Theo Đơn vị / PGD" if che_do_xem == "theo_pgd" else "Theo Xã/Phường"
    _set_font(p.add_run(f"{cd_txt} · Phạm vi: {tieu_de_pham_vi}"), bold=False, size=13, italic=True)
    doc.add_paragraph()

    # TÓM TẮT 2x2
    tong_to_tong_so = int(summary.get("tong_to_tong_so") or 0)
    tong_to_co_dl  = int(summary.get("tong_to") or 0)
    khong_co_dl    = int(summary.get("khong_co_du_lieu") or 0)
    tb_tuoi        = summary.get("tb_tuoi")
    min_tuoi       = summary.get("min_tuoi")
    max_tuoi       = summary.get("max_tuoi")
    median_tuoi    = summary.get("median_tuoi")
    def _fmt_age(v):
        return "—" if v is None else f"{v:g}"

    kpi_rows = [
        (
            "Tổng số Tổ trưởng trong dữ liệu",
            f"{tong_to_tong_so:,}  tổ".replace(",", "."),
        ),
        (
            "Số Tổ trưởng có dữ liệu độ tuổi",
            f"{tong_to_co_dl:,}  tổ".replace(",", "."),
        ),
        (
            "Số Tổ trưởng chưa có dữ liệu",
            f"{khong_co_dl:,}  tổ".replace(",", "."),
        ),
        (
            "Tuổi trung bình (Median / Min – Max)",
            f"{_fmt_age(tb_tuoi)} tuổi    (Median {_fmt_age(median_tuoi)} · {_fmt_age(min_tuoi)} – {_fmt_age(max_tuoi)} tuổi)",
        ),
    ]
    kpi_tbl = doc.add_table(rows=2, cols=2)
    _remove_borders(kpi_tbl)
    for col_idx, w_cm in [(0, 8.6), (1, 8.6)]:
        for cell in kpi_tbl.columns[col_idx].cells:
            cell.width = Cm(w_cm)
    for idx, (lbl, val) in enumerate(kpi_rows):
        r_idx, c_idx = divmod(idx, 2)
        cell = kpi_tbl.rows[r_idx].cells[c_idx]
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r1 = p1.add_run(lbl + ": ")
        _set_font(r1, bold=False, size=12, italic=True)
        r2 = p1.add_run(val)
        _set_font(r2, bold=True, size=14)
        _shade_cell(cell, fill="DEEBF7")
    _set_fixed_table_widths(kpi_tbl, [8.6, 8.6])

    doc.add_paragraph()
    p = doc.add_paragraph()
    _set_font(p.add_run("I. Bảng chi tiết phân bổ số lượng Tổ trưởng theo nhóm độ tuổi"), bold=True, size=14)
    doc.add_paragraph()

    if df_bins is None or df_bins.empty:
        p = doc.add_paragraph()
        _set_font(p.add_run("⚠️ Chưa có dòng nào có dữ liệu tuổi tổ trưởng hợp lệ."), italic=True, size=12)
    else:
        present_bins = [b for b in BIN_LABELS if b in df_bins.columns]
        id_cols  = [c for c in ["PGD", "Xã/Phường"] if c in df_bins.columns]
        sum_col  = "Tổng tổ trưởng có dữ liệu"
        all_cols = id_cols + present_bins + ([sum_col] if sum_col in df_bins.columns else [])
        df_show  = df_bins[all_cols].copy().reset_index(drop=True)

        if che_do_xem == "theo_pgd":
            widths_cm = [4.5] + [1.55] * len(present_bins) + [2.2]
        else:
            widths_cm = [3.0, 3.0] + [1.3] * len(present_bins) + [2.0]
        if len(widths_cm) > len(all_cols):
            widths_cm = widths_cm[: len(all_cols)]

        tbl = doc.add_table(rows=1, cols=len(all_cols))
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, w in enumerate(widths_cm):
            for cell in tbl.columns[i].cells:
                cell.width = Cm(w)

        hdr_row = tbl.rows[0]
        for i, h in enumerate(all_cols):
            cell = hdr_row.cells[i]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(h)
            _set_font(run, bold=True, size=10)
            _shade_cell(cell, HDR_FILL)

        is_xa_mode = "Xã/Phường" in df_show.columns
        n_id_cols = len(id_cols)
        n_bins = len(present_bins)
        for ridx in range(len(df_show)):
            row = tbl.add_row()
            rec = df_show.iloc[ridx]
            for c in range(n_id_cols):
                cell = row.cells[c]
                p = cell.paragraphs[0]
                p.alignment = (WD_ALIGN_PARAGRAPH.LEFT
                               if (c == 0 and not is_xa_mode) or (c == 1 and is_xa_mode)
                               else WD_ALIGN_PARAGRAPH.CENTER)
                v_text = rec[all_cols[c]]
                _set_font(p.add_run(str(v_text) if pd.notna(v_text) else ""), size=10)
            for offset in range(n_bins):
                c = n_id_cols + offset
                cell = row.cells[c]
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                try:
                    v = int(rec[all_cols[c]]) if pd.notna(rec[all_cols[c]]) else 0
                except Exception:
                    v = 0
                _set_font(p.add_run(f"{v:,}".replace(",", ".")), size=10)
            if sum_col in all_cols:
                c_sum = n_id_cols + n_bins
                cell = row.cells[c_sum]
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                try:
                    v_sum = int(rec[sum_col]) if pd.notna(rec[sum_col]) else 0
                except Exception:
                    v_sum = 0
                la_tong_cong = (
                    str(rec.get("PGD", "")).strip().startswith("🌐")
                    or str(rec.get("PGD", "")).strip().lower() in {"tổng cộng", "tong cong", "tong_cong"}
                )
                r = p.add_run(f"{v_sum:,}".replace(",", "."))
                _set_font(r, bold=la_tong_cong, size=10)
                if la_tong_cong:
                    for _c_c in range(len(all_cols)):
                        _shade_cell(row.cells[_c_c], TONG_FILL)
        _set_fixed_table_widths(tbl, widths_cm)

        doc.add_paragraph()
        p = doc.add_paragraph()
        _set_font(
            p.add_run(
                "Ghi chú: Số liệu đếm theo unique Tổ (mã tổ / tên tổ) để tránh đếm trùng "
                "khi file chấm điểm có dòng trùng lặp cùng một Tổ."
            ),
            italic=True, size=11,
        )

    doc.add_paragraph()
    doc.add_paragraph()
    foot = doc.add_table(rows=4, cols=2)
    _remove_borders(foot)
    foot_l = [
        f"Người lập báo cáo{(f' ({ten_pgd})' if ten_pgd else '')}",
        "",
        "(Ký, ghi rõ họ tên)",
        "",
    ]
    foot_r = [
        f"CHỦ TỊCH {'CHI NHÁNH' if not ten_pgd else 'PGD'}",
        "",
        "(Ký, đóng dấu, ghi rõ họ tên)",
        "",
    ]
    for i in range(4):
        pl = foot.rows[i].cells[0].paragraphs[0]
        pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(pl.add_run(foot_l[i]), bold=(i == 0), size=12)
        pr = foot.rows[i].cells[1].paragraphs[0]
        pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(pr.add_run(foot_r[i]), bold=(i == 0), size=12)
    _set_fixed_table_widths(foot, [8.6, 8.6])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
