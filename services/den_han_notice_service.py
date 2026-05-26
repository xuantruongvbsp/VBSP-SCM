"""Thông báo đến hạn — tạo thư nhắc nợ cho từng KH (ROADMAP §2.3)."""
from __future__ import annotations

import os
from datetime import date, datetime
from io import BytesIO

import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from data.core import _duckdb_query
from config import (
    CACHE_HSTD, TEN_CHI_NHANH_HIEN_THI,
    COT_TEN_PGD, COT_TEN_KH, COT_SO_KU, COT_TEN_CT,
    COT_TONG_DU_NO, COT_NGAY_DH, COT_DIA_CHI, COT_TEN_XA, COT_TEN_TO,
)


def _add_formatted_para(doc, text: str, bold: bool = False, size: int = 12, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    run.bold = bold
    return p


def tao_thu_nhac_no(
    ten_kh: str,
    dia_chi: str,
    so_ku: str,
    ten_ct: str,
    du_no: float,
    ngay_den_han,
    ten_pgd: str,
    ngay_lap: date | None = None,
) -> bytes:
    """Tạo 1 thư nhắc nợ cho 1 khách hàng."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    ngay_lap = ngay_lap or date.today()
    ngay_dh_str = ""
    if ngay_den_han is not None:
        if isinstance(ngay_den_han, str):
            ngay_dh_str = ngay_den_han[:10]
        elif hasattr(ngay_den_han, "strftime"):
            ngay_dh_str = ngay_den_han.strftime("%d/%m/%Y")
        else:
            ngay_dh_str = str(ngay_den_han)[:10]

    du_no_str = f"{du_no:,.0f}".replace(",", ".")

    # Header
    _add_formatted_para(doc, "NGÂN HÀNG CHÍNH SÁCH XÃ HỘI", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_formatted_para(doc, f"CHI NHÁNH {TEN_CHI_NHANH_HIEN_THI.upper()}", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_formatted_para(doc, f"PHÒNG GIAO DỊCH {ten_pgd.upper()}", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_formatted_para(doc, "—" * 40, align=WD_ALIGN_PARAGRAPH.CENTER)

    _add_formatted_para(doc, "")
    _add_formatted_para(doc, "THÔNG BÁO NỢ ĐẾN HẠN", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_formatted_para(doc, f"(Lần 1 · Ngày {ngay_lap.strftime('%d/%m/%Y')})", size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_formatted_para(doc, "")

    _add_formatted_para(doc, f"Kính gửi: Ông/Bà {ten_kh}")
    _add_formatted_para(doc, f"Địa chỉ: {dia_chi or '...'}")
    _add_formatted_para(doc, "")

    body = (
        f"Ngân hàng Chính sách xã hội — {TEN_CHI_NHANH_HIEN_THI} trân trọng thông báo "
        f"khoản vay của Ông/Bà sắp đến hạn thanh toán như sau:\n\n"
        f"  • Số khế ước: {so_ku}\n"
        f"  • Chương trình: {ten_ct}\n"
        f"  • Dư nợ hiện tại: {du_no_str} đồng\n"
        f"  • Ngày đến hạn: {ngay_dh_str}\n\n"
        f"Kính đề nghị Ông/Bà chủ động sắp xếp tài chính và đến Phòng Giao dịch "
        f"{ten_pgd} để thanh toán đúng hạn, tránh phát sinh nợ quá hạn ảnh hưởng "
        f"đến lịch sử tín dụng.\n\n"
        f"Mọi thắc mắc xin liên hệ CBTD phụ trách địa bàn hoặc Phòng Giao dịch "
        f"{ten_pgd} — ĐT: (0251) ... để được hỗ trợ."
    )
    for line in body.split("\n"):
        if line.strip():
            _add_formatted_para(doc, line, size=12)
        else:
            _add_formatted_para(doc, "", size=6)

    _add_formatted_para(doc, "")
    _add_formatted_para(doc, "")
    _add_formatted_para(doc, "Trân trọng thông báo!", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _add_formatted_para(doc, "")
    _add_formatted_para(doc, "GIÁM ĐỐC PHÒNG GIAO DỊCH", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.RIGHT)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def lay_ds_den_han(
    parquet_path: str | None = None,
    days_ahead: int = 30,
    max_rows: int = 200,
) -> pd.DataFrame:
    """Lấy danh sách khoản vay đến hạn trong N ngày tới."""
    parquet_path = parquet_path or str(CACHE_HSTD)
    if not os.path.exists(parquet_path):
        return pd.DataFrame()

    cutoff = (date.today() + __import__("datetime").timedelta(days=days_ahead)).isoformat()

    schema_cols = pd.read_parquet(parquet_path, engine='pyarrow').columns.tolist()
    select_cols = [
        COT_TEN_PGD, COT_TEN_KH, COT_SO_KU, COT_TEN_CT,
        COT_TONG_DU_NO, COT_NGAY_DH, COT_DIA_CHI, COT_TEN_XA,
    ]
    available = [c for c in select_cols if c in schema_cols]
    if COT_TEN_KH not in available or COT_TONG_DU_NO not in available:
        return pd.DataFrame()

    cols_sql = ", ".join(f'"{c}"' for c in available)
    sql = f"""
        SELECT {cols_sql}
        FROM read_parquet(?)
        WHERE "{COT_NGAY_DH}" IS NOT NULL
          AND "{COT_NGAY_DH}" >= CURRENT_DATE
          AND "{COT_NGAY_DH}" <= '{cutoff}'
          AND "{COT_TONG_DU_NO}" > 0
        ORDER BY "{COT_NGAY_DH}", "{COT_TONG_DU_NO}" DESC
        LIMIT {max_rows}
    """
    try:
        return _duckdb_query(sql, [parquet_path])
    except Exception:
        return pd.DataFrame()


def tao_thu_hang_loat(
    df: pd.DataFrame,
    ten_pgd_filter: str | None = None,
) -> list[dict]:
    """Từ DataFrame đến hạn, tạo danh sách thư nhắc nợ."""
    results = []
    df_loc = df.copy()
    if ten_pgd_filter and COT_TEN_PGD in df_loc.columns:
        df_loc = df_loc[df_loc[COT_TEN_PGD] == ten_pgd_filter]

    for _, row in df_loc.iterrows():
        du_no = float(row.get(COT_TONG_DU_NO, 0) or 0)
        ten_kh = str(row.get(COT_TEN_KH, ""))
        so_ku = str(row.get(COT_SO_KU, ""))
        ten_ct = str(row.get(COT_TEN_CT, ""))
        dia_chi = str(row.get(COT_DIA_CHI, ""))
        ten_pgd = str(row.get(COT_TEN_PGD, ""))
        ngay_dh = row.get(COT_NGAY_DH)

        try:
            doc_bytes = tao_thu_nhac_no(
                ten_kh=ten_kh, dia_chi=dia_chi, so_ku=so_ku,
                ten_ct=ten_ct, du_no=du_no, ngay_den_han=ngay_dh,
                ten_pgd=ten_pgd,
            )
            safe_name = "".join(c for c in ten_kh if c.isalnum() or c in " _-")[:30]
            filename = f"ThongBao_DenHan_{safe_name}_{so_ku[:8]}.docx"
            results.append({
                "ten_kh": ten_kh,
                "so_ku": so_ku,
                "du_no": du_no,
                "bytes": doc_bytes,
                "filename": filename,
            })
        except Exception:
            pass

    return results
