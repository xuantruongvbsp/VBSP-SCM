"""Dịch vụ xuất báo cáo Tổng hợp HSTD Word (.docx) — toàn Chi nhánh.

Bố cục file Word:
  Trang bìa: Header NHCSXH + Tiêu đề + Tháng/Năm + Người xuất
  Bảng 1: Tổng hợp theo chương trình (Số KH, Dư nợ, QH%, Khoanh, Lãi tồn)
  Bảng 2: Chi tiết 22 PGD (Số món, Số KH, Dư nợ, QH, TL QH%, Khoanh...)
  Biểu đồ: Bar chart QH theo PGD (dạng ảnh PNG nhúng)
  Khối chữ ký cuối
"""
from __future__ import annotations

from io import BytesIO
from datetime import datetime
from pathlib import Path
from typing import Sequence

import pandas as pd
import streamlit as st

from config import (
    COT_TEN_PGD, COT_MA_KH, COT_TONG_DU_NO, COT_DU_NO_QH, COT_DU_NO_KHOANH,
    COT_LAI_TON, COT_TEN_CT, DS_PGD, DON_VI_CHI_NHANH,
)
from utils import fmt, fmt_so
from logger import get_logger

logger = get_logger(__name__)

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    _DOCX_READY = True
except ImportError:
    _DOCX_READY = False

VBSP_GREEN = (0x2E, 0x7D, 0x32)
VBSP_GREEN_LIGHT = (0xE8, 0xF5, 0xE9)
RED_TEXT = (0xC6, 0x28, 0x28)
ORANGE_TEXT = (0xE6, 0x51, 0x00)
HEADER_BG = (0x2E, 0x7D, 0x32)
ROW_ALT = (0xF5, 0xF5, 0xF5)


def xuat_word_hstd_tong_hop(
    df: pd.DataFrame,
    thang: int | None = None,
    nam: int | None = None,
    username: str = "unknown",
    figs: list | None = None,
) -> bytes:
    """Tạo file Word Tổng hợp HSTD toàn Chi nhánh.

    Args:
        df: HSTD toàn CN (đã lọc kỳ)
        thang: Tháng báo cáo (None = lấy tháng hiện tại)
        nam: Năm báo cáo (None = lấy năm hiện tại)
        username: Người xuất
        figs: List[(bytes, caption)] — ảnh biểu đồ PNG (từ plotly)

    Returns:
        bytes: File .docx
    """
    if not _DOCX_READY:
        st.error("Chưa cài python-docx. Chạy: pip install python-docx")
        return b""
    if df is None or df.empty:
        st.warning("Không có dữ liệu HSTD để xuất Word.")
        return b""

    now = datetime.now()
    if thang is None:
        thang = now.month
    if nam is None:
        nam = now.year

    df_copy = df.copy()
    for c in [COT_TONG_DU_NO, COT_DU_NO_QH, COT_DU_NO_KHOANH, COT_LAI_TON]:
        if c in df_copy.columns:
            df_copy[c] = pd.to_numeric(df_copy[c], errors="coerce").fillna(0)

    doc = Document()

    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)

    _ve_trang_bia(doc, thang, nam, username)
    doc.add_page_break()

    _ve_bang_tong_hop_ct(doc, df_copy)
    doc.add_paragraph()
    doc.add_page_break()

    _ve_bang_chi_tiet_pgd(doc, df_copy)
    doc.add_paragraph()

    if figs:
        _ve_bieu_do(doc, figs)

    _ve_chu_ky(doc, username)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _p(
    doc: Document,
    text: str = "",
    bold: bool = False,
    size: int = 14,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    color: tuple = (0, 0, 0),
    space_before: float = 0,
    space_after: float = 4,
    font_name: str = "Times New Roman",
):
    para = doc.add_paragraph()
    para.alignment = align
    fmt_pf = para.paragraph_format
    fmt_pf.space_before = Pt(space_before)
    fmt_pf.space_after = Pt(space_after)
    if text:
        r = para.add_run(text)
        r.bold = bold
        r.font.name = font_name
        r.font.size = Pt(size)
        r.font.color.rgb = RGBColor(*color)
    return para


def _set_cell(cell, text: str, bold: bool = False, size: int = 9,
              align=WD_ALIGN_PARAGRAPH.CENTER, bg: tuple | None = None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    if bg:
        _set_cell_bg(cell, bg)


def _set_cell_bg(cell, color: tuple):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "%02X%02X%02X" % color)
    shading.set(qn("w:val"), "clear")
    tc_pr.append(shading)


def _ve_trang_bia(doc: Document, thang: int, nam: int, username: str):
    _p(doc, "", size=14, space_after=10)
    _p(doc, "NGÂN HÀNG CHÍNH SÁCH XÃ HỘI VIỆT NAM", bold=True, size=13,
       align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    _p(doc, "CHI NHÁNH TỈNH ĐỒNG NAI", bold=True, size=13,
       align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

    _p(doc, "—" * 50, size=9, align=WD_ALIGN_PARAGRAPH.CENTER,
       color=VBSP_GREEN, space_after=8)

    _p(doc, "", size=14, space_after=20)
    _p(doc, "BÁO CÁO TỔNG HỢP", bold=True, size=20,
       align=WD_ALIGN_PARAGRAPH.CENTER, color=VBSP_GREEN, space_after=4)
    _p(doc, "HỒ SƠ TÍN DỤNG TOÀN CHI NHÁNH", bold=True, size=16,
       align=WD_ALIGN_PARAGRAPH.CENTER, color=VBSP_GREEN, space_after=10)
    _p(doc, f"Tháng {thang}/{nam}", bold=True, size=14,
       align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

    _p(doc, "", size=14, space_after=40)
    now_str = datetime.now().strftime("%d/%m/%Y %H:%M")
    _p(doc, f"Ngày xuất: {now_str}", size=11,
       align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    _p(doc, f"Người xuất: {username}", size=11,
       align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    _p(doc, "Hệ thống VBSP-SCM", size=10,
       align=WD_ALIGN_PARAGRAPH.CENTER, color=(120, 120, 120))


def _ve_bang_tong_hop_ct(doc: Document, df: pd.DataFrame):
    _p(doc, "I. TỔNG HỢP THEO CHƯƠNG TRÌNH TÍN DỤNG", bold=True, size=13,
       color=VBSP_GREEN, space_after=8)

    dn = float(df[COT_TONG_DU_NO].sum())
    qh = float(df[COT_DU_NO_QH].sum())
    kh = float(df.get(COT_DU_NO_KHOANH, pd.Series([0])).sum())
    lt = float(df.get(COT_LAI_TON, pd.Series([0])).sum())
    n_kh = int(df[COT_MA_KH].nunique())
    n_mon = int(len(df))
    tl_qh = qh / dn * 100 if dn > 0 else 0.0
    tl_kh = kh / dn * 100 if dn > 0 else 0.0

    summary_text = (
        f"Tính đến hết tháng, toàn Chi nhánh quản lý {fmt_so(n_kh)} khách hàng "
        f"với {fmt_so(n_mon)} món vay. "
        f"Tổng dư nợ {fmt(int(dn))} triệu đồng. "
        f"Dư nợ quá hạn {fmt(int(qh))} triệu đồng, chiếm {tl_qh:.2f}%. "
        f"Dư nợ khoanh {fmt(int(kh))} triệu đồng, chiếm {tl_kh:.2f}%. "
        f"Lãi tồn {fmt(int(lt))} triệu đồng."
    )
    _p(doc, summary_text, size=11, space_after=8,
       align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    if COT_TEN_CT not in df.columns:
        _p(doc, "⚠️ Dữ liệu không có cột Chương trình tín dụng.", size=11,
           color=ORANGE_TEXT, space_after=8)
        return

    g = df.groupby(COT_TEN_CT).agg(
        so_mon=(COT_MA_KH, "count"),
        so_kh=(COT_MA_KH, "nunique"),
        du_no=(COT_TONG_DU_NO, "sum"),
        qh=(COT_DU_NO_QH, "sum"),
        khoanh=(COT_DU_NO_KHOANH, "sum") if COT_DU_NO_KHOANH in df.columns else (COT_MA_KH, lambda x: 0),
        lai_ton=(COT_LAI_TON, "sum") if COT_LAI_TON in df.columns else (COT_MA_KH, lambda x: 0),
    ).reset_index()
    g = g.sort_values("du_no", ascending=False).reset_index(drop=True)

    headers = ["Stt", "Chương trình tín dụng", "Số món", "Số KH",
               "Dư nợ (trđ)", "QH (trđ)", "QH%", "Khoanh (trđ)", "Lãi tồn (trđ)"]
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        _set_cell(tbl.rows[0].cells[i], h, bold=True, size=8, bg=HEADER_BG)
        tbl.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for idx, (_, row) in enumerate(g.iterrows()):
        tl_qh_ct = row["qh"] / row["du_no"] * 100 if row["du_no"] > 0 else 0
        vals = [
            str(idx + 1),
            str(row[COT_TEN_CT]),
            fmt_so(int(row["so_mon"])),
            fmt_so(int(row["so_kh"])),
            fmt(int(row["du_no"])),
            fmt(int(row["qh"])),
            f"{tl_qh_ct:.1f}%",
            fmt(int(row["khoanh"])),
            fmt(int(row["lai_ton"])),
        ]
        tr = tbl.add_row()
        for i, v in enumerate(vals):
            _set_cell(tr.cells[i], v, size=8,
                      align=WD_ALIGN_PARAGRAPH.LEFT if i == 1 else WD_ALIGN_PARAGRAPH.CENTER)
        if tl_qh_ct > 5:
            tr.cells[6].paragraphs[0].runs[0].font.color.rgb = RGBColor(*RED_TEXT)
        elif tl_qh_ct > 3:
            tr.cells[6].paragraphs[0].runs[0].font.color.rgb = RGBColor(*ORANGE_TEXT)
        if idx % 2 == 1:
            for i in range(len(headers)):
                _set_cell_bg(tr.cells[i], ROW_ALT)

    tr = tbl.add_row()
    cong_vals = ["", "TỔNG CỘNG",
                 fmt_so(int(g["so_mon"].sum())),
                 fmt_so(int(g["so_kh"].sum())),
                 fmt(int(g["du_no"].sum())),
                 fmt(int(g["qh"].sum())), "",
                 fmt(int(g["khoanh"].sum())),
                 fmt(int(g["lai_ton"].sum()))]
    for i, v in enumerate(cong_vals):
        _set_cell(tr.cells[i], v, bold=True, size=8,
                  align=WD_ALIGN_PARAGRAPH.LEFT if i == 1 else WD_ALIGN_PARAGRAPH.CENTER,
                  bg=VBSP_GREEN_LIGHT)


def _ve_bang_chi_tiet_pgd(doc: Document, df: pd.DataFrame):
    _p(doc, "II. CHI TIẾT THEO PHÒNG GIAO DỊCH", bold=True, size=13,
       color=VBSP_GREEN, space_after=8)

    pgd_stats = df.groupby(COT_TEN_PGD).agg(
        so_mon=(COT_MA_KH, "count"),
        so_kh=(COT_MA_KH, "nunique"),
        du_no=(COT_TONG_DU_NO, "sum"),
        qh=(COT_DU_NO_QH, "sum"),
        khoanh=(COT_DU_NO_KHOANH, "sum") if COT_DU_NO_KHOANH in df.columns else (COT_MA_KH, lambda x: 0),
        lai_ton=(COT_LAI_TON, "sum") if COT_LAI_TON in df.columns else (COT_MA_KH, lambda x: 0),
    ).reset_index()

    pgd_stats["tl_qh"] = pgd_stats.apply(
        lambda r: r["qh"] / r["du_no"] * 100 if r["du_no"] > 0 else 0, axis=1
    )
    pgd_stats = pgd_stats.sort_values("du_no", ascending=False).reset_index(drop=True)

    headers = ["Stt", "Phòng giao dịch", "Số món", "Số KH",
               "Dư nợ (trđ)", "QH (trđ)", "QH%", "Khoanh (trđ)", "Lãi tồn (trđ)"]
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        _set_cell(tbl.rows[0].cells[i], h, bold=True, size=8, bg=HEADER_BG)
        tbl.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for idx, (_, row) in enumerate(pgd_stats.iterrows()):
        vals = [
            str(idx + 1),
            str(row[COT_TEN_PGD]),
            fmt_so(int(row["so_mon"])),
            fmt_so(int(row["so_kh"])),
            fmt(int(row["du_no"])),
            fmt(int(row["qh"])),
            f"{row['tl_qh']:.1f}%",
            fmt(int(row["khoanh"])),
            fmt(int(row["lai_ton"])),
        ]
        tr = tbl.add_row()
        for i, v in enumerate(vals):
            _set_cell(tr.cells[i], v, size=8,
                      align=WD_ALIGN_PARAGRAPH.LEFT if i == 1 else WD_ALIGN_PARAGRAPH.CENTER)
        if row["tl_qh"] > 5:
            tr.cells[6].paragraphs[0].runs[0].font.color.rgb = RGBColor(*RED_TEXT)
        elif row["tl_qh"] > 3:
            tr.cells[6].paragraphs[0].runs[0].font.color.rgb = RGBColor(*ORANGE_TEXT)
        if idx % 2 == 1:
            for i in range(len(headers)):
                _set_cell_bg(tr.cells[i], ROW_ALT)

    tr = tbl.add_row()
    cong_vals = ["", "TOÀN CHI NHÁNH",
                 fmt_so(int(pgd_stats["so_mon"].sum())),
                 fmt_so(int(pgd_stats["so_kh"].sum())),
                 fmt(int(pgd_stats["du_no"].sum())),
                 fmt(int(pgd_stats["qh"].sum())), "",
                 fmt(int(pgd_stats["khoanh"].sum())),
                 fmt(int(pgd_stats["lai_ton"].sum()))]
    for i, v in enumerate(cong_vals):
        _set_cell(tr.cells[i], v, bold=True, size=8,
                  align=WD_ALIGN_PARAGRAPH.LEFT if i == 1 else WD_ALIGN_PARAGRAPH.CENTER,
                  bg=VBSP_GREEN_LIGHT)

    _p(doc, "", size=8, space_after=4)
    _p(doc, "Dữ liệu hiển thị: Dư nợ, QH, Khoanh, Lãi tồn tính bằng triệu đồng (trđ).",
       size=8, color=(120, 120, 120))


def _ve_bieu_do(doc: Document, figs: list):
    _p(doc, "", size=8, space_after=4)
    _p(doc, "III. BIỂU ĐỒ", bold=True, size=13,
       color=VBSP_GREEN, space_after=8)

    for i, (png_bytes, caption) in enumerate(figs):
        if png_bytes and len(png_bytes) > 100:
            try:
                img_stream = BytesIO(png_bytes)
                doc.add_picture(img_stream, width=Inches(5.5))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                if caption:
                    _p(doc, caption, size=9,
                       align=WD_ALIGN_PARAGRAPH.CENTER,
                       color=(120, 120, 120), space_after=8)
            except Exception as e:
                logger.error("Lỗi chèn biểu đồ %d: %s", i + 1, e, exc_info=True)
                _p(doc, f"⚠️ Không chèn được biểu đồ: {caption}", size=9,
                   color=RED_TEXT)


def _ve_chu_ky(doc: Document, username: str):
    _p(doc, "", size=8, space_after=12)
    _p(doc, "IV. KÝ DUYỆT", bold=True, size=13,
       color=VBSP_GREEN, space_after=8)

    tbl = doc.add_table(rows=2, cols=3)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    ky_labels = ["NGƯỜI LẬP BIỂU", "KIỂM SOÁT", "GIÁM ĐỐC"]
    ky_subs = ["(Ký, ghi rõ họ tên)", "(Ký, ghi rõ họ tên)", "(Ký, ghi rõ họ tên)"]

    for i, label in enumerate(ky_labels):
        _set_cell(tbl.rows[0].cells[i], label, bold=True, size=10)
    for i, sub in enumerate(ky_subs):
        _set_cell(tbl.rows[1].cells[i], sub, size=8)
        tbl.rows[1].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(120, 120, 120)

    _p(doc, "", size=8, space_after=4)
    now_str = datetime.now().strftime("%d/%m/%Y %H:%M")
    _p(doc, f"Xuất lúc {now_str} bởi {username} — Hệ thống VBSP-SCM",
       size=8, align=WD_ALIGN_PARAGRAPH.CENTER, color=(120, 120, 120))
