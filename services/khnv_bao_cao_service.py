"""Dịch vụ tổng hợp dữ liệu cho báo cáo KHNV hàng tháng.

Tích hợp 2 nguồn:
- HSTD (df_full): dữ liệu chi tiết từng món vay
- ĐIỆN BÁO: số liệu tổng hợp từ file Điện báo CN (matrix 22 đơn vị)

Cung cấp:
- tong_hop_so_lieu_thang():     Dict các metric tổng hợp từ HSTD
- tong_hop_tu_dienbao():        Đọc Điện báo → dict metric + bảng theo PGD
- so_sanh_hstd_vs_dienbao():    Đối chiếu 2 nguồn, cảnh báo chênh lệch
- xuat_excel_bao_cao_khnv():    Xuất file Excel
- xuat_word_bao_cao_khnv():     Xuất file Word
- lay_danh_sach_mau():          Quét thư mục mẫu MD
"""
from __future__ import annotations

from datetime import date, datetime
from io import BytesIO
from pathlib import Path
from typing import Optional
import os

import pandas as pd

from config import (
    COT_TEN_PGD, COT_TONG_DU_NO, COT_DU_NO_TH, COT_DU_NO_QH,
    COT_DU_NO_KHOANH, COT_TEN_CT, COT_NGUON_VON, COT_DVUT,
    COT_TEN_KH, COT_GIAI_NGAN_TRONG_THANG, COT_NGAY_VAY,
    COT_TEN_XA, DB_HT_CACHE, DB_PREV_CACHE, FILE_PATH_DB,
    DS_PGD, TEN_CHI_NHANH_HIEN_THI,
)
from logger import get_logger

logger = get_logger(__name__)

MAU_BAO_CAO_DIR = Path(__file__).parent.parent / "docs" / "MAU BAO CAO KHNV"


def _fmt(v) -> str:
    try:
        n = float(v)
        if abs(n) >= 1_000_000_000:
            return f"{n/1_000_000_000:,.2f} tỷ"
        elif abs(n) >= 1_000_000:
            return f"{n/1_000_000:,.0f} tr"
        return f"{n:,.0f}"
    except (ValueError, TypeError):
        return str(v)


# ══════════════════════════════════════════════════════════════════════════════
# PHẦN 1: TỔNG HỢP TỪ HSTD
# ══════════════════════════════════════════════════════════════════════════════

def tong_hop_so_lieu_thang(
    df: pd.DataFrame,
    thang: Optional[int] = None,
    nam: Optional[int] = None,
) -> dict:
    """Tổng hợp các chỉ số chính từ HSTD."""
    if df is None or df.empty:
        return {}

    today = date.today()
    thang = thang or today.month
    nam = nam or today.year

    df = df.copy()
    tong_du_no = pd.to_numeric(df[COT_TONG_DU_NO], errors="coerce").sum()
    du_no_th = pd.to_numeric(df[COT_DU_NO_TH], errors="coerce").sum()
    du_no_qh = pd.to_numeric(df[COT_DU_NO_QH], errors="coerce").sum()
    du_no_khoanh = pd.to_numeric(df[COT_DU_NO_KHOANH], errors="coerce").sum() if COT_DU_NO_KHOANH in df.columns else 0
    so_kh = df[COT_TEN_KH].nunique() if COT_TEN_KH in df.columns else len(df)
    so_mon = len(df)
    tl_qh = (du_no_qh / tong_du_no * 100) if tong_du_no > 0 else 0

    nguon_tw = 0.0
    nguon_dp = 0.0
    if COT_NGUON_VON in df.columns:
        mask_tw = df[COT_NGUON_VON].astype(str).str.strip().isin(["1", "TW", "Trung ương"])
        mask_dp = df[COT_NGUON_VON].astype(str).str.strip().isin(["2", "ĐP", "Địa phương"])
        nguon_tw = pd.to_numeric(df.loc[mask_tw, COT_TONG_DU_NO], errors="coerce").sum()
        nguon_dp = pd.to_numeric(df.loc[mask_dp, COT_TONG_DU_NO], errors="coerce").sum()

    giai_nggan = pd.to_numeric(df[COT_GIAI_NGAN_TRONG_THANG], errors="coerce").sum() if COT_GIAI_NGAN_TRONG_THANG in df.columns else 0

    return {
        "thang": thang, "nam": nam,
        "ngay_bao_cao": today.strftime("%d/%m/%Y"),
        "tong_du_no": tong_du_no, "du_no_trong_han": du_no_th,
        "du_no_qua_han": du_no_qh, "du_no_khoanh": du_no_khoanh,
        "so_khach_hang": so_kh, "so_mon_vay": so_mon,
        "ty_le_no_qua_han": round(tl_qh, 2),
        "nguon_von_tw": nguon_tw, "nguon_von_dp": nguon_dp,
        "giai_ngan_trong_thang": giai_nggan,
        "nguon": "HSTD",
        "bang_pgd": _bang_theo_pgd(df),
        "bang_chuong_trinh": _bang_theo_chuong_trinh(df),
        "bang_uy_thac": _bang_theo_uy_thac(df),
    }


def _bang_theo_pgd(df: pd.DataFrame) -> pd.DataFrame:
    if COT_TEN_PGD not in df.columns:
        return pd.DataFrame()
    grouped = df.groupby(COT_TEN_PGD).agg(
        Tổng_dư_nợ=(COT_TONG_DU_NO, lambda x: pd.to_numeric(x, errors="coerce").sum()),
        Dư_nợ_trong_hạn=(COT_DU_NO_TH, lambda x: pd.to_numeric(x, errors="coerce").sum()),
        Dư_nợ_quá_hạn=(COT_DU_NO_QH, lambda x: pd.to_numeric(x, errors="coerce").sum()),
        Số_khách_hàng=(COT_TEN_KH, "nunique") if COT_TEN_KH in df.columns else (COT_TONG_DU_NO, "count"),
    ).reset_index()
    grouped["Tỷ_lệ_QH_%"] = grouped.apply(
        lambda r: round(r["Dư_nợ_quá_hạn"] / r["Tổng_dư_nợ"] * 100, 2) if r["Tổng_dư_nợ"] > 0 else 0, axis=1
    )
    return grouped.sort_values("Tổng_dư_nợ", ascending=False)


def _bang_theo_chuong_trinh(df: pd.DataFrame) -> pd.DataFrame:
    if COT_TEN_CT not in df.columns:
        return pd.DataFrame()
    grouped = df.groupby(COT_TEN_CT).agg(
        Tổng_dư_nợ=(COT_TONG_DU_NO, lambda x: pd.to_numeric(x, errors="coerce").sum()),
        Số_món=(COT_TONG_DU_NO, "count"),
    ).reset_index()
    return grouped.sort_values("Tổng_dư_nợ", ascending=False)


def _bang_theo_uy_thac(df: pd.DataFrame) -> pd.DataFrame:
    if COT_DVUT not in df.columns:
        return pd.DataFrame()
    grouped = df.groupby(COT_DVUT).agg(
        Tổng_dư_nợ=(COT_TONG_DU_NO, lambda x: pd.to_numeric(x, errors="coerce").sum()),
        Số_khách_hàng=(COT_TEN_KH, "nunique") if COT_TEN_KH in df.columns else (COT_TONG_DU_NO, "count"),
    ).reset_index()
    return grouped.sort_values("Tổng_dư_nợ", ascending=False)


# ══════════════════════════════════════════════════════════════════════════════
# PHẦN 2: TỔNG HỢP TỪ ĐIỆN BÁO
# ══════════════════════════════════════════════════════════════════════════════

CAC_CHI_TIEU_DIEN_BAO = {
    "tong_du_no":            "Tổng dư nợ",
    "du_no_kha":             "Dư nợ Kế hoạch A",
    "du_no_khb":             "Dư nợ Kế hoạch B",
    "du_no_qua_han_kha":     "Dư nợ Quá hạn KHA",
    "du_no_qua_han_khb":     "Dư nợ Quá hạn KHB",
    "du_no_khoanh_kha":      "Dư nợ Khoanh KHA",
    "du_no_khoanh_khb":      "Dư nợ Khoanh KHB",
    "nguon_tw_kha":          "Nguồn vốn cân đối từ TW (KHA)",
    "huy_dong_von":          "Tổng huy động vốn",
    "utdt_dp":               "Nguồn vốn nhận UTĐT tại ĐP",
    "von_an_toan":           "Vốn An toàn",
    "ton_quy":               "Tồn quỹ tiền mặt",
    "tg_nhnn":               "Tiền gửi tại NHNN và TCTD",
}


def _tim_file_dienbao() -> str | None:
    """Tìm file Điện báo hiện tại: cache trước, data/ sau."""
    if os.path.exists(DB_HT_CACHE):
        return DB_HT_CACHE
    if os.path.exists(FILE_PATH_DB):
        return FILE_PATH_DB
    return None


def _tim_file_dienbao_prev() -> str | None:
    if os.path.exists(DB_PREV_CACHE):
        return DB_PREV_CACHE
    return None


def tong_hop_tu_dienbao(sheet_name: str = "DB1") -> dict:
    """Đọc file Điện báo → tổng hợp các chỉ số chính + bảng theo đơn vị.

    Args:
        sheet_name: Sheet để đọc (mặc định DB1 - format dọc tiện cho lookup)

    Returns:
        dict: {
            "tong_du_no", "du_no_kha", "du_no_khb",
            "du_no_qua_han_kha", "du_no_qua_han_khb",
            "nguon_tw_kha", "huy_dong_von", "utdt_dp",
            "nguon": "Điện báo",
            "ngay_bao_cao": str,
            "file_path": str,
            "matrix": {ten_ct: {ten_dv: float}},  # ma trận đầy đủ
            "units": list[str],
            "bang_theo_dv": pd.DataFrame,
        }
    """
    from data.hstd import doc_dienbao_matrix, db_lookup

    fp = _tim_file_dienbao()
    if not fp:
        return {"nguon": "Điện báo", "error": "Chưa có file Điện báo"}

    try:
        data = doc_dienbao_matrix(fp, 0, sheet_name=sheet_name)
    except Exception as e:
        # Fallback: đọc sheet đầu tiên có matrix format
        try:
            from data.hstd import liet_ke_sheet_dienbao
            sheets = liet_ke_sheet_dienbao(fp)
            matrix_sheets = [s["sheet"] for s in sheets if s["format"] == "matrix"]
            if not matrix_sheets:
                # Thử format cũ
                from data.hstd import doc_dienbao
                rows = doc_dienbao(fp, 0)
                return _tong_hop_tu_format_cu(rows, fp)
            data = doc_dienbao_matrix(fp, 0, sheet_name=matrix_sheets[0])
        except Exception as e2:
            logger.error("Điện báo: %s / %s", e, e2)
            return {"nguon": "Điện báo", "error": f"Lỗi đọc: {e}"}

    rows = data.get("rows", [])
    result = {"nguon": "Điện báo", "file_path": fp, "ngay_bao_cao": data.get("ngay_bao_cao", "")}

    for key, ten_ct in CAC_CHI_TIEU_DIEN_BAO.items():
        result[key] = db_lookup(rows, ten_ct)

    # Ma trận đầy đủ
    result["matrix"] = data.get("matrix", {})
    result["units"] = data.get("units", [])

    # Bảng tổng hợp theo đơn vị
    bang_dv = _bang_theo_dv_tu_matrix(data.get("matrix", {}), data.get("units", []),
                                       ["Tổng dư nợ", "Dư nợ Kế hoạch A", "Dư nợ Kế hoạch B",
                                        "Dư nợ Quá hạn KHA", "Dư nợ Quá hạn KHB"])
    result["bang_theo_dv"] = bang_dv

    return result


def _tong_hop_tu_format_cu(rows: list, fp: str) -> dict:
    """Fallback: đọc Điện báo format cũ (chỉ có cột Cộng)."""
    from data.hstd import db_lookup
    result = {"nguon": "Điện báo (format cũ)", "file_path": fp}
    for key, ten_ct in CAC_CHI_TIEU_DIEN_BAO.items():
        result[key] = db_lookup(rows, ten_ct)
    return result


def _bang_theo_dv_tu_matrix(
    matrix: dict,
    units: list[str],
    chi_tieu_quan_tam: list[str],
) -> pd.DataFrame:
    """Tạo bảng tổng hợp theo đơn vị từ ma trận Điện báo."""
    if not matrix or not units:
        return pd.DataFrame()
    rows_list = []
    for dv in units:
        row = {"Đơn_vị": dv}
        for ct in chi_tieu_quan_tam:
            row[ct] = matrix.get(ct, {}).get(dv, 0)
        rows_list.append(row)
    return pd.DataFrame(rows_list)


# ══════════════════════════════════════════════════════════════════════════════
# PHẦN 3: SO SÁNH HSTD vs ĐIỆN BÁO
# ══════════════════════════════════════════════════════════════════════════════

def so_sanh_hstd_vs_dienbao(so_lieu_hstd: dict, so_lieu_db: dict) -> list[dict]:
    """Đối chiếu chỉ số giữa HSTD và Điện báo, trả về danh sách chênh lệch."""
    if not so_lieu_hstd or not so_lieu_db:
        return []

    mapping = [
        ("Tổng dư nợ",         "tong_du_no",      "tong_du_no"),
        ("Dư nợ quá hạn KHA",  "du_no_qua_han",   "du_no_qua_han_kha"),
    ]

    chenh_lech = []
    for ten, key_hstd, key_db in mapping:
        val_hstd = so_lieu_hstd.get(key_hstd, 0)
        val_db = so_lieu_db.get(key_db, 0)
        if val_hstd and val_db:
            cl = val_hstd - val_db
            tl = (cl / val_db * 100) if val_db else 0
            chenh_lech.append({
                "Chỉ tiêu": ten, "HSTD": val_hstd, "Điện báo": val_db,
                "Chênh lệch": cl, "Tỷ lệ %": round(tl, 2),
                "Cảnh báo": "⚠️" if abs(tl) > 1 else "✅",
            })

    return chenh_lech


# ══════════════════════════════════════════════════════════════════════════════
# PHẦN 4: DANH SÁCH MẪU BÁO CÁO
# ══════════════════════════════════════════════════════════════════════════════

def lay_danh_sach_mau() -> list[dict]:
    """Quét thư mục MAU BAO CAO KHNV, trả về list các mẫu có sẵn."""
    if not MAU_BAO_CAO_DIR.exists():
        return []
    result = []
    for f in sorted(MAU_BAO_CAO_DIR.glob("*.md")):
        ten_hien_thi = f.stem.replace("_", " ")
        result.append({"ten_hien_thi": ten_hien_thi, "ten_file": f.name, "path": str(f)})
    return result


def doc_noi_dung_mau(ten_file: str) -> str:
    path = MAU_BAO_CAO_DIR / ten_file
    if not path.exists():
        return ""
    return path.read_text(encoding="utf-8")


# ══════════════════════════════════════════════════════════════════════════════
# PHẦN 5: XUẤT EXCEL
# ══════════════════════════════════════════════════════════════════════════════

def xuat_excel_bao_cao_khnv(
    so_lieu: dict,
    bang_pgd: pd.DataFrame,
    bang_ct: pd.DataFrame,
    bang_uy_thac: pd.DataFrame,
    bang_dienbao: pd.DataFrame | None = None,
    chenh_lech: list | None = None,
) -> bytes:
    """Xuất file Excel báo cáo KHNV."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment

    wb = Workbook()
    ws = wb.active
    ws.title = "Tổng quan"

    header_font = Font(name="Times New Roman", size=14, bold=True, color="2E75B6")
    bold_font = Font(name="Times New Roman", size=12, bold=True)
    normal_font = Font(name="Times New Roman", size=12)
    header_fill = PatternFill(start_color="2E75B6", end_color="2E75B6", fill_type="solid")

    ws.merge_cells("A1:D1")
    ws["A1"] = f"BÁO CÁO KHNV — Tháng {so_lieu.get('thang', '')}/{so_lieu.get('nam', '')}"
    ws["A1"].font = header_font
    ws["A1"].alignment = Alignment(horizontal="center")
    ws["A2"] = f"Nguồn: {so_lieu.get('nguon', '')} | Ngày: {so_lieu.get('ngay_bao_cao', '')}"
    ws["A2"].font = Font(name="Times New Roman", size=10, italic=True)

    metrics = [
        ("Tổng dư nợ", so_lieu.get("tong_du_no", 0)),
        ("Dư nợ trong hạn", so_lieu.get("du_no_trong_han", 0)),
        ("Dư nợ quá hạn", so_lieu.get("du_no_qua_han", 0)),
        ("Dư nợ khoanh", so_lieu.get("du_no_khoanh", 0)),
        ("Tỷ lệ nợ quá hạn", f"{so_lieu.get('ty_le_no_qua_han', 0)}%"),
        ("Số khách hàng", so_lieu.get("so_khach_hang", 0)),
        ("Số món vay", so_lieu.get("so_mon_vay", 0)),
        ("Nguồn vốn TW", so_lieu.get("nguon_von_tw", 0)),
        ("Nguồn vốn ĐP", so_lieu.get("nguon_von_dp", 0)),
        ("Giải ngân trong tháng", so_lieu.get("giai_ngan_trong_thang", 0)),
    ]

    for i, (label, value) in enumerate(metrics, start=4):
        ws.cell(row=i, column=1, value=label).font = bold_font
        if isinstance(value, (int, float)):
            ws.cell(row=i, column=2, value=value).font = normal_font
            ws.cell(row=i, column=2).number_format = "#,##0"
        else:
            ws.cell(row=i, column=2, value=value).font = normal_font

    ws.column_dimensions["A"].width = 30
    ws.column_dimensions["B"].width = 25

    # Dữ liệu Điện báo (nếu có)
    if so_lieu.get("nguon") == "Điện báo" and so_lieu.get("du_no_kha"):
        row_start = 15
        ws.cell(row=row_start, column=1, value="── Số liệu từ Điện báo ──").font = bold_font
        db_metrics = [
            ("Dư nợ KHA", so_lieu.get("du_no_kha", 0)),
            ("Dư nợ KHB", so_lieu.get("du_no_khb", 0)),
            ("NQH KHA", so_lieu.get("du_no_qua_han_kha", 0)),
            ("NQH KHB", so_lieu.get("du_no_qua_han_khb", 0)),
            ("Vốn TW (KHA)", so_lieu.get("nguon_tw_kha", 0)),
            ("Huy động vốn", so_lieu.get("huy_dong_von", 0)),
            ("UTĐT ĐP", so_lieu.get("utdt_dp", 0)),
        ]
        for j, (label, value) in enumerate(db_metrics, start=row_start + 1):
            ws.cell(row=j, column=1, value=label).font = bold_font
            ws.cell(row=j, column=2, value=float(value) if value else 0).font = normal_font
            ws.cell(row=j, column=2).number_format = "#,##0"

    # So sánh chênh lệch (nếu có)
    if chenh_lech:
        row_cl = 28
        ws.cell(row=row_cl, column=1, value="── Đối chiếu HSTD vs Điện báo ──").font = bold_font
        headers_cl = ["Chỉ tiêu", "HSTD", "Điện báo", "Chênh lệch", "Tỷ lệ %", "Cảnh báo"]
        for k, h in enumerate(headers_cl, 1):
            ws.cell(row=row_cl + 1, column=k, value=h).font = Font(color="FFFFFF", bold=True)
            ws.cell(row=row_cl + 1, column=k).fill = header_fill
        for m, item in enumerate(chenh_lech, row_cl + 2):
            for n, key in enumerate(["Chỉ tiêu", "HSTD", "Điện báo", "Chênh lệch", "Tỷ lệ %", "Cảnh báo"], 1):
                ws.cell(row=m, column=n, value=item.get(key, "")).font = normal_font

    def _add_df_sheet(wb, name, df):
        if df is None or df.empty:
            return
        ws2 = wb.create_sheet(name)
        for c_idx, col_name in enumerate(df.columns, 1):
            ws2.cell(row=1, column=c_idx, value=col_name).font = Font(bold=True, color="FFFFFF")
            ws2.cell(row=1, column=c_idx).fill = header_fill
        for r_idx, row in enumerate(df.itertuples(index=False), 2):
            for c_idx, value in enumerate(row, 1):
                ws2.cell(row=r_idx, column=c_idx, value=value).font = normal_font

    if not bang_pgd.empty:
        _add_df_sheet(wb, "Theo PGD (HSTD)", bang_pgd)
    if not bang_ct.empty:
        _add_df_sheet(wb, "Theo Chương trình", bang_ct)
    if not bang_uy_thac.empty:
        _add_df_sheet(wb, "Theo Ủy thác", bang_uy_thac)
    if bang_dienbao is not None and not bang_dienbao.empty:
        _add_df_sheet(wb, "Theo ĐV (Điện báo)", bang_dienbao)

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ══════════════════════════════════════════════════════════════════════════════
# PHẦN 6: XUẤT WORD
# ══════════════════════════════════════════════════════════════════════════════

def xuat_word_bao_cao_khnv(
    so_lieu: dict,
    ten_mau: str,
    bang_pgd: pd.DataFrame,
    bang_ct: pd.DataFrame,
    bang_dienbao: pd.DataFrame | None = None,
    chenh_lech: list | None = None,
) -> bytes:
    """Sinh file Word báo cáo KHNV với số liệu tự động."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(13)

    title = doc.add_heading(
        f"BÁO CÁO KHNV — Tháng {so_lieu.get('thang', '')}/{so_lieu.get('nam', '')}", level=1
    )
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

    doc.add_paragraph(f"Ngày báo cáo: {so_lieu.get('ngay_bao_cao', '')}")
    doc.add_paragraph(f"Nguồn dữ liệu: {so_lieu.get('nguon', '')}")
    doc.add_paragraph(f"Mẫu: {ten_mau}")
    doc.add_paragraph("")

    # ── Số liệu tổng quan ──
    doc.add_heading("I. TỔNG QUAN SỐ LIỆU", level=2)

    if so_lieu.get("nguon") == "Điện báo":
        metrics_text = [
            f"• Tổng dư nợ: {_fmt(so_lieu.get('tong_du_no', 0))} đồng",
            f"• Dư nợ Kế hoạch A: {_fmt(so_lieu.get('du_no_kha', 0))} đồng",
            f"• Dư nợ Kế hoạch B: {_fmt(so_lieu.get('du_no_khb', 0))} đồng",
            f"• Nợ quá hạn KHA: {_fmt(so_lieu.get('du_no_qua_han_kha', 0))} đồng",
            f"• Nợ quá hạn KHB: {_fmt(so_lieu.get('du_no_qua_han_khb', 0))} đồng",
            f"• Nợ khoanh KHA: {_fmt(so_lieu.get('du_no_khoanh_kha', 0))} đồng",
            f"• Nợ khoanh KHB: {_fmt(so_lieu.get('du_no_khoanh_khb', 0))} đồng",
            f"• Vốn TW (KHA): {_fmt(so_lieu.get('nguon_tw_kha', 0))} đồng",
            f"• Huy động vốn: {_fmt(so_lieu.get('huy_dong_von', 0))} đồng",
            f"• UTĐT ĐP: {_fmt(so_lieu.get('utdt_dp', 0))} đồng",
            f"• Vốn An toàn: {_fmt(so_lieu.get('von_an_toan', 0))} đồng",
        ]
    else:
        metrics_text = [
            f"• Tổng dư nợ: {_fmt(so_lieu.get('tong_du_no', 0))} đồng",
            f"• Dư nợ trong hạn: {_fmt(so_lieu.get('du_no_trong_han', 0))} đồng",
            f"• Dư nợ quá hạn: {_fmt(so_lieu.get('du_no_qua_han', 0))} đồng ({so_lieu.get('ty_le_no_qua_han', 0)}%)",
            f"• Dư nợ khoanh: {_fmt(so_lieu.get('du_no_khoanh', 0))} đồng",
            f"• Nguồn vốn TW: {_fmt(so_lieu.get('nguon_von_tw', 0))} đồng",
            f"• Nguồn vốn ĐP: {_fmt(so_lieu.get('nguon_von_dp', 0))} đồng",
            f"• Số khách hàng: {so_lieu.get('so_khach_hang', 0):,}",
            f"• Số món vay: {so_lieu.get('so_mon_vay', 0):,}",
            f"• Giải ngân trong tháng: {_fmt(so_lieu.get('giai_ngan_trong_thang', 0))} đồng",
        ]
    for t in metrics_text:
        doc.add_paragraph(t)

    # ── Đối chiếu HSTD vs Điện báo ──
    if chenh_lech:
        doc.add_heading("II. ĐỐI CHIẾU HSTD vs ĐIỆN BÁO", level=2)
        for item in chenh_lech:
            line = (f"• {item['Cảnh báo']} {item['Chỉ tiêu']}: "
                    f"HSTD={_fmt(item['HSTD'])} | "
                    f"Điện báo={_fmt(item['Điện báo'])} | "
                    f"Chênh lệch={_fmt(item['Chênh lệch'])} ({item['Tỷ lệ %']}%)")
            doc.add_paragraph(line)

    # ── Bảng theo PGD ──
    if not bang_pgd.empty:
        doc.add_heading(f"{'III' if chenh_lech else 'II'}. CHI TIẾT THEO PGD", level=2)
        _add_df_to_docx_table(doc, bang_pgd, 8)

    # ── Bảng theo Chương trình ──
    if not bang_ct.empty:
        doc.add_heading(f"{'IV' if chenh_lech else 'III'}. CHI TIẾT THEO CHƯƠNG TRÌNH", level=2)
        _add_df_to_docx_table(doc, bang_ct, 8)

    # ── Bảng Điện báo theo đơn vị ──
    if bang_dienbao is not None and not bang_dienbao.empty:
        doc.add_heading(f"{'V' if chenh_lech else 'IV'}. ĐIỆN BÁO THEO ĐƠN VỊ", level=2)
        _add_df_to_docx_table(doc, bang_dienbao, 8)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _add_df_to_docx_table(doc, df: pd.DataFrame, font_size: int = 8):
    """Thêm DataFrame vào Word dạng bảng."""
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    for i, col in enumerate(df.columns):
        cell = table.rows[0].cells[i]
        cell.text = str(col)
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(font_size)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = _fmt(val) if isinstance(val, (int, float)) else str(val)
            for p in cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(font_size)
