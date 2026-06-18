"""Dịch vụ xuất Lịch làm việc tuần KH-NV ra file Word.

Dùng mẫu MAU_LICH LAM VIEC TUAN KH-NV.md.
- Nhập tuần (từ ngày → đến ngày)
- Nhập nội dung cho từng mục
- Sinh file .docx hoàn chỉnh
"""
from __future__ import annotations

from datetime import date, timedelta
from io import BytesIO

from logger import get_logger

logger = get_logger(__name__)

SECTION_LABELS: list[tuple[str, str]] = [
    ("cong_tac_td", "1. Công tác tín dụng chi nhánh tỉnh"),
    ("dia_ban_hs", "2. Địa bàn 09 phường do Hội sở tỉnh kiêm nhiệm"),
    ("noi_dung_khac", "3. Nội dung công việc khác"),
]

MAU_NOI_DUNG_MAC_DINH = {
    "cong_tac_td": (
        "- Tham mưu Ban Giám đốc chi nhánh các nội dung theo chỉ đạo.\n"
        "- Đôn đốc các PGD:\n"
        "  + Làm việc với UBND các xã, phường để sớm được chuyển vốn ủy thác.\n"
        "  + Huy động tiền gửi; tổ chức giải ngân kịp thời.\n"
        "  + Kiểm soát, xử lý nợ đến hạn phân kỳ.\n"
        "  + Phân tích nguyên nhân, có giải pháp xử lý NQH, nợ khoanh, hộ KHĐ.\n"
        "  + Đôn đốc CT-XH nâng cao chất lượng ủy thác, bình xét, kiểm tra.\n"
        "- Chủ động rà soát, chỉnh sửa tồn tại theo kết luận kiểm tra.\n"
        "- Các nội dung công việc khác theo chỉ đạo của BGĐ CN tỉnh."
    ),
    "dia_ban_hs": (
        "- Làm việc với UBND các phường để sớm được chuyển vốn ủy thác.\n"
        "- Huy động tiền gửi; giải ngân các chỉ tiêu kế hoạch dư nợ.\n"
        "- Xử lý nợ đến hạn; đôn đốc thu hồi NQH, nợ khoanh.\n"
        "- Đôn đốc CT-XH cấp xã nâng cao chất lượng hoạt động ủy thác.\n"
        "- Các nội dung công việc khác theo chỉ đạo của BGĐ CN tỉnh."
    ),
    "noi_dung_khac": "- Các nội dung công việc khác theo chỉ đạo của BGĐ CN tỉnh.",
}


_THU_LABEL = {0: "Thứ Hai", 1: "Thứ Ba", 2: "Thứ Tư", 3: "Thứ Năm", 4: "Thứ Sáu", 5: "Thứ Bảy", 6: "Chủ Nhật"}
_LOAI_SHORT = {"hop": "Họp", "kiem_tra": "Kiểm tra", "cong_tac": "Công tác", "tap_huan": "Tập huấn", "khac": ""}

_MAC_DINH_CAN_BO = [
    {"ho_ten": "Nguyễn Xuân Trường", "chuc_vu": "vp1"},
    {"ho_ten": "Hoàng Trần Thành Nhân", "chuc_vu": "vp2"},
    {"ho_ten": "Phạm Thế Dinh", "chuc_vu": "cbtd"},
]


def xuat_lich_bang_tuan(
    tu_ngay: date,
    den_ngay: date,
    ds_lich: list,
    ds_phan_cong: list,
    can_bo_list: list | None = None,
    ten_truong_phong: str = "",
    ngay_ky: date | None = None,
) -> bytes:
    """Tạo file Word lịch làm việc tuần dạng BẢNG (A4 Landscape).
    Khớp mẫu: DN-Lich cong tac tuan KH-NV.docx
    """
    from docx import Document
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # ── Helper functions ──
    def _run(run, bold=False, size=12, italic=False):
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        try:
            rFonts = run._r.get_or_add_rPr().get_or_add_rFonts()
            rFonts.set(qn("w:eastAsia"), "Times New Roman")
        except Exception:
            pass

    def _clear_para(cell):
        p = cell.paragraphs[0]
        for child in list(p._element):
            p._element.remove(child)
        return p

    def _set_borders(cell, thin=True):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for old in tcPr.findall(qn("w:tcBorders")):
            tcPr.remove(old)
        tcB = OxmlElement("w:tcBorders")
        val, sz = ("single", "6") if thin else ("nil", "0")
        for side in ("top", "left", "bottom", "right"):
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), val)
            if thin:
                el.set(qn("w:sz"), sz)
                el.set(qn("w:space"), "0")
                el.set(qn("w:color"), "000000")
            tcB.append(el)
        tcPr.append(tcB)

    def _no_border(cell):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for old in tcPr.findall(qn("w:tcBorders")):
            tcPr.remove(old)
        tcB = OxmlElement("w:tcBorders")
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "nil")
            tcB.append(el)
        tcPr.append(tcB)

    def _valign(cell, val="center"):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for old in tcPr.findall(qn("w:vAlign")):
            tcPr.remove(old)
        el = OxmlElement("w:vAlign")
        el.set(qn("w:val"), val)
        tcPr.append(el)

    def _shading(cell, fill):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for old in tcPr.findall(qn("w:shd")):
            tcPr.remove(old)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        tcPr.append(shd)

    def _cell_w(cell, w_twips):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for old in tcPr.findall(qn("w:tcW")):
            tcPr.remove(old)
        tcW = OxmlElement("w:tcW")
        tcW.set(qn("w:w"), str(int(w_twips)))
        tcW.set(qn("w:type"), "dxa")
        tcPr.insert(0, tcW)

    def _vmerge(cell, restart=True):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for old in tcPr.findall(qn("w:vMerge")):
            tcPr.remove(old)
        vm = OxmlElement("w:vMerge")
        if restart:
            vm.set(qn("w:val"), "restart")
        tcPr.append(vm)

    def _row_h(row, h_cm):
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        for old in trPr.findall(qn("w:trHeight")):
            trPr.remove(old)
        trH = OxmlElement("w:trHeight")
        trH.set(qn("w:val"), str(int(Cm(h_cm).twips)))
        trH.set(qn("w:hRule"), "atLeast")
        trPr.append(trH)

    def _hmerge(row, start, end):
        """Merge cells[start..end] horizontally, return merged cell."""
        cells = row.cells
        c = cells[start]
        for i in range(start + 1, end + 1):
            c = c.merge(cells[i])
        return c

    def _get_content(staff_name: str, ngay: date) -> list:
        lines = []
        name_parts = [p for p in staff_name.split() if len(p) > 1]
        for ev in ds_lich:
            try:
                ev_date = date.fromisoformat(ev.get("ngay", ""))
            except (ValueError, TypeError):
                continue
            if ev_date != ngay:
                continue
            tv = ev.get("thanh_vien", "")
            if not tv or any(p.lower() in tv.lower() for p in name_parts):
                tieu_de = ev.get("tieu_de", "")
                loai = _LOAI_SHORT.get(ev.get("loai", ""), "")
                dia = ev.get("dia_diem", "")
                line = f"[{loai}] {tieu_de}" if loai else tieu_de
                if dia:
                    line += f" - {dia}"
                lines.append(line)
        for cv in ds_phan_cong:
            try:
                dl = date.fromisoformat(cv.get("ngay_deadline", ""))
            except (ValueError, TypeError):
                continue
            if dl != ngay:
                continue
            nguoi = cv.get("nguoi_thuc_hien", "")
            if any(p.lower() in nguoi.lower() for p in name_parts):
                lines.append(f"[Deadline] {cv.get('tieu_de', '')}")
        return lines

    # ── Setup ──
    staff = can_bo_list or _MAC_DINH_CAN_BO
    n_staff = len(staff)
    n_cols = 2 + n_staff
    ngay_ky = ngay_ky or (tu_ngay - timedelta(days=3))

    week_days = []
    d = tu_ngay
    while d <= den_ngay:
        week_days.append(d)
        d += timedelta(days=1)

    # Column widths (A4 landscape: 29.7cm - 1.5cm - 1.5cm = 26.7cm content)
    cw_time = int(Cm(2.5).twips)
    cw_date = int(Cm(2.5).twips)
    cw_total = int(Cm(26.7).twips)
    cw_remaining = cw_total - cw_time - cw_date
    cw_staff = cw_remaining // n_staff
    cw_last = cw_remaining - cw_staff * (n_staff - 1)
    col_ws = [cw_time, cw_date] + [cw_staff] * (n_staff - 1) + [cw_last]

    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(29.7)
    sec.page_height = Cm(21)
    sec.left_margin = Cm(1.5)
    sec.right_margin = Cm(1.5)
    sec.top_margin = Cm(1.0)
    sec.bottom_margin = Cm(1.0)
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(12)

    # ── BẢNG HEADER (letterhead + title, không viền) ──
    hdr = doc.add_table(rows=2, cols=2)
    for row in hdr.rows:
        for cell in row.cells:
            _no_border(cell)

    # Row 0: NHCSXH | CHXHCNVN
    _row_h(hdr.rows[0], 0.65)
    pl = _clear_para(hdr.rows[0].cells[0])
    pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(pl.add_run("NGÂN HÀNG CHÍNH SÁCH XÃ HỘI"), bold=True, size=11)
    pr = _clear_para(hdr.rows[0].cells[1])
    pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(pr.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM"), bold=True, size=11)

    # Row 1: PHÒNG KH-NV | Độc lập
    _row_h(hdr.rows[1], 0.55)
    pl1 = _clear_para(hdr.rows[1].cells[0])
    pl1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(pl1.add_run("PHÒNG KH-NV"), bold=True, size=11)
    pr1 = _clear_para(hdr.rows[1].cells[1])
    pr1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(pr1.add_run("Độc lập - Tự do - Hạnh phúc"), bold=True, size=11, italic=True)

    # Date line
    p_ngayky = doc.add_paragraph()
    p_ngayky.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _run(p_ngayky.add_run(
        f"Đồng Nai, ngày {ngay_ky.day} tháng {ngay_ky.month} năm {ngay_ky.year}"
    ), size=11, italic=True)

    # Title
    p_t1 = doc.add_paragraph()
    p_t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_t1.paragraph_format.space_before = Pt(2)
    _run(p_t1.add_run("THÔNG BÁO LỊCH LÀM VIỆC TUẦN"), bold=True, size=14)

    p_t2 = doc.add_paragraph()
    p_t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p_t2.add_run("CỦA PHÒNG KẾ HOẠCH - NGHIỆP VỤ TÍN DỤNG TỈNH ĐỒNG NAI"), bold=True, size=12)

    p_t3 = doc.add_paragraph()
    p_t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_t3.paragraph_format.space_after = Pt(4)
    _run(p_t3.add_run(
        f"(Từ ngày {tu_ngay.strftime('%d/%m/%Y')} đến ngày {den_ngay.strftime('%d/%m/%Y')})"
    ), bold=True, size=12)

    # ── BẢNG LỊCH (có viền) ──
    # Rows: 0=header-top (Thời gian/Ngày tháng vMerge start + LÃNH ĐẠO PHÒNG)
    #       1=header-bot (vMerge continue + staff names)
    #       2..n+1 = data rows
    #       n+2 = footer
    n_rows = 2 + len(week_days) + 1
    tbl = doc.add_table(rows=n_rows, cols=n_cols)

    # Set tblGrid
    tblGrid = OxmlElement("w:tblGrid")
    for w in col_ws:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        tblGrid.append(gc)
    existing = tbl._tbl.find(qn("w:tblGrid"))
    if existing is not None:
        tbl._tbl.remove(existing)
    tblPr = tbl._tbl.find(qn("w:tblPr"))
    if tblPr is not None:
        tblPr.addnext(tblGrid)
    else:
        tbl._tbl.insert(0, tblGrid)

    # Apply borders and widths to all cells
    for ri, row in enumerate(tbl.rows):
        for ci, cell in enumerate(row.cells):
            _cell_w(cell, col_ws[ci])
            _set_borders(cell)

    # ── Row 0: header top ──
    r0 = tbl.rows[0]
    _row_h(r0, 0.9)

    c_tg = r0.cells[0]
    _shading(c_tg, "BDD7EE")
    _valign(c_tg, "center")
    _vmerge(c_tg, restart=True)
    p = _clear_para(c_tg)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p.add_run("Thời gian"), bold=True, size=11)

    c_ng = r0.cells[1]
    _shading(c_ng, "BDD7EE")
    _valign(c_ng, "center")
    _vmerge(c_ng, restart=True)
    p = _clear_para(c_ng)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p.add_run("Ngày, tháng"), bold=True, size=11)

    c_ldp = _hmerge(r0, 2, n_cols - 1)
    _shading(c_ldp, "BDD7EE")
    _valign(c_ldp, "center")
    p = _clear_para(c_ldp)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p.add_run("LÃNH ĐẠO PHÒNG"), bold=True, size=12)

    # ── Row 1: header bot (staff names) ──
    r1 = tbl.rows[1]
    _row_h(r1, 1.0)

    c_tg2 = r1.cells[0]
    _shading(c_tg2, "BDD7EE")
    _vmerge(c_tg2, restart=False)
    _clear_para(c_tg2)

    c_ng2 = r1.cells[1]
    _shading(c_ng2, "BDD7EE")
    _vmerge(c_ng2, restart=False)
    _clear_para(c_ng2)

    for si, cb in enumerate(staff):
        c = r1.cells[2 + si]
        _shading(c, "BDD7EE")
        _valign(c, "center")
        p = _clear_para(c)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p.add_run(f"Đ/c {cb.get('ho_ten', '')}"), bold=True, size=11)

    # ── Data rows ──
    for di, ngay in enumerate(week_days):
        row = tbl.rows[2 + di]
        _row_h(row, 2.0)

        c0 = row.cells[0]
        _valign(c0, "center")
        p = _clear_para(c0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p.add_run(_THU_LABEL.get(ngay.weekday(), "")), bold=True, size=11)

        c1 = row.cells[1]
        _valign(c1, "center")
        p = _clear_para(c1)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p.add_run(ngay.strftime("%d/%m/%Y")), size=11)

        for si, cb in enumerate(staff):
            c = row.cells[2 + si]
            _valign(c, "top")
            lines = _get_content(cb.get("ho_ten", ""), ngay)
            first_p = _clear_para(c)
            first_p.paragraph_format.space_before = Pt(2)
            first_p.paragraph_format.space_after = Pt(1)
            for li, line in enumerate(lines):
                p_curr = first_p if li == 0 else c.add_paragraph()
                p_curr.paragraph_format.space_before = Pt(1)
                p_curr.paragraph_format.space_after = Pt(1)
                _run(p_curr.add_run(line), size=10)

    # ── Footer row ──
    r_foot = tbl.rows[2 + len(week_days)]
    _row_h(r_foot, 3.0)

    split = max(1, n_cols // 2 - 1)
    c_nr = _hmerge(r_foot, 0, split)
    _no_border(c_nr)
    _valign(c_nr, "top")
    p = _clear_para(c_nr)
    p.paragraph_format.space_before = Pt(4)
    _run(p.add_run("Nơi nhận:"), bold=True, size=11)
    for item in ("- Phòng HC-TC;", "- Các cán bộ phòng TD (biết, thực hiện);", "- Lưu TD."):
        pn = c_nr.add_paragraph()
        _run(pn.add_run(item), size=11)

    c_sig = _hmerge(r_foot, split + 1, n_cols - 1)
    _no_border(c_sig)
    _valign(c_sig, "center")
    p = _clear_para(c_sig)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p.add_run("TRƯỞNG PHÒNG"), bold=True, size=12)
    if ten_truong_phong:
        for _ in range(3):
            c_sig.add_paragraph()
        p_name = c_sig.add_paragraph()
        p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p_name.add_run(ten_truong_phong), bold=True, size=12)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def lay_tuan_tiep_theo() -> tuple[date, date]:
    """Ngày thứ 2 và thứ 6 của tuần sau."""
    today = date.today()
    days_until_monday = (7 - today.weekday()) % 7
    if days_until_monday == 0:
        days_until_monday = 7
    next_monday = today + timedelta(days=days_until_monday)
    next_friday = next_monday + timedelta(days=4)
    return next_monday, next_friday


def xuat_lich_lam_viec_tuan(
    tu_ngay: date,
    den_ngay: date,
    noi_dung: dict[str, str] | None = None,
    ten_truong_phong: str = "",
    ngay_ky: date | None = None,
) -> bytes:
    """Tạo file Word Lịch công tác tuần."""
    from docx import Document
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(3)
    sec.right_margin = Cm(2)
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(13)

    nd = noi_dung or {}
    ngay_ky = ngay_ky or (tu_ngay - timedelta(days=3))

    def fmt_date(d: date) -> str:
        return f"ngày {d.day} tháng {d.month} năm {d.year}"

    def _set(run, bold=False, size=13):
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.font.bold = bold
        try:
            rPr = run._r.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn("w:eastAsia"), "Times New Roman")
        except Exception:
            pass

    def _no_border(table):
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                for old in tcPr.findall(qn("w:tcBorders")):
                    tcPr.remove(old)
                tcBorders = OxmlElement("w:tcBorders")
                for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                    el = OxmlElement(f"w:{side}")
                    el.set(qn("w:val"), "nil")
                    tcBorders.append(el)
                tcPr.append(tcBorders)

    # ── Header ──
    hdr = doc.add_table(rows=1, cols=2)
    _no_border(hdr)
    cell_l = hdr.rows[0].cells[0]
    for txt, bold in [("NGÂN HÀNG CHÍNH SÁCH XÃ HỘI", True),
                       ("CHI NHÁNH TỈNH ĐỒNG NAI", True),
                       ("Phòng KHNVTD", False),
                       ("──────────────", False)]:
        p = cell_l.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set(p.add_run(txt), bold=bold, size=12 if bold else 11)
    p0 = cell_l.paragraphs[0]
    p0._element.getparent().remove(p0._element)

    cell_r = hdr.rows[0].cells[1]
    for txt, bold in [("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", True),
                       ("Độc lập - Tự do - Hạnh phúc", True),
                       ("──────────────────────────", False),
                       (f"Đồng Nai, {fmt_date(ngay_ky)}", False)]:
        p = cell_r.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set(p.add_run(txt), bold=bold, size=12 if bold else 11)
    p0r = cell_r.paragraphs[0]
    p0r._element.getparent().remove(p0r._element)

    doc.add_paragraph("")

    # ── Title ──
    p_t = doc.add_paragraph()
    p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set(p_t.add_run("LỊCH CÔNG TÁC"), bold=True, size=16)

    p_s = doc.add_paragraph()
    p_s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set(p_s.add_run(f"Từ {fmt_date(tu_ngay)} đến {fmt_date(den_ngay)}"), bold=True, size=13)

    doc.add_paragraph("")

    # ── NỘI DUNG ──
    p_h = doc.add_paragraph()
    _set(p_h.add_run("NỘI DUNG CÔNG VIỆC:"), bold=True, size=13)

    for key, label in SECTION_LABELS:
        doc.add_paragraph("")
        p_sec = doc.add_paragraph()
        _set(p_sec.add_run(label), bold=True, size=13)
        text = nd.get(key, MAU_NOI_DUNG_MAC_DINH.get(key, ""))
        for line in text.strip().split("\n"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            _set(p.add_run(line), bold=False, size=13)

    doc.add_paragraph("")
    doc.add_paragraph("")

    # ── Ký tên ──
    sig = doc.add_table(rows=1, cols=2)
    _no_border(sig)
    sl = sig.rows[0].cells[0]
    sp = sl.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set(sp.add_run("Nơi nhận:\n- Phòng HC-TC;\n- Các cán bộ phòng TD (biết, thực hiện);\n- Lưu TD."), bold=False, size=11)
    p0sl = sl.paragraphs[0]
    p0sl._element.getparent().remove(p0sl._element)

    sr = sig.rows[0].cells[1]
    for _ in range(4):
        sr.add_paragraph("")
    sp2 = sr.add_paragraph()
    sp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set(sp2.add_run("TRƯỞNG PHÒNG"), bold=True, size=13)
    if ten_truong_phong:
        for _ in range(3):
            sr.add_paragraph("")
        sp3 = sr.add_paragraph()
        sp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set(sp3.add_run(ten_truong_phong), bold=True, size=13)
    p0sr = sr.paragraphs[0]
    p0sr._element.getparent().remove(p0sr._element)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
