from __future__ import annotations

from datetime import date

import db


def doc_ds(key: str) -> list:
    val = db.doc_kv(key)
    return val if isinstance(val, list) else []


def ghi_ds(key: str, ds: list, username: str, action: str, mo_ta: str) -> None:
    db.ghi_kv(key, ds, username)
    db.ghi_audit(username, action, mo_ta)


# ── Word export — NĐ30/2020 ──────────────────────────────────────────────────

def _xuat_bc_phan_cong(ds: list, thang: int, nam: int, ten_truong_phong: str = "") -> bytes:
    """Tạo Word 'DANH SÁCH PHÂN CÔNG VÀ GIAO VIỆC' chuẩn NĐ30/2020."""
    import io
    from datetime import date as _date
    from docx import Document
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def _set_font(run, bold: bool = False, size: int = 13) -> None:
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.font.bold = bold
        try:
            rPr = run._r.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn("w:eastAsia"), "Times New Roman")
        except Exception:
            pass

    def _remove_borders(table) -> None:
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                for old in tcPr.findall(qn("w:tcBorders")):
                    tcPr.remove(old)
                tcBorders = OxmlElement("w:tcBorders")
                for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                    el = OxmlElement(f"w:{side}")
                    el.set(qn("w:val"), "nil")
                    tcBorders.append(el)
                tcPr.append(tcBorders)

    def _shade_cell(cell, fill: str = "BDD7EE") -> None:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        tcPr.append(shd)

    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(3)
    sec.right_margin = Cm(2)
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(13)

    today = _date.today()

    hdr = doc.add_table(rows=1, cols=2)
    _remove_borders(hdr)
    cell_l = hdr.rows[0].cells[0]
    for i, (txt, bold) in enumerate([
        ("NGÂN HÀNG CHÍNH SÁCH XÃ HỘI", True),
        ("Chi nhánh tỉnh Đồng Nai", False),
        ("Phòng KH-NV", False),
        ("──────────────", False),
        ("Số:      /BC-KHNV", False),
    ]):
        p = cell_l.paragraphs[0] if i == 0 else cell_l.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(p.add_run(txt), bold=bold, size=12 if i == 0 else 11)

    cell_r = hdr.rows[0].cells[1]
    for i, (txt, bold) in enumerate([
        ("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", True),
        ("Độc lập - Tự do - Hạnh phúc", True),
        ("──────────────────────────", False),
        (f"Đồng Nai, ngày {today.day} tháng {today.month} năm {today.year}", False),
    ]):
        p = cell_r.paragraphs[0] if i == 0 else cell_r.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(p.add_run(txt), bold=bold, size=11)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(p.add_run("DANH SÁCH PHÂN CÔNG VÀ GIAO VIỆC"), bold=True, size=14)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(p.add_run(f"Tháng {thang} năm {nam}"), size=13)
    doc.add_paragraph()

    _uu = {"khan_cap": "Khẩn cấp", "quan_trong": "Quan trọng", "binh_thuong": "Bình thường"}
    _cv = {"vp1": "Phó Phòng VT1", "vp2": "Phó Phòng VT2", "cbtd": "Cán bộ TD"}
    headers = ["STT", "Nhóm việc", "Đầu việc", "Người thực hiện",
               "Chức vụ", "Mức độ", "Ngày giao", "Thời gian\nhoàn thành"]
    widths_cm = [0.8, 2.5, 4.0, 2.5, 1.8, 1.6, 1.7, 1.7]

    tbl = doc.add_table(rows=1, cols=8)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, w in enumerate(widths_cm):
        for cell in tbl.columns[i].cells:
            cell.width = Cm(w)

    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(p.add_run(h), bold=True, size=11)
        _shade_cell(cell)

    for stt, item in enumerate(ds, 1):
        row = tbl.add_row()
        cv_code = item.get("chuc_vu", "")
        vals = [
            str(stt),
            item.get("nhom", ""),
            item.get("tieu_de", ""),
            item.get("nguoi_thuc_hien", ""),
            _cv.get(cv_code, cv_code),
            _uu.get(item.get("uu_tien", ""), ""),
            item.get("ngay_giao", ""),
            item.get("ngay_deadline", ""),
        ]
        CENTER_COLS = {0, 5, 6, 7}
        for i, val in enumerate(vals):
            p = row.cells[i].paragraphs[0]
            p.alignment = (WD_ALIGN_PARAGRAPH.CENTER if i in CENTER_COLS
                           else WD_ALIGN_PARAGRAPH.LEFT)
            _set_font(p.add_run(val), size=11)

    doc.add_paragraph()
    foot = doc.add_table(rows=3, cols=2)
    _remove_borders(foot)
    foot_data = [
        ("Người lập",          "TRƯỞNG PHÒNG KH-NV"),
        ("(Ký, ghi rõ họ tên)", "(Ký, ghi rõ họ tên)"),
        ("",                   ten_truong_phong),
    ]
    foot_bold = [(False, True), (False, False), (False, True)]
    for i, ((txt_l, txt_r), (bl, br)) in enumerate(zip(foot_data, foot_bold)):
        pl = foot.rows[i].cells[0].paragraphs[0]
        pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(pl.add_run(txt_l), bold=bl, size=12)
        pr = foot.rows[i].cells[1].paragraphs[0]
        pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(pr.add_run(txt_r), bold=br, size=12)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _xuat_bc_tien_do(ds: list, thang: int, nam: int, ten_truong_phong: str = "") -> bytes:
    """Tạo Word 'BÁO CÁO TIẾN ĐỘ THỰC HIỆN CÔNG VIỆC' chuẩn NĐ30/2020."""
    import io
    from datetime import date as _date
    from docx import Document
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def _set_font(run, bold: bool = False, size: int = 13) -> None:
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.font.bold = bold
        try:
            rPr = run._r.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn("w:eastAsia"), "Times New Roman")
        except Exception:
            pass

    def _remove_borders(table) -> None:
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                for old in tcPr.findall(qn("w:tcBorders")):
                    tcPr.remove(old)
                tcBorders = OxmlElement("w:tcBorders")
                for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                    el = OxmlElement(f"w:{side}")
                    el.set(qn("w:val"), "nil")
                    tcBorders.append(el)
                tcPr.append(tcBorders)

    def _shade_cell(cell, fill: str = "BDD7EE") -> None:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        tcPr.append(shd)

    _tt_map = {"chua_lam": "Chưa làm", "dang_lam": "Đang làm",
               "hoan_thanh": "Hoàn thành", "tre_han": "Trễ hạn"}
    _uu_map = {"khan_cap": "Khẩn cấp", "quan_trong": "Quan trọng", "binh_thuong": "Bình thường"}

    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(3)
    sec.right_margin = Cm(2)
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(13)

    today = _date.today()

    hdr = doc.add_table(rows=1, cols=2)
    _remove_borders(hdr)
    cell_l = hdr.rows[0].cells[0]
    for i, (txt, bold) in enumerate([
        ("NGÂN HÀNG CHÍNH SÁCH XÃ HỘI", True),
        ("Chi nhánh tỉnh Đồng Nai", False),
        ("Phòng KH-NV", False),
        ("──────────────", False),
        ("Số:      /BC-KHNV", False),
    ]):
        p = cell_l.paragraphs[0] if i == 0 else cell_l.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(p.add_run(txt), bold=bold, size=12 if i == 0 else 11)

    cell_r = hdr.rows[0].cells[1]
    for i, (txt, bold) in enumerate([
        ("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", True),
        ("Độc lập - Tự do - Hạnh phúc", True),
        ("──────────────────────────", False),
        (f"Đồng Nai, ngày {today.day} tháng {today.month} năm {today.year}", False),
    ]):
        p = cell_r.paragraphs[0] if i == 0 else cell_r.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(p.add_run(txt), bold=bold, size=11)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(p.add_run("BÁO CÁO TIẾN ĐỘ THỰC HIỆN CÔNG VIỆC"), bold=True, size=14)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(p.add_run(f"Tháng {thang} năm {nam}"), size=13)
    doc.add_paragraph()

    p = doc.add_paragraph()
    _set_font(p.add_run("I. TỔNG QUAN"), bold=True, size=13)

    tong = len(ds)
    ht   = sum(1 for x in ds if x.get("trang_thai") == "hoan_thanh")
    dang = sum(1 for x in ds if x.get("trang_thai") == "dang_lam")
    chua = sum(1 for x in ds if x.get("trang_thai") == "chua_lam")
    tre  = sum(1 for x in ds if x.get("trang_thai") == "tre_han")
    pct  = f"{ht / tong * 100:.0f}" if tong > 0 else "0"

    p = doc.add_paragraph()
    _set_font(p.add_run(
        f"Tổng số đầu việc: {tong}  |  Đã hoàn thành: {ht} ({pct}%)  |  "
        f"Đang thực hiện: {dang}  |  Chưa làm: {chua}  |  Trễ hạn: {tre}"
    ), size=13)
    doc.add_paragraph()

    p = doc.add_paragraph()
    _set_font(p.add_run("II. CHI TIẾT"), bold=True, size=13)

    headers = ["STT", "Đầu việc", "Người thực hiện", "Mức độ",
               "Ngày giao", "Thời gian\nhoàn thành", "Trạng thái", "Ghi chú"]
    widths_cm = [0.8, 3.6, 2.5, 1.8, 1.8, 2.0, 1.8, 2.3]

    tbl = doc.add_table(rows=1, cols=8)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, w in enumerate(widths_cm):
        for cell in tbl.columns[i].cells:
            cell.width = Cm(w)

    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(p.add_run(h), bold=True, size=11)
        _shade_cell(cell)

    for stt, item in enumerate(ds, 1):
        row = tbl.add_row()
        vals = [
            str(stt),
            item.get("tieu_de", ""),
            item.get("nguoi_thuc_hien", ""),
            _uu_map.get(item.get("uu_tien", ""), ""),
            item.get("ngay_giao", ""),
            item.get("ngay_deadline", ""),
            _tt_map.get(item.get("trang_thai", ""), ""),
            item.get("ghi_chu_ket_qua", ""),
        ]
        CENTER_COLS = {0, 3, 4, 5, 6}
        for i, val in enumerate(vals):
            p = row.cells[i].paragraphs[0]
            p.alignment = (WD_ALIGN_PARAGRAPH.CENTER if i in CENTER_COLS
                           else WD_ALIGN_PARAGRAPH.LEFT)
            _set_font(p.add_run(val), size=11)

    doc.add_paragraph()
    foot = doc.add_table(rows=3, cols=2)
    _remove_borders(foot)
    foot_data = [
        ("Người lập",           "TRƯỞNG PHÒNG KH-NV"),
        ("(Ký, ghi rõ họ tên)", "(Ký, ghi rõ họ tên)"),
        ("",                    ten_truong_phong),
    ]
    foot_bold = [(False, True), (False, False), (False, True)]
    for i, ((txt_l, txt_r), (bl, br)) in enumerate(zip(foot_data, foot_bold)):
        pl = foot.rows[i].cells[0].paragraphs[0]
        pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(pl.add_run(txt_l), bold=bl, size=12)
        pr = foot.rows[i].cells[1].paragraphs[0]
        pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(pr.add_run(txt_r), bold=br, size=12)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Constants — chức vụ & mẫu giao việc ─────────────────────────────────────
# (moved from tab_khnv_noi_bo.py)

_CHUC_VU_MAP = {
    "vp1":  "Phó phòng (VT 1)",
    "vp2":  "Phó phòng (VT 2)",
    "cbtd": "Cán bộ TD",
}
_CHUC_VU_LABEL = {
    "vp1":  "👔 Phó phòng Vị trí 1",
    "vp2":  "👔 Phó phòng Vị trí 2",
    "cbtd": "🧑‍💼 Cán bộ Tín dụng",
}
_CHUC_VU_SHORT = {
    "vp1":  "Phó Phòng VT1",
    "vp2":  "Phó Phòng VT2",
    "cbtd": "Cán bộ Tín dụng",
}
_CHUC_VU_TASK_FILTER = {
    "vp1":  {"Phó phòng (VT 1)", "Phó phòng (VT 1 & VT 2)",
             "Phó phòng (VT 1) + Cán bộ TD", "Phó phòng (VT 1 & VT 2), Cán bộ TD",
             "Tất cả cán bộ"},
    "vp2":  {"Phó phòng (VT 2)", "Phó phòng (VT 1 & VT 2)",
             "Phó phòng (VT 2) + Cán bộ TD", "Phó phòng (VT 1 & VT 2), Cán bộ TD",
             "Tất cả cán bộ"},
    "cbtd": {"Cán bộ TD", "Cán bộ TD (theo địa bàn)",
             "Phó phòng (VT 1) + Cán bộ TD", "Phó phòng (VT 2) + Cán bộ TD",
             "Phó phòng (VT 1 & VT 2), Cán bộ TD", "Tất cả cán bộ"},
}

_MAU_GIAO_VIEC = [
    # I. QUẢN LÝ CHUNG & HÀNH CHÍNH (5 việc)
    {"nhom": "I. Quản lý chung & Điều hành",
     "tieu_de": "Tổng hợp kế hoạch công tác tháng của Phòng, báo cáo Trưởng phòng",
     "nguoi_thuc_hien": "Phó phòng (VT 1)",
     "mo_ta": "⏱ 25 hàng tháng · 📄 Dự thảo kế hoạch"},
    {"nhom": "I. Quản lý chung & Điều hành",
     "tieu_de": "Theo dõi, đôn đốc tiến độ công việc; tổng hợp phiếu giao việc",
     "nguoi_thuc_hien": "Phó phòng (VT 1)",
     "mo_ta": "⏱ Hàng tuần · 📄 Báo cáo chiều thứ Sáu"},
    {"nhom": "I. Quản lý chung & Điều hành",
     "tieu_de": "Phân công cán bộ đi giao dịch xã theo lịch cố định",
     "nguoi_thuc_hien": "Phó phòng (VT 1)",
     "mo_ta": "⏱ Trước 20 hàng tháng · 📄 Danh sách phân công"},
    {"nhom": "I. Quản lý chung & Điều hành",
     "tieu_de": "Kiểm soát, ký nháy các văn bản do Phòng soạn thảo",
     "nguoi_thuc_hien": "Phó phòng (VT 1 & VT 2)",
     "mo_ta": "⏱ Trong ngày · 📄 Văn bản trình Giám đốc"},
    {"nhom": "I. Quản lý chung & Điều hành",
     "tieu_de": "Theo dõi chấm công, nghỉ phép, nghỉ bù của cán bộ Phòng",
     "nguoi_thuc_hien": "Phó phòng (VT 2)",
     "mo_ta": "⏱ Hàng tháng · 📄 Bảng chấm công"},
    # II. TÍN DỤNG & CHO VAY (6 việc)
    {"nhom": "II. Tín dụng & Cho vay",
     "tieu_de": "Thẩm định và phê duyệt các khoản vay theo phân quyền trên hệ thống Intellect",
     "nguoi_thuc_hien": "Phó phòng (VT 1)",
     "mo_ta": "⏱ Trong 2 ngày kể từ khi nhận hồ sơ · 📄 Kết quả thẩm định"},
    {"nhom": "II. Tín dụng & Cho vay",
     "tieu_de": "Tổng hợp nhu cầu vay vốn hộ nghèo, cận nghèo, đối tượng chính sách tại địa bàn TP",
     "nguoi_thuc_hien": "Phó phòng (VT 1)",
     "mo_ta": "⏱ Hàng quý · 📄 Báo cáo đề xuất bổ sung vốn"},
    {"nhom": "II. Tín dụng & Cho vay",
     "tieu_de": "Triển khai cho vay: HTTVL, hộ nghèo, cận nghèo, nhà ở xã hội, HS-SV, XKLĐ, NĐ75, QĐ755, 2085,...",
     "nguoi_thuc_hien": "Cán bộ TD (theo địa bàn)",
     "mo_ta": "⏱ Theo kế hoạch giải ngân · 📄 Hồ sơ giải ngân đúng quy trình"},
    {"nhom": "II. Tín dụng & Cho vay",
     "tieu_de": "Tập hợp, kiểm tra hồ sơ vay từ Tổ TK&VV, trình lãnh đạo phê duyệt",
     "nguoi_thuc_hien": "Cán bộ TD",
     "mo_ta": "⏱ Hàng tuần · 📄 Hồ sơ hợp lệ"},
    {"nhom": "II. Tín dụng & Cho vay",
     "tieu_de": "Xây dựng kế hoạch tín dụng năm của các xã, phường, trình Trưởng BĐD HĐQT TP phê duyệt",
     "nguoi_thuc_hien": "Phó phòng (VT 1)",
     "mo_ta": "⏱ Quý IV hàng năm · 📄 Kế hoạch được phê duyệt"},
    {"nhom": "II. Tín dụng & Cho vay",
     "tieu_de": "Xây dựng kế hoạch tín dụng toàn tỉnh, trình Trưởng BĐD HĐQT tỉnh",
     "nguoi_thuc_hien": "Phó phòng (VT 2)",
     "mo_ta": "⏱ Theo chỉ đạo · 📄 Kế hoạch trình Trung ương"},
    # III. NGUỒN VỐN & QUỸ (3 việc)
    {"nhom": "III. Nguồn vốn & Quỹ",
     "tieu_de": "Theo dõi biến động quỹ an toàn chi trả, đề xuất bổ sung hoặc điều chuyển",
     "nguoi_thuc_hien": "Phó phòng (VT 2)",
     "mo_ta": "⏱ Hàng ngày · 📄 Điện chuyển vốn kịp thời"},
    {"nhom": "III. Nguồn vốn & Quỹ",
     "tieu_de": "Huy động tiền gửi tiết kiệm từ tổ chức, cá nhân trên địa bàn",
     "nguoi_thuc_hien": "Phó phòng (VT 2) + Cán bộ TD",
     "mo_ta": "⏱ Hàng tháng · 📄 Báo cáo kết quả huy động"},
    {"nhom": "III. Nguồn vốn & Quỹ",
     "tieu_de": "Quản lý nguồn vốn nhận ủy thác từ UBND tỉnh, các tổ chức; theo dõi Quỹ QGVL",
     "nguoi_thuc_hien": "Phó phòng (VT 2)",
     "mo_ta": "⏱ Thường xuyên · 📄 Sổ theo dõi, báo cáo đối chiếu"},
    # IV. NỢ RỦI RO & QUẢN LÝ NỢ (5 việc)
    {"nhom": "IV. Nợ rủi ro & Quản lý nợ",
     "tieu_de": "Hướng dẫn, đôn đốc các đơn vị lập hồ sơ xử lý nợ rủi ro",
     "nguoi_thuc_hien": "Phó phòng (VT 2)",
     "mo_ta": "⏱ Theo phát sinh · 📄 Hồ sơ đầy đủ, đúng quy định"},
    {"nhom": "IV. Nợ rủi ro & Quản lý nợ",
     "tieu_de": "Kiểm tra, tổng hợp hồ sơ nợ rủi ro toàn tỉnh, trình cấp thẩm quyền",
     "nguoi_thuc_hien": "Phó phòng (VT 2)",
     "mo_ta": "⏱ Hàng quý · 📄 Tờ trình kèm hồ sơ"},
    {"nhom": "IV. Nợ rủi ro & Quản lý nợ",
     "tieu_de": "Thông báo công khai kết quả xử lý nợ rủi ro tại địa bàn thành phố",
     "nguoi_thuc_hien": "Phó phòng (VT 2)",
     "mo_ta": "⏱ Sau khi được phê duyệt · 📄 Biên bản, thông báo"},
    {"nhom": "IV. Nợ rủi ro & Quản lý nợ",
     "tieu_de": "Lưu giữ hồ sơ nợ rủi ro theo quy định",
     "nguoi_thuc_hien": "Phó phòng (VT 2)",
     "mo_ta": "⏱ Thường xuyên · 📄 Hồ sơ đầy đủ"},
    {"nhom": "IV. Nợ rủi ro & Quản lý nợ",
     "tieu_de": "Đôn đốc thu nợ đến hạn, quá hạn; lập danh sách nợ chây ỳ",
     "nguoi_thuc_hien": "Cán bộ TD",
     "mo_ta": "⏱ Hàng tháng · 📄 Báo cáo nợ chi tiết"},
    # V. ỦY THÁC CT-XH & TỔ TK&VV (4 việc)
    {"nhom": "V. Ủy thác CT-XH & Tổ TK&VV",
     "tieu_de": "Tham mưu ký Văn bản liên tịch, Hợp đồng ủy thác với các tổ chức CT-XH",
     "nguoi_thuc_hien": "Phó phòng (VT 1)",
     "mo_ta": "⏱ Khi có thay đổi · 📄 Hợp đồng đã ký"},
    {"nhom": "V. Ủy thác CT-XH & Tổ TK&VV",
     "tieu_de": "Tổ chức họp giao ban với các tổ chức CT-XH cấp tỉnh, thành phố",
     "nguoi_thuc_hien": "Phó phòng (VT 2)",
     "mo_ta": "⏱ Định kỳ (2 tháng/lần) · 📄 Biên bản, thông báo kết luận"},
    {"nhom": "V. Ủy thác CT-XH & Tổ TK&VV",
     "tieu_de": "Đánh giá chất lượng hoạt động Tổ TK&VV; đề xuất củng cố tổ yếu kém",
     "nguoi_thuc_hien": "Cán bộ TD (theo địa bàn)",
     "mo_ta": "⏱ Hàng tháng · 📄 Bảng xếp loại Tổ"},
    {"nhom": "V. Ủy thác CT-XH & Tổ TK&VV",
     "tieu_de": "Tham gia sinh hoạt Tổ TK&VV theo lịch; kiểm tra sổ sách của Tổ",
     "nguoi_thuc_hien": "Cán bộ TD",
     "mo_ta": "⏱ Hàng tháng · 📄 Biên bản kiểm tra"},
    # VI. GIAO DỊCH XÃ & KIỂM TRA CƠ SỞ (4 việc)
    {"nhom": "VI. Giao dịch xã & Kiểm tra cơ sở",
     "tieu_de": "Tham mưu tổ chức phiên giao dịch xã đúng lịch, an toàn",
     "nguoi_thuc_hien": "Phó phòng (VT 1)",
     "mo_ta": "⏱ Theo lịch cố định · 📄 Báo cáo sau phiên giao dịch"},
    {"nhom": "VI. Giao dịch xã & Kiểm tra cơ sở",
     "tieu_de": "Kiểm tra, giám sát hoạt động tại Điểm giao dịch xã; tỷ lệ giải ngân, thu nợ, thu lãi",
     "nguoi_thuc_hien": "Phó phòng (VT 1) + Cán bộ TD",
     "mo_ta": "⏱ Hàng quý · 📄 Báo cáo đánh giá"},
    {"nhom": "VI. Giao dịch xã & Kiểm tra cơ sở",
     "tieu_de": "Kiểm tra sử dụng vốn vay 100% món vay trong vòng 30 ngày sau giải ngân",
     "nguoi_thuc_hien": "Cán bộ TD",
     "mo_ta": "⏱ Hàng tháng · 📄 Mẫu 06/TD"},
    {"nhom": "VI. Giao dịch xã & Kiểm tra cơ sở",
     "tieu_de": "Mở hòm thư góp ý, tổng hợp ý kiến khách hàng tại Điểm giao dịch xã",
     "nguoi_thuc_hien": "Cán bộ TD",
     "mo_ta": "⏱ Hàng tháng · 📄 Báo cáo tham mưu giải quyết"},
    # VII. BÁO CÁO THỐNG KÊ & TỔNG HỢP (6 việc)
    {"nhom": "VII. Báo cáo & Thống kê",
     "tieu_de": "Tổng hợp báo cáo thống kê tín dụng toàn tỉnh gửi NHCSXH TW, NHNN, các sở ngành",
     "nguoi_thuc_hien": "Phó phòng (VT 2)",
     "mo_ta": "⏱ Trước ngày 10 hàng tháng · 📄 Báo cáo đầy đủ biểu"},
    {"nhom": "VII. Báo cáo & Thống kê",
     "tieu_de": "Dự thảo Nghị quyết, báo cáo kết quả hoạt động của BĐD HĐQT tỉnh",
     "nguoi_thuc_hien": "Phó phòng (VT 2)",
     "mo_ta": "⏱ Theo kỳ họp · 📄 Nghị quyết trình ký"},
    {"nhom": "VII. Báo cáo & Thống kê",
     "tieu_de": "Xây dựng dự thảo báo cáo kết quả hoạt động chi nhánh tháng, quý, năm",
     "nguoi_thuc_hien": "Phó phòng (VT 2)",
     "mo_ta": "⏱ Theo định kỳ · 📄 Báo cáo trình Giám đốc"},
    {"nhom": "VII. Báo cáo & Thống kê",
     "tieu_de": "Lập dự toán, tờ trình văn phòng phẩm theo tháng; thanh toán các khoản của Phòng",
     "nguoi_thuc_hien": "Phó phòng (VT 2)",
     "mo_ta": "⏱ Hàng tháng · 📄 Dự toán, chứng từ"},
    {"nhom": "VII. Báo cáo & Thống kê",
     "tieu_de": "Kiểm soát, chỉnh sửa các cảnh báo trên chương trình TTBC-IMS",
     "nguoi_thuc_hien": "Cán bộ TD",
     "mo_ta": "⏱ Hàng tuần · 📄 Báo cáo kết quả chỉnh sửa"},
    {"nhom": "VII. Báo cáo & Thống kê",
     "tieu_de": "Chấm điểm 05 chuyên đề thi đua của chi nhánh",
     "nguoi_thuc_hien": "Phó phòng (VT 1)",
     "mo_ta": "⏱ Hàng quý · 📄 Bảng chấm điểm"},
    # VIII. ĐÀO TẠO & CÔNG TÁC KHÁC (5 việc)
    {"nhom": "VIII. Đào tạo & Công tác khác",
     "tieu_de": "Dự thảo kế hoạch, tài liệu tập huấn nghiệp vụ cho cán bộ trong và ngoài ngành",
     "nguoi_thuc_hien": "Phó phòng (VT 2)",
     "mo_ta": "⏱ Quý II hàng năm · 📄 Kế hoạch, tài liệu"},
    {"nhom": "VIII. Đào tạo & Công tác khác",
     "tieu_de": "Tham gia tập huấn, bồi dưỡng nghiệp vụ khi được phân công",
     "nguoi_thuc_hien": "Cán bộ TD",
     "mo_ta": "⏱ Theo yêu cầu · 📄 Giấy chứng nhận (nếu có)"},
    {"nhom": "VIII. Đào tạo & Công tác khác",
     "tieu_de": "Thành viên Tổ giao dịch lưu động tại xã, phường",
     "nguoi_thuc_hien": "Phó phòng (VT 1 & VT 2), Cán bộ TD",
     "mo_ta": "⏱ Theo lịch phân công · 📄 Thực hiện giao dịch"},
    {"nhom": "VIII. Đào tạo & Công tác khác",
     "tieu_de": "Làm thư ký giúp việc cho thành viên BĐD HĐQT các cấp khi kiểm tra địa bàn",
     "nguoi_thuc_hien": "Phó phòng (VT 1 & VT 2), Cán bộ TD",
     "mo_ta": "⏱ Theo phân công · 📄 Biên bản kiểm tra"},
    {"nhom": "VIII. Đào tạo & Công tác khác",
     "tieu_de": "Thực hiện nhiệm vụ đột xuất do Trưởng phòng / Ban Giám đốc giao",
     "nguoi_thuc_hien": "Tất cả cán bộ",
     "mo_ta": "⏱ Theo yêu cầu · 📄 Báo cáo hoàn thành"},
]

_MAU_GIAO_VIEC_TP = [
    {"ma": "TP01", "tieu_de": "Xây dựng chương trình, kế hoạch công tác của Phòng",
     "mo_ta": "Lập kế hoạch tháng, quý, năm; tổng hợp, đánh giá kết quả thực hiện; báo cáo Ban Giám đốc",
     "tan_suat": "Tháng/Quý/Năm"},
    {"ma": "TP02", "tieu_de": "Quản lý, phân công, giám sát và đánh giá cán bộ",
     "mo_ta": "Phân công nhiệm vụ cụ thể; theo dõi, đôn đốc, kiểm tra, nhận xét, đánh giá kết quả",
     "tan_suat": "Hàng tuần/Tháng"},
    {"ma": "TP03", "tieu_de": "Kiểm soát, ký nháy văn bản do Phòng soạn thảo",
     "mo_ta": "Kiểm soát văn bản trước khi trình Giám đốc tỉnh phê duyệt hoặc ban hành",
     "tan_suat": "Hàng ngày"},
    {"ma": "TP04", "tieu_de": "Đầu mối triển khai tín dụng chính sách trên địa bàn",
     "mo_ta": "Hướng dẫn, triển khai các chương trình tín dụng; phát hiện vướng mắc, đề xuất giải pháp",
     "tan_suat": "Thường xuyên"},
    {"ma": "TP05", "tieu_de": "Tham mưu Ban đại diện HĐQT tỉnh",
     "mo_ta": "Tham mưu tổ chức họp, ban hành Nghị quyết; giao chỉ tiêu tín dụng; đề xuất bổ sung vốn",
     "tan_suat": "Theo định kỳ/Đột xuất"},
    {"ma": "TP06", "tieu_de": "Điều hành công tác nguồn vốn",
     "mo_ta": "Giao hạn mức quỹ; điều hành quỹ hàng ngày; chỉ đạo huy động tiền gửi; quản lý vốn ủy thác",
     "tan_suat": "Hàng ngày/Tháng"},
    {"ma": "TP07", "tieu_de": "Chỉ đạo thực hiện và điều chỉnh kế hoạch tín dụng",
     "mo_ta": "Chỉ đạo xây dựng kế hoạch; tổng hợp trình phê duyệt; điều chỉnh chỉ tiêu; đảm bảo tăng trưởng",
     "tan_suat": "Quý/Năm"},
    {"ma": "TP08", "tieu_de": "Chỉ đạo rà soát nhu cầu vay vốn",
     "mo_ta": "Đảm bảo 100% hộ nghèo, cận nghèo, đối tượng chính sách có nhu cầu được tiếp cận vốn",
     "tan_suat": "Năm/Đột xuất"},
    {"ma": "TP09", "tieu_de": "Tham mưu ký kết và giám sát ủy thác qua tổ chức CT-XH",
     "mo_ta": "Ký văn bản liên tịch, hợp đồng ủy thác; tổ chức triển khai; duy trì giao ban, sơ kết, tập huấn",
     "tan_suat": "Theo định kỳ"},
    {"ma": "TP10", "tieu_de": "Chỉ đạo xử lý nợ rủi ro",
     "mo_ta": "Tham mưu chỉ đạo xử lý nợ rủi ro; thành lập đoàn kiểm tra; kiểm soát hồ sơ; tổng hợp trình cấp thẩm quyền",
     "tan_suat": "Thường xuyên"},
    {"ma": "TP11", "tieu_de": "Giám sát hoạt động giao dịch xã",
     "mo_ta": "Chỉ đạo tổ chức giao dịch xã; kiểm tra mạng lưới điểm giao dịch; đánh giá chất lượng; đề xuất chấn chỉnh",
     "tan_suat": "Tháng/Quý"},
    {"ma": "TP12", "tieu_de": "Tổ chức đào tạo, tập huấn nghiệp vụ",
     "mo_ta": "Dự thảo kế hoạch, chương trình, tài liệu; tham mưu tổ chức tập huấn cho thành viên BĐD HĐQT huyện",
     "tan_suat": "Năm/Đột xuất"},
    {"ma": "TP13", "tieu_de": "Kiểm soát và phê duyệt báo cáo thống kê",
     "mo_ta": "Lập báo cáo định tính tập thể Phòng; kiểm soát, phê duyệt chỉ tiêu được phân quyền; tổng hợp theo quy định",
     "tan_suat": "Tháng"},
    {"ma": "TP14", "tieu_de": "Dự thảo báo cáo định kỳ và đột xuất",
     "mo_ta": "Dự thảo báo cáo tín dụng, tham luận, giải trình, trả lời kiến nghị, góp ý dự thảo văn bản",
     "tan_suat": "Theo yêu cầu"},
    {"ma": "TP15", "tieu_de": "Làm thư ký cho thành viên BĐD HĐQT tỉnh",
     "mo_ta": "Thư ký giúp việc khi thành viên BĐD HĐQT kiểm tra, giám sát địa bàn xã được phân công",
     "tan_suat": "Theo phân công"},
    {"ma": "TP16", "tieu_de": "Đầu mối giao dịch với Sở, ban, ngành, tổ chức CT-XH",
     "mo_ta": "Phối hợp triển khai các hoạt động liên quan đến tín dụng chính sách",
     "tan_suat": "Thường xuyên"},
    {"ma": "TP17", "tieu_de": "Thực hiện nhiệm vụ khác do Ban Giám đốc giao",
     "mo_ta": "Triển khai các nhiệm vụ phát sinh ngoài kế hoạch",
     "tan_suat": "Đột xuất"},
]


def _guess_chuc_vu(cv: dict) -> str:
    """Đoán chức vụ từ task (field chuc_vu mới hoặc fallback từ nguoi_thuc_hien cũ)."""
    cv_field = cv.get("chuc_vu")
    if cv_field in _CHUC_VU_LABEL:
        return cv_field
    nguoi = cv.get("nguoi_thuc_hien", "")
    if "VT 1" in nguoi and "VT 2" not in nguoi:
        return "vp1"
    if "VT 2" in nguoi:
        return "vp2"
    if "Phó phòng" in nguoi:
        return "vp1"
    return "cbtd"


def _safe_date_lt(date_str: str, ref) -> bool:
    """True nếu date_str < ref, bắt lỗi parse."""
    from datetime import date as _date
    try:
        return _date.fromisoformat(date_str) < ref
    except ValueError:
        return False
