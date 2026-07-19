"""Service cho Mẫu 07 — Giao/Điều chỉnh KHTD theo Ấp/Thôn.

Chứa các hàm thuần logic/xuất Word không phụ thuộc streamlit:
  - Slug/key helpers
  - KV-store read/write helpers
  - Business logic: lay_so_goc_cho_ap, _sync_khtd_xa, _lay_ds_ma_key_co_du_lieu
  - DataFrame helpers: _build_table_data, _extract_data_from_edited_df, _update_total_row
  - Word export helpers + xuat_mau07_word
"""
from __future__ import annotations

import json
import re
import unicodedata
from datetime import datetime
from io import BytesIO

import pandas as pd

import db
try:
    from logger import get_logger
    logger = get_logger(__name__)
except Exception as e:
    logger.error("Lỗi trong khối except: %s", e, exc_info=True)
    import logging
    logger = logging.getLogger(__name__)
from config import (
    CHUONG_TRINH_KHTD,
    COT_MA_CHUONG_TRINH,
    COT_NGUON_VON,
    COT_TEN_THON,
    COT_TEN_XA,
    COT_TONG_DU_NO,
)
from tabs.tab_khtd import MAKEY_BY_MACT_NV

# ── Derived constants (mirrors tab) ────────────────────────────────────────────
CT_TW = [(mk, ten) for mk, _, ten, nv, _ in CHUONG_TRINH_KHTD if nv == "TW"]
CT_DP = [(mk, ten) for mk, _, ten, nv, _ in CHUONG_TRINH_KHTD if nv == "DP"]
TEN_BY_MAKEY = {mk: ten for mk, _, ten, _, _ in CHUONG_TRINH_KHTD}


# ══════════════════════════════════════════════════════════════════════════════
# SLUG / KEY HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def _slug(text: str) -> str:
    s = unicodedata.normalize("NFD", text.strip().lower())
    s = "".join(c for c in s if unicodedata.category(c) != "Mn")
    return re.sub(r"[^a-z0-9]+", "_", s).strip("_")


def _chuan_hoa_ten(ten: str) -> str:
    """Loại bỏ dấu, viết thường, bỏ 'xã'/'phường'/'thị trấn'/'TT' để so sánh."""
    s = unicodedata.normalize('NFD', str(ten).lower())
    s = ''.join(c for c in s if unicodedata.category(c) != 'Mn')
    s = re.sub(r'\b(xa|phuong|phuong|thi tran|tt)\b', '', s)
    s = re.sub(r'\s+', ' ', s).strip()
    return s


def _kv_key(pgd: str, xa: str) -> str:
    return f"khtd_ap_{_slug(pgd)}_{_slug(xa)}"


def _kv_key_ls(pgd: str, xa: str) -> str:
    return f"khtd_ap_lich_su_{_slug(pgd)}_{_slug(xa)}"


# ══════════════════════════════════════════════════════════════════════════════
# KV-STORE HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def _doc_kv_dict(key: str) -> dict:
    try:
        with db.get_conn() as conn:
            row = conn.execute("SELECT value FROM kv_store WHERE key=?", (key,)).fetchone()
        return json.loads(row["value"]) if row else {}
    except Exception as e:
        logger.error("_doc_kv_dict: lỗi đọc kv key=%s — %s", key, e, exc_info=True)
        return {}


def _doc_kv_list(key: str) -> list:
    try:
        with db.get_conn() as conn:
            row = conn.execute("SELECT value FROM kv_store WHERE key=?", (key,)).fetchone()
        val = json.loads(row["value"]) if row else []
        return val if isinstance(val, list) else []
    except Exception as e:
        logger.error("_doc_kv_list: lỗi đọc kv key=%s — %s", key, e, exc_info=True)
        return []


def _luu_kv(key: str, data, username: str) -> bool:
    try:
        with db.get_conn() as conn:
            conn.execute(
                "INSERT OR REPLACE INTO kv_store (key, value, updated_at, updated_by) VALUES (?,?,?,?)",
                (key, json.dumps(data, ensure_ascii=False), datetime.now().isoformat(), username),
            )
            conn.commit()
        return True
    except Exception as e:
        logger.error("_luu_kv: lỗi ghi kv key=%s — %s", key, e, exc_info=True)
        return False


# ══════════════════════════════════════════════════════════════════════════════
# BUSINESS LOGIC
# ══════════════════════════════════════════════════════════════════════════════

def tinh_du_no_ap_baseline(df_baseline: pd.DataFrame, ten_xa: str) -> dict:
    """Từ HSTD 31/12 tính dư nợ theo ấp × chương trình × nguồn vốn.

    Returns: {"{ten_thon}|{ma_key}": du_no_trieu} hoặc {"_err": "..."}.
    """
    if df_baseline is None or df_baseline.empty:
        return {}
    try:
        df = df_baseline.copy()
        col_xa = COT_TEN_XA if COT_TEN_XA in df.columns else None
        col_thon = COT_TEN_THON if COT_TEN_THON in df.columns else None
        col_mact = COT_MA_CHUONG_TRINH if COT_MA_CHUONG_TRINH in df.columns else None
        col_nv = COT_NGUON_VON if COT_NGUON_VON in df.columns else None
        col_dn = COT_TONG_DU_NO if COT_TONG_DU_NO in df.columns else None

        missing = [
            name
            for name, val in [
                ("Tên xã", col_xa),
                ("Tên thôn", col_thon),
                ("Mã chương trình", col_mact),
                ("Nguồn vốn", col_nv),
                ("Tổng dư nợ", col_dn),
            ]
            if not val
        ]
        if missing:
            return {"_err": f"Thiếu cột: {missing}"}

        df = df[df[col_xa] == ten_xa].copy()
        if df.empty:
            return {"_err": f"Xã '{ten_xa}' không có trong baseline"}

        df = df.dropna(subset=[col_thon, col_mact, col_nv])
        df[col_dn] = pd.to_numeric(df[col_dn], errors="coerce").fillna(0)
        df[col_mact] = pd.to_numeric(df[col_mact], errors="coerce").fillna(0).astype(int)
        df[col_nv] = pd.to_numeric(df[col_nv], errors="coerce").fillna(0).astype(int)

        result: dict[str, float] = {}
        for (ten_thon, ma_ct, nv_int), grp in df.groupby([col_thon, col_mact, col_nv]):
            du_no = float(grp[col_dn].sum())
            if du_no <= 0:
                continue
            mk_list = MAKEY_BY_MACT_NV.get((int(ma_ct), int(nv_int)), [])
            for ma_key in mk_list:
                result[f"{str(ten_thon).strip()}|{ma_key}"] = round(du_no / 1_000_000, 1)
        return result
    except Exception as e:  # conv: skip
        return {"_err": str(e)}


def lay_so_goc_cho_ap(ten_ap: str, ma_key: str, du_no_baseline: dict, lich_su: list) -> float:
    """Số gốc lũy kế: lần 1 = baseline 31/12, lần 2+ = KH lần trước."""
    composite = f"{ten_ap}|{ma_key}"
    if not lich_su:
        return du_no_baseline.get(composite, 0.0)
    lan_truoc_data = lich_su[-1].get("data", {})
    return lan_truoc_data.get(composite, du_no_baseline.get(composite, 0.0))


def _sync_khtd_xa(xa: str, data_dict: dict, username: str) -> None:
    """Sum theo xã × ma_key → ghi vào khtd_xa để các tab khác không bị break."""
    kv_xa = _doc_kv_dict("khtd_xa")
    tong: dict = {}
    for composite, gia_tri in data_dict.items():
        parts = composite.split("|", 1)
        if len(parts) == 2:
            mk = parts[1]
            k = f"{xa}|{mk}"
            tong[k] = round(tong.get(k, 0.0) + float(gia_tri), 1)
    kv_xa.update(tong)
    _luu_kv("khtd_xa", kv_xa, username)


def _lay_ds_ma_key_co_du_lieu(
    ten_xa: str,
    df_full: pd.DataFrame | None,
    du_no_baseline: dict,
    lich_su: list,
) -> set:
    """Lấy danh sách ma_key có dữ liệu từ baseline + HSTD hiện tại + lịch sử."""
    ds_ma_key: set = set()

    # Từ baseline
    if du_no_baseline:
        for key in du_no_baseline.keys():
            if "|" in key:
                ma_key = key.split("|")[1]
                ds_ma_key.add(ma_key)

    # Từ df_full (HSTD hiện tại)
    if df_full is not None and not df_full.empty:
        col_xa   = COT_TEN_XA if COT_TEN_XA in df_full.columns else None
        col_thon = COT_TEN_THON if COT_TEN_THON in df_full.columns else None
        col_mact = COT_MA_CHUONG_TRINH if COT_MA_CHUONG_TRINH in df_full.columns else None
        col_nv   = COT_NGUON_VON if COT_NGUON_VON in df_full.columns else None
        if col_xa and col_mact and col_nv:
            df_xa = df_full[df_full[col_xa] == ten_xa].copy()
            df_xa = df_xa.dropna(subset=[col_mact, col_nv])
            for _, row in df_xa.iterrows():
                ma_ct = int(row[col_mact])
                nguon_von_code = int(row[col_nv])
                for ma_key in MAKEY_BY_MACT_NV.get((ma_ct, nguon_von_code), []):
                    ds_ma_key.add(ma_key)

    # Từ lịch sử
    for item in lich_su:
        for key in item.get("data", {}).keys():
            if "|" in key:
                ma_key = key.split("|")[1]
                ds_ma_key.add(ma_key)

    # Fallback: nếu không có CT nào, dùng list mặc định
    if not ds_ma_key:
        ds_ma_key = {"1_TW", "19_TW", "9_TW", "3_TW", "2_TW"}  # Các CT chính

    return ds_ma_key


def _build_table_data(
    ds_ap: list,
    ds_ma_key: set,
    du_no_baseline: dict,
    lich_su: list,
    data_hien_tai: dict,
    nam: int,
) -> pd.DataFrame:
    """Build DataFrame cho st.data_editor — mỗi dòng = 1 CT trong 1 ấp."""
    rows = []

    for ten_ap in ds_ap:
        # Dòng header ấp
        rows.append({
            "_type": "header_ap",
            "_key": "",
            "Ấp/Thôn": f"📍 {ten_ap}",
            "Chương trình": "",
            "Nguồn vốn": "",
            f"Số gốc (triệu)": None,
            "Giao tăng/giảm (triệu)": None,
            f"Chỉ tiêu KH {nam} (triệu)": None,
        })

        # Lấy các CT có dữ liệu cho ấp này
        ds_ct_ap = []
        for ma_key in sorted(ds_ma_key):
            key = f"{ten_ap}|{ma_key}"
            # Luôn hiển thị nếu có trong baseline hoặc lịch sử
            co_du_lieu = (
                key in du_no_baseline or
                (lich_su and any(key in h["data"] for h in lich_su))
            )
            # Hoặc: luôn hiển thị các CT chính
            if co_du_lieu or ma_key in {"1_TW", "19_TW", "9_TW", "3_TW"}:
                ds_ct_ap.append(ma_key)

        # Sort: TW trước, ĐP sau
        ds_ct_ap = sorted(ds_ct_ap, key=lambda x: (x.split("_")[1], int(x.split("_")[0])))

        # Dòng dữ liệu CT
        for ma_key in ds_ct_ap:
            key = f"{ten_ap}|{ma_key}"
            so_goc = lay_so_goc_cho_ap(ten_ap, ma_key, du_no_baseline, lich_su)

            # Lấy giá trị từ lịch sử hoặc data hiện tại
            if data_hien_tai and key in data_hien_tai:
                chi_tieu_kh = data_hien_tai[key]
            elif lich_su and key in lich_su[-1].get("data", {}):
                chi_tieu_kh = lich_su[-1]["data"][key]
            else:
                chi_tieu_kh = so_goc

            giao_tang_giam = round(chi_tieu_kh - so_goc, 1)

            # Lấy tên CT
            ma_ct_num = int(ma_key.split("_")[0])
            ten_ct = TEN_BY_MAKEY.get(ma_key, f"CT {ma_ct_num}")
            nguon_von = ma_key.split("_")[1]

            rows.append({
                "_type": "data",
                "_key": key,
                "Ấp/Thôn": "",
                "Chương trình": ten_ct,
                "Nguồn vốn": nguon_von,
                f"Số gốc (triệu)": so_goc,
                "Giao tăng/giảm (triệu)": giao_tang_giam,
                f"Chỉ tiêu KH {nam} (triệu)": chi_tieu_kh,
            })

    # Thêm dòng tổng (tính sau)
    rows.append({
        "_type": "total",
        "_key": "",
        "Ấp/Thôn": "TỔNG CỘNG",
        "Chương trình": "",
        "Nguồn vốn": "",
        f"Số gốc (triệu)": 0.0,
        "Giao tăng/giảm (triệu)": 0.0,
        f"Chỉ tiêu KH {nam} (triệu)": 0.0,
    })

    return pd.DataFrame(rows)


def _extract_data_from_edited_df(edited_df: pd.DataFrame, nam: int) -> dict:
    """Extract data_dict từ edited_df sau khi user nhập."""
    data_dict = {}
    col_kh = f"Chỉ tiêu KH {nam} (triệu)"

    for _, row in edited_df[edited_df["_type"] == "data"].iterrows():
        key = row.get("_key")
        chi_tieu = row.get(col_kh)
        if key and chi_tieu and chi_tieu != 0:
            data_dict[key] = round(float(chi_tieu), 1)

    return data_dict


def _update_total_row(df: pd.DataFrame, nam: int) -> pd.DataFrame:
    """Cập nhật dòng tổng cộng."""
    col_so_goc = f"Số gốc (triệu)"
    col_giao = "Giao tăng/giảm (triệu)"
    col_kh = f"Chỉ tiêu KH {nam} (triệu)"

    data_rows = df[df["_type"] == "data"]

    tong_so_goc = data_rows[col_so_goc].sum() if not data_rows.empty else 0
    tong_giao = data_rows[col_giao].sum() if not data_rows.empty else 0
    tong_kh = data_rows[col_kh].sum() if not data_rows.empty else 0

    idx_total = df[df["_type"] == "total"].index
    if not idx_total.empty:
        df.loc[idx_total[0], col_so_goc] = round(tong_so_goc)
        df.loc[idx_total[0], col_giao] = round(tong_giao)
        df.loc[idx_total[0], col_kh] = round(tong_kh)

    return df


# ══════════════════════════════════════════════════════════════════════════════
# WORD EXPORT — HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def _remove_table_borders(table) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblBorders = OxmlElement("w:tblBorders")
    for name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:val"), "none")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _set_cell_bg(cell, color_hex: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  color_hex.lstrip("#"))
    cell._tc.get_or_add_tcPr().append(shd)


def _set_cell_border_top_double(cell) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    top = OxmlElement("w:top")
    top.set(qn("w:val"),   "double")
    top.set(qn("w:sz"),    "8")
    top.set(qn("w:color"), "000000")
    borders.append(top)
    tcPr.append(borders)


def _cell_write(cell, text: str, bold: bool = False, italic: bool = False,
                align: str = "left", size: int = 13, indent_cm: float = 0) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, Cm
    p = cell.paragraphs[0]
    p.alignment = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right":  WD_ALIGN_PARAGRAPH.RIGHT,
        "left":   WD_ALIGN_PARAGRAPH.LEFT,
    }.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    run = p.add_run(text)
    run.font.name  = "Times New Roman"
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic


def _fmt_delta(val: float) -> str:
    """Định dạng tăng/giảm: +50 / −10 / 0 (chuẩn VN: . nghìn, 0 dp)."""
    s = f"{abs(val):,.0f}".replace(",","X").replace(".",",").replace("X",".")
    if val > 0:
        return f"+{s}"
    if val < 0:
        return f"−{s}"
    return "0"


def _fmt_money(val: float) -> str:
    """Định dạng tiền (chuẩn VN: . nghìn, 0 chữ số thập phân)."""
    return f"{val:,.0f}".replace(",","X").replace(".",",").replace("X",".")


# ══════════════════════════════════════════════════════════════════════════════
# WORD EXPORT — MAIN
# ══════════════════════════════════════════════════════════════════════════════

def xuat_mau07_word(
    xa: str,
    nam: int,
    loai_van_ban: str,
    data_dict: dict,
    du_no_baseline: dict,
    lich_su: list,
    lan_chon,
) -> bytes:
    """Xuất file Word Mẫu 07 theo CV 7064. Số QĐ & ngày tháng để TRỐNG.

    Returns: bytes .docx
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, Cm
    except ImportError:
        raise ImportError("Cần cài: pip install python-docx")

    # Xác định data và lich_su_truoc cho hàm so_goc
    if lan_chon == "dang_nhap":
        xuat_data    = data_dict
        ls_truoc     = lich_su
    else:
        idx          = int(lan_chon) - 1
        xuat_data    = lich_su[idx]["data"] if idx < len(lich_su) else data_dict
        ls_truoc     = lich_su[:idx]

    # Thu thập danh sách ấp theo thứ tự
    ds_ap = sorted({composite.split("|")[0] for composite in xuat_data if "|" in composite})

    # Chuẩn bị dữ liệu bảng
    bang_rows = []
    tong_delta = 0.0
    tong_kh    = 0.0
    stt = 0

    for ten_ap in ds_ap:
        ct_rows = []
        for mk, ten_ct in CT_TW + CT_DP:
            composite = f"{ten_ap}|{mk}"
            kh_val  = float(xuat_data.get(composite, 0.0))
            so_goc  = lay_so_goc_cho_ap(ten_ap, mk, du_no_baseline, ls_truoc)
            delta   = round(kh_val - so_goc, 1)
            if kh_val == 0 and so_goc == 0:
                continue
            nv_label = "TW" if mk.endswith("_TW") else "ĐP"
            ct_rows.append({"ten_ct": f"{ten_ct} ({nv_label})", "delta": delta, "kh": kh_val})
        if ct_rows:
            bang_rows.append({"is_ap": True, "ten_ap": ten_ap})
            for r in ct_rows:
                stt += 1
                r["stt"] = stt
                bang_rows.append(r)
                tong_delta += r["delta"]
                tong_kh    += r["kh"]

    loai_upper = "GIAO" if loai_van_ban == "giao" else "ĐIỀU CHỈNH"
    loai_lower = "giao" if loai_van_ban == "giao" else "điều chỉnh"

    # ── Tạo Document ──────────────────────────────────────────────────────────
    doc = Document()

    # Page setup A4, lề chuẩn công văn
    sec = doc.sections[0]
    sec.page_height   = Cm(29.7)
    sec.page_width    = Cm(21.0)
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin   = Cm(3.0)
    sec.right_margin  = Cm(2.0)

    # Font mặc định
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(13)

    def _para(text: str = "", bold: bool = False, italic: bool = False,
              align: str = "left", size: int = 13,
              space_before: int = 0, space_after: int = 0,
              first_indent_cm: float = 0, line_spacing: float = 0) -> None:
        p = doc.add_paragraph()
        p.alignment = {
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right":  WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }.get(align, WD_ALIGN_PARAGRAPH.LEFT)
        if space_before:
            p.paragraph_format.space_before = Pt(space_before)
        if space_after:
            p.paragraph_format.space_after = Pt(space_after)
        if first_indent_cm:
            p.paragraph_format.first_line_indent = Cm(first_indent_cm)
        if line_spacing:
            p.paragraph_format.line_spacing = line_spacing
        if text:
            run = p.add_run(text)
            run.font.name  = "Times New Roman"
            run.font.size  = Pt(size)
            run.font.bold  = bold
            run.font.italic = italic
        return p

    # ══════════════════════════════════════════════════════════════════════════
    # TRANG 1 — HEADER VĂN BẢN
    # ══════════════════════════════════════════════════════════════════════════

    # Bảng header 2 cột (ẩn border)
    hdr_tbl = doc.add_table(rows=1, cols=2)
    hdr_tbl.autofit = False
    hdr_tbl.columns[0].width = Cm(9)
    hdr_tbl.columns[1].width = Cm(9)
    _remove_table_borders(hdr_tbl)

    # Cột trái
    lc = hdr_tbl.rows[0].cells[0]
    lc.paragraphs[0].clear()
    p_l = lc.paragraphs[0]
    r = p_l.add_run(f"ỦY BAN NHÂN DÂN\n{xa.upper()}")
    r.font.name = "Times New Roman"; r.font.size = Pt(13); r.font.bold = True
    p_l2 = lc.add_paragraph("Số:       /TB-UBND")
    p_l2.runs[0].font.name = "Times New Roman"; p_l2.runs[0].font.size = Pt(13)

    # Cột phải
    rc = hdr_tbl.rows[0].cells[1]
    rc.paragraphs[0].clear()
    p_r = rc.paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p_r.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    r1.font.name = "Times New Roman"; r1.font.size = Pt(13); r1.font.bold = True
    p_r2 = rc.add_paragraph("Độc lập - Tự do - Hạnh phúc")
    p_r2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_r2.runs[0].font.name = "Times New Roman"
    p_r2.runs[0].font.size = Pt(13); p_r2.runs[0].font.bold = True
    p_r3 = rc.add_paragraph("─────────────────")
    p_r3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_r3.runs[0].font.name = "Times New Roman"; p_r3.runs[0].font.size = Pt(11)
    p_r4 = rc.add_paragraph("........., ngày     tháng     năm 202....")
    p_r4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_r4.runs[0].font.name = "Times New Roman"
    p_r4.runs[0].font.size = Pt(13); p_r4.runs[0].font.italic = True

    doc.add_paragraph()

    # Tiêu đề
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after  = Pt(6)
    r_t = p_title.add_run(f"THÔNG BÁO {loai_upper} KẾ HOẠCH DƯ NỢ\nNĂM {nam}")
    r_t.font.name = "Times New Roman"; r_t.font.size = Pt(14); r_t.font.bold = True

    if loai_van_ban == "dieu_chinh" and lich_su:
        p_lan = doc.add_paragraph()
        p_lan.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_lan = p_lan.add_run(f"(LẦN {len(lich_su) + 1})")
        r_lan.font.name = "Times New Roman"; r_lan.font.size = Pt(14); r_lan.font.bold = True

    doc.add_paragraph()

    # Kính gửi
    p_kg = doc.add_paragraph()
    p_kg.paragraph_format.first_line_indent = Cm(1.5)
    r_kg = p_kg.add_run(f"Kính gửi: Ông (bà) Trưởng các thôn, ấp thuộc xã {xa}")
    r_kg.font.name = "Times New Roman"; r_kg.font.size = Pt(13); r_kg.font.italic = True

    # Các đoạn nội dung chuẩn
    body_paras = [
        (f"Căn cứ Quyết định số        /QĐ-NHCS ngày     /     /      của Trưởng Ban đại "
         f"diện HĐQT NHCSXH tỉnh về việc {loai_lower} kế hoạch tín dụng năm {nam} cho các xã;"),
        "Căn cứ tổng hợp nhu cầu vốn tín dụng chính sách tại các thôn, ấp.",
        (f"Ủy ban nhân dân xã {xa} thông báo {loai_lower} kế hoạch dư nợ năm {nam} "
         f"của các thôn, ấp theo danh mục đính kèm."),
        ("Đề nghị ông (bà) Trưởng thôn, ấp căn cứ Thông báo này để chỉ đạo các Tổ TK&VV "
         "tại thôn tổ chức bình xét cho vay và lập hồ sơ vay vốn theo quy định của NHCSXH. "
         "Việc bình xét cho vay phải đảm bảo công khai, dân chủ, đúng đối tượng và có sự "
         "tham gia giám sát của Trưởng thôn, ấp, đại diện lãnh đạo của tổ chức chính trị "
         "- xã hội nhận ủy thác."),
        (f"Trên đây là chỉ tiêu kế hoạch dư nợ năm {nam} được Ủy ban nhân dân xã {xa} phê "
         f"duyệt, đề nghị ông (bà) Trưởng thôn, ấp nghiêm túc thực hiện."),
    ]
    for txt in body_paras:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(1.5)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(txt)
        r.font.name = "Times New Roman"; r.font.size = Pt(13)

    # Ký tên
    doc.add_paragraph()
    doc.add_paragraph()
    p_ky = doc.add_paragraph()
    p_ky.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_ky.paragraph_format.space_before = Pt(24)
    for line, bold, italic in [
        ("TM. ỦY BAN NHÂN DÂN XÃ", True, False),
        ("CHỦ TỊCH",               True, False),
        ("(Ký tên, đóng dấu)",     False, True),
    ]:
        r_ky = p_ky.add_run(f"{line}\n")
        r_ky.font.name = "Times New Roman"
        r_ky.font.size = Pt(13)
        r_ky.font.bold   = bold
        r_ky.font.italic = italic

    # ══════════════════════════════════════════════════════════════════════════
    # TRANG 2 — BẢNG DỮ LIỆU
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_page_break()

    p_h1 = doc.add_paragraph()
    p_h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_h1 = p_h1.add_run(f"ỦY BAN NHÂN DÂN XÃ {xa.upper()}")
    r_h1.font.name = "Times New Roman"; r_h1.font.size = Pt(14); r_h1.font.bold = True

    p_h2 = doc.add_paragraph()
    p_h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_h2 = p_h2.add_run(
        f"DANH MỤC {loai_upper} CHỈ TIÊU KẾ HOẠCH DƯ NỢ\n"
        f"NGUỒN VỐN TRUNG ƯƠNG & ĐỊA PHƯƠNG NĂM {nam}"
    )
    r_h2.font.name = "Times New Roman"; r_h2.font.size = Pt(14); r_h2.font.bold = True

    p_km = doc.add_paragraph()
    p_km.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_km = p_km.add_run(
        f"(Kèm theo Thông báo số       /TB-UBND ngày     /     /      của\n"
        f"Chủ tịch UBND xã {xa})"
    )
    r_km.font.name = "Times New Roman"; r_km.font.size = Pt(12); r_km.font.italic = True

    p_dv = doc.add_paragraph()
    p_dv.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_dv.paragraph_format.space_after = Pt(6)
    r_dv = p_dv.add_run("Đơn vị: triệu đồng")
    r_dv.font.name = "Times New Roman"; r_dv.font.size = Pt(12); r_dv.font.italic = True

    # Bảng chính
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    tbl.autofit = False
    tbl.columns[0].width = Cm(1.5)
    tbl.columns[1].width = Cm(9.5)
    tbl.columns[2].width = Cm(3.5)
    tbl.columns[3].width = Cm(3.5)

    # Header row
    hdr = tbl.rows[0].cells
    _set_cell_bg(hdr[0], "E8E8E8"); _cell_write(hdr[0], "STT",                    bold=True, align="center", size=12)
    _set_cell_bg(hdr[1], "E8E8E8"); _cell_write(hdr[1], "Chỉ tiêu",               bold=True, align="center", size=12)
    _set_cell_bg(hdr[2], "E8E8E8"); _cell_write(hdr[2], "Giao tăng/giảm",         bold=True, align="center", size=12)
    _set_cell_bg(hdr[3], "E8E8E8"); _cell_write(hdr[3], f"Chỉ tiêu KH\nnăm {nam}", bold=True, align="center", size=12)

    # Data rows
    for item in bang_rows:
        row = tbl.add_row()
        cells = row.cells
        if item.get("is_ap"):
            merged = cells[0].merge(cells[3])
            _set_cell_bg(merged, "E3F2FD")
            _cell_write(merged, f"  {item['ten_ap']}", bold=True, align="left", size=12)
        else:
            _cell_write(cells[0], str(item["stt"]),           align="center", size=12)
            _cell_write(cells[1], item["ten_ct"],              align="left",   size=12, indent_cm=0.3)
            _cell_write(cells[2], _fmt_delta(item["delta"]),   align="right",  size=12)
            _cell_write(cells[3], _fmt_money(item["kh"]),      align="right",  size=12)

    # Dòng Tổng cộng (double border-top)
    row_tc = tbl.add_row()
    tc = row_tc.cells
    merged_tc = tc[0].merge(tc[1])
    _set_cell_bg(merged_tc, "F5F5F5")
    _cell_write(merged_tc, "TỔNG CỘNG", bold=True, align="center", size=12)
    _set_cell_bg(tc[2], "F5F5F5")
    _cell_write(tc[2], _fmt_delta(round(tong_delta)), bold=True, align="right", size=12)
    _set_cell_bg(tc[3], "F5F5F5")
    _cell_write(tc[3], _fmt_money(round(tong_kh)),    bold=True, align="right", size=12)
    for c in [merged_tc, tc[2], tc[3]]:
        _set_cell_border_top_double(c)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
