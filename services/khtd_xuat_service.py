"""Xuất Excel / Word cho tab Xây dựng KHTD tương lai.

Functions:
  xuat_excel_1pgd()       → Excel 3 sheet (Biểu01C, Biểu02C, Thuyết minh) cho 1 PGD
  xuat_excel_tong_hop_cn()→ Excel tổng hợp toàn CN (1 sheet/năm)
  xuat_word_bao_cao_pgd() → Word tờ trình kế hoạch 1 PGD
  xuat_word_tong_hop_cn() → Word tổng hợp toàn CN
"""
from __future__ import annotations

from io import BytesIO
from datetime import date

import pandas as pd

from logger import get_logger

logger = get_logger(__name__)

from config import (
    CHUONG_TRINH_KHTD,
    DS_PGD,
    TEN_CHI_NHANH_HIEN_THI,
    THUYET_MINH_LABELS,
)
from data.pgd import pgd_slug as _slug
from services.khtd_import_service import (
    doc_bieu_01c_xd,
    doc_bieu_02c,
    doc_thuyet_minh,
    doc_trang_thai_approval,
    tong_hop_bieu_01c_cn,
    tong_hop_bieu_02c_cn,
    trang_thai_approval_cn,
)


# ── Excel: 1 PGD ─────────────────────────────────────────────────────────────

def xuat_excel_1pgd(pgd_ten: str, ds_nam: list[int], loai: str) -> bytes:
    """Xuất Excel 3 sheet cho 1 PGD: Biểu01C, Biểu02C, ThuyetMinh."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = Workbook()
    wb.remove(wb.active)

    header_fill = PatternFill("solid", fgColor="1F4E79")
    header_font = Font(color="FFFFFF", bold=True)
    sub_fill    = PatternFill("solid", fgColor="D6E4F0")
    thin = Side(border_style="thin", color="AAAAAA")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    def _hdr(ws, row, col, val, fill=None, bold=False):
        c = ws.cell(row=row, column=col, value=val)
        if fill:
            c.fill = fill
            c.font = Font(color="FFFFFF", bold=True) if fill == header_fill else Font(bold=bold)
        c.alignment = Alignment(wrap_text=True, vertical="center")
        c.border = border
        return c

    # ── Sheet 1: Biểu 02C (Dư nợ theo chương trình) ──────────────────────────
    ws1 = wb.create_sheet("Biểu 02C — Dư nợ")
    ws1.freeze_panes = "C3"
    headers = ["Năm", "Chương trình", "Nguồn vốn", "Dư nợ (triệu đồng)"]
    for ci, h in enumerate(headers, 1):
        _hdr(ws1, 1, ci, h, header_fill)
    ws1.column_dimensions["B"].width = 35
    ws1.column_dimensions["C"].width = 28

    row = 2
    for nam in ds_nam:
        data = doc_bieu_02c(pgd_ten, nam, loai)
        if not data:
            continue
        for nguon, ct_dict in data.items():
            for ma_key, vnd in ct_dict.items():
                ten_ct = CHUONG_TRINH_KHTD.get(ma_key, ma_key)
                ws1.cell(row=row, column=1, value=nam)
                ws1.cell(row=row, column=2, value=ten_ct)
                ws1.cell(row=row, column=3, value=nguon)
                ws1.cell(row=row, column=4, value=round(vnd / 1_000_000, 3))
                for ci in range(1, 5):
                    ws1.cell(row=row, column=ci).border = border
                row += 1

    # ── Sheet 2: Thuyết minh ──────────────────────────────────────────────────
    ws2 = wb.create_sheet("Thuyết minh")
    _hdr(ws2, 1, 1, "Chỉ tiêu", header_fill)
    _hdr(ws2, 1, 2, "Năm / Giá trị", header_fill)
    ws2.column_dimensions["A"].width = 40

    row = 2
    for nam in ds_nam:
        data = doc_thuyet_minh(pgd_ten, nam, loai)
        if not data:
            continue
        ws2.cell(row=row, column=1, value=f"── Năm {nam} ──")
        ws2.cell(row=row, column=1).fill = sub_fill
        ws2.cell(row=row, column=1).font = Font(bold=True)
        row += 1
        for key, label in THUYET_MINH_LABELS.items():
            ws2.cell(row=row, column=1, value=label)
            ws2.cell(row=row, column=2, value=int(data.get(key, 0) or 0))
            for ci in range(1, 3):
                ws2.cell(row=row, column=ci).border = border
            row += 1

    # ── Sheet 3: Trạng thái phê duyệt ────────────────────────────────────────
    ws3 = wb.create_sheet("Trạng thái")
    approval = doc_trang_thai_approval(pgd_ten, ds_nam, loai)
    ws3.cell(row=1, column=1, value="Đơn vị").font = Font(bold=True)
    ws3.cell(row=1, column=2, value=pgd_ten)
    ws3.cell(row=2, column=1, value="Loại kế hoạch").font = Font(bold=True)
    ws3.cell(row=2, column=2, value=loai)
    ws3.cell(row=3, column=1, value="Trạng thái").font = Font(bold=True)
    ws3.cell(row=3, column=2, value=approval.get("trang_thai", "nhap_lieu"))
    ws3.cell(row=4, column=1, value="Ngày nộp").font = Font(bold=True)
    ws3.cell(row=4, column=2, value=approval.get("ngay_nop") or "")
    ws3.cell(row=5, column=1, value="Người nộp").font = Font(bold=True)
    ws3.cell(row=5, column=2, value=approval.get("nguoi_nop") or "")
    ws3.cell(row=6, column=1, value="Ngày duyệt").font = Font(bold=True)
    ws3.cell(row=6, column=2, value=approval.get("ngay_duyet") or "")
    ws3.cell(row=7, column=1, value="Người duyệt").font = Font(bold=True)
    ws3.cell(row=7, column=2, value=approval.get("nguoi_duyet") or "")
    ws3.column_dimensions["A"].width = 20
    ws3.column_dimensions["B"].width = 30

    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ── Excel: Tổng hợp CN ────────────────────────────────────────────────────────

def xuat_excel_tong_hop_cn(ds_nam: list[int], loai: str) -> bytes:
    """Xuất Excel tổng hợp toàn CN: 1 sheet/năm + sheet trạng thái phê duyệt."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = Workbook()
    wb.remove(wb.active)

    header_fill = PatternFill("solid", fgColor="1F4E79")
    thin = Side(border_style="thin", color="AAAAAA")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    def _hdr(ws, row, col, val):
        c = ws.cell(row=row, column=col, value=val)
        c.fill = header_fill
        c.font = Font(color="FFFFFF", bold=True)
        c.alignment = Alignment(wrap_text=True, vertical="center")
        c.border = border

    for nam in ds_nam:
        ws = wb.create_sheet(f"Năm {nam}")
        ws.freeze_panes = "C2"
        hdrs = ["PGD", "Chương trình", "Nguồn vốn", "Dư nợ (triệu đồng)"]
        for ci, h in enumerate(hdrs, 1):
            _hdr(ws, 1, ci, h)
        ws.column_dimensions["A"].width = 22
        ws.column_dimensions["B"].width = 35
        ws.column_dimensions["C"].width = 25

        df = tong_hop_bieu_02c_cn(nam, loai)
        if df.empty:
            ws.cell(row=2, column=1, value="(Chưa có dữ liệu)")
            continue

        for ri, (_, r) in enumerate(df.iterrows(), 2):
            ws.cell(row=ri, column=1, value=r.get("PGD", ""))
            ws.cell(row=ri, column=2, value=r.get("ten_ct", ""))
            ws.cell(row=ri, column=3, value=r.get("nguon_von", ""))
            du_no = r.get("du_no_vnd", 0)
            ws.cell(row=ri, column=4, value=round(float(du_no) / 1_000_000, 3))
            for ci in range(1, 5):
                ws.cell(row=ri, column=ci).border = border

    # Sheet phê duyệt
    ws_pd = wb.create_sheet("Phê duyệt")
    hdrs_pd = ["PGD", "Trạng thái", "Ngày nộp", "Người nộp", "Ngày duyệt", "Người duyệt", "Ý kiến"]
    for ci, h in enumerate(hdrs_pd, 1):
        _hdr(ws_pd, 1, ci, h)
    ws_pd.column_dimensions["A"].width = 22
    ws_pd.column_dimensions["B"].width = 16
    ws_pd.column_dimensions["G"].width = 40

    approval_map = trang_thai_approval_cn(ds_nam, loai)
    for ri, pgd in enumerate(DS_PGD, 2):
        ap = approval_map.get(pgd, {})
        ws_pd.cell(row=ri, column=1, value=pgd)
        ws_pd.cell(row=ri, column=2, value=ap.get("trang_thai", "nhap_lieu"))
        ws_pd.cell(row=ri, column=3, value=ap.get("ngay_nop") or "")
        ws_pd.cell(row=ri, column=4, value=ap.get("nguoi_nop") or "")
        ws_pd.cell(row=ri, column=5, value=ap.get("ngay_duyet") or "")
        ws_pd.cell(row=ri, column=6, value=ap.get("nguoi_duyet") or "")
        ws_pd.cell(row=ri, column=7, value=ap.get("y_kien") or "")
        for ci in range(1, 8):
            ws_pd.cell(row=ri, column=ci).border = Border(
                left=Side(border_style="thin", color="AAAAAA"),
                right=Side(border_style="thin", color="AAAAAA"),
                top=Side(border_style="thin", color="AAAAAA"),
                bottom=Side(border_style="thin", color="AAAAAA"),
            )

    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ── Word: Tờ trình 1 PGD ──────────────────────────────────────────────────────

def xuat_word_bao_cao_pgd(pgd_ten: str, ds_nam: list[int], loai: str) -> bytes:
    """Tờ trình Kế hoạch Tín dụng cho 1 PGD — định dạng Word."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()

    # Lề trang
    for sec in doc.sections:
        sec.top_margin = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2)

    def _heading(text: str, level: int = 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14 if level == 1 else 12)
        if level == 1:
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    def _para(text: str, bold: bool = False):
        p = doc.add_paragraph(text)
        if bold:
            for run in p.runs:
                run.bold = True
        return p

    giai_doan = str(ds_nam[0]) if len(ds_nam) == 1 else f"{ds_nam[0]}–{ds_nam[-1]}"
    loai_str = {"1n": "1 năm", "3n": "3 năm", "5n": "5 năm (2026–2030)"}.get(loai, loai)

    _heading(f"TỜ TRÌNH KẾ HOẠCH TÍN DỤNG {loai_str.upper()}", level=1)
    _heading(f"Giai đoạn: {giai_doan}", level=1)

    doc.add_paragraph()
    _para(f"Đơn vị lập: {pgd_ten}", bold=True)
    _para(f"Trực thuộc: {TEN_CHI_NHANH_HIEN_THI}", bold=True)
    _para(f"Ngày lập: {date.today().strftime('%d/%m/%Y')}")

    doc.add_paragraph()
    approval = doc_trang_thai_approval(pgd_ten, ds_nam, loai)
    tt_map = {
        "nhap_lieu": "Đang nhập liệu",
        "da_nop": "Đã nộp — Chờ duyệt",
        "da_duyet": "Đã duyệt",
        "tu_choi": "Bị trả lại",
    }
    _para(f"Trạng thái phê duyệt: {tt_map.get(approval.get('trang_thai',''), 'Không rõ')}", bold=True)
    if approval.get("ngay_duyet"):
        _para(f"Ngày duyệt: {approval['ngay_duyet'][:10]}")
    if approval.get("y_kien"):
        _para(f"Ý kiến Chi nhánh: {approval['y_kien']}")

    # Bảng dư nợ theo năm
    for nam in ds_nam:
        doc.add_paragraph()
        _heading(f"I. Dư nợ dự kiến — Năm {nam}", level=2)

        data_02c = doc_bieu_02c(pgd_ten, nam, loai)
        if not data_02c:
            _para("(Chưa có dữ liệu Biểu 02C)")
            continue

        rows_data: list[tuple[str, str, float]] = []
        tong = 0.0
        for nguon, ct_dict in data_02c.items():
            for ma_key, vnd in ct_dict.items():
                ten_ct = CHUONG_TRINH_KHTD.get(ma_key, ma_key)
                trieu = round(float(vnd) / 1_000_000, 3)
                rows_data.append((ten_ct, nguon, trieu))
                tong += trieu

        tbl = doc.add_table(rows=1, cols=4)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.style = "Table Grid"
        hdrs = ["Chương trình", "Nguồn vốn", "Dư nợ (tr.đ)", ""]
        for ci, h in enumerate(hdrs[:3]):
            cell = tbl.rows[0].cells[ci]
            cell.text = h
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for ten_ct, nguon, trieu in rows_data:
            row = tbl.add_row()
            row.cells[0].text = ten_ct
            row.cells[1].text = nguon
            row.cells[2].text = f"{trieu:,.3f}".replace(",", ".")
            row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        tong_row = tbl.add_row()
        tong_row.cells[0].text = "TỔNG CỘNG"
        tong_row.cells[0].paragraphs[0].runs[0].bold = True
        tong_row.cells[2].text = f"{tong:,.3f}".replace(",", ".")
        tong_row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Thuyết minh tóm tắt
    for nam in ds_nam:
        tm = doc_thuyet_minh(pgd_ten, nam, loai)
        if not tm:
            continue
        doc.add_paragraph()
        _heading(f"II. Thuyết minh chỉ tiêu — Năm {nam}", level=2)
        for key, label in THUYET_MINH_LABELS.items():
            val = int(tm.get(key, 0) or 0)
            if val:
                _para(f"• {label}: {val:,}")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Word: Tổng hợp CN ────────────────────────────────────────────────────────

def xuat_word_tong_hop_cn(ds_nam: list[int], loai: str) -> bytes:
    """Báo cáo tổng hợp toàn Chi nhánh — Word."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2)

    def _heading(text: str, level: int = 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14 if level == 1 else 12)
        if level == 1:
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    giai_doan = str(ds_nam[0]) if len(ds_nam) == 1 else f"{ds_nam[0]}–{ds_nam[-1]}"
    loai_str = {"1n": "1 năm", "3n": "3 năm", "5n": "5 năm (2026–2030)"}.get(loai, loai)

    _heading(f"BÁO CÁO TỔNG HỢP KẾ HOẠCH TÍN DỤNG {loai_str.upper()}", level=1)
    _heading(f"Giai đoạn: {giai_doan}", level=1)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(TEN_CHI_NHANH_HIEN_THI).bold = True

    doc.add_paragraph(f"Ngày lập: {date.today().strftime('%d/%m/%Y')}")

    # Trạng thái phê duyệt
    doc.add_paragraph()
    _heading("I. Trạng thái phê duyệt các PGD", level=2)

    approval_map = trang_thai_approval_cn(ds_nam, loai)
    tt_label = {
        "nhap_lieu": "Đang nhập liệu",
        "da_nop": "Đã nộp",
        "da_duyet": "Đã duyệt",
        "tu_choi": "Bị trả lại",
    }

    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    for ci, h in enumerate(["PGD", "Trạng thái", "Ngày nộp"]):
        cell = tbl.rows[0].cells[ci]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True

    da_duyet = 0
    for pgd in DS_PGD:
        ap = approval_map.get(pgd, {})
        tt = ap.get("trang_thai", "nhap_lieu")
        if tt == "da_duyet":
            da_duyet += 1
        row = tbl.add_row()
        row.cells[0].text = pgd
        row.cells[1].text = tt_label.get(tt, tt)
        row.cells[2].text = (ap.get("ngay_nop") or "")[:10]

    doc.add_paragraph()
    doc.add_paragraph(f"Tổng đã duyệt: {da_duyet}/{len(DS_PGD)} PGD")

    # Tổng hợp dư nợ toàn CN
    for nam in ds_nam:
        doc.add_paragraph()
        _heading(f"II. Tổng hợp dư nợ dự kiến — Năm {nam}", level=2)

        df = tong_hop_bieu_02c_cn(nam, loai)
        if df.empty:
            doc.add_paragraph("(Chưa có dữ liệu)")
            continue

        df_ct = (
            df.groupby(["ten_ct", "nguon_von"], as_index=False)
              .agg(du_no_vnd=("du_no_vnd", "sum"))
        )
        tong = df["du_no_vnd"].sum()

        tbl2 = doc.add_table(rows=1, cols=3)
        tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl2.style = "Table Grid"
        for ci, h in enumerate(["Chương trình", "Nguồn vốn", "Dư nợ (tr.đ)"]):
            tbl2.rows[0].cells[ci].text = h
            tbl2.rows[0].cells[ci].paragraphs[0].runs[0].bold = True

        for _, r in df_ct.iterrows():
            row = tbl2.add_row()
            row.cells[0].text = str(r.get("ten_ct", ""))
            row.cells[1].text = str(r.get("nguon_von", ""))
            trieu = round(float(r.get("du_no_vnd", 0)) / 1_000_000, 3)
            row.cells[2].text = f"{trieu:,.3f}".replace(",", ".")
            row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        tong_row = tbl2.add_row()
        tong_row.cells[0].text = "TỔNG CỘNG"
        tong_row.cells[0].paragraphs[0].runs[0].bold = True
        tong_trieu = round(float(tong) / 1_000_000, 3)
        tong_row.cells[2].text = f"{tong_trieu:,.3f}".replace(",", ".")
        tong_row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
