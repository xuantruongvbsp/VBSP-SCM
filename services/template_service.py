"""Template-based document generation cho VBSP-SCM."""
from __future__ import annotations
import io, os, tempfile, logging
from datetime import date, datetime
from pathlib import Path
from typing import Any

from docxtpl import DocxTemplate

from config import (
    BASE_DIR,
    COT_DU_NO_QH,
    COT_DVUT,
    COT_LAI_TON,
    COT_MUC_VAY,
    COT_SO_DU_TG,
    COT_SO_KU,
    COT_TEN_CT,
    COT_TEN_KH,
    COT_TEN_TO,
    COT_TEN_XA,
    COT_TONG_DU_NO,
)

TEMPLATE_DIR = BASE_DIR / "templates"
logger = logging.getLogger(__name__)

# ── Tên file template chuẩn ──────────────────────────
TMPL_MAU06    = "mau_06td.docx"
TMPL_MAU06A   = "mau_06atd.docx"
TMPL_MAU15    = "mau_15td.docx"
TMPL_MAU16    = "mau_16td.docx"   # BB kiểm tra Tổ TK&VV
TMPL_KH_KT    = "ke_hoach_kt.docx"
TMPL_BB_XMN   = "bb_xac_minh_no.docx"

# ── Mẫu xử lý nợ rủi ro QĐ 62/2015/QĐ-TTg ──────────────────────────────────
TMPL_13XLN     = "mau_13xln.docx"     # Đề nghị khoanh nợ
TMPL_14XLN     = "mau_14xln.docx"     # Đề nghị xóa nợ
TMPL_TT_KHOANH = "to_trinh_khoanh_no.docx"  # Tờ trình khoanh nợ
TMPL_TT_XOA    = "to_trinh_xoa_no.docx"     # Tờ trình xóa nợ


def co_template(ten_mau: str) -> bool:
    """Kiểm tra file template có tồn tại không."""
    return (TEMPLATE_DIR / ten_mau).exists()


def dien_template(ten_mau: str, context: dict[str, Any]) -> bytes:
    """
    Điền dữ liệu vào template .docx → trả về bytes.
    Raise FileNotFoundError nếu template chưa có.
    """
    path = TEMPLATE_DIR / ten_mau
    if not path.exists():
        raise FileNotFoundError(
            f"Template '{ten_mau}' chưa có trong {TEMPLATE_DIR}. "
            f"Vui lòng upload template vào thư mục templates/."
        )
    tpl = DocxTemplate(str(path))
    tpl.render(context)
    buf = io.BytesIO()
    tpl.save(buf)
    return buf.getvalue()


def docx_to_pdf(docx_bytes: bytes) -> bytes | None:
    """
    Convert .docx → .pdf bằng docx2pdf (cần MS Word trên Windows).
    Trả về None nếu không convert được — caller tự xử lý fallback.
    """
    try:
        from docx2pdf import convert
        with tempfile.TemporaryDirectory() as tmp:
            inp = os.path.join(tmp, "input.docx")
            out = os.path.join(tmp, "input.pdf")
            with open(inp, "wb") as f:
                f.write(docx_bytes)
            convert(inp, out)
            if os.path.exists(out):
                with open(out, "rb") as f:
                    return f.read()
        return None
    except Exception as e:  # conv: skip
        logger.warning(f"docx_to_pdf failed: {e}")
        return None


def docx_bytes_to_pdf(docx_bytes: bytes) -> bytes | None:
    """Convert .docx bytes → .pdf bytes dùng MS Word (Windows)."""
    try:
        from docx2pdf import convert
        import tempfile, os
        with tempfile.TemporaryDirectory() as tmp:
            inp = os.path.join(tmp, "input.docx")
            out = os.path.join(tmp, "input.pdf")
            with open(inp, "wb") as f:
                f.write(docx_bytes)
            convert(inp, out)
            if os.path.exists(out):
                with open(out, "rb") as f:
                    return f.read()
        return None
    except Exception as e:  # conv: skip
        import logging
        logging.warning(f"docx_to_pdf failed: {e}")
        return None


def nut_tai_word_va_pdf(
    docx_bytes: bytes,
    ten_file_goc: str,
    key_prefix: str,
) -> None:
    """
    Lưu Word + PDF bytes vào session_state.
    Gọi bên trong if st.button().
    """
    import streamlit as st

    st.session_state[f"_w_bytes_{key_prefix}"] = docx_bytes
    st.session_state[f"_f_name_{key_prefix}"] = ten_file_goc
    pdf_bytes = docx_to_pdf(docx_bytes)
    if pdf_bytes:
        st.session_state[f"_p_bytes_{key_prefix}"] = pdf_bytes


def hien_thi_nut_tai(key_prefix: str) -> None:
    """
    Hiển thị 2 nút download: Word + PDF từ session_state.
    Gọi NGOÀI if st.button() — luôn render.
    """
    import streamlit as st

    docx_bytes = st.session_state.get(f"_w_bytes_{key_prefix}")
    if not docx_bytes:
        return

    col1, col2 = st.columns(2)
    with col1:
        st.download_button(
            "⬇️ Tải Word (.docx)",
            data=docx_bytes,
            file_name=f"{st.session_state.get(f'_f_name_{key_prefix}', 'file')}.docx",
            mime="application/vnd.openxmlformats-officedocument"
                 ".wordprocessingml.document",
            key=f"{key_prefix}_dl_docx",
            use_container_width=True,
        )
    with col2:
        pdf_bytes = st.session_state.get(f"_p_bytes_{key_prefix}")
        if pdf_bytes:
            st.download_button(
                "⬇️ Tải PDF",
                data=pdf_bytes,
                file_name=f"{st.session_state.get(f'_f_name_{key_prefix}', 'file')}.pdf",
                mime="application/pdf",
                key=f"{key_prefix}_dl_pdf",
                use_container_width=True,
            )
        else:
            st.caption("⚠️ PDF: cần MS Word trên server")


def _fmt_bang_ty(x, so_le: int = 0) -> str:
    try:
        x = float(x)
        if x == 0:
            return "—"
        trieu = x / 1_000_000
        s = f"{trieu:,.{so_le}f}".replace(",", "X").replace(".", ",").replace("X", ".")
        return s
    except Exception:
        return "—"


def _fmt_so(x) -> str:
    try:
        return f"{int(x):,}".replace(",", ".")
    except Exception:
        return "—"


def _uyt_style_doc(doc) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = __import__("docx").shared.Pt(12)


def _uyt_xoa_border_table(tbl) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    for cell in tbl.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for bn in ["top", "left", "bottom", "right"]:
            b = OxmlElement(f"w:{bn}")
            b.set(qn("w:val"), "none")
            tcBorders.append(b)
        tcPr.append(tcBorders)


def _uyt_parse_date(s: str | None) -> date:
    if not s:
        return date.today()
    try:
        return datetime.strptime(s, "%Y-%m-%d").date()
    except Exception:
        return date.today()


def _uyt_add_header_quoc_hieu(doc, don_vi: str, so_vb: str, dia_danh: str, ngay_ky: date) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    for cell in t.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for border_name in ["top", "left", "bottom", "right"]:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "none")
            tcBorders.append(border)
        tcPr.append(tcBorders)

    cell_l = t.rows[0].cells[0]
    p = cell_l.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run((don_vi or "").upper())
    run.bold = True
    run.font.size = Pt(12)
    p2 = cell_l.add_paragraph(f"Số: {so_vb}")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell_r = t.rows[0].cells[1]
    p3 = cell_r.paragraphs[0]
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    run3.bold = True
    run3.font.size = Pt(12)
    p4 = cell_r.add_paragraph("Độc lập - Tự do - Hạnh phúc")
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.runs[0] if p4.runs else p4.add_run("Độc lập - Tự do - Hạnh phúc")
    run4.bold = True

    p5 = cell_r.add_paragraph(
        f"{dia_danh}, ngày {ngay_ky.day} tháng {ngay_ky.month} năm {ngay_ky.year}"
    )
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def tao_word_uythac_ke_hoach(du_lieu: dict, ds_to: list) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    doc = Document()
    _uyt_style_doc(doc)
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    _uyt_add_header_quoc_hieu(
        doc,
        don_vi=du_lieu.get("don_vi_kt", ""),
        so_vb=du_lieu.get("so_vb", ""),
        dia_danh=du_lieu.get("dia_danh", ""),
        ngay_ky=du_lieu.get("ngay_ky", date.today()),
    )

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = title.add_run(
        "KẾ HOẠCH\nKiểm tra giám sát hoạt động ủy thác "
        f"năm {du_lieu.get('nam_kh', date.today().year)}"
    )
    run_t.bold = True
    run_t.font.size = Pt(13)
    doc.add_paragraph("Căn cứ văn bản số 727/HD-NHCS ngày 11/02/2026 của Tổng Giám đốc NHCSXH;")

    doc.add_paragraph("I. MỤC ĐÍCH, YÊU CẦU").runs[0].bold = True
    doc.add_paragraph(f"1. Mục đích\n{du_lieu.get('muc_dich','')}")
    doc.add_paragraph(f"2. Yêu cầu\n{du_lieu.get('yeu_cau','')}")

    doc.add_paragraph("II. KẾ HOẠCH KIỂM TRA").runs[0].bold = True
    doc.add_paragraph(f"1. Nội dung, thời hiệu kiểm tra\n{du_lieu.get('noi_dung_kt','')}")

    doc.add_paragraph("2. Đối tượng được kiểm tra")
    if ds_to:
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Table Grid"
        hdrs = ["STT", "Hội đoàn thể", "Xã/Phường", "Tên Tổ TK&VV"]
        for i, h in enumerate(hdrs):
            cell = tbl.rows[0].cells[i]
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for idx, to in enumerate(ds_to, 1):
            row = tbl.add_row()
            row.cells[0].text = str(idx)
            row.cells[1].text = str(to.get(COT_DVUT, ""))
            row.cells[2].text = str(to.get(COT_TEN_XA, ""))
            row.cells[3].text = str(to.get(COT_TEN_TO, ""))
        doc.add_paragraph()

    doc.add_paragraph(f"3. Thành phần Đoàn kiểm tra\n{du_lieu.get('thanh_phan','')}")

    doc.add_paragraph("III. KẾ HOẠCH GIÁM SÁT").runs[0].bold = True
    doc.add_paragraph(f"1. Nội dung, thời hiệu giám sát\n{du_lieu.get('noi_dung_gs','')}")
    doc.add_paragraph(f"2. Phân công cán bộ giám sát\n{du_lieu.get('phan_cong_gs','')}")

    doc.add_paragraph("IV. TỔ CHỨC THỰC HIỆN").runs[0].bold = True
    doc.add_paragraph(du_lieu.get("to_chuc", ""))

    doc.add_paragraph()
    ky = doc.add_table(rows=1, cols=2)
    ky.style = "Table Grid"
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    for cell in ky.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for border_name in ["top", "left", "bottom", "right"]:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "none")
            tcBorders.append(border)
        tcPr.append(tcBorders)
    ky.rows[0].cells[0].text = "Nơi nhận:\n- NHCSXH;\n- Lưu: VT."
    p_ky = ky.rows[0].cells[1].paragraphs[0]
    p_ky.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ky.add_run("CHỦ TỊCH\n(Ký, ghi rõ họ tên)\n\n\n\n").bold = True
    p_ky.add_run(du_lieu.get("chu_tich", ""))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def tao_word_uythac_mau15(du_lieu: dict, df_to) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    doc = Document()
    _uyt_style_doc(doc)
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(1.5)
        section.page_width = Cm(29.7)
        section.page_height = Cm(21)

    _uyt_add_header_quoc_hieu(
        doc,
        don_vi=du_lieu.get("pgd", ""),
        so_vb="15/TD",
        dia_danh=du_lieu.get("ten_xa", ""),
        ngay_ky=du_lieu.get("ngay_chot", date.today()),
    )

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = title.add_run("DANH SÁCH ĐỐI CHIẾU SỐ DƯ TIỀN VAY VÀ TIỀN GỬI")
    run_t.bold = True
    run_t.font.size = Pt(13)
    doc.add_paragraph(
        f"Đến ngày {du_lieu.get('ngay_chot', date.today()).strftime('%d/%m/%Y')}"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        f"Tổ trưởng: {du_lieu.get('to_truong','')}    "
        f"Mã Tổ: {du_lieu.get('ma_to','')}    "
        f"Địa chỉ: {du_lieu.get('dia_chi','')}"
    )
    doc.add_paragraph("Đơn vị tính: đồng")

    headers = [
        "STT",
        "Họ và tên KH / Chương trình",
        "Mã khoản vay",
        "Nợ gốc (NH)",
        "Nợ lãi (NH)",
        "Số dư TG (NH)",
        "Nợ gốc (KH)",
        "Nợ lãi (KH)",
        "Số dư TG (KH)",
        "CL Nợ gốc",
        "CL Nợ lãi",
        "CL Số dư TG",
        "Nguyên nhân CL",
        "Chữ ký KH",
    ]
    tbl = doc.add_table(rows=2, cols=len(headers))
    tbl.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].font.size = Pt(9)

    stt_row = tbl.rows[1]
    for i in range(len(headers)):
        stt_row.cells[i].text = str(i + 1)
        stt_row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        stt_row.cells[i].paragraphs[0].runs[0].font.size = Pt(9)

    tong_goc = tong_lai = tong_tg = 0
    for idx, row in df_to.iterrows():
        r = tbl.add_row()
        ten_kh = str(row.get(COT_TEN_KH, ""))
        ten_ct = str(row.get(COT_TEN_CT, ""))
        ma_ku = str(row.get(COT_SO_KU, ""))
        goc = float(row.get(COT_TONG_DU_NO, 0) or 0)
        lai = float(row.get("Nợ lãi", 0) or 0)
        tg = float(row.get(COT_SO_DU_TG, 0) or 0)
        tong_goc += goc
        tong_lai += lai
        tong_tg += tg

        vals = [
            str(idx + 1),
            f"{ten_kh}\n{ten_ct}",
            ma_ku,
            f"{goc:,.0f}",
            f"{lai:,.0f}",
            f"{tg:,.0f}",
            "",
            "",
            "",
            "",
            "",
            "",
            "",
            "",
        ]
        for i, v in enumerate(vals):
            r.cells[i].text = v
            r.cells[i].paragraphs[0].runs[0].font.size = Pt(9)

    r_tong = tbl.add_row()
    r_tong.cells[0].merge(r_tong.cells[2])
    r_tong.cells[0].text = "Tổng cộng"
    r_tong.cells[0].paragraphs[0].runs[0].bold = True
    r_tong.cells[3].text = f"{tong_goc:,.0f}"
    r_tong.cells[4].text = f"{tong_lai:,.0f}"
    r_tong.cells[5].text = f"{tong_tg:,.0f}"

    doc.add_paragraph()
    p_kt = doc.add_paragraph()
    p_kt.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_kt.add_run(
        "CÁN BỘ ĐỐI CHIẾU\n(Ký, ghi rõ họ tên)\n\n\n\n" f"{du_lieu.get('can_bo_kt','')}"
    ).bold = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def tao_word_uythac_mau06(du_lieu: dict, df_m06, loai: str = "06") -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    _uyt_style_doc(doc)
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.3)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(2.0)
        section.page_width = Cm(29.7)
        section.page_height = Cm(21)

    def _xoa_border_row(tbl):
        for cell in tbl.rows[0].cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for bn in ["top", "left", "bottom", "right"]:
                b = OxmlElement(f"w:{bn}")
                b.set(qn("w:val"), "none")
                tcBorders.append(b)
            tcPr.append(tcBorders)

    def _cell_text(cell, text, bold=False, size=10, align=None):
        p = cell.paragraphs[0]
        if align:
            p.alignment = align
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)

    tbl_h = doc.add_table(rows=1, cols=3)
    tbl_h.style = "Table Grid"
    _xoa_border_row(tbl_h)

    _cell_text(
        tbl_h.rows[0].cells[0],
        f"Đơn vị kiểm tra: {du_lieu.get('don_vi_kt', '')}",
        size=11,
    )

    c1 = tbl_h.rows[0].cells[1]
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    r1.bold = True
    r1.font.size = Pt(12)
    p1b = c1.add_paragraph("Độc lập - Tự do - Hạnh phúc")
    p1b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if p1b.runs:
        p1b.runs[0].bold = True
    p1c = c1.add_paragraph("PHIẾU KIỂM TRA SỬ DỤNG VỐN VAY")
    p1c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if p1c.runs:
        p1c.runs[0].bold = True
        p1c.runs[0].font.size = Pt(13)

    c2 = tbl_h.rows[0].cells[2]
    p2 = c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = p2.add_run(f"Mẫu số {loai}/TD")
    r2.bold = True
    r2.font.size = Pt(11)
    for txt in ["    Lập 02 liên:", "- 01 liên chính lưu NH;", "- 01 liên phô tô lưu Đ.v k.tra."]:
        p_tmp = c2.add_paragraph(txt)
        if p_tmp.runs:
            p_tmp.runs[0].font.size = Pt(10)

    ngay_kt = du_lieu.get("ngay_kt", date.today())
    p_cb1 = doc.add_paragraph()
    p_cb1.add_run("Họ và tên cán bộ kiểm tra: \t1. Ông (bà): ")
    p_cb1.add_run(f"{du_lieu.get('can_bo_1', '')}   Chức vụ: {du_lieu.get('chuc_vu_1', '')}")
    doc.add_paragraph(
        f"\t2. Ông (bà): {du_lieu.get('can_bo_2', '')}   Chức vụ: {du_lieu.get('chuc_vu_2', '')}"
    )

    doc.add_paragraph(
        "Thời điểm kiểm tra: ............   "
        f"Địa bàn kiểm tra: {du_lieu.get('dia_ban', '')}   "
        f"Tổ TK&VV: {du_lieu.get('ten_to', '')}"
    )
    p_dvt = doc.add_paragraph("Đơn vị tính: triệu đồng")
    p_dvt.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    if loai == "06":
        n_data = len(df_m06)
        tbl = doc.add_table(rows=3 + n_data + 1, cols=15)
        tbl.style = "Table Grid"

        r0 = tbl.rows[0]
        r0.cells[0].merge(r0.cells[6])
        r0.cells[7].merge(r0.cells[14])
        for ci, txt in [(0, "PHẦN GHI THEO HỒ SƠ CHO VAY"), (7, "PHẦN KIỂM TRA THỰC TẾ TẠI KHÁCH HÀNG")]:
            r0.cells[ci].text = txt
            r0.cells[ci].paragraphs[0].runs[0].bold = True
            r0.cells[ci].paragraphs[0].runs[0].font.size = Pt(8)
            r0.cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        r1 = tbl.rows[1]
        r1.cells[9].merge(r1.cells[11])
        col1_names = [
            (0, "STT"),
            (1, "Họ và tên người vay"),
            (2, "Mã khoản vay"),
            (3, "Chương trình cho vay"),
            (4, "Số tiền giải ngân (tr.đ)"),
            (5, "Dư nợ đến ngày KT (tr.đ)"),
            (6, "Mục đích sử dụng vốn"),
            (7, "Tổng tiền thực nhận (tr.đ)"),
            (8, "Dư nợ thực tế (tr.đ)"),
            (9, "Thực tế sử dụng vốn"),
            (12, "Hiệu quả ĐT"),
            (13, "Nợ lãi (tr.đ)"),
            (14, "Chữ ký KH"),
        ]
        for ci, name in col1_names:
            r1.cells[ci].text = name
            r1.cells[ci].paragraphs[0].runs[0].bold = True
            r1.cells[ci].paragraphs[0].runs[0].font.size = Pt(8)
            r1.cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        r2r = tbl.rows[2]
        for ci, name in [(9, "Vào việc"), (10, "Đúng MĐ"), (11, "Sai MĐ")]:
            r2r.cells[ci].text = name
            r2r.cells[ci].paragraphs[0].runs[0].bold = True
            r2r.cells[ci].paragraphs[0].runs[0].font.size = Pt(8)
            r2r.cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        tong_gn = tong_dn = 0
        for seq, (_, row) in enumerate(df_m06.iterrows()):
            r = tbl.rows[3 + seq]
            gn = float(row.get(COT_MUC_VAY, 0) or 0)
            dn = float(row.get(COT_TONG_DU_NO, 0) or 0)
            tong_gn += gn
            tong_dn += dn
            vals = {
                0: str(seq + 1),
                1: str(row.get(COT_TEN_KH, "")),
                2: str(row.get(COT_SO_KU, "")),
                3: str(row.get(COT_TEN_CT, "")),
                4: _fmt_bang_ty(gn),
                5: _fmt_bang_ty(dn),
                6: str(row.get("Mục đích sử dụng vốn vay", "")),
            }
            for ci, v in vals.items():
                r.cells[ci].text = v
                if r.cells[ci].paragraphs[0].runs:
                    r.cells[ci].paragraphs[0].runs[0].font.size = Pt(8)

        r_c = tbl.rows[3 + n_data]
        r_c.cells[0].merge(r_c.cells[3])
        r_c.cells[0].text = "Cộng"
        r_c.cells[0].paragraphs[0].runs[0].bold = True
        r_c.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for ci, v in [(4, _fmt_bang_ty(tong_gn)), (5, _fmt_bang_ty(tong_dn))]:
            r_c.cells[ci].text = v
            if r_c.cells[ci].paragraphs[0].runs:
                r_c.cells[ci].paragraphs[0].runs[0].font.size = Pt(9)
    else:
        if not df_m06.empty:
            row = df_m06.iloc[0]
            gn = float(row.get(COT_MUC_VAY, 0) or 0)
            dn = float(row.get(COT_TONG_DU_NO, 0) or 0)
            lai = float(row.get("Nợ lãi", 0) or 0)
            tg = float(row.get(COT_SO_DU_TG, 0) or 0)
            doc.add_paragraph(
                f"Họ và tên: {row.get(COT_TEN_KH, '')}\n"
                f"Mã khoản vay: {row.get(COT_SO_KU, '')}\n"
                f"Chương trình: {row.get(COT_TEN_CT, '')}\n"
                f"Số tiền giải ngân: {_fmt_bang_ty(gn)} triệu đồng\n"
                f"Dư nợ gốc: {_fmt_bang_ty(dn)} triệu đồng\n"
                f"Nợ lãi: {_fmt_bang_ty(lai)} triệu đồng\n"
                f"Số dư tiền gửi TK: {_fmt_bang_ty(tg)} triệu đồng"
            )
            doc.add_paragraph("Mục đích vay vốn: " + str(row.get("Mục đích sử dụng vốn vay", "")))
            doc.add_paragraph(
                "Thực tế sử dụng vốn:\n"
                "- Sử dụng đúng mục đích: ............... triệu đồng\n"
                "- Sử dụng sai mục đích: ................. triệu đồng"
            )
            doc.add_paragraph("Hiệu quả đầu tư: ................................")
            doc.add_paragraph("Khả năng trả nợ: ................................")

    p_nx = doc.add_paragraph()
    p_nx.add_run("Nhận xét:").bold = True

    nhan_xet_chung = du_lieu.get("nhan_xet_chung", "")
    so_kh_kt = du_lieu.get("so_kh_kt", len(df_m06))
    tong_tien_kt = du_lieu.get(
        "tong_tien_kt",
        df_m06[COT_TONG_DU_NO].sum() if COT_TONG_DU_NO in df_m06.columns else 0,
    )
    so_kh_dung = du_lieu.get("so_kh_dung", "")
    so_tien_dung = du_lieu.get("so_tien_dung", "")
    ty_trong_dung = du_lieu.get("ty_trong_dung", "")
    so_kh_sai = du_lieu.get("so_kh_sai", "")
    so_tien_sai = du_lieu.get("so_tien_sai", "")
    ty_trong_sai = du_lieu.get("ty_trong_sai", "")
    bien_phap = du_lieu.get("bien_phap", "")

    doc.add_paragraph(f"1. Tình hình thực hiện phương án vay vốn: {nhan_xet_chung}")
    doc.add_paragraph(
        f"2. Kiểm tra, đối chiếu thực tế được {so_kh_kt} KH, "
        f"số tiền {_fmt_bang_ty(float(tong_tien_kt or 0))} triệu đồng. Trong đó:"
    )
    doc.add_paragraph(
        f"- Số KH đúng MĐ: {so_kh_dung} KH, số tiền: {so_tien_dung}, tỷ trọng: {ty_trong_dung}%."
    )
    doc.add_paragraph(
        f"- Số KH sai MĐ: {so_kh_sai} KH, số tiền: {so_tien_sai}, tỷ trọng: {ty_trong_sai}%."
    )
    p_bp = doc.add_paragraph()
    p_bp.add_run("Biện pháp xử lý: ").bold = True
    p_bp.add_run(bien_phap)

    doc.add_paragraph()
    ky = doc.add_table(rows=1, cols=2)
    ky.style = "Table Grid"
    _xoa_border_row(ky)

    p_ck = ky.rows[0].cells[0].paragraphs[0]
    p_ck.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ck.add_run(
        ".............................................\n"
        "CÁN BỘ CHỨNG KIẾN (nếu có)\n"
        "(Ký, ghi rõ họ tên)\n\n\n"
    )

    p_cbkt = ky.rows[0].cells[1].paragraphs[0]
    p_cbkt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cbkt.add_run(
        f"Ngày {ngay_kt.day} tháng {ngay_kt.month} năm {ngay_kt.year}\n"
        "CÁN BỘ KIỂM TRA\n"
        "(Ký, ghi rõ họ tên)\n\n\n"
    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def tao_word_uythac_mau16(du_lieu: dict, df_to) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    doc = Document()
    _uyt_style_doc(doc)
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)

    tbl_h = doc.add_table(rows=1, cols=2)
    tbl_h.style = "Table Grid"
    _uyt_xoa_border_table(tbl_h)

    cell_l = tbl_h.rows[0].cells[0]
    p_l = cell_l.paragraphs[0]
    p_l.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rl = p_l.add_run("ĐƠN VỊ KIỂM TRA")
    rl.bold = True
    rl.font.size = Pt(13)
    p_l2 = cell_l.add_paragraph(du_lieu.get("don_vi_kt", ""))
    p_l2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if p_l2.runs:
        p_l2.runs[0].bold = True
        p_l2.runs[0].font.size = Pt(13)

    cell_r = tbl_h.rows[0].cells[1]
    p_r = cell_r.paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = p_r.add_run("CỘNG HOÀ XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    rr.bold = True
    rr.font.size = Pt(13)
    p_r2 = cell_r.add_paragraph("Độc lập - Tự do - Hạnh phúc")
    p_r2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if p_r2.runs:
        p_r2.runs[0].bold = True
        p_r2.runs[0].font.size = Pt(13)
    doc.add_paragraph()

    t1 = doc.add_paragraph()
    t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t1.add_run("BIÊN BẢN KIỂM TRA")
    r.bold = True
    r.font.size = Pt(13.5)

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t2.add_run("Hoạt động của Tổ Tiết kiệm và vay vốn")
    r.bold = True
    r.font.size = Pt(13.5)
    doc.add_paragraph()

    ngay_kt = du_lieu.get("ngay_kt", date.today())
    if isinstance(ngay_kt, str):
        ngay_kt = _uyt_parse_date(ngay_kt)

    ten_thon = du_lieu.get("ten_thon", "")
    ten_xa = du_lieu.get("ten_xa", "")
    don_vi_kt = du_lieu.get("don_vi_kt", "")
    can_bo_1 = du_lieu.get("can_bo_1", "")
    chuc_vu_1 = du_lieu.get("chuc_vu_1", "")
    can_bo_2 = du_lieu.get("can_bo_2", "")
    chuc_vu_2 = du_lieu.get("chuc_vu_2", "")
    to_truong = du_lieu.get("to_truong", "")
    to_pho = du_lieu.get("to_pho", "")
    hoi_dt = du_lieu.get("hoi_doan_the", "")

    def _pj(bold_prefix: str = "", rest: str = "", size: float = 13.5):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if bold_prefix:
            r0 = p.add_run(bold_prefix)
            r0.bold = True
            r0.font.size = Pt(size)
        if rest:
            r1 = p.add_run(rest)
            r1.font.size = Pt(size)
        return p

    _pj(
        rest=(
            f"Hôm nay, ngày {ngay_kt.day} tháng {ngay_kt.month} năm {ngay_kt.year}, "
            "tại Tổ Tiết kiệm và vay vốn (Tổ) thôn/tổ dân phố "
            f"{ten_thon or '.....................................'}, "
            f"xã/phường {ten_xa or '.......................'}"
        )
    )
    _pj(bold_prefix="ĐOÀN KIỂM TRA: ", rest=don_vi_kt)
    _pj(
        rest=(
            f"- Ông (bà): {can_bo_1 or '........................................'}   "
            f"Chức vụ: {chuc_vu_1 or '.............................'}"
        )
    )
    if can_bo_2:
        _pj(rest=f"- Ông (bà): {can_bo_2}   Chức vụ: {chuc_vu_2 or '.............................'}")
    _pj(bold_prefix="ĐƠN VỊ ĐƯỢC KIỂM TRA: ", rest=f"Tổ thuộc Hội {hoi_dt or '.....................................'}")
    _pj(rest=f"- Ông (bà): {to_truong or '........................................'}   Chức vụ: Tổ trưởng")
    _pj(rest=f"- Ông (bà): {to_pho or '........................................'}   Chức vụ: Tổ phó")
    _pj(
        rest=(
            "Cùng tiến hành kiểm tra việc thực hiện Hợp đồng ủy nhiệm của Tổ đã ký "
            "với NHCSXH, thống nhất kết quả kiểm tra như sau:"
        )
    )

    p_i = doc.add_paragraph()
    p_i.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    ri = p_i.add_run("I. TÌNH HÌNH CHUNG CỦA TỔ")
    ri.bold = True
    ri.font.size = Pt(12)
    ri2 = p_i.add_run(f" (đến thời điểm {ngay_kt.strftime('%d/%m/%Y')})")
    ri2.font.size = Pt(13.5)

    tong_dn = df_to[COT_TONG_DU_NO].sum() if df_to is not None and COT_TONG_DU_NO in df_to.columns else 0
    so_tv = len(df_to) if df_to is not None else 0
    nqh_val = df_to[COT_DU_NO_QH].sum() if df_to is not None and COT_DU_NO_QH in df_to.columns else 0
    lai_val = df_to[COT_LAI_TON].sum() if df_to is not None and COT_LAI_TON in df_to.columns else 0
    tg_val = df_to[COT_SO_DU_TG].sum() if df_to is not None and COT_SO_DU_TG in df_to.columns else 0
    ty_le = du_lieu.get("ty_le_nqh", "")
    xep_loai = du_lieu.get("xep_loai_to", "")

    _pj(
        rest=(
            f"1. Tổng dư nợ của Tổ: {_fmt_bang_ty(tong_dn)} triệu đồng, "
            f"{_fmt_so(so_tv)} tổ viên. Trong đó, nợ quá hạn "
            f"{_fmt_bang_ty(nqh_val)} triệu đồng (tỷ lệ {ty_le or '......'}%)"
        )
    )
    _pj(rest=f"2. Tổng lãi tồn của Tổ: {_fmt_bang_ty(lai_val)} triệu đồng.")
    _pj(rest=f"3. Số dư tiền gửi của Tổ: {_fmt_bang_ty(tg_val)} triệu đồng.")
    _pj(
        rest=(
            "4. Kết quả chấm điểm xếp loại Tổ "
            f"(tháng {ngay_kt.month}/{ngay_kt.year}): {xep_loai or '............................................'}"
        )
    )

    p_ii = doc.add_paragraph()
    p_ii.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    rii = p_ii.add_run("II. HOẠT ĐỘNG CỦA TỔ VÀ BAN QUẢN LÝ TỔ")
    rii.bold = True
    rii.font.size = Pt(12)

    _pj(
        rest=(
            "Đoàn/cán bộ kiểm tra thực hiện kiểm tra theo nội dung kiểm tra, "
            "giám sát tại khoản 3 Phụ lục I văn bản số 727/HD-NHCS ngày 11/02/2026. Cụ thể:"
        )
    )

    CL1 = [
        (True, "1. Thành lập Tổ"),
        (False, "- Tổ có được thành lập theo cụm dân cư liền kề?"),
        (False, "- Số lượng đảm bảo quy định (từ 05 đến 60 tổ viên)"),
        (True, "2. Ban quản lý Tổ"),
        (False, "- Ban quản lý Tổ gồm mấy người? Phân công nhiệm vụ Tổ trưởng, Tổ phó?"),
        (False, "- Tổ trưởng, Tổ phó có mối quan hệ vợ, chồng, cha, mẹ, con hoặc anh, chị, em ruột?"),
        (False, "- Ban quản lý Tổ có tham gia Ban Thường vụ tổ chức CT-XH cấp xã?"),
        (False, "- Ban quản lý Tổ có kiêm Tổ trưởng Tổ vay vốn của ngân hàng khác?"),
    ]
    tbl1 = doc.add_table(rows=1 + len(CL1), cols=2)
    tbl1.style = "Table Grid"
    for ci, txt in enumerate(["Nội dung", "Kết quả kiểm tra"]):
        c = tbl1.rows[0].cells[ci]
        c.text = txt
        if c.paragraphs[0].runs:
            c.paragraphs[0].runs[0].bold = True
            c.paragraphs[0].runs[0].font.size = Pt(12)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, (is_hdr, txt) in enumerate(CL1, 1):
        cell = tbl1.rows[i].cells[0]
        cell.text = txt
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].bold = is_hdr
            cell.paragraphs[0].runs[0].font.size = Pt(12)

    CL2 = [
        (True, "3. Thực hiện nhiệm vụ của Ban quản lý Tổ"),
        (False, "- Duy trì sinh hoạt Tổ định kỳ như thế nào (địa điểm, thời gian theo tháng/quý,….)?"),
        (False, "- Tuyên truyền, phổ biến chính sách đến tổ viên? Vận động tổ viên thực hành tiết kiệm như thế nào?"),
        (False, "- Tiếp nhận đề nghị vay vốn, tổ chức họp bình xét cho vay? Thành phần tham dự họp có đầy đủ?"),
        (False, "- Tham gia giao dịch, chứng kiến giải ngân, thu nợ, họp giao ban tại điểm giao dịch?"),
        (False, "- Có thu phí, thu nợ gốc, vay ké, chiếm dụng vốn, giữ sổ vay vốn của tổ viên?"),
        (
            False,
            "- Trường hợp được ủy nhiệm thu lãi, có trả biên lai cho tổ viên sau khi thu tiền? "
            "Nộp về ngân hàng đầy đủ số tiền đã thu và biên lai chưa thu (nếu có)?",
        ),
        (
            False,
            "- Chứng kiến việc kiểm tra sử dụng vốn vay, đối chiếu nợ vay và số dư tiền gửi tổ viên "
            "của các tổ chức, cá nhân có thẩm quyền khi được yêu cầu?",
        ),
        (
            False,
            "- Giám sát, đôn đốc tổ viên sử dụng vốn vay đúng mục đích; trả nợ gốc, lãi đúng hạn? "
            "Số tổ viên có lãi tồn (không bao gồm lãi tồn ân hạn), nợ quá hạn, sử dụng vốn vay sai mục đích?",
        ),
        (
            False,
            "- Thông báo, phối hợp xác minh và xử lý các trường hợp gia hạn nợ, điều chỉnh kỳ hạn nợ, "
            "nợ quá hạn, nợ bị rủi ro, nợ bị chiếm dụng, sử dụng vốn sai mục đích, người vay đi khỏi nơi cư trú?",
        ),
        (False, "- Sắp xếp, lưu giữ hồ sơ hoạt động Tổ theo quy định?"),
        (False, "- ……………………………………………………"),
    ]
    doc.add_paragraph()
    tbl2 = doc.add_table(rows=1 + len(CL2), cols=2)
    tbl2.style = "Table Grid"
    for ci, txt in enumerate(["Nội dung", "Kết quả kiểm tra"]):
        c = tbl2.rows[0].cells[ci]
        c.text = txt
        if c.paragraphs[0].runs:
            c.paragraphs[0].runs[0].bold = True
            c.paragraphs[0].runs[0].font.size = Pt(12)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, (is_hdr, txt) in enumerate(CL2, 1):
        cell = tbl2.rows[i].cells[0]
        cell.text = txt
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].bold = is_hdr
            cell.paragraphs[0].runs[0].font.size = Pt(12)

    doc.add_paragraph()
    p_iii = doc.add_paragraph()
    p_iii.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    riii = p_iii.add_run("III. ĐÁNH GIÁ, NHẬN XÉT CỦA ĐOÀN KIỂM TRA")
    riii.bold = True
    riii.font.size = Pt(12)

    so_kh_tt = du_lieu.get("so_kh_kt_thuc_te", "")
    _pj(rest=f"Qua kiểm tra tại Tổ và thực tế tại {so_kh_tt or '…..'} khách hàng, Đoàn có nhận xét như sau:")

    uu_diem = du_lieu.get("uu_diem", "") or ".........................................."
    ton_tai = du_lieu.get("ton_tai", "") or ".........................................."
    kien_nghi = du_lieu.get("kien_nghi", "") or ".........................................."

    p_uu = doc.add_paragraph()
    p_uu.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r0 = p_uu.add_run("1. Ưu điểm: ")
    r0.bold = True
    r0.font.size = Pt(13.5)
    p_uu.add_run(uu_diem).font.size = Pt(13.5)

    p_tt = doc.add_paragraph()
    p_tt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r0 = p_tt.add_run("2. Tồn tại: ")
    r0.bold = True
    r0.font.size = Pt(13.5)
    p_tt.add_run(ton_tai).font.size = Pt(13.5)

    p_kn = doc.add_paragraph()
    p_kn.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r0 = p_kn.add_run("3. Kiến nghị (nếu có): ")
    r0.bold = True
    r0.font.size = Pt(13.5)
    p_kn.add_run(kien_nghi).font.size = Pt(13.5)

    so_phieu = du_lieu.get("so_phieu_kem_theo", "")
    _pj(rest=f"Kèm theo Biên bản này là {so_phieu or '…'} Phiếu kiểm tra sử dụng vốn vay của khách hàng.")
    _pj(rest="Biên bản được lập thành 02 bản (01 bản lưu Đoàn kiểm tra, 01 bản lưu Tổ).")

    doc.add_paragraph()
    ky = doc.add_table(rows=1, cols=2)
    ky.style = "Table Grid"
    _uyt_xoa_border_table(ky)

    p_kl = ky.rows[0].cells[0].paragraphs[0]
    p_kl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p_kl.add_run("TRƯỞNG ĐOÀN KIỂM TRA\n(Ký, ghi rõ họ tên)\n\n\n\n")
    r0.bold = True
    r0.font.size = Pt(12)
    p_kl.add_run(can_bo_1).font.size = Pt(12)

    p_kr = ky.rows[0].cells[1].paragraphs[0]
    p_kr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p_kr.add_run("TỔ TRƯỞNG TỔ TK&VV\n(Ký, ghi rõ họ tên)\n\n\n\n")
    r0.bold = True
    r0.font.size = Pt(12)
    p_kr.add_run(to_truong).font.size = Pt(12)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def tao_word_uythac_bb_xac_minh(du_lieu: dict) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    _uyt_style_doc(doc)
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)

    ngay_lap = du_lieu.get("ngay_lap", date.today())
    _uyt_add_header_quoc_hieu(doc, don_vi=du_lieu.get("pgd_user", ""), so_vb="", dia_danh="", ngay_ky=ngay_lap)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = title.add_run("BIÊN BẢN\nXác minh nợ chiếm dụng")
    run_t.bold = True
    run_t.font.size = Pt(13)

    doc.add_paragraph(
        f"Hôm nay, ngày {ngay_lap.day} tháng {ngay_lap.month} năm {ngay_lap.year}\n"
        f"Cán bộ lập biên bản: {du_lieu.get('can_bo_lap', '')}\n"
        f"Làm việc với: Khách hàng {du_lieu.get('ten_kh', '')}\n"
        f"Số khế ước: {du_lieu.get('so_ku', '')}\n"
        "NỘI DUNG XÁC MINH:\n\n"
        f"Số tiền chiếm dụng: {du_lieu.get('so_tien', '')} triệu đồng\n"
        "Lý do / Hoàn cảnh:\n"
        f"{du_lieu.get('ly_do', '')}\n"
        "Biện pháp xử lý đã thống nhất:\n"
        f"{du_lieu.get('bien_phap', '')}\n\n"
        "Biên bản được lập thành 02 liên, mỗi bên giữ 01 liên."
    )

    doc.add_paragraph()
    ky = doc.add_table(rows=1, cols=2)
    ky.style = "Table Grid"
    for cell in ky.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for border_name in ["top", "left", "bottom", "right"]:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "none")
            tcBorders.append(border)
        tcPr.append(tcBorders)

    p_l = ky.rows[0].cells[0].paragraphs[0]
    p_l.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_l.add_run("KHÁCH HÀNG\n(Ký, ghi rõ họ tên)\n\n\n\n").bold = True
    p_l.add_run(str(du_lieu.get("ten_kh", "")))

    p_r = ky.rows[0].cells[1].paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_r.add_run("CÁN BỘ LẬP BIÊN BẢN\n(Ký, ghi rõ họ tên)\n\n\n\n").bold = True
    p_r.add_run(str(du_lieu.get("can_bo_lap", "")))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def tao_word_uythac_bb_ct_cx(du_lieu: dict, cap: str = "tinh") -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    doc = Document()
    _uyt_style_doc(doc)
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)

    tbl_h = doc.add_table(rows=1, cols=2)
    tbl_h.style = "Table Grid"
    _uyt_xoa_border_table(tbl_h)
    cell_l = tbl_h.rows[0].cells[0]
    p_l = cell_l.paragraphs[0]
    p_l.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_l = p_l.add_run("ĐƠN VỊ KIỂM TRA\n")
    r_l.bold = True
    r_l2 = p_l.add_run((du_lieu.get("don_vi_kt") or "").upper())
    r_l2.bold = True

    cell_r = tbl_h.rows[0].cells[1]
    p_r = cell_r.paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_r.add_run("CỘNG HOÀ XÃ HỘI CHỦ NGHĨA VIỆT NAM").bold = True
    p_r2 = cell_r.add_paragraph("Độc lập - Tự do - Hạnh phúc")
    p_r2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if p_r2.runs:
        p_r2.runs[0].bold = True
    ngay_kt = du_lieu.get("ngay_kt", date.today())
    if isinstance(ngay_kt, str):
        ngay_kt = _uyt_parse_date(ngay_kt)
    p_r3 = cell_r.add_paragraph(
        f"{du_lieu.get('dia_danh', '')}, ngày {ngay_kt.day} tháng {ngay_kt.month} năm {ngay_kt.year}"
    )
    p_r3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    ten_cap = "cấp tỉnh" if cap == "tinh" else "cấp xã"
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = title.add_run(
        "BIÊN BẢN KIỂM TRA\n"
        f"Hoạt động nhận ủy thác với NHCSXH của tổ chức CT-XH {ten_cap}"
    )
    run_t.bold = True
    run_t.font.size = Pt(13)

    doc.add_paragraph(f"Hôm nay, ngày {ngay_kt.day} tháng {ngay_kt.month} năm {ngay_kt.year}, chúng tôi gồm:")
    doc.add_paragraph(
        f"ĐOÀN KIỂM TRA: {du_lieu.get('don_vi_kt', '')}\n"
        f"- Ông (bà): {du_lieu.get('truong_doan', '')}   Chức vụ: Trưởng đoàn\n"
        f"- Ông (bà): {du_lieu.get('can_bo_2', '')}   Chức vụ: {du_lieu.get('chuc_vu_2', '')}"
    )
    doc.add_paragraph(
        f"ĐƠN VỊ ĐƯỢC KIỂM TRA: {du_lieu.get('ten_don_vi', '')}\n"
        f"- Ông (bà): {du_lieu.get('dai_dien_dc', '')}   Chức vụ: {du_lieu.get('chuc_vu_dc', '')}"
    )
    doc.add_paragraph(
        "Cùng tiến hành kiểm tra việc thực hiện các nội dung công việc được ủy thác của "
        f"{du_lieu.get('ten_don_vi', '')}, thống nhất kết quả kiểm tra như sau:"
    )

    p_i = doc.add_paragraph("I. KẾT QUẢ HOẠT ĐỘNG ỦY THÁC")
    p_i.runs[0].bold = True
    doc.add_paragraph(
        f"(Đến thời điểm {ngay_kt.strftime('%d/%m/%Y')})\n"
        "- Tổng số Tổ TK&VV do Hội quản lý: ........... tổ\n"
        "- Tổng số khách hàng vay vốn: ........... người\n"
        "- Tổng dư nợ nhận ủy thác: ........... triệu đồng\n"
        "  Trong đó: Nợ quá hạn ........... (tỷ lệ .......%), Nợ khoanh ..........."
    )

    ten_khoan = "khoản 1" if cap == "tinh" else "khoản 2"
    p_ii = doc.add_paragraph(
        f"II. KẾT QUẢ THỰC HIỆN CÁC NỘI DUNG NHẬN ỦY THÁC CỦA TỔ CHỨC CT-XH {ten_cap.upper()}"
    )
    p_ii.runs[0].bold = True
    doc.add_paragraph(
        f"Đoàn kiểm tra thực hiện kiểm tra theo nội dung kiểm tra, giám sát tại {ten_khoan} "
        "Phụ lục I văn bản số 727/HD-NHCS ngày 11/02/2026. Cụ thể:"
    )

    muc_noi_dung = [
        ("tuyen_truyen", "1. Công tác tuyên truyền, vận động"),
        ("kiem_tra_giam_sat", "2. Công tác kiểm tra, giám sát hoạt động ủy thác"),
        ("tap_huan", "3. Công tác tập huấn"),
        ("phoi_hop_nhcs", "4. Hoạt động phối hợp thực hiện cùng NHCSXH"),
    ]
    if cap == "tinh":
        muc_noi_dung.append(("trach_nhiem", "5. Trách nhiệm của tổ chức CT-XH cấp tỉnh"))

    for field_key, ten_muc in muc_noi_dung:
        p_muc = doc.add_paragraph(ten_muc)
        p_muc.runs[0].bold = True
        data = du_lieu.get(field_key) or {}
        doc.add_paragraph(f"a) Kết quả đạt được\n{data.get('ket_qua', '.....')}")
        doc.add_paragraph(f"b) Tồn tại\n{data.get('ton_tai', '.....')}")

    p_iii = doc.add_paragraph("III. ĐÁNH GIÁ, NHẬN XÉT CỦA ĐOÀN KIỂM TRA")
    p_iii.runs[0].bold = True
    han_str = du_lieu.get("han_hoan_thanh", "....../....../20......")
    if han_str and len(han_str) == 10 and han_str[4] == "-":
        try:
            han_str = datetime.strptime(han_str, "%Y-%m-%d").strftime("%d/%m/%Y")
        except Exception:
            pass
    doc.add_paragraph(
        f"1. Ưu điểm\n{du_lieu.get('uu_diem', '.....')}\n\n"
        f"2. Tồn tại\n{du_lieu.get('ton_tai_chung', '.....')}\n\n"
        f"3. Kiến nghị của Đoàn kiểm tra\n{du_lieu.get('kien_nghi', '.....')}\n\n"
        "Đơn vị được kiểm tra hoàn thành các kiến nghị và báo cáo kết quả "
        f"trước ngày {han_str}."
    )

    p_iv = doc.add_paragraph("IV. Ý KIẾN CỦA ĐƠN VỊ ĐƯỢC KIỂM TRA")
    p_iv.runs[0].bold = True
    doc.add_paragraph(du_lieu.get("y_kien_don_vi_dc", ".....") or ".....")

    doc.add_paragraph(
        "Biên bản được lập thành 02 bản (01 bản lưu Đoàn kiểm tra, 01 bản lưu Đơn vị được kiểm tra)."
    )

    doc.add_paragraph()
    ky = doc.add_table(rows=1, cols=2)
    ky.style = "Table Grid"
    _uyt_xoa_border_table(ky)
    p_kl = ky.rows[0].cells[0].paragraphs[0]
    p_kl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_kl.add_run("TRƯỞNG ĐOÀN KIỂM TRA\n(Ký, ghi rõ họ tên)\n\n\n\n").bold = True
    p_kl.add_run(du_lieu.get("truong_doan", ""))
    p_kr = ky.rows[0].cells[1].paragraphs[0]
    p_kr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_kr.add_run("ĐƠN VỊ ĐƯỢC KIỂM TRA\n(Ký tên, đóng dấu)\n\n\n\n").bold = True
    p_kr.add_run(du_lieu.get("dai_dien_dc", ""))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def tao_word_uythac_bc_th(du_lieu: dict, ds_bien_ban: list) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    doc = Document()
    _uyt_style_doc(doc)
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)

    ngay_bc = du_lieu.get("ngay_bc", date.today())
    if isinstance(ngay_bc, str):
        ngay_bc = _uyt_parse_date(ngay_bc)

    tbl_h = doc.add_table(rows=1, cols=2)
    tbl_h.style = "Table Grid"
    _uyt_xoa_border_table(tbl_h)
    cell_l = tbl_h.rows[0].cells[0]
    cell_l.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_l.paragraphs[0].add_run(f"ĐƠN VỊ KIỂM TRA\n{(du_lieu.get('don_vi_kt') or '').upper()}").bold = True
    cell_r = tbl_h.rows[0].cells[1]
    p_r = cell_r.paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_r.add_run("CỘNG HOÀ XÃ HỘI CHỦ NGHĨA VIỆT NAM").bold = True
    p_r2 = cell_r.add_paragraph("Độc lập - Tự do - Hạnh phúc")
    p_r2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if p_r2.runs:
        p_r2.runs[0].bold = True
    p_r3 = cell_r.add_paragraph(
        f"{du_lieu.get('dia_danh', '')}, ngày {ngay_bc.day} tháng {ngay_bc.month} năm {ngay_bc.year}"
    )
    p_r3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = title.add_run("BÁO CÁO TỔNG HỢP\nKết quả kiểm tra hoạt động nhận ủy thác cho vay")
    run_t.bold = True
    run_t.font.size = Pt(13)
    doc.add_paragraph()

    p_i = doc.add_paragraph("I. THÀNH PHẦN")
    p_i.runs[0].bold = True
    doc.add_paragraph(
        f"1. Đoàn kiểm tra: {du_lieu.get('don_vi_kt', '')}\n"
        f"   Trưởng đoàn: {du_lieu.get('truong_doan', '')}\n"
        f"2. Cấp ủy, chính quyền địa phương (nếu có): {du_lieu.get('cap_uy', '')}"
    )

    p_ii = doc.add_paragraph("II. THỜI GIAN, ĐỊA ĐIỂM, ĐƠN VỊ ĐƯỢC KIỂM TRA")
    p_ii.runs[0].bold = True
    tbl_dv = doc.add_table(rows=1, cols=4)
    tbl_dv.style = "Table Grid"
    for i, h in enumerate(["STT", "Thời gian", "Đơn vị được kiểm tra", "Địa điểm"]):
        c = tbl_dv.rows[0].cells[i]
        c.text = h
        if c.paragraphs[0].runs:
            c.paragraphs[0].runs[0].bold = True
            c.paragraphs[0].runs[0].font.size = Pt(10)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for idx, bb in enumerate(ds_bien_ban, 1):
        ngay_str = bb.get("ngay_kt", "")
        try:
            ngay_str = datetime.strptime(ngay_str, "%Y-%m-%d").strftime("%d/%m/%Y")
        except Exception:
            pass
        r = tbl_dv.add_row()
        r.cells[0].text = str(idx)
        r.cells[1].text = ngay_str
        r.cells[2].text = bb.get("ten_don_vi", "")
        r.cells[3].text = bb.get("dia_danh", "")
        for cell in r.cells:
            if cell.paragraphs and cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_paragraph()

    p_iii = doc.add_paragraph("III. NỘI DUNG KIỂM TRA")
    p_iii.runs[0].bold = True
    doc.add_paragraph(du_lieu.get("noi_dung_kt", "Theo Phụ lục I văn bản số 727/HD-NHCS ngày 11/02/2026."))

    p_iv = doc.add_paragraph("IV. KẾT QUẢ KIỂM TRA")
    p_iv.runs[0].bold = True
    doc.add_paragraph("1. Đánh giá, nhận xét của Đoàn kiểm tra").runs[0].bold = True
    doc.add_paragraph(
        f"a) Đối với tổ chức CT-XH được kiểm tra\n{du_lieu.get('nx_ctxh', '.....')}\n\n"
        f"b) Đối với Tổ TK&VV\n{du_lieu.get('nx_to', '.....')}\n\n"
        f"c) Đối với tổ viên Tổ TK&VV\n{du_lieu.get('nx_to_vien', '.....')}"
    )
    doc.add_paragraph("2. Kiến nghị của Đoàn kiểm tra").runs[0].bold = True
    doc.add_paragraph(
        f"a) Đối với tổ chức CT-XH được kiểm tra\n{du_lieu.get('kn_ctxh', '.....')}\n\n"
        f"b) Đối với Tổ TK&VV\n{du_lieu.get('kn_to', '.....')}\n\n"
        f"c) Đối với tổ viên Tổ TK&VV\n{du_lieu.get('kn_to_vien', '.....')}\n\n"
        f"d) Đối với NHCSXH\n{du_lieu.get('kn_nhcs', '.....')}\n\n"
        f"đ) Đối với tổ chức CT-XH cấp trên\n{du_lieu.get('kn_cap_tren', '.....')}"
    )
    doc.add_paragraph("3. Kiến nghị của Đơn vị được kiểm tra").runs[0].bold = True
    doc.add_paragraph("a) Đối với NHCSXH\n.....\n\nb) Đối với tổ chức CT-XH cấp trên\n.....")

    p_vi = doc.add_paragraph("VI. TÀI LIỆU KÈM THEO (nếu có)")
    p_vi.runs[0].bold = True
    doc.add_paragraph(
        "1. Phiếu kiểm tra sử dụng vốn vay (mẫu 06/TD, 06A/TD): ............... phiếu.\n"
        "2. Danh sách đối chiếu (mẫu 15/TD): ...................... danh sách.\n"
        f"3. Biên bản kiểm tra tổ chức CT-XH: {len(ds_bien_ban)} biên bản."
    )

    doc.add_paragraph()
    ky = doc.add_table(rows=1, cols=1)
    ky.style = "Table Grid"
    _uyt_xoa_border_table(ky)
    p_ky = ky.rows[0].cells[0].paragraphs[0]
    p_ky.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_ky.add_run("TRƯỞNG ĐOÀN KIỂM TRA\n(Ký, ghi rõ họ tên)\n\n\n\n").bold = True
    p_ky.add_run(du_lieu.get("truong_doan", ""))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
