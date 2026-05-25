"""Word generation helpers cho biểu mẫu XLN (Xử Lý Nợ)."""
from __future__ import annotations

import io
from datetime import date

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from utils import fmt
from services.rui_ro_aggregation import _tong_hop_no


# ── Helpers dùng chung ────────────────────────────────────────────────

def _pgd_plain(ten: str) -> str:
    s = str(ten or "").strip()
    if s.lower().startswith("pgd "):
        return s[4:].strip()
    return s


def _pgd_line(ten: str) -> str:
    s = str(ten or "").strip()
    if not s:
        return "PGD"
    if s.lower().startswith("pgd "):
        return s
    return f"PGD {s}"


def _style_doc_xln(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)


def _bo_border_cell(cell) -> None:
    """Xóa border khỏi 4 cạnh của một cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for bn in ["top", "left", "bottom", "right"]:
        b = OxmlElement(f"w:{bn}")
        b.set(qn("w:val"), "none")
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _set_cell(
    cell, text: str, *,
    bold: bool = False, italic: bool = False,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    v_align=WD_ALIGN_VERTICAL.CENTER,
    font_size: int = 10,
) -> None:
    """Gán text + format cho một cell bảng."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(font_size)
    cell.vertical_alignment = v_align


def _set_row_font(row, font_size: int = 10) -> None:
    """Đồng bộ font cho toàn bộ row."""
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(font_size)


def _num(v) -> float:
    """Chuyển đổi an toàn string → float."""
    if v is None:
        return 0.0
    if isinstance(v, (int, float)):
        return float(v)
    s = str(v).strip()
    if not s:
        return 0.0
    s = s.replace(" ", "").replace(".", "").replace(",", ".")
    try:
        return float(s)
    except Exception:
        return 0.0


def _add_header_xln(doc: Document, dia_danh: str, ngay_ky: date) -> None:
    """Header Quốc hiệu + ngày tháng, không có số VB (mẫu đơn cá nhân)."""
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    for cell in t.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for bn in ["top", "left", "bottom", "right"]:
            b = OxmlElement(f"w:{bn}")
            b.set(qn("w:val"), "none")
            tcBorders.append(b)
        tcPr.append(tcBorders)

    cell_l = t.rows[0].cells[0]
    cell_l.paragraphs[0].add_run("Mẫu số 01/XLN")

    cell_r = t.rows[0].cells[1]
    p3 = cell_r.paragraphs[0]
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run("CỘNG HOÀ XÃ HỘI CHỦ NGHĨA VIỆT NAM").bold = True
    p4 = cell_r.add_paragraph("Độc lập - Tự do - Hạnh phúc")
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.runs[0].bold = True
    p5 = cell_r.add_paragraph(
        f"{dia_danh}, ngày {ngay_ky.day} tháng {ngay_ky.month} năm {ngay_ky.year}"
    )
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def _set_margins(
    doc: Document,
    left_cm: float = 3.0,
    right_cm: float = 2.0,
    top_cm: float = 2.5,
    bottom_cm: float = 2.5,
) -> None:
    for section in doc.sections:
        section.left_margin = Cm(left_cm)
        section.right_margin = Cm(right_cm)
        section.top_margin = Cm(top_cm)
        section.bottom_margin = Cm(bottom_cm)


# ── Tạo biểu mẫu Word ────────────────────────────────────────────────

def _tao_word_01xln(du_lieu: dict) -> bytes:
    """Mẫu 01/XLN — Đơn đề nghị xử lý nợ (KH tự viết)."""
    doc = Document()
    _style_doc_xln(doc)
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    _add_header_xln(
        doc,
        dia_danh=du_lieu.get("dia_danh", ""),
        ngay_ky=du_lieu.get("ngay_ky", date.today()),
    )

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("ĐƠN ĐỀ NGHỊ XỬ LÝ NỢ")
    r.bold = True
    r.font.size = Pt(14)

    doc.add_paragraph(
        f"Kính gửi: Ngân hàng Chính sách xã hội {du_lieu.get('ten_nhcsxh','')}"
    )
    doc.add_paragraph()

    doc.add_paragraph(
        f"Tên tôi là: {du_lieu.get('ten_kh','')}\n"
        f"Hiện cư trú tại: {du_lieu.get('dia_chi','')}\n"
        f"Là thành viên của Tổ TK&VV: {du_lieu.get('ten_to','')} "
        f"do ông (bà): {du_lieu.get('to_truong','')} làm Tổ trưởng"
    )
    doc.add_paragraph(
        f"1. Theo HĐTD (sổ Vay vốn) số {du_lieu.get('so_ku','')}, "
        f"ngày {du_lieu.get('ngay_vay','')}, tôi có đứng tên vay vốn "
        f"chương trình {du_lieu.get('ten_ct','')} tại NHCSXH "
        f"{du_lieu.get('ten_nhcsxh','')}.\n"
        f"    Số tiền vay: {du_lieu.get('muc_vay','')} đồng; "
        f"Hạn trả nợ: {du_lieu.get('ngay_dh','')}; "
        f"Mục đích vay vốn: {du_lieu.get('muc_dich_vay','')}\n"
        f"    Hiện nay, tôi còn nợ Ngân hàng số tiền: "
        f"{du_lieu.get('tong_du_no','')} đồng\n"
        f"    (Trong đó: Nợ gốc: {du_lieu.get('du_no_goc','')} đồng; "
        f"Nợ lãi: {du_lieu.get('lai_ton','')} đồng)"
    )
    doc.add_paragraph(f"2. Trong thời gian vừa qua do:\n{du_lieu.get('nguyen_nhan','')}")
    doc.add_paragraph(
        "3. Số vốn, tài sản của dự án bị thiệt hại:\n"
        f"    - Số vốn và tài sản bị thiệt hại: {du_lieu.get('so_tien_thiet_hai','')} đồng\n"
        f"    - Tổng số vốn thực hiện dự án: {du_lieu.get('muc_vay','')} đồng\n"
        f"    - Mức độ thiệt hại: {du_lieu.get('muc_do_thiet_hai','')}%"
    )
    doc.add_paragraph(
        "4. Tình hình kinh tế, khả năng trả nợ sau khi gặp rủi ro:\n"
        f"{du_lieu.get('kha_nang_tra_no','')}"
    )
    doc.add_paragraph(
        f"Vậy tôi làm đơn này đề nghị NHCSXH {du_lieu.get('ten_nhcsxh','')} "
        f"xem xét {du_lieu.get('bien_phap','')} số nợ bị rủi ro, cụ thể:\n"
        f"    - Số tiền đề nghị: {du_lieu.get('so_tien_de_nghi','')} đồng\n"
        f"      (Nợ gốc: {du_lieu.get('du_no_goc','')} đồng; "
        f"Nợ lãi: {du_lieu.get('lai_ton','')} đồng)\n"
        f"    - Thời gian đề nghị: {du_lieu.get('so_thang','')} tháng\n"
        f"    - Kế hoạch trả nợ: {du_lieu.get('ke_hoach_tra_no','')}"
    )
    doc.add_paragraph(
        "Tôi xin cam đoan và chịu trách nhiệm trước pháp luật về nội dung "
        "kê khai trên đơn và các hồ sơ giấy tờ chứng minh là đúng."
    )

    doc.add_paragraph()
    p_ky = doc.add_paragraph()
    p_ky.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    d_ky = du_lieu.get("ngay_ky", date.today())
    p_ky.add_run(
        f"Ngày {d_ky.day} tháng {d_ky.month} năm {d_ky.year}\n"
        "Người làm đơn\n(Ký, ghi rõ họ tên)\n\n\n\n"
        f"{du_lieu.get('ten_kh','')}"
    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _tao_word_02xln(du_lieu: dict) -> bytes:
    """Mẫu 02/XLN — Biên bản đề nghị xử lý nợ bị rủi ro (nhiều bên ký)."""
    doc = Document()
    _style_doc_xln(doc)
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    _add_header_xln(
        doc, dia_danh=du_lieu.get("dia_danh", ""), ngay_ky=du_lieu.get("ngay_lap", date.today())
    )
    if doc.tables:
        cell = doc.tables[0].rows[0].cells[0]
        if cell.paragraphs and cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].text = "Mẫu số 02/XLN"
            for rr in cell.paragraphs[0].runs[1:]:
                rr.text = ""
        else:
            cell.text = "Mẫu số 02/XLN"

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.add_run("BIÊN BẢN\nĐề nghị xử lý nợ bị rủi ro").bold = True
    p_ct = doc.add_paragraph(f"(Chương trình {du_lieu.get('ten_ct','')})")
    p_ct.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ngay_lap = du_lieu.get("ngay_lap", date.today())
    doc.add_paragraph(
        f"Hôm nay, ngày {ngay_lap.day} tháng {ngay_lap.month} "
        f"năm {ngay_lap.year}, tại {du_lieu.get('dia_diem','')}, "
        "chúng tôi gồm có:"
    )

    thanh_phan = du_lieu.get("thanh_phan", [])
    for i, tp in enumerate(thanh_phan, 1):
        doc.add_paragraph(
            f"{i}. Ông (bà) {tp.get('ho_ten','')}   "
            f"Chức vụ: {tp.get('chuc_vu','')}   "
            f"Đại diện: {tp.get('dai_dien','')}"
        )
    for i in range(len(thanh_phan) + 1, 8):
        doc.add_paragraph(
            f"{i}. Ông (bà) ....................................   "
            "Chức vụ: ....................   "
            "Đại diện: ......................"
        )

    doc.add_paragraph(
        "Đã tiến hành thẩm tra và lập biên bản đề nghị xử lý nợ bị rủi ro "
        f"của ông (bà): {du_lieu.get('ten_kh','')}  "
        f"địa chỉ: {du_lieu.get('dia_chi','')}\n"
        "Là đại diện hộ gia đình vay vốn NHCSXH theo HĐTD số "
        f"{du_lieu.get('so_ku','')} ngày {du_lieu.get('ngay_vay','')}, "
        f"có mã món vay: {du_lieu.get('so_ku','')}. Cụ thể như sau:"
    )

    p = doc.add_paragraph("I. Nguyên nhân khách hàng bị rủi ro:")
    p.runs[0].bold = True
    doc.add_paragraph(du_lieu.get("nguyen_nhan", ""))

    p = doc.add_paragraph("II. Xác định mức độ thiệt hại về vốn và tài sản:")
    p.runs[0].bold = True
    doc.add_paragraph(
        "1. Số vốn và tài sản bị thiệt hại: "
        f"{du_lieu.get('so_tien_thiet_hai','')} đồng\n"
        f"2. Tổng số vốn thực hiện dự án: {du_lieu.get('muc_vay','')} đồng\n"
        f"3. Đánh giá mức độ thiệt hại: {du_lieu.get('muc_do_thiet_hai','')}%"
    )

    p = doc.add_paragraph("III. Dư nợ tại NHCSXH đến ngày lập biên bản:")
    p.runs[0].bold = True
    doc.add_paragraph(
        f"Tổng số nợ còn phải trả: {du_lieu.get('tong_du_no','')} đồng\n"
        f"    Trong đó:  + Nợ gốc: {du_lieu.get('du_no_goc','')} đồng\n"
        f"               + Nợ lãi: {du_lieu.get('lai_ton','')} đồng"
    )

    p = doc.add_paragraph("IV. Đánh giá thực trạng dự án, tài sản và khả năng trả nợ:")
    p.runs[0].bold = True
    doc.add_paragraph(
        "1. Đánh giá thực trạng dự án / phương án khôi phục:\n"
        f"{du_lieu.get('thuc_trang_du_an','')}\n\n"
        "2. Tài sản hiện tại của khách hàng:\n"
        f"{du_lieu.get('tai_san_hien_tai','')}\n\n"
        "3. Đánh giá khả năng trả nợ:\n"
        f"{du_lieu.get('kha_nang_tra_no','')}\n\n"
        "4. Về việc áp dụng biện pháp thu hồi nợ:\n"
        f"{du_lieu.get('bien_phap_thu_hoi','')}"
    )

    p = doc.add_paragraph("V. Đề xuất biện pháp xử lý:")
    p.runs[0].bold = True
    doc.add_paragraph(
        f"Chúng tôi nhất trí đề nghị NHCSXH xem xét {du_lieu.get('bien_phap','')} "
        f"cho ông (bà) {du_lieu.get('ten_kh','')} với thời gian "
        f"{du_lieu.get('so_thang','')} tháng, số tiền "
        f"{du_lieu.get('so_tien_de_nghi','')} đồng.\n"
        f"    Trong đó:  + Nợ gốc: {du_lieu.get('du_no_goc','')} đồng\n"
        f"               + Nợ lãi: {du_lieu.get('lai_ton','')} đồng\n"
        "Biên bản này lập thành 02 bản có giá trị pháp lý như nhau."
    )

    doc.add_paragraph()
    ky = doc.add_table(rows=1, cols=3)
    ky.style = "Table Grid"
    for cell in ky.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for bn in ["top", "left", "bottom", "right"]:
            b = OxmlElement(f"w:{bn}")
            b.set(qn("w:val"), "none")
            tcBorders.append(b)
        tcPr.append(tcBorders)

    def _ky(cell, nhan: str, ten: str = ""):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"{nhan}\n(Ký, ghi rõ họ tên)\n\n\n\n{ten}")

    _ky(
        ky.rows[0].cells[0],
        "ĐẠI DIỆN KHÁCH HÀNG\nĐẠI DIỆN UBND CẤP XÃ",
        du_lieu.get("ten_kh", ""),
    )
    _ky(ky.rows[0].cells[1], "TỔ TRƯỞNG TỔ TK&VV\nĐẠI DIỆN HỘI ĐOÀN THỂ")
    _ky(
        ky.rows[0].cells[2],
        "CÁN BỘ TÍN DỤNG\nĐẠI DIỆN NHCSXH",
        du_lieu.get("can_bo_td", ""),
    )

    ky2 = doc.add_table(rows=1, cols=2)
    ky2.style = "Table Grid"
    for cell in ky2.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for bn in ["top", "left", "bottom", "right"]:
            b = OxmlElement(f"w:{bn}")
            b.set(qn("w:val"), "none")
            tcBorders.append(b)
        tcPr.append(tcBorders)
    _ky(ky2.rows[0].cells[0], "ĐẠI DIỆN CƠ QUAN CÔNG AN CẤP XÃ\n(Xác nhận, ký tên, đóng dấu)")
    _ky(ky2.rows[0].cells[1], "ĐẠI DIỆN TỔ CHỨC, CÁ NHÂN LIÊN QUAN (nếu có)")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _tao_word_xln_bao_cao(
    tong_hop: dict,
    ten_pgd: str,
    nguon_label: str,
    so_qd: str,
    ngay_qd: date,
    ngay_bat_dau: date,
    ngay_ket_thuc: date,
    mau_so: str,
    tieu_de: str,
) -> bytes:
    doc = Document()
    _style_doc_xln(doc)
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width = Cm(29.7)
    sec.page_height = Cm(21.0)
    _set_margins(doc, left_cm=2.0, right_cm=2.0, top_cm=2.0, bottom_cm=2.0)

    p_ms = doc.add_paragraph()
    p_ms.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = p_ms.add_run(f"Mẫu số: {mau_so}\n")
    r1.italic = True
    r2 = p_ms.add_run("PGD gửi tỉnh, tỉnh tổng hợp gửi TW")
    r2.italic = True

    p = doc.add_paragraph("NGÂN HÀNG CHÍNH SÁCH XÃ HỘI")
    if p.runs:
        p.runs[0].bold = True
    if ten_pgd:
        p2 = doc.add_paragraph(_pgd_line(ten_pgd))
        if p2.runs:
            p2.runs[0].bold = True
    else:
        doc.add_paragraph("PGD")

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(tieu_de)
    r.bold = True

    ngay_qd_txt = ngay_qd.strftime("%d/%m/%Y")
    t2 = doc.add_paragraph(
        f"(theo Quyết định số {so_qd} ngày {ngay_qd_txt} của Chủ tịch Hội đồng quản trị NHCSXH)"
    )
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    tu_ngay = ngay_bat_dau.strftime("%d/%m/%Y")
    den_ngay = ngay_ket_thuc.strftime("%d/%m/%Y")
    t3 = doc.add_paragraph(f"Từ ngày {tu_ngay} đến ngày {den_ngay}")
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Nguồn vốn: {nguon_label}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Đơn vị tính: hộ, đồng").alignment = WD_ALIGN_PARAGRAPH.RIGHT

    tbl = doc.add_table(rows=3, cols=15)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr0 = tbl.rows[0].cells
    hdr1 = tbl.rows[1].cells
    hdr2 = tbl.rows[2].cells

    hdr0[0].merge(hdr2[0])
    hdr0[1].merge(hdr2[1])
    hdr0[14].merge(hdr2[14])

    hdr0[2].merge(hdr0[5])
    hdr0[6].merge(hdr0[9])
    hdr0[10].merge(hdr0[13])

    hdr1[2].merge(hdr2[2])
    hdr1[6].merge(hdr2[6])
    hdr1[10].merge(hdr2[10])

    hdr1[3].merge(hdr1[5])
    hdr1[7].merge(hdr1[9])
    hdr1[11].merge(hdr1[13])

    _set_cell(hdr0[0], "STT\n(1)", bold=True)
    _set_cell(hdr0[1], "Chi nhánh tỉnh, thành phố\n(2)", bold=True)
    _set_cell(hdr0[2], "Số được thông báo", bold=True)
    _set_cell(hdr0[6], "Số hạch toán thực tế", bold=True)
    _set_cell(hdr0[10], "Chênh lệch", bold=True)
    _set_cell(hdr0[14], "Nguyên nhân\nchênh lệch\n(15)", bold=True)

    _set_cell(hdr1[2], "Số hộ\n(3)", bold=True)
    _set_cell(hdr1[3], "Số tiền", bold=True)
    _set_cell(hdr1[6], "Số hộ\n(7)", bold=True)
    _set_cell(hdr1[7], "Số tiền", bold=True)
    _set_cell(hdr1[10], "Số hộ\n(11)", bold=True)
    _set_cell(hdr1[11], "Số tiền", bold=True)

    _set_cell(hdr2[3], "Tổng\n(4)", bold=True)
    _set_cell(hdr2[4], "Gốc\n(5)", bold=True)
    _set_cell(hdr2[5], "Lãi\n(6)", bold=True)
    _set_cell(hdr2[7], "Tổng\n(8)", bold=True)
    _set_cell(hdr2[8], "Gốc\n(9)", bold=True)
    _set_cell(hdr2[9], "Lãi\n(10)", bold=True)
    _set_cell(hdr2[11], "Tổng\n(12)", bold=True)
    _set_cell(hdr2[12], "Gốc\n(13)", bold=True)
    _set_cell(hdr2[13], "Lãi\n(14)", bold=True)

    _set_row_font(tbl.rows[0], 10)
    _set_row_font(tbl.rows[1], 10)
    _set_row_font(tbl.rows[2], 10)

    stt = 0
    nhom_ct = tong_hop.get("nhom_ct", {}) or {}
    for ten_ct, v in nhom_ct.items():
        stt += 1

        row_nhom = tbl.add_row()
        _set_row_font(row_nhom, 10)
        _set_cell(row_nhom.cells[0], str(stt), bold=True)
        cell_nhom = row_nhom.cells[1].merge(row_nhom.cells[-1])
        _set_cell(cell_nhom, str(ten_ct), bold=True, italic=True, align=WD_ALIGN_PARAGRAPH.LEFT)
        so_ho = int(v.get("so_ho", 0) or 0)
        goc = float(v.get("goc", 0) or 0)
        lai = float(v.get("lai", 0) or 0)
        tong = goc + lai

        row = tbl.add_row()
        _set_row_font(row, 10)
        _set_cell(row.cells[0], str(stt), align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell(row.cells[1], ten_pgd or "", align=WD_ALIGN_PARAGRAPH.LEFT)

        _set_cell(row.cells[2], str(so_ho), align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell(row.cells[3], fmt(tong), align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell(row.cells[4], fmt(goc), align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell(row.cells[5], fmt(lai), align=WD_ALIGN_PARAGRAPH.RIGHT)

        _set_cell(row.cells[6], str(so_ho), align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell(row.cells[7], fmt(tong), align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell(row.cells[8], fmt(goc), align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell(row.cells[9], fmt(lai), align=WD_ALIGN_PARAGRAPH.RIGHT)

        _set_cell(row.cells[10], "0", align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell(row.cells[11], "0", align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell(row.cells[12], "0", align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell(row.cells[13], "0", align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell(row.cells[14], "", align=WD_ALIGN_PARAGRAPH.LEFT)

    row_tong = tbl.add_row()
    _set_row_font(row_tong, 10)
    _set_cell(row_tong.cells[0], "", bold=True)
    _set_cell(row_tong.cells[1], "Tổng cộng:", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)

    tong_ho = int(tong_hop.get("tong_ho", 0) or 0)
    tong_goc = float(tong_hop.get("tong_goc", 0) or 0)
    tong_lai = float(tong_hop.get("tong_lai", 0) or 0)
    tong_tien = tong_goc + tong_lai

    _set_cell(row_tong.cells[2], str(tong_ho), bold=True)
    _set_cell(row_tong.cells[3], fmt(tong_tien), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell(row_tong.cells[4], fmt(tong_goc), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell(row_tong.cells[5], fmt(tong_lai), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)

    _set_cell(row_tong.cells[6], str(tong_ho), bold=True)
    _set_cell(row_tong.cells[7], fmt(tong_tien), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell(row_tong.cells[8], fmt(tong_goc), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell(row_tong.cells[9], fmt(tong_lai), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)

    _set_cell(row_tong.cells[10], "0", bold=True)
    _set_cell(row_tong.cells[11], "0", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell(row_tong.cells[12], "0", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell(row_tong.cells[13], "0", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell(row_tong.cells[14], "", bold=True)

    doc.add_paragraph()
    ky = doc.add_table(rows=1, cols=3)
    ky.style = "Table Grid"
    for cell in ky.rows[0].cells:
        _bo_border_cell(cell)
        _set_cell(cell, "", font_size=11)

    _set_cell(ky.rows[0].cells[0], "LẬP BIỂU\n\n\n\n", bold=True, font_size=11)
    _set_cell(ky.rows[0].cells[1], "KIỂM SOÁT\n\n\n\n", bold=True, font_size=11)
    _set_cell(ky.rows[0].cells[2], "GIÁM ĐỐC\n\n\n\n", bold=True, font_size=11)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _tao_word_13xln(
    tong_hop: dict,
    ten_pgd: str,
    nguon_label: str,
    so_qd: str,
    ngay_qd: date,
    ngay_bat_dau: date,
    ngay_ket_thuc: date,
) -> bytes:
    return _tao_word_xln_bao_cao(
        tong_hop=tong_hop,
        ten_pgd=ten_pgd,
        nguon_label=nguon_label,
        so_qd=so_qd,
        ngay_qd=ngay_qd,
        ngay_bat_dau=ngay_bat_dau,
        ngay_ket_thuc=ngay_ket_thuc,
        mau_so="13/XLN",
        tieu_de="BÁO CÁO CÁC KHOẢN NỢ SAU KHI THỰC HIỆN HẠCH TOÁN KHOANH NỢ",
    )


def _tao_word_14xln(
    tong_hop: dict,
    ten_pgd: str,
    nguon_label: str,
    so_qd: str,
    ngay_qd: date,
    ngay_bat_dau: date,
    ngay_ket_thuc: date,
) -> bytes:
    return _tao_word_xln_bao_cao(
        tong_hop=tong_hop,
        ten_pgd=ten_pgd,
        nguon_label=nguon_label,
        so_qd=so_qd,
        ngay_qd=ngay_qd,
        ngay_bat_dau=ngay_bat_dau,
        ngay_ket_thuc=ngay_ket_thuc,
        mau_so="14/XLN",
        tieu_de="BÁO CÁO CÁC KHOẢN NỢ SAU KHI THỰC HIỆN HẠCH TOÁN XÓA NỢ",
    )


def _tao_word_04xln(
    tong_hop: dict,
    ten_pgd: str,
    nguon_label: str,
    dot: int,
    nam: int,
) -> bytes:
    doc = Document()
    _style_doc_xln(doc)
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    _set_margins(doc, left_cm=3.0, right_cm=1.5, top_cm=2.0, bottom_cm=2.0)

    hdr = doc.add_table(rows=2, cols=2)
    hdr.style = "Table Grid"
    for row in hdr.rows:
        for cell in row.cells:
            _bo_border_cell(cell)

    p0l = hdr.rows[0].cells[0].paragraphs[0]
    p0l.add_run("NGÂN HÀNG CHÍNH SÁCH XÃ HỘI")
    p0r = hdr.rows[0].cells[1].paragraphs[0]
    p0r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r0r = p0r.add_run("Mẫu số: 04/XLN")
    r0r.italic = True

    hdr.rows[1].cells[0].paragraphs[0].add_run("Chi nhánh NHCSXH tỉnh Đồng Nai")
    p1r = hdr.rows[1].cells[1].paragraphs[0]
    p1r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1r = p1r.add_run("PGD gửi tỉnh, tỉnh tổng hợp gửi TW")
    r1r.italic = True

    doc.add_paragraph(_pgd_line(ten_pgd))

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("BIỂU TỔNG HỢP ĐỀ NGHỊ KHOANH NỢ ĐỐI VỚI KHÁCH HÀNG")
    r.bold = True
    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("VAY VỐN BỊ RỦI RO DO NGUYÊN NHÂN KHÁCH QUAN")
    r2.bold = True
    doc.add_paragraph(f"Đợt {dot} năm {nam} — Nguồn vốn: {nguon_label}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Đơn vị tính: đồng").alignment = WD_ALIGN_PARAGRAPH.RIGHT

    tbl = doc.add_table(rows=2, cols=13)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    h0 = tbl.rows[0].cells
    h1 = tbl.rows[1].cells
    h0[0].merge(h1[0])
    h0[1].merge(h1[1])
    h0[2].merge(h1[2])
    h0[3].merge(h1[3])
    h0[4].merge(h1[4])
    h0[5].merge(h1[5])
    h0[12].merge(h1[12])

    h0[6].merge(h0[8])
    h0[9].merge(h0[11])

    _set_cell(h0[0], "STT", bold=True)
    _set_cell(h0[1], "Chương trình; Huyện, thị xã;\nHọ và tên", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell(h0[2], "Địa chỉ (Xã, phường)", bold=True)
    _set_cell(h0[3], "Mã món vay", bold=True)
    _set_cell(h0[4], "Ngày vay", bold=True)
    _set_cell(h0[5], "Mức độ thiệt hại (%)", bold=True)
    _set_cell(h0[6], "Số dư nợ tại NHCS", bold=True)
    _set_cell(h0[9], "Số nợ đề nghị xử lý", bold=True)
    _set_cell(h0[12], "Ghi chú", bold=True)

    _set_cell(h1[6], "Số tiền", bold=True)
    _set_cell(h1[7], "Gốc", bold=True)
    _set_cell(h1[8], "Lãi", bold=True)
    _set_cell(h1[9], "Số tiền", bold=True)
    _set_cell(h1[10], "Gốc", bold=True)
    _set_cell(h1[11], "Lãi", bold=True)

    ds_all: list[dict] = []
    for v in (tong_hop.get("nhom_ct", {}) or {}).values():
        ds_all.extend(list(v.get("ds", []) or []))

    nhom_3nam = [r for r in ds_all if int(r.get("so_thang", 0) or 0) <= 36]
    nhom_5nam = [r for r in ds_all if int(r.get("so_thang", 0) or 0) > 36]
    th_3 = _tong_hop_no(nhom_3nam)
    th_5 = _tong_hop_no(nhom_5nam)

    def _roman(n: int) -> str:
        pairs = [
            (1000, "M"), (900, "CM"), (500, "D"), (400, "CD"),
            (100, "C"), (90, "XC"), (50, "L"), (40, "XL"),
            (10, "X"), (9, "IX"), (5, "V"), (4, "IV"), (1, "I"),
        ]
        out = ""
        x = n
        for v, s in pairs:
            while x >= v:
                out += s
                x -= v
        return out

    def _muc_do_txt(v: str) -> str:
        s = (v or "").strip()
        if not s or "Không" in s:
            return ""
        if "40%" in s and "80%" in s:
            return "40-<80"
        if "80%" in s and "100%" in s:
            return "80-100"
        return s

    def _render_nhom(tong_hop_nhom: dict, label: str) -> None:
        row_t = tbl.add_row()
        cell = row_t.cells[0].merge(row_t.cells[-1])
        _set_cell(
            cell,
            f"{label}  —  Số tiền {fmt(float(tong_hop_nhom.get('tong_tien', 0) or 0))}  "
            f"(Gốc {fmt(float(tong_hop_nhom.get('tong_goc', 0) or 0))}, "
            f"Lãi {fmt(float(tong_hop_nhom.get('tong_lai', 0) or 0))})",
            bold=True,
            align=WD_ALIGN_PARAGRAPH.LEFT,
        )

        row_a = tbl.add_row()
        cell_a = row_a.cells[0].merge(row_a.cells[-1])
        _set_cell(cell_a, "A  Khoanh nợ lần đầu", italic=True, align=WD_ALIGN_PARAGRAPH.LEFT)

        for i_ct, (ten_ct, v) in enumerate((tong_hop_nhom.get("nhom_ct", {}) or {}).items(), 1):
            goc_ct = float(v.get("goc", 0) or 0)
            lai_ct = float(v.get("lai", 0) or 0)
            tong_ct = goc_ct + lai_ct

            row_ct = tbl.add_row()
            _set_cell(row_ct.cells[0], _roman(i_ct), italic=True)
            _set_cell(row_ct.cells[1], str(ten_ct), italic=True, align=WD_ALIGN_PARAGRAPH.LEFT)
            for j in range(2, 6):
                _set_cell(row_ct.cells[j], "", align=WD_ALIGN_PARAGRAPH.LEFT)
            _set_cell(row_ct.cells[6], fmt(tong_ct), align=WD_ALIGN_PARAGRAPH.RIGHT)
            _set_cell(row_ct.cells[7], fmt(goc_ct), align=WD_ALIGN_PARAGRAPH.RIGHT)
            _set_cell(row_ct.cells[8], fmt(lai_ct), align=WD_ALIGN_PARAGRAPH.RIGHT)
            _set_cell(row_ct.cells[9], fmt(tong_ct), align=WD_ALIGN_PARAGRAPH.RIGHT)
            _set_cell(row_ct.cells[10], fmt(goc_ct), align=WD_ALIGN_PARAGRAPH.RIGHT)
            _set_cell(row_ct.cells[11], fmt(lai_ct), align=WD_ALIGN_PARAGRAPH.RIGHT)
            _set_cell(row_ct.cells[12], "", align=WD_ALIGN_PARAGRAPH.LEFT)

            ds_ct = list(v.get("ds", []) or [])
            for idx, r0 in enumerate(ds_ct, 1):
                goc = float(r0.get("du_no_goc", 0) or 0)
                lai = float(r0.get("lai_ton", 0) or 0)
                tong = goc + lai
                row = tbl.add_row()
                _set_cell(row.cells[0], str(idx))
                _set_cell(row.cells[1], str(r0.get("ten_kh", "")), align=WD_ALIGN_PARAGRAPH.LEFT)
                _set_cell(row.cells[2], str(r0.get("dia_chi", "")), align=WD_ALIGN_PARAGRAPH.LEFT)
                _set_cell(row.cells[3], str(r0.get("so_ku", "")))
                _set_cell(row.cells[4], str(r0.get("ngay_vay", "")))
                _set_cell(row.cells[5], _muc_do_txt(str(r0.get("muc_do", ""))))
                _set_cell(row.cells[6], fmt(tong), align=WD_ALIGN_PARAGRAPH.RIGHT)
                _set_cell(row.cells[7], fmt(goc), align=WD_ALIGN_PARAGRAPH.RIGHT)
                _set_cell(row.cells[8], fmt(lai), align=WD_ALIGN_PARAGRAPH.RIGHT)
                _set_cell(row.cells[9], fmt(tong), align=WD_ALIGN_PARAGRAPH.RIGHT)
                _set_cell(row.cells[10], fmt(goc), align=WD_ALIGN_PARAGRAPH.RIGHT)
                _set_cell(row.cells[11], fmt(lai), align=WD_ALIGN_PARAGRAPH.RIGHT)
                _set_cell(row.cells[12], str(r0.get("ghi_chu", "")), align=WD_ALIGN_PARAGRAPH.LEFT)

    _render_nhom(th_3, "Khoanh nợ tối đa 3 năm")
    _render_nhom(th_5, "Khoanh nợ tối đa 5 năm")

    row_tong = tbl.add_row()
    cell_t = row_tong.cells[0].merge(row_tong.cells[5])
    _set_cell(cell_t, "TỔNG CỘNG", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    tong_goc = float(tong_hop.get("tong_goc", 0) or 0)
    tong_lai = float(tong_hop.get("tong_lai", 0) or 0)
    tong_tien = tong_goc + tong_lai
    _set_cell(row_tong.cells[6], fmt(tong_tien), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell(row_tong.cells[7], fmt(tong_goc), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell(row_tong.cells[8], fmt(tong_lai), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell(row_tong.cells[9], fmt(tong_tien), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell(row_tong.cells[10], fmt(tong_goc), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell(row_tong.cells[11], fmt(tong_lai), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell(row_tong.cells[12], "", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)

    doc.add_paragraph()
    ky = doc.add_table(rows=2, cols=3)
    ky.style = "Table Grid"
    for row in ky.rows:
        for cell in row.cells:
            _bo_border_cell(cell)

    cell_ngay = ky.rows[0].cells[0].merge(ky.rows[0].cells[2])
    _set_cell(cell_ngay, f"{_pgd_plain(ten_pgd)}, ngày ... tháng ... năm ...", align=WD_ALIGN_PARAGRAPH.RIGHT, font_size=11)

    _set_cell(ky.rows[1].cells[0], "LẬP BIỂU\n(Ký, ghi rõ họ tên)", bold=True, font_size=11)
    _set_cell(ky.rows[1].cells[1], "KIỂM SOÁT\n(Ký, ghi rõ họ tên)", bold=True, font_size=11)
    _set_cell(ky.rows[1].cells[2], "GIÁM ĐỐC\n(Ký tên, đóng dấu)", bold=True, font_size=11)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _tao_word_05xln(
    tong_hop: dict,
    ten_pgd: str,
    nguon_label: str,
    dot: int,
    nam: int,
) -> bytes:
    doc = Document()
    _style_doc_xln(doc)
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    _set_margins(doc, left_cm=3.0, right_cm=1.5, top_cm=2.0, bottom_cm=2.0)

    hdr = doc.add_table(rows=2, cols=2)
    hdr.style = "Table Grid"
    for row in hdr.rows:
        for cell in row.cells:
            _bo_border_cell(cell)

    hdr.rows[0].cells[0].paragraphs[0].add_run("NGÂN HÀNG CHÍNH SÁCH XÃ HỘI")
    p0r = hdr.rows[0].cells[1].paragraphs[0]
    p0r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r0r = p0r.add_run("Mẫu số: 05/XLN")
    r0r.italic = True

    hdr.rows[1].cells[0].paragraphs[0].add_run("Chi nhánh NHCSXH tỉnh Đồng Nai")
    p1r = hdr.rows[1].cells[1].paragraphs[0]
    p1r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1r = p1r.add_run("PGD gửi tỉnh, tỉnh tổng hợp gửi TW")
    r1r.italic = True

    doc.add_paragraph(_pgd_line(ten_pgd))

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("BIỂU TỔNG HỢP ĐỀ NGHỊ XÓA NỢ ĐỐI VỚI KHÁCH HÀNG")
    r.bold = True
    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("VAY VỐN BỊ RỦI RO DO NGUYÊN NHÂN KHÁCH QUAN")
    r2.bold = True
    doc.add_paragraph(f"Đợt {dot} năm {nam} — Nguồn vốn: {nguon_label}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Đơn vị tính: đồng").alignment = WD_ALIGN_PARAGRAPH.RIGHT

    tbl = doc.add_table(rows=2, cols=13)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    h0 = tbl.rows[0].cells
    h1 = tbl.rows[1].cells
    h0[0].merge(h1[0])
    h0[1].merge(h1[1])
    h0[2].merge(h1[2])
    h0[3].merge(h1[3])
    h0[4].merge(h1[4])
    h0[5].merge(h1[5])
    h0[12].merge(h1[12])

    h0[6].merge(h0[8])
    h0[9].merge(h0[11])

    _set_cell(h0[0], "STT", bold=True)
    _set_cell(h0[1], "Chương trình; Huyện, thị xã;\nHọ và tên", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell(h0[2], "Địa chỉ (Xã, phường)", bold=True)
    _set_cell(h0[3], "Mã món vay", bold=True)
    _set_cell(h0[4], "Ngày vay", bold=True)
    _set_cell(h0[5], "Mức độ thiệt hại (%)", bold=True)
    _set_cell(h0[6], "Số dư nợ tại NHCS", bold=True)
    _set_cell(h0[9], "Số nợ đề nghị xử lý", bold=True)
    _set_cell(h0[12], "Ghi chú", bold=True)

    _set_cell(h1[6], "Số tiền", bold=True)
    _set_cell(h1[7], "Gốc", bold=True)
    _set_cell(h1[8], "Lãi", bold=True)
    _set_cell(h1[9], "Số tiền", bold=True)
    _set_cell(h1[10], "Gốc", bold=True)
    _set_cell(h1[11], "Lãi", bold=True)

    def _muc_do_txt(v: str) -> str:
        s = (v or "").strip()
        if not s or "Không" in s:
            return ""
        if "40%" in s and "80%" in s:
            return "40-<80"
        if "80%" in s and "100%" in s:
            return "80-100"
        return s

    for i_ct, (ten_ct, v) in enumerate((tong_hop.get("nhom_ct", {}) or {}).items(), 1):
        goc_ct = float(v.get("goc", 0) or 0)
        lai_ct = float(v.get("lai", 0) or 0)
        tong_ct = goc_ct + lai_ct
        row_ct = tbl.add_row()
        _set_cell(row_ct.cells[0], str(i_ct))
        _set_cell(row_ct.cells[1], str(ten_ct), italic=True, align=WD_ALIGN_PARAGRAPH.LEFT)
        for j in range(2, 6):
            _set_cell(row_ct.cells[j], "", align=WD_ALIGN_PARAGRAPH.LEFT)
        _set_cell(row_ct.cells[6], fmt(tong_ct), align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell(row_ct.cells[7], fmt(goc_ct), align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell(row_ct.cells[8], fmt(lai_ct), align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell(row_ct.cells[9], fmt(tong_ct), align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell(row_ct.cells[10], fmt(goc_ct), align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell(row_ct.cells[11], fmt(lai_ct), align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell(row_ct.cells[12], "", align=WD_ALIGN_PARAGRAPH.LEFT)

        ds_ct = list(v.get("ds", []) or [])
        for idx, r0 in enumerate(ds_ct, 1):
            goc = float(r0.get("du_no_goc", 0) or 0)
            lai = float(r0.get("lai_ton", 0) or 0)
            tong = goc + lai
            row = tbl.add_row()
            _set_cell(row.cells[0], str(idx))
            _set_cell(row.cells[1], str(r0.get("ten_kh", "")), align=WD_ALIGN_PARAGRAPH.LEFT)
            _set_cell(row.cells[2], str(r0.get("dia_chi", "")), align=WD_ALIGN_PARAGRAPH.LEFT)
            _set_cell(row.cells[3], str(r0.get("so_ku", "")))
            _set_cell(row.cells[4], str(r0.get("ngay_vay", "")))
            _set_cell(row.cells[5], _muc_do_txt(str(r0.get("muc_do", ""))))
            _set_cell(row.cells[6], fmt(tong), align=WD_ALIGN_PARAGRAPH.RIGHT)
            _set_cell(row.cells[7], fmt(goc), align=WD_ALIGN_PARAGRAPH.RIGHT)
            _set_cell(row.cells[8], fmt(lai), align=WD_ALIGN_PARAGRAPH.RIGHT)
            _set_cell(row.cells[9], fmt(tong), align=WD_ALIGN_PARAGRAPH.RIGHT)
            _set_cell(row.cells[10], fmt(goc), align=WD_ALIGN_PARAGRAPH.RIGHT)
            _set_cell(row.cells[11], fmt(lai), align=WD_ALIGN_PARAGRAPH.RIGHT)
            _set_cell(row.cells[12], str(r0.get("ghi_chu", "")), align=WD_ALIGN_PARAGRAPH.LEFT)

    row_tong = tbl.add_row()
    cell_t = row_tong.cells[0].merge(row_tong.cells[5])
    _set_cell(cell_t, "TỔNG CỘNG", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    tong_goc = float(tong_hop.get("tong_goc", 0) or 0)
    tong_lai = float(tong_hop.get("tong_lai", 0) or 0)
    tong_tien = tong_goc + tong_lai
    _set_cell(row_tong.cells[6], fmt(tong_tien), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell(row_tong.cells[7], fmt(tong_goc), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell(row_tong.cells[8], fmt(tong_lai), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell(row_tong.cells[9], fmt(tong_tien), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell(row_tong.cells[10], fmt(tong_goc), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell(row_tong.cells[11], fmt(tong_lai), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell(row_tong.cells[12], "", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)

    doc.add_paragraph()
    ky = doc.add_table(rows=2, cols=3)
    ky.style = "Table Grid"
    for row in ky.rows:
        for cell in row.cells:
            _bo_border_cell(cell)

    cell_ngay = ky.rows[0].cells[0].merge(ky.rows[0].cells[2])
    _set_cell(cell_ngay, f"{_pgd_plain(ten_pgd)}, ngày ... tháng ... năm ...", align=WD_ALIGN_PARAGRAPH.RIGHT, font_size=11)

    _set_cell(ky.rows[1].cells[0], "LẬP BIỂU\n(Ký, ghi rõ họ tên)", bold=True, font_size=11)
    _set_cell(ky.rows[1].cells[1], "KIỂM SOÁT\n(Ký, ghi rõ họ tên)", bold=True, font_size=11)
    _set_cell(ky.rows[1].cells[2], "GIÁM ĐỐC\n(Ký tên, đóng dấu)", bold=True, font_size=11)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _tao_word_to_trinh_pgd(
    tong_hop_khoanh: dict,
    tong_hop_xoa: dict,
    ds_khoanh: list[dict],
    ten_pgd: str,
    nguon_label: str,
    dot: int,
    nam: int,
) -> bytes:
    def _set_run(p, text: str, *, bold: bool = False, italic: bool = False) -> None:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    def _p(doc_or_cell, text: str = "", *, bold: bool = False, align=None) -> None:
        if hasattr(doc_or_cell, "add_paragraph"):
            p = doc_or_cell.add_paragraph()
        else:
            p = doc_or_cell.paragraphs[0]
            p.text = ""
        if align is not None:
            p.alignment = align
        _set_run(p, text, bold=bold)

    ten_pgd_plain = _pgd_plain(ten_pgd)
    ten_pgd_line = _pgd_line(ten_pgd)

    nhom_3nam = [r for r in ds_khoanh if int(r.get("so_thang", 0) or 0) <= 36]
    nhom_5nam = [r for r in ds_khoanh if int(r.get("so_thang", 0) or 0) > 36]
    th_3 = _tong_hop_no(nhom_3nam)
    th_5 = _tong_hop_no(nhom_5nam)

    doc = Document()
    _style_doc_xln(doc)
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    _set_margins(doc, left_cm=3.0, right_cm=2.0, top_cm=2.5, bottom_cm=2.5)

    ngay_hn = date.today()
    hdr = doc.add_table(rows=1, cols=2)
    hdr.style = "Table Grid"
    for cell in hdr.rows[0].cells:
        _bo_border_cell(cell)

    cell_l = hdr.rows[0].cells[0]
    cell_r = hdr.rows[0].cells[1]

    p_l = cell_l.paragraphs[0]
    p_l.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_run(p_l, "CHI NHÁNH NHCSXH TỈNH ĐỒNG NAI", bold=True)
    p_l2 = cell_l.add_paragraph()
    _set_run(p_l2, ten_pgd_line, bold=True)
    p_l3 = cell_l.add_paragraph()
    _set_run(p_l3, "Số: .../TTr-NHCS", bold=False)

    p_r = cell_r.paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(p_r, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", bold=True)
    p_r2 = cell_r.add_paragraph()
    p_r2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(p_r2, "Độc lập - Tự do - Hạnh phúc", bold=True)
    p_r3 = cell_r.add_paragraph()
    p_r3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(p_r3, f"{ten_pgd_plain}, ngày ... tháng ... năm ...", bold=False)

    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(t, "TỜ TRÌNH", bold=True)
    t.runs[0].font.size = Pt(14)
    doc.add_paragraph(f"Về việc đề nghị xử lý nợ bị rủi ro đợt {dot} năm {nam}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Nguồn vốn: {nguon_label}").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph("Kính gửi: Giám đốc Chi nhánh Ngân hàng Chính sách xã hội tỉnh Đồng Nai")

    for line in [
        "Thực hiện Quyết định số 62/QĐ-HĐQT ngày 27/9/2021 của Chủ tịch Hội đồng",
        "quản trị NHCSXH về việc Ban hành Quy định xử lý nợ bị rủi ro trong hệ thống",
        "NHCSXH và các Quyết định sửa đổi, bổ sung Quyết định số 62/QĐ-HĐQT của Hội",
        "đồng quản trị NHCSXH.",
        "",
        "Căn cứ tình hình nợ bị rủi ro tại Phòng giao dịch; Sau khi nhận được hồ sơ",
        "đề nghị xử lý rủi ro của khách hàng, Phòng giao dịch NHCSXH "
        f"{ten_pgd_plain} đã",
        "phối hợp với khách hàng, cá nhân, tổ chức có thẩm quyền theo quy định tại",
        "Điều 7 Quyết định số 62/QĐ-HĐQT tổ chức kiểm tra thực tế, kiểm tra điều",
        "kiện, tính đầy đủ, chính xác, hợp pháp, hợp lệ của toàn bộ khoản nợ đề nghị",
        f"xử lý rủi ro đợt {dot} năm {nam}. Phòng giao dịch NHCSXH {ten_pgd_plain} báo cáo",
        "và trình Giám đốc Chi nhánh NHCSXH tỉnh Đồng Nai các nội dung như sau:",
    ]:
        if line:
            doc.add_paragraph(line)
        else:
            doc.add_paragraph()

    so_mon_khoanh = int(tong_hop_khoanh.get("tong_ho", 0) or 0)
    so_mon_xoa = int(tong_hop_xoa.get("tong_ho", 0) or 0)
    so_mon_tong = so_mon_khoanh + so_mon_xoa

    goc_khoanh = float(tong_hop_khoanh.get("tong_goc", 0) or 0)
    lai_khoanh = float(tong_hop_khoanh.get("tong_lai", 0) or 0)
    tien_khoanh = goc_khoanh + lai_khoanh

    goc_xoa = float(tong_hop_xoa.get("tong_goc", 0) or 0)
    lai_xoa = float(tong_hop_xoa.get("tong_lai", 0) or 0)
    tien_xoa = goc_xoa + lai_xoa

    goc_tong = goc_khoanh + goc_xoa
    lai_tong = lai_khoanh + lai_xoa
    tien_tong = goc_tong + lai_tong

    doc.add_paragraph(
        f"1. Tổng số đề nghị xử lý rủi ro đợt {dot} năm {nam}: {so_mon_tong} món vay, số"
        f" tiền {fmt(tien_tong)} đồng (gốc {fmt(goc_tong)} đồng, lãi {fmt(lai_tong)} đồng). Cụ thể:"
    )

    if so_mon_khoanh:
        so_mon_3 = int(th_3.get("tong_ho", 0) or 0)
        so_mon_5 = int(th_5.get("tong_ho", 0) or 0)
        goc_3 = float(th_3.get("tong_goc", 0) or 0)
        lai_3 = float(th_3.get("tong_lai", 0) or 0)
        goc_5 = float(th_5.get("tong_goc", 0) or 0)
        lai_5 = float(th_5.get("tong_lai", 0) or 0)
        doc.add_paragraph(
            f"a) Khoanh nợ: {so_mon_khoanh} món vay, số tiền {fmt(tien_khoanh)} đồng, gốc"
            f" {fmt(goc_khoanh)} đồng, lãi {fmt(lai_khoanh)} đồng. Trong đó:"
        )
        doc.add_paragraph(
            f"+ Khoanh nợ tối đa 3 năm: {so_mon_3} món vay, số tiền {fmt(goc_3 + lai_3)}"
            f" đồng, gốc {fmt(goc_3)} đồng, lãi {fmt(lai_3)} đồng."
        )
        doc.add_paragraph(
            f"+ Khoanh nợ tối đa 5 năm: {so_mon_5} món vay, số tiền {fmt(goc_5 + lai_5)}"
            f" đồng, gốc {fmt(goc_5)} đồng, lãi {fmt(lai_5)} đồng."
        )

    if so_mon_xoa:
        doc.add_paragraph(
            f"b) Xóa nợ: {so_mon_xoa} món vay, số tiền {fmt(tien_xoa)} đồng, gốc"
            f" {fmt(goc_xoa)} đồng, lãi {fmt(lai_xoa)} đồng."
        )

    doc.add_paragraph("(Biểu tổng hợp đề nghị xử lý nợ mẫu số 03, 04, 05/XLN đính kèm)")

    doc.add_paragraph(
        "2. Các khoản nợ bị rủi ro đề nghị xử lý đợt "
        f"{dot} năm {nam} của Phòng giao"
        " dịch đảm bảo đủ điều kiện, đúng thực tế, hồ sơ được thiết lập đúng và đầy đủ"
        " theo quy định."
    )
    doc.add_paragraph(
        "3. Phòng giao dịch chịu trách nhiệm trước pháp luật, trước Tổng Giám đốc về"
        " tính đầy đủ, chính xác, hợp pháp, hợp lệ, đúng thực tế của toàn bộ hồ sơ và"
        f" số liệu đề nghị xử lý rủi ro đợt {dot} năm {nam}."
    )

    doc.add_paragraph()
    doc.add_paragraph(
        f"Phòng giao dịch NHCSXH {ten_pgd_plain} kính trình Giám đốc chi nhánh NHCSXH tỉnh"
        " Đồng Nai trình cấp có thẩm quyền xem xét, quyết định./."
    )

    doc.add_paragraph()
    ky = doc.add_table(rows=1, cols=2)
    ky.style = "Table Grid"
    for cell in ky.rows[0].cells:
        _bo_border_cell(cell)

    ky.rows[0].cells[0].paragraphs[0].add_run("Nơi nhận:\n- Như kính gửi;\n- Lưu VT")
    pky = ky.rows[0].cells[1].paragraphs[0]
    pky.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(pky, "GIÁM ĐỐC", bold=True)
    pky2 = ky.rows[0].cells[1].add_paragraph()
    pky2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(pky2, "(Ký tên, đóng dấu)", bold=False)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _tao_word_to_trinh_cn(
    tong_hop_khoanh: dict,
    tong_hop_xoa: dict,
    ds_khoanh: list[dict],
    ten_tinh: str,
    nguon_label: str,
    dot: int,
    nam: int,
) -> bytes:
    def _set_run(p, text: str, *, bold: bool = False, italic: bool = False) -> None:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    nhom_3nam = [r for r in ds_khoanh if int(r.get("so_thang", 0) or 0) <= 36]
    nhom_5nam = [r for r in ds_khoanh if int(r.get("so_thang", 0) or 0) > 36]
    th_3 = _tong_hop_no(nhom_3nam)
    th_5 = _tong_hop_no(nhom_5nam)

    doc = Document()
    _style_doc_xln(doc)
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    _set_margins(doc, left_cm=3.0, right_cm=2.0, top_cm=2.5, bottom_cm=2.5)

    hdr = doc.add_table(rows=1, cols=2)
    hdr.style = "Table Grid"
    for cell in hdr.rows[0].cells:
        _bo_border_cell(cell)

    cell_l = hdr.rows[0].cells[0]
    cell_r = hdr.rows[0].cells[1]

    p_l = cell_l.paragraphs[0]
    p_l.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_run(p_l, "NGÂN HÀNG CHÍNH SÁCH XÃ HỘI", bold=True)
    p_l2 = cell_l.add_paragraph()
    _set_run(p_l2, f"CHI NHÁNH TỈNH {ten_tinh.upper()}", bold=True)
    p_l3 = cell_l.add_paragraph()
    _set_run(p_l3, "Số: .../TTr-NHCS", bold=False)

    p_r = cell_r.paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(p_r, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", bold=True)
    p_r2 = cell_r.add_paragraph()
    p_r2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(p_r2, "Độc lập - Tự do - Hạnh phúc", bold=True)
    p_r3 = cell_r.add_paragraph()
    p_r3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(p_r3, f"{ten_tinh}, ngày ... tháng ... năm ...", bold=False)

    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(t, "TỜ TRÌNH", bold=True)
    t.runs[0].font.size = Pt(14)
    doc.add_paragraph(f"Về việc đề nghị xử lý nợ bị rủi ro đợt {dot} năm {nam}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Nguồn vốn: {nguon_label}").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph("Kính gửi: Tổng Giám đốc Ngân hàng Chính sách xã hội")

    for line in [
        "Thực hiện Quyết định số 62/QĐ-HĐQT ngày 27/9/2021 của Chủ tịch Hội đồng",
        "quản trị NHCSXH về việc Ban hành Quy định xử lý nợ bị rủi ro trong hệ thống",
        "NHCSXH và các Quyết định sửa đổi, bổ sung Quyết định số 62/QĐ-HĐQT của Hội",
        "đồng quản trị NHCSXH.",
        "",
        "Căn cứ tình hình nợ bị rủi ro tại chi nhánh; Sau khi nhận được hồ sơ",
        "đề nghị xử lý rủi ro của khách hàng, chi nhánh NHCSXH tỉnh đã tổ chức",
        "kiểm tra thực tế khách hàng, kiểm tra hồ sơ của toàn bộ khoản nợ... "
        f"Chi nhánh NHCSXH tỉnh {ten_tinh} kính trình Tổng Giám đốc các nội dung như sau:",
    ]:
        if line:
            doc.add_paragraph(line)
        else:
            doc.add_paragraph()

    so_mon_khoanh = int(tong_hop_khoanh.get("tong_ho", 0) or 0)
    so_mon_xoa = int(tong_hop_xoa.get("tong_ho", 0) or 0)
    so_mon_tong = so_mon_khoanh + so_mon_xoa

    goc_khoanh = float(tong_hop_khoanh.get("tong_goc", 0) or 0)
    lai_khoanh = float(tong_hop_khoanh.get("tong_lai", 0) or 0)
    tien_khoanh = goc_khoanh + lai_khoanh

    goc_xoa = float(tong_hop_xoa.get("tong_goc", 0) or 0)
    lai_xoa = float(tong_hop_xoa.get("tong_lai", 0) or 0)
    tien_xoa = goc_xoa + lai_xoa

    goc_tong = goc_khoanh + goc_xoa
    lai_tong = lai_khoanh + lai_xoa
    tien_tong = goc_tong + lai_tong

    doc.add_paragraph(
        f"1. Tổng số đề nghị xử lý rủi ro đợt {dot} năm {nam}: {so_mon_tong} món vay, số"
        f" tiền {fmt(tien_tong)} đồng (gốc {fmt(goc_tong)} đồng, lãi {fmt(lai_tong)} đồng). Cụ thể:"
    )

    if so_mon_khoanh:
        so_mon_3 = int(th_3.get("tong_ho", 0) or 0)
        so_mon_5 = int(th_5.get("tong_ho", 0) or 0)
        goc_3 = float(th_3.get("tong_goc", 0) or 0)
        lai_3 = float(th_3.get("tong_lai", 0) or 0)
        goc_5 = float(th_5.get("tong_goc", 0) or 0)
        lai_5 = float(th_5.get("tong_lai", 0) or 0)
        doc.add_paragraph(
            f"a) Khoanh nợ: {so_mon_khoanh} món vay, số tiền {fmt(tien_khoanh)} đồng, gốc"
            f" {fmt(goc_khoanh)} đồng, lãi {fmt(lai_khoanh)} đồng. Trong đó:"
        )
        doc.add_paragraph(
            f"+ Khoanh nợ tối đa 3 năm: {so_mon_3} món vay, số tiền {fmt(goc_3 + lai_3)}"
            f" đồng, gốc {fmt(goc_3)} đồng, lãi {fmt(lai_3)} đồng."
        )
        doc.add_paragraph(
            f"+ Khoanh nợ tối đa 5 năm: {so_mon_5} món vay, số tiền {fmt(goc_5 + lai_5)}"
            f" đồng, gốc {fmt(goc_5)} đồng, lãi {fmt(lai_5)} đồng."
        )

    if so_mon_xoa:
        doc.add_paragraph(
            f"b) Xóa nợ: {so_mon_xoa} món vay, số tiền {fmt(tien_xoa)} đồng, gốc"
            f" {fmt(goc_xoa)} đồng, lãi {fmt(lai_xoa)} đồng."
        )

    doc.add_paragraph("(Biểu tổng hợp đề nghị xử lý nợ mẫu số 03, 04, 05/XLN đính kèm)")

    doc.add_paragraph(
        "2. Các khoản nợ bị rủi ro đề nghị xử lý đợt "
        f"{dot} năm {nam} của chi nhánh đảm bảo đủ điều kiện, đúng thực tế, hồ sơ được"
        " thiết lập đúng và đầy đủ theo quy định."
    )
    doc.add_paragraph(
        "3. Chi nhánh chịu trách nhiệm trước pháp luật, trước Tổng Giám đốc về tính"
        " đầy đủ, chính xác, hợp pháp, hợp lệ, đúng thực tế của toàn bộ hồ sơ và số"
        f" liệu đề nghị xử lý rủi ro đợt {dot} năm {nam}."
    )

    doc.add_paragraph()
    doc.add_paragraph(
        f"Chi nhánh NHCSXH tỉnh {ten_tinh} kính trình Tổng Giám đốc trình cấp có thẩm quyền"
        " xem xét, quyết định./."
    )

    doc.add_paragraph()
    ky = doc.add_table(rows=1, cols=2)
    ky.style = "Table Grid"
    for cell in ky.rows[0].cells:
        _bo_border_cell(cell)

    ky.rows[0].cells[0].paragraphs[0].add_run("Nơi nhận:\n- Như kính gửi;\n- Lưu VT")
    pky = ky.rows[0].cells[1].paragraphs[0]
    pky.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(pky, "GIÁM ĐỐC CHI NHÁNH", bold=True)
    pky2 = ky.rows[0].cells[1].add_paragraph()
    pky2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(pky2, "(Ký tên, đóng dấu)", bold=False)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════════════════
# MẪU 01/XLN và 02/XLN PHIÊN BẢN MỚI (v2) — Theo mẫu thực tế QĐ62
# ═══════════════════════════════════════════════════════════════════════════════

def _add_header_xln_v2(doc: Document, ten_nhcsxh: str, dia_danh: str, ngay: date, mau_so: str) -> None:
    """Header chuẩn: Bảng 2 cột không border."""
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    
    for cell in table.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for bn in ["top", "left", "bottom", "right"]:
            b = OxmlElement(f"w:{bn}")
            b.set(qn("w:val"), "none")
            tcBorders.append(b)
        tcPr.append(tcBorders)
    
    cell_left = table.rows[0].cells[0]
    cell_left.paragraphs[0].add_run(f"Mẫu số {mau_so}")
    
    cell_right = table.rows[0].cells[1]
    p1 = cell_right.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run("CỘNG HOÀ XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    r1.bold = True
    
    p2 = cell_right.add_paragraph("Độc lập - Tự do - Hạnh phúc")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.runs[0]
    r2.bold = True
    
    p3 = cell_right.add_paragraph(f"{dia_danh}, ngày {ngay.day} tháng {ngay.month} năm {ngay.year}")
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()


def _add_tieu_de_xln(doc: Document, line1: str, line2: str = None, line3: str = None) -> None:
    """Tiêu đề căn giữa, in đậm."""
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(line1)
    r1.bold = True
    r1.font.size = Pt(14)
    r1.font.name = "Times New Roman"
    
    if line2:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(line2)
        r2.font.size = Pt(13)
        r2.font.name = "Times New Roman"
    
    if line3:
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = p3.add_run(line3)
        r3.font.size = Pt(12)
        r3.font.name = "Times New Roman"
    
    doc.add_paragraph()


def _add_noi_dung_01xln(doc: Document, du_lieu: dict) -> None:
    """Thêm nội dung 4 điểm cho 01/XLN — text tự do."""
    doc.add_paragraph(f"Kính gửi: Ngân hàng Chính sách xã hội {du_lieu.get('ten_nhcsxh', '')}")
    doc.add_paragraph()
    
    doc.add_paragraph(f"Tên tôi là: {du_lieu.get('ten_kh', '')}")
    doc.add_paragraph(f"Hiện cư trú tại: {du_lieu.get('dia_chi', '')}")
    doc.add_paragraph(f"Là thành viên của Tổ TK&VV {du_lieu.get('ma_to', '')} do ông (bà): {du_lieu.get('ten_to_truong', '')} làm Tổ trưởng")
    doc.add_paragraph()
    
    p2 = doc.add_paragraph()
    p2.add_run(f"Theo HĐTD (sổ Vay vốn) số {du_lieu.get('so_ku', '')}, ngày {du_lieu.get('ngay_vay', '')}, tôi có đứng tên vay vốn chương trình {du_lieu.get('ten_ct', '')} tại NHCSXH {du_lieu.get('ten_nhcsxh', '')}.")
    
    p2_1 = doc.add_paragraph()
    p2_1.paragraph_format.left_indent = Cm(1)
    p2_1.add_run(f"Số tiền vay {du_lieu.get('muc_vay', '')} đồng; Thời điểm nhận tiền vay {du_lieu.get('ngay_vay', '')};\tHạn trả nợ: {du_lieu.get('ngay_dh', '')} ; Mục đích vay vốn {du_lieu.get('muc_dich_vay', '')}")
    
    p2_2 = doc.add_paragraph()
    p2_2.paragraph_format.left_indent = Cm(1)
    p2_2.add_run(f"Hiện nay, tôi còn nợ Ngân hàng số tiền: {du_lieu.get('tong_du_no', '')} đồng\n(Trong đó: Nợ gốc: {du_lieu.get('du_no_goc', '')} đồng; Nợ lãi: {du_lieu.get('lai_ton', '')} đồng)")
    doc.add_paragraph()
    
    p3 = doc.add_paragraph()
    p3.add_run(f"Trong thời gian vừa qua do: {du_lieu.get('nguyen_nhan', '')}")
    doc.add_paragraph("........................................................................................")
    doc.add_paragraph("3. Số vốn, tài sản của dự án, phương án vay vốn bị thiệt hại như sau:")
    
    p3_2 = doc.add_paragraph()
    p3_2.paragraph_format.left_indent = Cm(1)
    p3_2.add_run(f"- Số vốn và tài sản bị thiệt hại {du_lieu.get('so_tien_thiet_hai', '')} đồng.")
    
    p3_3 = doc.add_paragraph()
    p3_3.paragraph_format.left_indent = Cm(1)
    p3_3.add_run("(Ghi rõ tên, số lượng hiện vật bị thiệt hại): ..................................)")
    
    p3_4 = doc.add_paragraph()
    p3_4.paragraph_format.left_indent = Cm(1)
    p3_4.add_run(f"- Tổng số vốn thực hiện dự án (phương án vay vốn) {du_lieu.get('muc_vay', '')} đồng.")
    
    p3_5 = doc.add_paragraph()
    p3_5.paragraph_format.left_indent = Cm(1)
    p3_5.add_run(f"- Mức độ thiệt hại về vốn và tài sản {du_lieu.get('muc_do_thiet_hai', '')}%.")
    doc.add_paragraph()
    
    p4 = doc.add_paragraph()
    p4.add_run("4. Tình hình kinh tế, khả năng trả nợ Ngân hàng hiện nay của cá nhân và các thành viên cùng tham gia ký kết hợp đồng vay vốn sau khi gặp rủi ro:")
    
    p4_1 = doc.add_paragraph()
    p4_1.paragraph_format.left_indent = Cm(1)
    p4_1.add_run(du_lieu.get('kha_nang_tra_no', ''))
    doc.add_paragraph("........................................................................................")
    doc.add_paragraph()
    
    p_de_nghi = doc.add_paragraph()
    p_de_nghi.add_run(f"Vậy tôi làm đơn này đề nghị NHCSXH {du_lieu.get('ten_nhcsxh', '')} và các cơ quan chức năng xem xét {du_lieu.get('bien_phap', '')} số nợ bị rủi ro do nguyên nhân khách quan của tôi, cụ thể:")
    
    p_de_nghi_1 = doc.add_paragraph()
    p_de_nghi_1.paragraph_format.left_indent = Cm(1)
    p_de_nghi_1.add_run(f"- Số tiền đề nghị là {du_lieu.get('so_tien_de_nghi', '')} đồng")
    
    p_de_nghi_2 = doc.add_paragraph()
    p_de_nghi_2.paragraph_format.left_indent = Cm(1)
    p_de_nghi_2.add_run(f"(Trong đó: Nợ gốc: {du_lieu.get('du_no_goc', '')} đồng; Nợ lãi: {du_lieu.get('lai_ton', '')} đồng)")
    
    p_de_nghi_3 = doc.add_paragraph()
    p_de_nghi_3.paragraph_format.left_indent = Cm(1)
    p_de_nghi_3.add_run(f"- Thời gian đề nghị {du_lieu.get('so_thang', '')} tháng")
    
    p_de_nghi_4 = doc.add_paragraph()
    p_de_nghi_4.paragraph_format.left_indent = Cm(1)
    p_de_nghi_4.add_run(f"- Kế hoạch trả nợ: {du_lieu.get('ke_hoach_tra_no', '')}")
    doc.add_paragraph()
    
    doc.add_paragraph("Tôi xin cam đoan và chịu trách nhiệm trước pháp luật về nội dung kê khai trên đơn và các hồ sơ giấy tờ chứng minh là đúng.")
    doc.add_paragraph()


def _add_phan_ky_1cot(doc: Document, ngay: date, ten_kh: str, dia_danh: str) -> None:
    """Phần ký 1 cột — căn phải (cho 01/XLN)."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    
    for cell in table.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for bn in ["top", "left", "bottom", "right"]:
            b = OxmlElement(f"w:{bn}")
            b.set(qn("w:val"), "none")
            tcBorders.append(b)
        tcPr.append(tcBorders)
    
    cell = table.rows[0].cells[0]
    
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p1.add_run(f"{dia_danh}, ngày {ngay.day} tháng {ngay.month} năm {ngay.year}")
    
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.add_run("Người làm đơn")
    
    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p3.add_run("(Ký ghi rõ họ tên, hoặc điểm chỉ)")
    p3.runs[0].italic = True
    
    cell.add_paragraph()
    cell.add_paragraph()
    
    p4 = cell.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p4.add_run(ten_kh)


def _add_phan_ky_3cot_v2(doc: Document, ngay: date, dia_danh: str) -> None:
    """Phần ký bảng 3 cột có border (cho 02/XLN)."""
    doc.add_paragraph()
    
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    
    cell1 = table.rows[0].cells[0]
    p1_1 = cell1.paragraphs[0]
    p1_1.add_run("Đại diện khách hàng vay vốn\n")
    p1_1.add_run("(Ký ghi rõ họ tên hoặc điểm chỉ,\nđóng dấu nếu là đại diện pháp nhân)")
    p1_1.runs[1].italic = True
    cell1.add_paragraph("\n\nx\n\n\n")
    p1_2 = cell1.add_paragraph()
    p1_2.add_run("Đại diện UBND cấp xã\n")
    p1_2.add_run("(Xác nhận, ký tên, đóng dấu)")
    p1_2.runs[1].italic = True
    cell1.add_paragraph("\n\nx")
    
    cell2 = table.rows[0].cells[1]
    p2_1 = cell2.paragraphs[0]
    p2_1.add_run("Tổ trưởng Tổ TK&VV\n")
    p2_1.add_run("(Ký, ghi rõ họ tên)")
    p2_1.runs[1].italic = True
    cell2.add_paragraph("\n\nx\n\n\n")
    p2_2 = cell2.add_paragraph()
    p2_2.add_run("Đại diện tổ chức\nHội đoàn thể nhận\nủy thác\n")
    p2_2.add_run("(Ký tên, đóng dấu)")
    p2_2.runs[1].italic = True
    cell2.add_paragraph("\n\nx")
    
    cell3 = table.rows[0].cells[2]
    p3_1 = cell3.paragraphs[0]
    p3_1.add_run("Cán bộ tín dụng\n")
    p3_1.add_run("(Ký tên, ghi rõ họ tên)")
    p3_1.runs[1].italic = True
    cell3.add_paragraph("\n\nx\n\n\n")
    p3_2 = cell3.add_paragraph()
    p3_2.add_run("Đại diện NHCSXH\n")
    p3_2.add_run("(Ký tên, đóng dấu)")
    p3_2.runs[1].italic = True
    cell3.add_paragraph("\n\nx")


def _add_thanh_phan_tham_du_02xln(doc: Document, du_lieu: dict) -> None:
    """Thêm thành phần tham dự (6 người) cho 02/XLN."""
    ngay = du_lieu.get('ngay_lap', date.today())
    
    p = doc.add_paragraph()
    p.add_run(f"Hôm nay, ngày {ngay.day} tháng {ngay.month} năm {ngay.year}, tại {du_lieu.get('dia_diem', '')} chúng tôi gồm có:")
    doc.add_paragraph()
    
    thanh_phan = [
        (f"1. {du_lieu.get('ten_pgd', '')}", "Phó Giám đốc", f"NHCSXH {du_lieu.get('ten_nhcsxh', '')}"),
        (f"2. {du_lieu.get('ten_ubnd', '')}", "Phó Chủ tịch", "UBND xã"),
        (f"3. {du_lieu.get('ten_hoi_nd', '')}", "Chủ tịch", "Hội Nông dân xã"),
        (f"4. {du_lieu.get('ten_cbtd', '')}", "CBTD", f"NHCSXH {du_lieu.get('ten_nhcsxh', '')}"),
        (f"5. {du_lieu.get('ten_to_truong', '')}", "Tổ trưởng", "Tổ TKVV"),
        (f"6. {du_lieu.get('ten_kh', '')}", "", "là khách hàng vay vốn."),
    ]
    
    for ten, chuc_vu, dai_dien in thanh_phan:
        if chuc_vu:
            doc.add_paragraph(f"{ten}, Chức vụ: {chuc_vu}, Đại diện: {dai_dien}")
        else:
            doc.add_paragraph(f"{ten} {dai_dien}")
    
    doc.add_paragraph()
    
    p_mo_dau = doc.add_paragraph()
    p_mo_dau.add_run(f"Đã tiến hành thẩm tra và lập biên bản đề nghị xử lý nợ bị rủi ro của ông (bà): {du_lieu.get('ten_kh', '')}, địa chỉ: {du_lieu.get('dia_chi', '')} là đại diện hộ gia đình vay vốn Ngân hàng Chính sách xã hội (NHCSXH) theo HĐTD (Sổ Vay vốn) số {du_lieu.get('so_ku', '')}, ngày {du_lieu.get('ngay_vay', '')}, có mã món vay {du_lieu.get('ma_mon_vay', '')}. Cụ thể như sau:")
    doc.add_paragraph()


def _add_noi_dung_02xln(doc: Document, du_lieu: dict) -> None:
    """Thêm nội dung 5 mục La Mã cho 02/XLN."""
    
    p_i = doc.add_paragraph()
    run_i = p_i.add_run("I. Nguyên nhân khách hàng bị rủi ro:")
    run_i.bold = True
    doc.add_paragraph(du_lieu.get('nguyen_nhan', ''))
    doc.add_paragraph()
    
    p_ii = doc.add_paragraph()
    run_ii = p_ii.add_run("II. Xác định mức độ thiệt hại về vốn và tài sản:")
    run_ii.bold = True
    
    for i, text in enumerate([
        f"1. Số vốn và tài sản bị thiệt hại: {du_lieu.get('so_tien_thiet_hai', '')} đồng.",
        f"(Ghi rõ tên, số lượng hiện vật bị thiệt hại): {du_lieu.get('chi_tiet_thiet_hai', '')}",
        f"2. Tổng số vốn thực hiện dự án (phương án vay vốn): {du_lieu.get('muc_vay', '')} đồng.",
        f"3. Đánh giá mức độ thiệt hại về vốn và tài sản: {du_lieu.get('danh_gia_thiet_hai', '')}"
    ], 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.add_run(text)
    doc.add_paragraph()
    
    p_iii = doc.add_paragraph()
    run_iii = p_iii.add_run("III. Dư nợ tại NHCSXH đến ngày lập biên bản:")
    run_iii.bold = True
    doc.add_paragraph(f"Tổng số nợ còn phải trả ngân hàng đến ngày lập biên bản: {du_lieu.get('tong_du_no', '')} đồng.")
    
    for text in [f"Trong đó:  + Nợ gốc: {du_lieu.get('du_no_goc', '')} đồng.", f"+ Nợ lãi: {du_lieu.get('lai_ton', '')} đồng."]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.add_run(text)
    doc.add_paragraph()
    
    p_iv = doc.add_paragraph()
    run_iv = p_iv.add_run("IV. Đánh giá thực trạng dự án, phương án vay vốn, tài sản và khả năng trả nợ của khách hàng sau khi bị thiệt hại:")
    run_iv.bold = True
    
    for title, content in [
        ("1. Đánh giá thực trạng dự án, phương án vay vốn; phương án khôi phục dự án vay vốn:", du_lieu.get('danh_gia_du_an', '')),
        ("2. Tài sản hiện tại của khách hàng:", du_lieu.get('tai_san_hien_tai', '')),
        ("3. Đánh giá tình trạng khả năng trả nợ của khách hàng và tất cả các thành viên:", du_lieu.get('kha_nang_tra_no', '')),
        ("4. Về việc áp dụng mọi biện pháp thu hồi nợ:", du_lieu.get('bien_phap_thu_hoi', '...................................................................................................................................'))
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.add_run(title)
        
        p_nd = doc.add_paragraph()
        p_nd.paragraph_format.left_indent = Cm(1)
        p_nd.add_run(content)
    doc.add_paragraph()
    
    p_v = doc.add_paragraph()
    run_v = p_v.add_run("V. Đề xuất biện pháp xử lý:")
    run_v.bold = True
    
    p_v_1 = doc.add_paragraph()
    p_v_1.add_run(f"Căn cứ vào quy chế xử lý nợ bị rủi ro do nguyên nhân khách quan tại NHCSXH, thẩm tra tình hình thực tế của khách hàng, chúng tôi nhất trí xác nhận các nội dung trên là đúng và thống nhất đề nghị NHCSXH, các cơ quan có thẩm quyền xem xét {du_lieu.get('bien_phap', '')} cho {du_lieu.get('ten_kh', '')} với thời gian {du_lieu.get('so_thang', '')} tháng, số tiền {du_lieu.get('so_tien_de_nghi', '')} đồng.")
    
    for text in [f"Trong đó:  + Nợ gốc: {du_lieu.get('du_no_goc', '')} đồng.", f"+ Nợ lãi: {du_lieu.get('lai_ton', '')} đồng."]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.add_run(text)
    doc.add_paragraph()
    
    doc.add_paragraph("Biên bản này lập thành 02 bản có giá trị pháp lý như nhau và được các thành viên thống nhất thông qua ký tên dưới đây.")


def _tao_word_01xln_v2(du_lieu: dict) -> bytes:
    """Mẫu 01/XLN — Đơn đề nghị xử lý nợ (theo mẫu thực tế QĐ62)."""
    doc = Document()
    _style_doc_xln(doc)
    _set_margins(doc, left_cm=3.0, right_cm=2.0, top_cm=2.5, bottom_cm=2.5)
    
    _add_header_xln_v2(doc, du_lieu.get('ten_nhcsxh', ''), du_lieu.get('dia_danh', ''), du_lieu.get('ngay_ky', date.today()), "01/XLN")
    _add_tieu_de_xln(doc, "ĐƠN ĐỀ NGHỊ XỬ LÝ NỢ")
    _add_noi_dung_01xln(doc, du_lieu)
    _add_phan_ky_1cot(doc, du_lieu.get('ngay_ky', date.today()), du_lieu.get('ten_kh', ''), du_lieu.get('dia_danh', ''))
    
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _tao_word_02xln_v2(du_lieu: dict) -> bytes:
    """Mẫu 02/XLN — Biên bản đề nghị xử lý nợ (theo mẫu thực tế QĐ62)."""
    doc = Document()
    _style_doc_xln(doc)
    _set_margins(doc, left_cm=3.0, right_cm=2.0, top_cm=2.5, bottom_cm=2.5)
    
    ngay_lap = du_lieu.get('ngay_lap', date.today())
    
    _add_header_xln_v2(doc, du_lieu.get('ten_nhcsxh', ''), du_lieu.get('dia_danh', ''), ngay_lap, "02/XLN")
    _add_tieu_de_xln(doc, "BIÊN BẢN", "Đề nghị xử lý nợ bị rủi ro", f"(Chương trình {du_lieu.get('ten_ct', '')})")
    _add_thanh_phan_tham_du_02xln(doc, du_lieu)
    _add_noi_dung_02xln(doc, du_lieu)
    _add_phan_ky_3cot_v2(doc, ngay_lap, du_lieu.get('dia_danh', ''))
    
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
