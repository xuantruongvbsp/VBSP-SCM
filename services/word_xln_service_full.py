"""Word generation helpers cho biểu mẫu XLN (Xử Lý Nợ) — phần mẫu tổng hợp.

Tách từ services/word_xln_service.py (giữ nguyên logic): mẫu 04/05 XLN v2
(biên bản tổng hợp) và thông báo kết quả XLRR.

Lưu ý: KHÔNG import trực tiếp module này — import qua services.word_xln_service
(re-export ở cuối file đó) để tránh circular import.
"""
from __future__ import annotations

import io
from datetime import date

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from config import TEN_CHI_NHANH_HIEN_THI
from utils import fmt
from services.rui_ro_aggregation import _tong_hop_no

from services.word_xln_service import (
    _add_header_xln_v2,
    _add_tieu_de_xln,
    _bo_border_cell,
    _pgd_line,
    _pgd_plain,
    _set_cell,
    _set_margins,
    _style_doc_xln,
)


# ═════════════════════════════════════════════════════════════════════════════
# MẪU 04/05 - TỔNG HỢP
# ═════════════════════════════════════════════════════════════════════════════

def _tao_word_04xln_v2(
    ds_hs: list,
    thong_tin: dict,
) -> bytes:
    """Mẫu 04/XLN — Biên bản tổng hợp đề nghị KHOANH NỢ.
    
    Args:
        ds_hs: Danh sách HoSoRuiRo (đã lọc biện pháp khoanh)
        thong_tin: Dict chứa ten_nhcsxh, dia_danh, ngay_lap, 
                   ten_pgd, ten_ubnd, ten_hoi_nd, ten_cbtd
    
    Returns:
        File Word dạng bytes
    """
    from services.xlrr_service import HoSoRuiRo
    
    doc = Document()
    _style_doc_xln(doc)
    _set_margins(doc, left_cm=2.5, right_cm=2.0, top_cm=2.0, bottom_cm=2.0)
    
    ngay_lap = thong_tin.get('ngay_lap', date.today())
    ten_nhcsxh = thong_tin.get('ten_nhcsxh', TEN_CHI_NHANH_HIEN_THI)
    dia_danh = thong_tin.get('dia_danh', 'TP. Biên Hòa')
    
    # Header
    _add_header_xln_v2(doc, ten_nhcsxh, dia_danh, ngay_lap, "04/XLN")
    
    # Tiêu đề
    _add_tieu_de_xln(
        doc, 
        "BIÊN BẢN TỔNG HỢP", 
        "Đề nghị khoanh nợ bị rủi ro",
        f"(Tổng hợp {len(ds_hs)} hộ vay)"
    )
    
    # Thành phần tham dự
    _add_thanh_phan_tong_hop(doc, thong_tin)
    
    # Nội dung - Danh sách các hộ
    p = doc.add_paragraph()
    p.add_run("Căn cứ các biên bản kiểm tra, chúng tôi nhất trí đề nghị khoanh nợ cho các hộ vay sau:").italic = True
    
    # Bảng tổng hợp
    _add_bang_tong_hop_04_05(doc, ds_hs)
    
    # Tổng kết
    tong_goc = sum(hs.du_no_goc for hs in ds_hs)
    tong_lai = sum(hs.lai_ton for hs in ds_hs)
    
    p_tong = doc.add_paragraph()
    p_tong.add_run(f"\nTổng cộng: {len(ds_hs)} hộ vay, tổng dư nợ gốc {fmt(tong_goc)} đồng, lãi {fmt(tong_lai)} đồng.")
    p_tong.runs[0].bold = True
    
    # Phần ký
    _add_phan_ky_tong_hop(doc, ngay_lap, dia_danh, thong_tin)
    
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _tao_word_05xln_v2(
    ds_hs: list,
    thong_tin: dict,
) -> bytes:
    """Mẫu 05/XLN — Biên bản tổng hợp đề nghị XÓA NỢ.
    
    Args:
        ds_hs: Danh sách HoSoRuiRo (đã lọc biện pháp xoa)
        thong_tin: Dict chứa ten_nhcsxh, dia_danh, ngay_lap,
                   ten_pgd, ten_ubnd, ten_hoi_nd, ten_cbtd
    
    Returns:
        File Word dạng bytes
    """
    doc = Document()
    _style_doc_xln(doc)
    _set_margins(doc, left_cm=2.5, right_cm=2.0, top_cm=2.0, bottom_cm=2.0)
    
    ngay_lap = thong_tin.get('ngay_lap', date.today())
    ten_nhcsxh = thong_tin.get('ten_nhcsxh', TEN_CHI_NHANH_HIEN_THI)
    dia_danh = thong_tin.get('dia_danh', 'TP. Biên Hòa')
    
    # Header
    _add_header_xln_v2(doc, ten_nhcsxh, dia_danh, ngay_lap, "05/XLN")
    
    # Tiêu đề
    _add_tieu_de_xln(
        doc,
        "BIÊN BẢN TỔNG HỢP",
        "Đề nghị xóa nợ bị rủi ro",
        f"(Tổng hợp {len(ds_hs)} hộ vay)"
    )
    
    # Thành phần tham dự
    _add_thanh_phan_tong_hop(doc, thong_tin)
    
    # Nội dung
    p = doc.add_paragraph()
    p.add_run("Căn cứ các biên bản kiểm tra và đề nghị của PGD, chúng tôi nhất trí đề nghị xóa nợ cho các hộ vay sau:").italic = True
    
    # Bảng tổng hợp
    _add_bang_tong_hop_04_05(doc, ds_hs)
    
    # Tổng kết
    tong_goc = sum(hs.du_no_goc for hs in ds_hs)
    tong_lai = sum(hs.lai_ton for hs in ds_hs)
    
    p_tong = doc.add_paragraph()
    p_tong.add_run(f"\nTổng cộng: {len(ds_hs)} hộ vay, tổng dư nợ gốc {fmt(tong_goc)} đồng, lãi {fmt(tong_lai)} đồng.")
    p_tong.runs[0].bold = True
    
    # Phần ký
    _add_phan_ky_tong_hop(doc, ngay_lap, dia_danh, thong_tin)
    
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_thanh_phan_tong_hop(doc: Document, thong_tin: dict) -> None:
    """Thêm phần thành phần tham dự cho mẫu tổng hợp."""
    p = doc.add_paragraph()
    p.add_run("I. THÀNH PHẦN THAM DỰ:").bold = True
    
    thanh_phan = [
        ("1", "Phó Giám đốc NHCSXH", thong_tin.get('ten_pgd', '...')),
        ("2", "Phó Chủ tịch UBND", thong_tin.get('ten_ubnd', '...')),
        ("3", "Chủ tịch Hội Nông dân", thong_tin.get('ten_hoi_nd', '...')),
        ("4", "CBTD NHCSXH", thong_tin.get('ten_cbtd', '...')),
    ]
    
    for stt, chuc_danh, ten in thanh_phan:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.add_run(f"{stt}. {chuc_danh}: {ten}")


def _add_bang_tong_hop_04_05(
    doc: Document,
    ds_hs: list,
) -> None:
    """Thêm bảng tổng hợp danh sách hộ vay (dùng chung cho mẫu 04 và 05)."""
    # Tạo bảng
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    hdr_cells = table.rows[0].cells
    headers = ["STT", "Tên khách hàng", "Số KU", "Xã", "Dư nợ gốc", "Dư nợ lãi", "Tổng", "Nguyên nhân rủi ro"]
    for i, text in enumerate(headers):
        _set_cell(hdr_cells[i], text, bold=True, font_size=10)
    
    # Data rows
    for idx, hs in enumerate(ds_hs, 1):
        row = table.add_row().cells
        _set_cell(row[0], str(idx), font_size=10)
        _set_cell(row[1], hs.ten_kh, font_size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
        _set_cell(row[2], hs.so_ku, font_size=10)
        _set_cell(row[3], hs.xa, font_size=10)
        _set_cell(row[4], fmt(hs.du_no_goc), font_size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell(row[5], fmt(hs.lai_ton), font_size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell(row[6], fmt(hs.tong_du_no), font_size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell(row[7], hs.nguyen_nhan, font_size=9, align=WD_ALIGN_PARAGRAPH.LEFT)


def _add_phan_ky_tong_hop(
    doc: Document,
    ngay_lap: date,
    dia_danh: str,
    thong_tin: dict,
) -> None:
    """Thêm phần ký cho mẫu tổng hợp."""
    doc.add_paragraph()

    # Dòng địa danh + ngày tháng (bắt buộc với văn bản hành chính VN)
    p_ngay = doc.add_paragraph()
    p_ngay.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_ngay.add_run(
        f"{dia_danh}, ngày {ngay_lap.day:02d} tháng {ngay_lap.month:02d} năm {ngay_lap.year}"
    ).italic = True

    # Tạo bảng 3 cột cho phần ký
    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Row 1: Chức danh
    row1 = tbl.rows[0].cells
    _set_cell(row1[0], "PHÓ GIÁM ĐỐC\nNHCSXH", bold=True, font_size=11)
    _set_cell(row1[1], "PHÓ CHỦ TỊCH\nUBND", bold=True, font_size=11)
    _set_cell(row1[2], "CHỦ TỊCH\nHỘI NÔNG DÂN", bold=True, font_size=11)
    
    # Row 2: Khoảng trống ký
    row2 = tbl.add_row().cells
    for cell in row2:
        _set_cell(cell, "\n\n\n", font_size=11)
    
    # Row 3: Tên người ký
    row3 = tbl.add_row().cells
    _set_cell(row3[0], thong_tin.get('ten_pgd', ''), bold=True, font_size=11)
    _set_cell(row3[1], thong_tin.get('ten_ubnd', ''), bold=True, font_size=11)
    _set_cell(row3[2], thong_tin.get('ten_hoi_nd', ''), bold=True, font_size=11)


# ── Thông báo kết quả XLRR ──────────────────────────────────────────────────

def _tao_word_thong_bao_ket_qua_cn(
    ds_ket_qua: list[dict],
    so_quyet_dinh: str,
    ngay_quyet_dinh: date,
    dot: int,
    nam: int,
) -> bytes:
    """Thông báo kết quả xử lý nợ rủi ro toàn CN gửi các PGD.

    Bảng tổng hợp theo PGD: PGD | Số HS | Khoanh | Xóa | Không duyệt | Tổng tiền duyệt.
    """
    from collections import defaultdict

    doc = Document()
    _style_doc_xln(doc)
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    _set_margins(doc, left_cm=3.0, right_cm=2.0, top_cm=2.5, bottom_cm=2.5)

    ngay_hom_nay = date.today()

    # Header 2 cột
    hdr = doc.add_table(rows=1, cols=2)
    hdr.style = "Table Grid"
    for cell in hdr.rows[0].cells:
        _bo_border_cell(cell)

    cell_l = hdr.rows[0].cells[0]
    cell_r = hdr.rows[0].cells[1]

    p_l = cell_l.paragraphs[0]
    p_l.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p_l.add_run(TEN_CHI_NHANH_HIEN_THI.upper())
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    p_l2 = cell_l.add_paragraph()
    r2 = p_l2.add_run(f"Số: .../TB-NHCS")
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(12)

    p_r = cell_r.paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = p_r.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    rr.bold = True
    rr.font.name = "Times New Roman"
    rr.font.size = Pt(12)
    p_r2 = cell_r.add_paragraph()
    p_r2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr2 = p_r2.add_run("Độc lập - Tự do - Hạnh phúc")
    rr2.bold = True
    rr2.font.name = "Times New Roman"
    rr2.font.size = Pt(12)
    p_r3 = cell_r.add_paragraph()
    p_r3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr3 = p_r3.add_run(
        f"Đồng Nai, ngày {ngay_hom_nay.day:02d} tháng {ngay_hom_nay.month:02d} năm {ngay_hom_nay.year}"
    )
    rr3.font.name = "Times New Roman"
    rr3.font.size = Pt(12)

    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run("THÔNG BÁO")
    tr.bold = True
    tr.font.name = "Times New Roman"
    tr.font.size = Pt(14)

    sub = doc.add_paragraph(
        f"Kết quả xử lý nợ bị rủi ro đợt {dot} năm {nam}"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.name = "Times New Roman"
    sub.runs[0].font.size = Pt(12)

    doc.add_paragraph()
    kq = doc.add_paragraph(
        f"Căn cứ Quyết định số {so_quyet_dinh} ngày "
        f"{ngay_quyet_dinh.day:02d}/{ngay_quyet_dinh.month:02d}/{ngay_quyet_dinh.year} "
        "của NHCSXH về việc xử lý nợ bị rủi ro, Chi nhánh thông báo kết quả như sau:"
    )
    kq.runs[0].font.name = "Times New Roman"
    kq.runs[0].font.size = Pt(12)

    # Tổng hợp theo PGD
    agg: dict[str, dict] = defaultdict(lambda: {
        "so_hs": 0, "khoanh": 0, "xoa": 0, "khong_duyet": 0, "tong_tien": 0.0
    })
    for r_item in ds_ket_qua:
        pgd = r_item.get("ten_pgd", "Không rõ")
        agg[pgd]["so_hs"] += 1
        ket_qua = r_item.get("ket_qua", "")
        if ket_qua == "da_khoanh":
            agg[pgd]["khoanh"] += 1
        elif ket_qua == "da_xoa":
            agg[pgd]["xoa"] += 1
        elif ket_qua == "khong_duyet":
            agg[pgd]["khong_duyet"] += 1
        agg[pgd]["tong_tien"] += float(r_item.get("so_tien_duoc_duyet", 0) or 0)

    headers = ["STT", "Đơn vị", "Số HS", "Khoanh", "Xóa", "Không duyệt", "Tổng tiền (đồng)"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        _set_cell(hdr_cells[i], h, bold=True, font_size=10)

    tong_hs = tong_kh = tong_xoa = tong_kd = 0
    tong_tien_all = 0.0
    for idx, (pgd_name, v) in enumerate(sorted(agg.items()), 1):
        row_cells = table.add_row().cells
        _set_cell(row_cells[0], str(idx), font_size=10)
        _set_cell(row_cells[1], pgd_name, font_size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
        _set_cell(row_cells[2], str(v["so_hs"]), font_size=10)
        _set_cell(row_cells[3], str(v["khoanh"]), font_size=10)
        _set_cell(row_cells[4], str(v["xoa"]), font_size=10)
        _set_cell(row_cells[5], str(v["khong_duyet"]), font_size=10)
        _set_cell(row_cells[6], fmt(v["tong_tien"]), font_size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)
        tong_hs += v["so_hs"]
        tong_kh += v["khoanh"]
        tong_xoa += v["xoa"]
        tong_kd += v["khong_duyet"]
        tong_tien_all += v["tong_tien"]

    # Dòng tổng
    tong_row = table.add_row().cells
    _set_cell(tong_row[0], "", font_size=10)
    _set_cell(tong_row[1], "Tổng cộng", bold=True, font_size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
    _set_cell(tong_row[2], str(tong_hs), bold=True, font_size=10)
    _set_cell(tong_row[3], str(tong_kh), bold=True, font_size=10)
    _set_cell(tong_row[4], str(tong_xoa), bold=True, font_size=10)
    _set_cell(tong_row[5], str(tong_kd), bold=True, font_size=10)
    _set_cell(tong_row[6], fmt(tong_tien_all), bold=True, font_size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)

    doc.add_paragraph()
    p_note = doc.add_paragraph(
        "Đề nghị các Phòng giao dịch triển khai thực hiện theo quy định."
    )
    p_note.runs[0].font.name = "Times New Roman"
    p_note.runs[0].font.size = Pt(12)

    # Phần ký Giám đốc CN
    doc.add_paragraph()
    p_ngay = doc.add_paragraph()
    p_ngay.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_ngay = p_ngay.add_run(
        f"Đồng Nai, ngày {ngay_hom_nay.day:02d} tháng {ngay_hom_nay.month:02d} năm {ngay_hom_nay.year}"
    )
    r_ngay.italic = True
    r_ngay.font.name = "Times New Roman"
    r_ngay.font.size = Pt(12)

    ky_tbl = doc.add_table(rows=1, cols=2)
    ky_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell in ky_tbl.rows[0].cells:
        _bo_border_cell(cell)

    _set_cell(ky_tbl.rows[0].cells[0], "NƠI NHẬN\n- Các PGD trực thuộc;\n- Lưu NHCS.", font_size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
    p_gd = ky_tbl.rows[0].cells[1].paragraphs[0]
    p_gd.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_gd = p_gd.add_run("GIÁM ĐỐC")
    r_gd.bold = True
    r_gd.font.name = "Times New Roman"
    r_gd.font.size = Pt(12)
    p_gd2 = ky_tbl.rows[0].cells[1].add_paragraph()
    p_gd2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_gd2 = p_gd2.add_run("\n\n\n\n(Ký, đóng dấu)")
    r_gd2.italic = True
    r_gd2.font.name = "Times New Roman"
    r_gd2.font.size = Pt(12)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _tao_word_thong_bao_ket_qua_pgd(
    ds_ket_qua_pgd: list[dict],
    ten_pgd: str,
    so_quyet_dinh: str,
    ngay_quyet_dinh: date,
    dot: int,
    nam: int,
) -> bytes:
    """Thông báo kết quả xử lý nợ rủi ro chi tiết cho 1 PGD.

    Bảng chi tiết từng KH: KH | Số KU | Biện pháp | Kết quả | Tiền duyệt | Ghi chú.
    """
    ten_pgd_plain = _pgd_plain(ten_pgd)
    ten_pgd_line = _pgd_line(ten_pgd)

    doc = Document()
    _style_doc_xln(doc)
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    _set_margins(doc, left_cm=3.0, right_cm=2.0, top_cm=2.5, bottom_cm=2.5)

    ngay_hom_nay = date.today()

    # Header
    hdr = doc.add_table(rows=1, cols=2)
    hdr.style = "Table Grid"
    for cell in hdr.rows[0].cells:
        _bo_border_cell(cell)

    cell_l = hdr.rows[0].cells[0]
    cell_r = hdr.rows[0].cells[1]

    p_l = cell_l.paragraphs[0]
    p_l.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rl = p_l.add_run(TEN_CHI_NHANH_HIEN_THI.upper())
    rl.bold = True
    rl.font.name = "Times New Roman"
    rl.font.size = Pt(12)
    p_l2 = cell_l.add_paragraph()
    rl2 = p_l2.add_run(ten_pgd_line)
    rl2.bold = True
    rl2.font.name = "Times New Roman"
    rl2.font.size = Pt(12)
    p_l3 = cell_l.add_paragraph()
    rl3 = p_l3.add_run("Số: .../TB-NHCS")
    rl3.font.name = "Times New Roman"
    rl3.font.size = Pt(12)

    p_r = cell_r.paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = p_r.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    rr.bold = True
    rr.font.name = "Times New Roman"
    rr.font.size = Pt(12)
    p_r2 = cell_r.add_paragraph()
    p_r2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr2 = p_r2.add_run("Độc lập - Tự do - Hạnh phúc")
    rr2.bold = True
    rr2.font.name = "Times New Roman"
    rr2.font.size = Pt(12)
    p_r3 = cell_r.add_paragraph()
    p_r3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr3 = p_r3.add_run(
        f"{ten_pgd_plain}, ngày {ngay_hom_nay.day:02d} tháng {ngay_hom_nay.month:02d} năm {ngay_hom_nay.year}"
    )
    rr3.font.name = "Times New Roman"
    rr3.font.size = Pt(12)

    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run("THÔNG BÁO")
    tr.bold = True
    tr.font.name = "Times New Roman"
    tr.font.size = Pt(14)

    sub = doc.add_paragraph(
        f"Kết quả xử lý nợ bị rủi ro đợt {dot} năm {nam} — {ten_pgd_line}"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.name = "Times New Roman"
    sub.runs[0].font.size = Pt(12)

    doc.add_paragraph()
    kq = doc.add_paragraph(
        f"Căn cứ Quyết định số {so_quyet_dinh} ngày "
        f"{ngay_quyet_dinh.day:02d}/{ngay_quyet_dinh.month:02d}/{ngay_quyet_dinh.year} "
        "của NHCSXH, Chi nhánh thông báo kết quả xử lý nợ bị rủi ro cho "
        f"{ten_pgd_line} như sau:"
    )
    kq.runs[0].font.name = "Times New Roman"
    kq.runs[0].font.size = Pt(12)

    # Bảng chi tiết từng KH
    ket_qua_label = {
        "da_khoanh": "Đã khoanh",
        "da_xoa": "Đã xóa",
        "khong_duyet": "Không duyệt",
        "cho_xu_ly": "Chờ xử lý",
    }
    headers = ["STT", "Tên KH", "Số KU", "Biện pháp", "Kết quả", "Tiền duyệt (đồng)", "Ghi chú"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        _set_cell(hdr_cells[i], h, bold=True, font_size=10)

    tong_tien = 0.0
    for idx, r_item in enumerate(ds_ket_qua_pgd, 1):
        row_cells = table.add_row().cells
        _set_cell(row_cells[0], str(idx), font_size=10)
        _set_cell(row_cells[1], r_item.get("ten_kh", ""), font_size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
        _set_cell(row_cells[2], r_item.get("so_ku", ""), font_size=10)
        bp = "Khoanh" if r_item.get("bien_phap") == "khoanh" else "Xóa"
        _set_cell(row_cells[3], bp, font_size=10)
        kq_text = ket_qua_label.get(r_item.get("ket_qua", ""), r_item.get("ket_qua", ""))
        _set_cell(row_cells[4], kq_text, font_size=10)
        tien = float(r_item.get("so_tien_duoc_duyet", 0) or 0)
        _set_cell(row_cells[5], fmt(tien), font_size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell(row_cells[6], r_item.get("ghi_chu", ""), font_size=9, align=WD_ALIGN_PARAGRAPH.LEFT)
        tong_tien += tien

    # Dòng tổng
    tong_row = table.add_row().cells
    _set_cell(tong_row[0], "", font_size=10)
    _set_cell(tong_row[1], "Tổng cộng", bold=True, font_size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
    for i in range(2, 5):
        _set_cell(tong_row[i], "", font_size=10)
    _set_cell(tong_row[5], fmt(tong_tien), bold=True, font_size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell(tong_row[6], "", font_size=10)

    doc.add_paragraph()
    p_note = doc.add_paragraph(
        f"Đề nghị {ten_pgd_line} thực hiện theo quy định và thông báo kết quả đến "
        "khách hàng trong thời hạn quy định."
    )
    p_note.runs[0].font.name = "Times New Roman"
    p_note.runs[0].font.size = Pt(12)

    # Phần ký 2 cột: PGD nhận + Giám đốc CN
    doc.add_paragraph()
    p_ngay = doc.add_paragraph()
    p_ngay.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_ngay = p_ngay.add_run(
        f"Đồng Nai, ngày {ngay_hom_nay.day:02d} tháng {ngay_hom_nay.month:02d} năm {ngay_hom_nay.year}"
    )
    r_ngay.italic = True
    r_ngay.font.name = "Times New Roman"
    r_ngay.font.size = Pt(12)

    ky_tbl = doc.add_table(rows=1, cols=2)
    ky_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell in ky_tbl.rows[0].cells:
        _bo_border_cell(cell)

    p_pgd = ky_tbl.rows[0].cells[0].paragraphs[0]
    p_pgd.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_pgd = p_pgd.add_run(f"TRƯỞNG {ten_pgd_line.upper()}")
    r_pgd.bold = True
    r_pgd.font.name = "Times New Roman"
    r_pgd.font.size = Pt(12)
    p_pgd2 = ky_tbl.rows[0].cells[0].add_paragraph()
    p_pgd2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_pgd2 = p_pgd2.add_run("\n\n\n\n(Ký, đóng dấu)")
    r_pgd2.italic = True
    r_pgd2.font.name = "Times New Roman"
    r_pgd2.font.size = Pt(12)

    p_gd = ky_tbl.rows[0].cells[1].paragraphs[0]
    p_gd.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_gd = p_gd.add_run("GIÁM ĐỐC CHI NHÁNH")
    r_gd.bold = True
    r_gd.font.name = "Times New Roman"
    r_gd.font.size = Pt(12)
    p_gd2 = ky_tbl.rows[0].cells[1].add_paragraph()
    p_gd2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_gd2 = p_gd2.add_run("\n\n\n\n(Ký, đóng dấu)")
    r_gd2.italic = True
    r_gd2.font.name = "Times New Roman"
    r_gd2.font.size = Pt(12)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
