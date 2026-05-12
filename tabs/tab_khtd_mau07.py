"""Tab Mẫu 07 — Giao/Điều chỉnh Kế hoạch Tín dụng theo Ấp/Thôn.

Biểu số 07/NHCS-KH · Theo CV 7064
KV store:
  khtd_ap_{pgd_slug}_{xa_slug}          → dict  {"{ten_ap}|{ma_key}": float}
  khtd_ap_lich_su_{pgd_slug}_{xa_slug}  → list  [{lan, loai, ngay, username, data}]

Ghi chú:
  - Số QĐ và ngày tháng KHÔNG lưu — để TRỐNG cho UBND xã điền tay khi ký.
  - Số gốc lũy kế: lần 1 = baseline HSTD 31/12; lần 2+ = KH lần trước.
  - Data DB chỉ lưu "Chỉ tiêu KH" (float triệu đồng), không lưu giao +/−.
"""
import json
import os
import re
import unicodedata
from datetime import datetime
from io import BytesIO

import pandas as pd
import streamlit as st

import db
from config import (
    CHUONG_TRINH_KHTD,
    COT_MA_CHUONG_TRINH,
    COT_NGUON_VON,
    COT_TEN_THON,
    COT_TEN_XA,
    COT_TONG_DU_NO,
    DON_VI_CHI_NHANH,
    DS_PGD,
    NAM_HT,
    PGD_XA_MAP,
    tim_ten_xa_trong_hstd,
)
from tabs.tab_khtd import MAKEY_BY_MACT_NV
from utils import vn, hien_thi_dataframe_phan_trang

# ── Hằng số ───────────────────────────────────────────────────────────────────
_NAM_KH_DEFAULT = int(NAM_HT)

CT_TW = [(mk, ten) for mk, _, ten, nv, _ in CHUONG_TRINH_KHTD if nv == "TW"]
CT_DP = [(mk, ten) for mk, _, ten, nv, _ in CHUONG_TRINH_KHTD if nv == "DP"]
TEN_BY_MAKEY = {mk: ten for mk, _, ten, _, _ in CHUONG_TRINH_KHTD}

_BG_TW = "#e3f2fd"
_BG_DP = "#fff8e1"


# ══════════════════════════════════════════════════════════════════════════════
# HELPERS
# ══════════════════════════════════════════════════════════════════════════════
def _slug(text: str) -> str:
    s = unicodedata.normalize("NFD", text.strip().lower())
    s = "".join(c for c in s if unicodedata.category(c) != "Mn")
    return re.sub(r"[^a-z0-9]+", "_", s).strip("_")


def _chuan_hoa_ten(ten: str) -> str:
    """Loại bỏ dấu, viết thường, bỏ 'xã'/'phường'/'thị trấn'/'TT' để so sánh."""
    s = unicodedata.normalize('NFD', str(ten).lower())
    s = ''.join(c for c in s if unicodedata.category(c) != 'Mn')
    s = re.sub(r'\b(xa|phuong|phuong|thi tran|tt)\b', '', s)
    s = re.sub(r'\s+', ' ', s).strip()
    return s


def _kv_key(pgd: str, xa: str) -> str:
    return f"khtd_ap_{_slug(pgd)}_{_slug(xa)}"


def _kv_key_ls(pgd: str, xa: str) -> str:
    return f"khtd_ap_lich_su_{_slug(pgd)}_{_slug(xa)}"


def _doc_kv_dict(key: str) -> dict:
    try:
        with db.get_conn() as conn:
            row = conn.execute("SELECT value FROM kv_store WHERE key=?", (key,)).fetchone()
        return json.loads(row["value"]) if row else {}
    except Exception:
        return {}


def _doc_kv_list(key: str) -> list:
    try:
        with db.get_conn() as conn:
            row = conn.execute("SELECT value FROM kv_store WHERE key=?", (key,)).fetchone()
        val = json.loads(row["value"]) if row else []
        return val if isinstance(val, list) else []
    except Exception:
        return []


def _luu_kv(key: str, data, username: str) -> bool:
    try:
        with db.get_conn() as conn:
            conn.execute(
                "INSERT OR REPLACE INTO kv_store (key, value, updated_at, updated_by) VALUES (?,?,?,?)",
                (key, json.dumps(data, ensure_ascii=False), datetime.now().isoformat(), username),
            )
            conn.commit()
        return True
    except Exception:
        return False


# ══════════════════════════════════════════════════════════════════════════════
# BUSINESS LOGIC
# ══════════════════════════════════════════════════════════════════════════════
@st.cache_data(show_spinner=False)
def tinh_du_no_ap_baseline(_df_baseline: pd.DataFrame, ten_xa: str) -> dict:
    """Từ HSTD 31/12 tính dư nợ theo ấp × chương trình × nguồn vốn.

    Returns: {"{ten_thon}|{ma_key}": du_no_trieu}
    """
    if _df_baseline is None or _df_baseline.empty:
        return {}
    try:
        df = _df_baseline.copy()
        col_xa   = next((c for c in [COT_TEN_XA,   "Tên xã"]   if c in df.columns), None)
        col_thon = next((c for c in [COT_TEN_THON,  "Tên thôn"] if c in df.columns), None)
        col_mact = next((c for c in [COT_MA_CHUONG_TRINH, "Mã chương trình"] if c in df.columns), None)
        col_nv   = next((c for c in [COT_NGUON_VON, "Nguồn vốn"] if c in df.columns), None)
        col_dn   = next((c for c in [COT_TONG_DU_NO, "Tổng dư nợ"] if c in df.columns), None)

        missing = [name for name, val in [
            ("Tên xã", col_xa), ("Tên thôn", col_thon),
            ("Mã chương trình", col_mact), ("Nguồn vốn", col_nv), ("Tổng dư nợ", col_dn)
        ] if not val]
        if missing:
            return {"_err": f"Thiếu cột: {missing}"}

        df = df[df[col_xa] == ten_xa].copy()
        if df.empty:
            return {"_err": f"Xã '{ten_xa}' không có trong baseline"}

        df = df.dropna(subset=[col_thon, col_mact, col_nv])
        df[col_dn]   = pd.to_numeric(df[col_dn],   errors="coerce").fillna(0)
        df[col_mact] = pd.to_numeric(df[col_mact], errors="coerce").fillna(0).astype(int)
        df[col_nv]   = pd.to_numeric(df[col_nv],   errors="coerce").fillna(0).astype(int)

        result = {}
        for (ten_thon, ma_ct, nv_int), grp in df.groupby([col_thon, col_mact, col_nv]):
            du_no = grp[col_dn].sum()
            if du_no <= 0:
                continue
            ma_key = MAKEY_BY_MACT_NV.get((int(ma_ct), int(nv_int)))
            if ma_key:
                result[f"{str(ten_thon).strip()}|{ma_key}"] = round(du_no / 1_000_000, 1)
        return result
    except Exception as e:
        return {"_err": str(e)}


def lay_so_goc_cho_ap(ten_ap: str, ma_key: str, du_no_baseline: dict, lich_su: list) -> float:
    """Số gốc lũy kế: lần 1 = baseline 31/12, lần 2+ = KH lần trước."""
    composite = f"{ten_ap}|{ma_key}"
    if not lich_su:
        return du_no_baseline.get(composite, 0.0)
    lan_truoc_data = lich_su[-1].get("data", {})
    return lan_truoc_data.get(composite, du_no_baseline.get(composite, 0.0))


def _sync_khtd_xa(xa: str, data_dict: dict, username: str) -> None:
    """Sum theo xã × ma_key → ghi vào khtd_xa để các tab khác không bị break."""
    kv_xa = _doc_kv_dict("khtd_xa")
    tong: dict = {}
    for composite, gia_tri in data_dict.items():
        parts = composite.split("|", 1)
        if len(parts) == 2:
            mk = parts[1]
            k = f"{xa}|{mk}"
            tong[k] = round(tong.get(k, 0.0) + float(gia_tri), 1)
    kv_xa.update(tong)
    _luu_kv("khtd_xa", kv_xa, username)


# ══════════════════════════════════════════════════════════════════════════════
# WORD EXPORT — HELPERS
# ══════════════════════════════════════════════════════════════════════════════
def _remove_table_borders(table) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblBorders = OxmlElement("w:tblBorders")
    for name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:val"), "none")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _set_cell_bg(cell, color_hex: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  color_hex.lstrip("#"))
    cell._tc.get_or_add_tcPr().append(shd)


def _set_cell_border_top_double(cell) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    top = OxmlElement("w:top")
    top.set(qn("w:val"),   "double")
    top.set(qn("w:sz"),    "8")
    top.set(qn("w:color"), "000000")
    borders.append(top)
    tcPr.append(borders)


def _cell_write(cell, text: str, bold: bool = False, italic: bool = False,
                align: str = "left", size: int = 13, indent_cm: float = 0) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, Cm
    p = cell.paragraphs[0]
    p.alignment = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right":  WD_ALIGN_PARAGRAPH.RIGHT,
        "left":   WD_ALIGN_PARAGRAPH.LEFT,
    }.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    run = p.add_run(text)
    run.font.name  = "Times New Roman"
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic


def _fmt_delta(val: float) -> str:
    """Định dạng tăng/giảm: +50,0 / −10,5 / 0 (chuẩn VN: . nghìn, , thập phân)."""
    s = f"{abs(val):,.1f}".replace(",","X").replace(".",",").replace("X",".")
    if val > 0:
        return f"+{s}"
    if val < 0:
        return f"\u2212{s}"
    return "0"


def _fmt_money(val: float) -> str:
    """Định dạng tiền (chuẩn VN: . nghìn, , thập phân). VD: 2.500 hoặc 2.500,5"""
    if val == int(val):
        return f"{int(val):,}".replace(",",".")
    return f"{val:,.1f}".replace(",","X").replace(".",",").replace("X",".")


# ══════════════════════════════════════════════════════════════════════════════
# WORD EXPORT — MAIN
# ══════════════════════════════════════════════════════════════════════════════
def xuat_mau07_word(
    xa: str,
    nam: int,
    loai_van_ban: str,
    data_dict: dict,
    du_no_baseline: dict,
    lich_su: list,
    lan_chon,
) -> bytes:
    """Xuất file Word Mẫu 07 theo CV 7064. Số QĐ & ngày tháng để TRỐNG.

    Returns: bytes .docx
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, Cm
    except ImportError:
        raise ImportError("Cần cài: pip install python-docx")

    # Xác định data và lich_su_truoc cho hàm so_goc
    if lan_chon == "dang_nhap":
        xuat_data    = data_dict
        ls_truoc     = lich_su
    else:
        idx          = int(lan_chon) - 1
        xuat_data    = lich_su[idx]["data"] if idx < len(lich_su) else data_dict
        ls_truoc     = lich_su[:idx]

    # Thu thập danh sách ấp theo thứ tự
    ds_ap = sorted({composite.split("|")[0] for composite in xuat_data if "|" in composite})

    # Chuẩn bị dữ liệu bảng
    bang_rows = []
    tong_delta = 0.0
    tong_kh    = 0.0
    stt = 0

    for ten_ap in ds_ap:
        ct_rows = []
        for mk, ten_ct in CT_TW + CT_DP:
            composite = f"{ten_ap}|{mk}"
            kh_val  = float(xuat_data.get(composite, 0.0))
            so_goc  = lay_so_goc_cho_ap(ten_ap, mk, du_no_baseline, ls_truoc)
            delta   = round(kh_val - so_goc, 1)
            if kh_val == 0 and so_goc == 0:
                continue
            nv_label = "TW" if mk.endswith("_TW") else "ĐP"
            ct_rows.append({"ten_ct": f"{ten_ct} ({nv_label})", "delta": delta, "kh": kh_val})
        if ct_rows:
            bang_rows.append({"is_ap": True, "ten_ap": ten_ap})
            for r in ct_rows:
                stt += 1
                r["stt"] = stt
                bang_rows.append(r)
                tong_delta += r["delta"]
                tong_kh    += r["kh"]

    loai_upper = "GIAO" if loai_van_ban == "giao" else "ĐIỀU CHỈNH"
    loai_lower = "giao" if loai_van_ban == "giao" else "điều chỉnh"

    # ── Tạo Document ──────────────────────────────────────────────────────────
    doc = Document()

    # Page setup A4, lề chuẩn công văn
    sec = doc.sections[0]
    sec.page_height   = Cm(29.7)
    sec.page_width    = Cm(21.0)
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin   = Cm(3.0)
    sec.right_margin  = Cm(2.0)

    # Font mặc định
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(13)

    def _para(text: str = "", bold: bool = False, italic: bool = False,
              align: str = "left", size: int = 13,
              space_before: int = 0, space_after: int = 0,
              first_indent_cm: float = 0, line_spacing: float = 0) -> None:
        p = doc.add_paragraph()
        p.alignment = {
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right":  WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }.get(align, WD_ALIGN_PARAGRAPH.LEFT)
        if space_before:
            p.paragraph_format.space_before = Pt(space_before)
        if space_after:
            p.paragraph_format.space_after = Pt(space_after)
        if first_indent_cm:
            p.paragraph_format.first_line_indent = Cm(first_indent_cm)
        if line_spacing:
            p.paragraph_format.line_spacing = line_spacing
        if text:
            run = p.add_run(text)
            run.font.name  = "Times New Roman"
            run.font.size  = Pt(size)
            run.font.bold  = bold
            run.font.italic = italic
        return p

    # ══════════════════════════════════════════════════════════════════════════
    # TRANG 1 — HEADER VĂN BẢN
    # ══════════════════════════════════════════════════════════════════════════

    # Bảng header 2 cột (ẩn border)
    hdr_tbl = doc.add_table(rows=1, cols=2)
    hdr_tbl.autofit = False
    hdr_tbl.columns[0].width = Cm(9)
    hdr_tbl.columns[1].width = Cm(9)
    _remove_table_borders(hdr_tbl)

    # Cột trái
    lc = hdr_tbl.rows[0].cells[0]
    lc.paragraphs[0].clear()
    p_l = lc.paragraphs[0]
    r = p_l.add_run(f"ỦY BAN NHÂN DÂN\n{xa.upper()}")
    r.font.name = "Times New Roman"; r.font.size = Pt(13); r.font.bold = True
    p_l2 = lc.add_paragraph("Số:       /TB-UBND")
    p_l2.runs[0].font.name = "Times New Roman"; p_l2.runs[0].font.size = Pt(13)

    # Cột phải
    rc = hdr_tbl.rows[0].cells[1]
    rc.paragraphs[0].clear()
    p_r = rc.paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p_r.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    r1.font.name = "Times New Roman"; r1.font.size = Pt(13); r1.font.bold = True
    p_r2 = rc.add_paragraph("Độc lập - Tự do - Hạnh phúc")
    p_r2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_r2.runs[0].font.name = "Times New Roman"
    p_r2.runs[0].font.size = Pt(13); p_r2.runs[0].font.bold = True
    p_r3 = rc.add_paragraph("─────────────────")
    p_r3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_r3.runs[0].font.name = "Times New Roman"; p_r3.runs[0].font.size = Pt(11)
    p_r4 = rc.add_paragraph("........., ngày     tháng     năm 202....")
    p_r4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_r4.runs[0].font.name = "Times New Roman"
    p_r4.runs[0].font.size = Pt(13); p_r4.runs[0].font.italic = True

    doc.add_paragraph()

    # Tiêu đề
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after  = Pt(6)
    r_t = p_title.add_run(f"THÔNG BÁO {loai_upper} KẾ HOẠCH DƯ NỢ\nNĂM {nam}")
    r_t.font.name = "Times New Roman"; r_t.font.size = Pt(14); r_t.font.bold = True

    if loai_van_ban == "dieu_chinh" and lich_su:
        p_lan = doc.add_paragraph()
        p_lan.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_lan = p_lan.add_run(f"(LẦN {len(lich_su) + 1})")
        r_lan.font.name = "Times New Roman"; r_lan.font.size = Pt(14); r_lan.font.bold = True

    doc.add_paragraph()

    # Kính gửi
    p_kg = doc.add_paragraph()
    p_kg.paragraph_format.first_line_indent = Cm(1.5)
    r_kg = p_kg.add_run(f"Kính gửi: Ông (bà) Trưởng các thôn, ấp thuộc xã {xa}")
    r_kg.font.name = "Times New Roman"; r_kg.font.size = Pt(13); r_kg.font.italic = True

    # Các đoạn nội dung chuẩn
    body_paras = [
        (f"Căn cứ Quyết định số        /QĐ-NHCS ngày     /     /      của Trưởng Ban đại "
         f"diện HĐQT NHCSXH tỉnh về việc {loai_lower} kế hoạch tín dụng năm {nam} cho các xã;"),
        "Căn cứ tổng hợp nhu cầu vốn tín dụng chính sách tại các thôn, ấp.",
        (f"Ủy ban nhân dân xã {xa} thông báo {loai_lower} kế hoạch dư nợ năm {nam} "
         f"của các thôn, ấp theo danh mục đính kèm."),
        ("Đề nghị ông (bà) Trưởng thôn, ấp căn cứ Thông báo này để chỉ đạo các Tổ TK&VV "
         "tại thôn tổ chức bình xét cho vay và lập hồ sơ vay vốn theo quy định của NHCSXH. "
         "Việc bình xét cho vay phải đảm bảo công khai, dân chủ, đúng đối tượng và có sự "
         "tham gia giám sát của Trưởng thôn, ấp, đại diện lãnh đạo của tổ chức chính trị "
         "- xã hội nhận ủy thác."),
        (f"Trên đây là chỉ tiêu kế hoạch dư nợ năm {nam} được Ủy ban nhân dân xã {xa} phê "
         f"duyệt, đề nghị ông (bà) Trưởng thôn, ấp nghiêm túc thực hiện."),
    ]
    for txt in body_paras:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(1.5)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(txt)
        r.font.name = "Times New Roman"; r.font.size = Pt(13)

    # Ký tên
    doc.add_paragraph()
    doc.add_paragraph()
    p_ky = doc.add_paragraph()
    p_ky.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_ky.paragraph_format.space_before = Pt(24)
    for line, bold, italic in [
        ("TM. ỦY BAN NHÂN DÂN XÃ", True, False),
        ("CHỦ TỊCH",               True, False),
        ("(Ký tên, đóng dấu)",     False, True),
    ]:
        r_ky = p_ky.add_run(f"{line}\n")
        r_ky.font.name = "Times New Roman"
        r_ky.font.size = Pt(13)
        r_ky.font.bold   = bold
        r_ky.font.italic = italic

    # ══════════════════════════════════════════════════════════════════════════
    # TRANG 2 — BẢNG DỮ LIỆU
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_page_break()

    p_h1 = doc.add_paragraph()
    p_h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_h1 = p_h1.add_run(f"ỦY BAN NHÂN DÂN XÃ {xa.upper()}")
    r_h1.font.name = "Times New Roman"; r_h1.font.size = Pt(14); r_h1.font.bold = True

    p_h2 = doc.add_paragraph()
    p_h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_h2 = p_h2.add_run(
        f"DANH MỤC {loai_upper} CHỈ TIÊU KẾ HOẠCH DƯ NỢ\n"
        f"NGUỒN VỐN TRUNG ƯƠNG & ĐỊA PHƯƠNG NĂM {nam}"
    )
    r_h2.font.name = "Times New Roman"; r_h2.font.size = Pt(14); r_h2.font.bold = True

    p_km = doc.add_paragraph()
    p_km.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_km = p_km.add_run(
        f"(Kèm theo Thông báo số       /TB-UBND ngày     /     /      của\n"
        f"Chủ tịch UBND xã {xa})"
    )
    r_km.font.name = "Times New Roman"; r_km.font.size = Pt(12); r_km.font.italic = True

    p_dv = doc.add_paragraph()
    p_dv.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_dv.paragraph_format.space_after = Pt(6)
    r_dv = p_dv.add_run("Đơn vị: triệu đồng")
    r_dv.font.name = "Times New Roman"; r_dv.font.size = Pt(12); r_dv.font.italic = True

    # Bảng chính
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    tbl.autofit = False
    tbl.columns[0].width = Cm(1.5)
    tbl.columns[1].width = Cm(9.5)
    tbl.columns[2].width = Cm(3.5)
    tbl.columns[3].width = Cm(3.5)

    # Header row
    hdr = tbl.rows[0].cells
    _set_cell_bg(hdr[0], "E8E8E8"); _cell_write(hdr[0], "STT",                    bold=True, align="center", size=12)
    _set_cell_bg(hdr[1], "E8E8E8"); _cell_write(hdr[1], "Chỉ tiêu",               bold=True, align="center", size=12)
    _set_cell_bg(hdr[2], "E8E8E8"); _cell_write(hdr[2], "Giao tăng/giảm",         bold=True, align="center", size=12)
    _set_cell_bg(hdr[3], "E8E8E8"); _cell_write(hdr[3], f"Chỉ tiêu KH\nnăm {nam}", bold=True, align="center", size=12)

    # Data rows
    for item in bang_rows:
        row = tbl.add_row()
        cells = row.cells
        if item.get("is_ap"):
            merged = cells[0].merge(cells[3])
            _set_cell_bg(merged, "E3F2FD")
            _cell_write(merged, f"  {item['ten_ap']}", bold=True, align="left", size=12)
        else:
            _cell_write(cells[0], str(item["stt"]),           align="center", size=12)
            _cell_write(cells[1], item["ten_ct"],              align="left",   size=12, indent_cm=0.3)
            _cell_write(cells[2], _fmt_delta(item["delta"]),   align="right",  size=12)
            _cell_write(cells[3], _fmt_money(item["kh"]),      align="right",  size=12)

    # Dòng Tổng cộng (double border-top)
    row_tc = tbl.add_row()
    tc = row_tc.cells
    merged_tc = tc[0].merge(tc[1])
    _set_cell_bg(merged_tc, "F5F5F5")
    _cell_write(merged_tc, "TỔNG CỘNG", bold=True, align="center", size=12)
    _set_cell_bg(tc[2], "F5F5F5")
    _cell_write(tc[2], _fmt_delta(round(tong_delta, 1)), bold=True, align="right", size=12)
    _set_cell_bg(tc[3], "F5F5F5")
    _cell_write(tc[3], _fmt_money(round(tong_kh, 1)),    bold=True, align="right", size=12)
    for c in [merged_tc, tc[2], tc[3]]:
        _set_cell_border_top_double(c)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ══════════════════════════════════════════════════════════════════════════════
# UI — BẢNG NHẬP TỔNG HỢP (1 BẢNG DUY NHẤT THAY VÌ EXPANDER)
# ══════════════════════════════════════════════════════════════════════════════

def _lay_ds_ma_key_co_du_lieu(
    ten_xa: str,
    df_full: pd.DataFrame | None,
    du_no_baseline: dict,
    lich_su: list,
) -> set:
    """Lấy danh sách ma_key có dữ liệu từ baseline + HSTD hiện tại + lịch sử."""
    ds_ma_key: set = set()

    # Từ baseline
    if du_no_baseline:
        for key in du_no_baseline.keys():
            if "|" in key:
                ma_key = key.split("|")[1]
                ds_ma_key.add(ma_key)

    # Từ df_full (HSTD hiện tại)
    if df_full is not None and not df_full.empty:
        required_cols = ["Tên xã", "Tên thôn", "Mã chương trình", "Nguồn vốn"]
        if all(col in df_full.columns for col in required_cols):
            df_xa = df_full[df_full["Tên xã"] == ten_xa].copy()
            df_xa = df_xa.dropna(subset=["Mã chương trình", "Nguồn vốn"])
            for _, row in df_xa.iterrows():
                ma_ct = int(row["Mã chương trình"])
                nguon_von_code = int(row["Nguồn vốn"])
                ma_key = MAKEY_BY_MACT_NV.get((ma_ct, nguon_von_code))
                if ma_key:
                    ds_ma_key.add(ma_key)

    # Từ lịch sử
    for item in lich_su:
        for key in item.get("data", {}).keys():
            if "|" in key:
                ma_key = key.split("|")[1]
                ds_ma_key.add(ma_key)

    # Fallback: nếu không có CT nào, dùng list mặc định
    if not ds_ma_key:
        ds_ma_key = {"1_TW", "19_TW", "9_TW", "3_TW", "2_TW"}  # Các CT chính

    return ds_ma_key


def _build_table_data(
    ds_ap: list,
    ds_ma_key: set,
    du_no_baseline: dict,
    lich_su: list,
    data_hien_tai: dict,
    nam: int,
) -> pd.DataFrame:
    """Build DataFrame cho st.data_editor — mỗi dòng = 1 CT trong 1 ấp."""
    rows = []

    for ten_ap in ds_ap:
        # Dòng header ấp
        rows.append({
            "_type": "header_ap",
            "_key": "",
            "Ấp/Thôn": f"📍 {ten_ap}",
            "Chương trình": "",
            "Nguồn vốn": "",
            f"Số gốc (triệu)": None,
            "Giao tăng/giảm (triệu)": None,
            f"Chỉ tiêu KH {nam} (triệu)": None,
        })

        # Lấy các CT có dữ liệu cho ấp này
        ds_ct_ap = []
        for ma_key in sorted(ds_ma_key):
            key = f"{ten_ap}|{ma_key}"
            # Luôn hiển thị nếu có trong baseline hoặc lịch sử
            co_du_lieu = (
                key in du_no_baseline or
                (lich_su and any(key in h["data"] for h in lich_su))
            )
            # Hoặc: luôn hiển thị các CT chính
            if co_du_lieu or ma_key in {"1_TW", "19_TW", "9_TW", "3_TW"}:
                ds_ct_ap.append(ma_key)

        # Sort: TW trước, ĐP sau
        ds_ct_ap = sorted(ds_ct_ap, key=lambda x: (x.split("_")[1], int(x.split("_")[0])))

        # Dòng dữ liệu CT
        for ma_key in ds_ct_ap:
            key = f"{ten_ap}|{ma_key}"
            so_goc = lay_so_goc_cho_ap(ten_ap, ma_key, du_no_baseline, lich_su)

            # Lấy giá trị từ lịch sử hoặc data hiện tại
            if data_hien_tai and key in data_hien_tai:
                chi_tieu_kh = data_hien_tai[key]
            elif lich_su and key in lich_su[-1].get("data", {}):
                chi_tieu_kh = lich_su[-1]["data"][key]
            else:
                chi_tieu_kh = so_goc

            giao_tang_giam = round(chi_tieu_kh - so_goc, 1)

            # Lấy tên CT
            ma_ct_num = int(ma_key.split("_")[0])
            ten_ct = TEN_BY_MAKEY.get(ma_key, f"CT {ma_ct_num}")
            nguon_von = ma_key.split("_")[1]

            rows.append({
                "_type": "data",
                "_key": key,
                "Ấp/Thôn": "",
                "Chương trình": ten_ct,
                "Nguồn vốn": nguon_von,
                f"Số gốc (triệu)": so_goc,
                "Giao tăng/giảm (triệu)": giao_tang_giam,
                f"Chỉ tiêu KH {nam} (triệu)": chi_tieu_kh,
            })

    # Thêm dòng tổng (tính sau)
    rows.append({
        "_type": "total",
        "_key": "",
        "Ấp/Thôn": "TỔNG CỘNG",
        "Chương trình": "",
        "Nguồn vốn": "",
        f"Số gốc (triệu)": 0.0,
        "Giao tăng/giảm (triệu)": 0.0,
        f"Chỉ tiêu KH {nam} (triệu)": 0.0,
    })

    return pd.DataFrame(rows)


def _extract_data_from_edited_df(edited_df: pd.DataFrame, nam: int) -> dict:
    """Extract data_dict từ edited_df sau khi user nhập."""
    data_dict = {}
    col_kh = f"Chỉ tiêu KH {nam} (triệu)"

    for _, row in edited_df[edited_df["_type"] == "data"].iterrows():
        key = row.get("_key")
        chi_tieu = row.get(col_kh)
        if key and chi_tieu and chi_tieu != 0:
            data_dict[key] = round(float(chi_tieu), 1)

    return data_dict


def _update_total_row(df: pd.DataFrame, nam: int) -> pd.DataFrame:
    """Cập nhật dòng tổng cộng."""
    col_so_goc = f"Số gốc (triệu)"
    col_giao = "Giao tăng/giảm (triệu)"
    col_kh = f"Chỉ tiêu KH {nam} (triệu)"

    data_rows = df[df["_type"] == "data"]

    tong_so_goc = data_rows[col_so_goc].sum() if not data_rows.empty else 0
    tong_giao = data_rows[col_giao].sum() if not data_rows.empty else 0
    tong_kh = data_rows[col_kh].sum() if not data_rows.empty else 0

    idx_total = df[df["_type"] == "total"].index
    if not idx_total.empty:
        df.loc[idx_total[0], col_so_goc] = round(tong_so_goc, 1)
        df.loc[idx_total[0], col_giao] = round(tong_giao, 1)
        df.loc[idx_total[0], col_kh] = round(tong_kh, 1)

    return df


# ══════════════════════════════════════════════════════════════════════════════
# UI — LỊCH SỬ
# ══════════════════════════════════════════════════════════════════════════════
def _render_lich_su(lich_su: list) -> None:
    if not lich_su:
        st.info("Chưa có lịch sử giao/điều chỉnh nào.")
        return

    rows = [{
        "Lần":    item.get("lan", "?"),
        "Loại":   "Giao" if item.get("loai") == "giao" else "Điều chỉnh",
        "Ngày":   item.get("ngay", "")[:10],
        "Người":  item.get("username", ""),
        "Tổng KH (triệu)": vn(sum(float(v) for v in item.get("data", {}).values()), 1),
    } for item in lich_su]
    hien_thi_dataframe_phan_trang(
        pd.DataFrame(rows),
        key="mau07_lich_su",
    )

    lan_xem = st.selectbox(
        "Xem chi tiết lần",
        [f"Lần {i+1} — {item.get('ngay','')[:10]}" for i, item in enumerate(lich_su)],
        key="m07_ls_xem",
    )
    idx_xem  = int(lan_xem.split()[1]) - 1
    item_xem = lich_su[idx_xem]
    st.markdown(
        f"**Loại:** {'Giao' if item_xem.get('loai')=='giao' else 'Điều chỉnh'}"
        f" &nbsp;|&nbsp; **Người lưu:** {item_xem.get('username','')}"
    )
    df_ct = pd.DataFrame([
        {
            "Ấp/Thôn":              composite.split("|")[0] if "|" in composite else composite,
            "Chương trình":         TEN_BY_MAKEY.get(composite.split("|")[1] if "|" in composite else "", composite),
            "Chỉ tiêu KH (triệu)": v,
        }
        for composite, v in item_xem.get("data", {}).items()
    ])
    if not df_ct.empty:
        hien_thi_dataframe_phan_trang(df_ct, key="mau07_chi_tiet_lan")


# ══════════════════════════════════════════════════════════════════════════════
# ENTRY POINT
# ══════════════════════════════════════════════════════════════════════════════
def render(tab, **kwargs) -> None:
    """Entry point — render tab Mẫu 07."""
    role     = kwargs.get("role", "user")
    username = kwargs.get("username", "system")
    pgd_user = kwargs.get("pgd_user")

    _tab_ctx = tab if tab is not None else __import__('streamlit').container()
    with _tab_ctx:
        st.subheader("📋 Mẫu 07 — Giao/Điều chỉnh Chỉ tiêu KHTD theo Ấp/Thôn")
        st.caption("Biểu số 07/NHCS-KH · Theo CV 7064 · Số QĐ & ngày tháng để trống cho UBND xã điền khi ký")

        # ── ① PGD / Xã / Năm ──────────────────────────────────────────────────────
        col_pgd, col_xa, col_nam = st.columns([2, 2, 1])

        with col_pgd:
            if role == "user" and pgd_user:
                st.info(f"PGD: **{pgd_user}**")
                pgd_chon = pgd_user
            else:
                # Danh sách đơn vị = Hội sở + 21 PGD
                ds_don_vi = [DON_VI_CHI_NHANH] + DS_PGD
                pgd_chon = st.selectbox(
                    "Chọn Đơn vị",
                    ds_don_vi,
                    format_func=lambda x: f"🏦 {x}" if x == DON_VI_CHI_NHANH else f"🏢 {x}",
                    key="m07_pgd",
                )

        ds_xa_pgd = PGD_XA_MAP.get(pgd_chon, [])

        # ═══ MAPPING VÀ FUZZY MATCHING TÊNN XÃ ═══
        df_full = kwargs.get("df_full")
        ten_xa_mapping: dict = {}

        # Mapping từ XA_NAME_MAP + tự động bỏ prefix (config -> HSTD)
        for xa_pgd in ds_xa_pgd:
            mapped = tim_ten_xa_trong_hstd(xa_pgd)
            if mapped != xa_pgd:
                ten_xa_mapping[xa_pgd] = mapped

        # Fuzzy matching (chuẩn hóa) cho các xã chưa có mapping
        if df_full is not None and isinstance(df_full, pd.DataFrame) and not df_full.empty:
            if "Tên xã" in df_full.columns:
                ds_xa_co_data = df_full["Tên xã"].dropna().unique()
                for xa_pgd in ds_xa_pgd:
                    if xa_pgd in ten_xa_mapping:
                        continue
                    key_pgd = _chuan_hoa_ten(xa_pgd)
                    for xa_data in ds_xa_co_data:
                        if _chuan_hoa_ten(xa_data) == key_pgd:
                            ten_xa_mapping[xa_pgd] = xa_data
                            break

        with col_xa:
            if not ds_xa_pgd:
                st.warning(f"⚠️ Đơn vị **{pgd_chon}** chưa có danh sách xã trong cấu hình PGD_XA_MAP.")
                return

            if st.query_params.get("debug") == "1":
                with st.expander("🔍 Debug: Mapping xã", expanded=False):
                    st.write("**Từ PGD_XA_MAP:**", ds_xa_pgd)
                    if df_full is not None and "Tên xã" in df_full.columns:
                        st.write("**Trong HSTD:**", sorted(df_full["Tên xã"].dropna().unique().tolist())[:20])
                    st.write("**Mapping:**", ten_xa_mapping if ten_xa_mapping else "(Tất cả tên khớp trực tiếp)")

            xa_chon_raw = st.selectbox("Chọn Xã/Phường", ds_xa_pgd, key="m07_xa")

            # xa_chon_raw: tên gốc từ config (dùng cho slug/key lưu)
            # xa_chon:     tên trong HSTD (dùng cho filter dữ liệu)
            xa_chon = ten_xa_mapping.get(xa_chon_raw, xa_chon_raw)
            if st.query_params.get("debug") == "1" and xa_chon != xa_chon_raw:
                st.caption(f"📌 Mapping: '{xa_chon_raw}' → '{xa_chon}' (trong HSTD)")
            elif df_full is not None and "Tên xã" in df_full.columns:
                ds_xa_co_data = df_full["Tên xã"].dropna().unique()
                if xa_chon not in ds_xa_co_data:
                    st.warning(f"⚠️ Xã '{xa_chon}' không tìm thấy trong HSTD. Kiểm tra XA_NAME_MAP hoặc tên xã.")

        with col_nam:
            nam_kh = st.number_input(
                "Năm KH", min_value=2024, max_value=2030,
                value=_NAM_KH_DEFAULT, step=1, key="m07_nam",
            )

        pgd_slug_key = _slug(pgd_chon)
        xa_slug_key  = _slug(xa_chon_raw)  # Dùng xa_chon_raw (tên gốc) cho slug để ổn định key lưu
        nam_baseline = nam_kh - 1

        # ── ② Loại văn bản ────────────────────────────────────────────────────
        loai_chon = st.radio(
            "Loại văn bản",
            ["🆕 Giao lần đầu", "✏️ Điều chỉnh"],
            horizontal=True,
            key="m07_loai",
            label_visibility="collapsed",
        )
        loai_van_ban = "giao" if "Giao" in loai_chon else "dieu_chinh"

        # ── ③ Đọc dữ liệu ─────────────────────────────────────────────────────
        kv_key_ht = _kv_key(pgd_chon, xa_chon_raw)
        kv_key_ls = _kv_key_ls(pgd_chon, xa_chon_raw)
        data_hien_tai: dict = _doc_kv_dict(kv_key_ht)
        lich_su: list       = _doc_kv_list(kv_key_ls)

        # Đọc baseline HSTD 31/12
        df_baseline = None
        co_baseline = False
        try:
            from config import (
                baseline_pgd_path, baseline_path,
                danh_sach_nam_baseline_pgd,
            )
            from data.hstd import doc_baseline_merged, doc_baseline

            ds_nam_bl = danh_sach_nam_baseline_pgd()
            if nam_baseline in ds_nam_bl:
                fp = baseline_pgd_path(pgd_chon, nam_baseline)
                _ts = os.path.getmtime(fp) if os.path.exists(fp) else 0
                df_baseline = doc_baseline_merged(nam_baseline, _ts=_ts)
                co_baseline = df_baseline is not None and not df_baseline.empty

            if not co_baseline:
                fp2  = baseline_path(nam_baseline)
                _ts2 = os.path.getmtime(fp2) if os.path.exists(fp2) else 0
                df_baseline = doc_baseline(nam_baseline, _ts=_ts2)
                co_baseline = df_baseline is not None and not df_baseline.empty
        except Exception:
            co_baseline = False

        if co_baseline:
            _raw = tinh_du_no_ap_baseline(df_baseline, xa_chon)
            # Kiểm tra lỗi trả về từ hàm
            if "_err" in _raw:
                _err_msg = _raw["_err"]
                st.warning(f"⚠️ Baseline 31/12/{nam_baseline} có nhưng không đọc được dư nợ: {_err_msg}")
                # Vẫn hiển thị debug để chẩn đoán
                with st.expander("🔍 Debug baseline", expanded=True):
                    st.write(f"**Cột trong file:** {list(df_baseline.columns[:15])}")
                    if COT_TEN_XA in df_baseline.columns or "Tên xã" in df_baseline.columns:
                        col_xa_debug = COT_TEN_XA if COT_TEN_XA in df_baseline.columns else "Tên xã"
                        ds_xa_debug = sorted(df_baseline[col_xa_debug].dropna().unique().tolist())
                        st.write(f"**Danh sách xã trong baseline ({len(ds_xa_debug)} xã):** {ds_xa_debug[:10]}")
                        st.caption(f"Đang tìm xã: **'{xa_chon}'** — kiểm tra tên có khớp không")
                du_no_baseline = {}
            else:
                so_ap = len({k.split("|")[0] for k in _raw})
                so_ct = len(_raw)
                st.success(
                    f"✅ Baseline 31/12/{nam_baseline}: {so_ap} ấp/thôn, {so_ct} chỉ tiêu "
                    f"— dùng làm số gốc lần đầu"
                )
                du_no_baseline = _raw
        else:
            st.warning(
                f"⚠️ Chưa có dữ liệu HSTD 31/12/{nam_baseline} — "
                "số gốc lần đầu = 0. Vui lòng upload baseline trước."
            )
            du_no_baseline = {}

        if lich_su:
            st.info(
                f"📜 Đã có **{len(lich_su)}** lần giao/điều chỉnh. "
                f"Số gốc hiện tại = Chỉ tiêu KH lần {len(lich_su)}."
            )

        st.divider()

        # ── ④⑤ BẢNG NHẬP DUY NHẤT (THAY CHO EXPANDER TỪNG ẤP) ────────────────

        # ═══ LẤY DANH SÁCH ẤP TỰ ĐỘNG ═══
        # Nguồn 1: Từ baseline (ưu tiên)
        ds_ap_from_baseline = set()
        if du_no_baseline:
            ds_ap_from_baseline = set(key.split("|")[0] for key in du_no_baseline.keys() if "|" in key)

        # Nguồn 2: Từ df_baseline (HSTD 31/12)
        ds_ap_from_baseline_df = set()
        if co_baseline and df_baseline is not None and not df_baseline.empty:
            col_xa_bl = next((c for c in [COT_TEN_XA, "Tên xã"] if c in df_baseline.columns), None)
            col_thon_bl = next((c for c in [COT_TEN_THON, "Tên thôn"] if c in df_baseline.columns), None)
            if col_xa_bl and col_thon_bl:
                df_xa_bl = df_baseline[df_baseline[col_xa_bl] == xa_chon]
                ap_list = df_xa_bl[col_thon_bl].dropna().astype(str).str.strip()
                ds_ap_from_baseline_df = set(ap_list[ap_list != ""].unique())

        # Nguồn 3: Từ HSTD hiện tại (df_full) - đã lấy từ kwargs ở trên
        ds_ap_from_hstd_current = set()
        if df_full is not None and isinstance(df_full, pd.DataFrame) and not df_full.empty:
            if "Tên xã" in df_full.columns and "Tên thôn" in df_full.columns:
                df_xa_full = df_full[df_full["Tên xã"] == xa_chon]
                ap_list = df_xa_full["Tên thôn"].dropna().astype(str).str.strip()
                ds_ap_from_hstd_current = set(ap_list[ap_list != ""].unique())

        # Nguồn 4: Từ lịch sử
        ds_ap_from_history = set()
        for item in lich_su:
            for key in item.get("data", {}).keys():
                if "|" in key:
                    ds_ap_from_history.add(key.split("|")[0])
        for key in data_hien_tai.keys():
            if "|" in key:
                ds_ap_from_history.add(key.split("|")[0])

        # Gộp tất cả nguồn
        ds_ap = sorted(ds_ap_from_baseline | ds_ap_from_baseline_df | ds_ap_from_hstd_current | ds_ap_from_history)

        # Debug info
        if ds_ap:
            with st.expander("🔍 Debug: Nguồn dữ liệu ấp", expanded=False):
                st.caption(f"✅ Tìm thấy **{len(ds_ap)}** ấp/thôn cho xã **{xa_chon}**")
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.markdown("**Từ baseline:**")
                    st.write(sorted(ds_ap_from_baseline) if ds_ap_from_baseline else "(Chưa có)")
                with col2:
                    st.markdown("**Từ baseline df:**")
                    st.write(sorted(ds_ap_from_baseline_df) if ds_ap_from_baseline_df else "(Chưa có)")
                with col3:
                    st.markdown("**Từ HSTD hiện tại:**")
                    st.write(sorted(ds_ap_from_hstd_current) if ds_ap_from_hstd_current else "(Chưa có)")
                with col4:
                    st.markdown("**Từ lịch sử:**")
                    st.write(sorted(ds_ap_from_history) if ds_ap_from_history else "(Chưa có)")
        else:
            st.error(
                f"❌ Không tìm thấy dữ liệu ấp/thôn cho xã **{xa_chon}**\n\n"
                "**Nguyên nhân có thể:**\n"
                "- Tên xã không khớp (kiểm tra dấu cách, dấu thanh)\n"
                "- Chưa có dữ liệu HSTD cho xã này\n"
                "- File HSTD chưa được upload\n\n"
                "**Giải pháp:** Upload file HSTD có chứa dữ liệu xã này"
            )
            st.stop()

        # ═══ LẤY DANH SÁCH CHƯƠNG TRÌNH CÓ DỮ LIỆU ═══
        ds_ma_key = _lay_ds_ma_key_co_du_lieu(xa_chon, df_full, du_no_baseline, lich_su)

        # ═══ BUILD VÀ HIỂN THỊ BẢNG ═══
        st.markdown("### 📊 Nhập chỉ tiêu theo ấp/thôn")

        lan = len(lich_su) + 1 if lich_su else 1
        st.caption(
            f"Xã: **{xa_chon}** · Năm: **{nam_kh}** · "
            f"Số gốc từ: **{'Dư nợ 31/12/' + str(nam_kh-1) if lan == 1 else 'KH lũy kế lần ' + str(lan-1)}**"
        )

        # Build DataFrame
        df_edit = _build_table_data(ds_ap, ds_ma_key, du_no_baseline, lich_su, data_hien_tai, nam_kh)

        # st.data_editor
        col_so_goc = f"Số gốc (triệu)"
        col_giao = "Giao tăng/giảm (triệu)"
        col_kh = f"Chỉ tiêu KH {nam_kh} (triệu)"

        edited_df = st.data_editor(
            df_edit,
            disabled=["Ấp/Thôn", "Chương trình", "Nguồn vốn", col_so_goc, col_kh],
            column_config={
                "_type": None,  # Ẩn cột internal
                "_key": None,
                "Ấp/Thôn": st.column_config.TextColumn("Ấp/Thôn", width="medium"),
                "Chương trình": st.column_config.TextColumn("Chương trình", width="large"),
                "Nguồn vốn": st.column_config.TextColumn("NV", width="small"),
                col_so_goc: st.column_config.NumberColumn("Số gốc\n(triệu)", format="%.1f", width="small"),
                col_giao: st.column_config.NumberColumn("Giao tăng/giảm\n(triệu)", format="%.1f", width="medium", step=0.1),
                col_kh: st.column_config.NumberColumn(f"Chỉ tiêu KH {nam_kh}\n(triệu)", format="%.1f", width="medium"),
            },
            hide_index=True,
            use_container_width=True,
            height=600,
            key=f"m07_table_edit_{pgd_slug_key}_{xa_slug_key}_{nam_kh}",
        )

        # Auto-calculate cột "Chỉ tiêu KH" khi user nhập "Giao tăng/giảm"
        if edited_df is not None and not edited_df.empty:
            for idx, row in edited_df.iterrows():
                if row.get("_type") == "data":
                    so_goc = row.get(col_so_goc) or 0
                    giao = row.get(col_giao) or 0
                    edited_df.at[idx, col_kh] = round(so_goc + giao, 1)

            # Update tổng
            edited_df = _update_total_row(edited_df, nam_kh)

            # Hiển thị dòng tổng nổi bật
            total_row = edited_df[edited_df["_type"] == "total"]
            if not total_row.empty:
                tong_so_goc = total_row.iloc[0][col_so_goc]
                tong_giao = total_row.iloc[0][col_giao]
                tong_kh = total_row.iloc[0][col_kh]
                _giao_str = (f"+{vn(tong_giao, 1)}" if tong_giao > 0 else vn(tong_giao, 1))
                st.markdown(
                    f"<div style='background:#e3f2fd;padding:10px 14px;border-radius:6px;"
                    f"font-weight:700;font-size:14px;margin-top:8px'>"
                    f"📊 TỔNG CỘNG TOÀN XÃ: Số gốc {vn(tong_so_goc, 1)} · "
                    f"Giao tăng/giảm {_giao_str} · "
                    f"Chỉ tiêu KH {vn(tong_kh, 1)} triệu đồng</div>",
                    unsafe_allow_html=True,
                )

        # Extract data để lưu
        data_nhap = _extract_data_from_edited_df(edited_df, nam_kh) if edited_df is not None else {}

        st.divider()

        # ── ⑥ Hành động: Lưu + Xuất Word ─────────────────────────────────────
        col_luu, col_word_label, col_word_sel, col_word_btn = st.columns([1, 1, 2, 1])

        with col_luu:
            btn_luu = st.button("💾 Lưu", type="primary", key="m07_btn_luu")

        with col_word_label:
            st.markdown("<div style='padding-top:8px'>📄 Xuất Word lần:</div>",
                        unsafe_allow_html=True)

        with col_word_sel:
            opts_xuat = ["Đang nhập"] + [f"Lần {i+1}" for i in range(len(lich_su))]
            lan_xuat = st.selectbox(
                "Xuất lần", opts_xuat,
                key="m07_lan_xuat",
                label_visibility="collapsed",
            )

        with col_word_btn:
            btn_word = st.button("Xuất ▶", key="m07_btn_word")

        # ── Xử lý Lưu ─────────────────────────────────────────────────────────
        if btn_luu:
            co_du_lieu = any(float(v) != 0 for v in data_nhap.values()) if data_nhap else False
            if not co_du_lieu:
                st.error("❌ Chưa có chỉ tiêu nào khác 0. Vui lòng nhập dữ liệu trước khi lưu.")
            else:
                # Cảnh báo nếu tổng thay đổi > 20% so với lần trước
                if lich_su:
                    tong_cu = sum(float(v) for v in lich_su[-1].get("data", {}).values())
                    tong_moi = sum(float(v) for v in data_nhap.values())
                    if tong_cu > 0:
                        pct_thay_doi = abs(tong_moi - tong_cu) / tong_cu * 100
                        if pct_thay_doi > 20:
                            st.warning(
                                f"⚠️ Tổng xã thay đổi **{pct_thay_doi:.1f}%** so với lần trước "
                                f"({vn(tong_cu,1)} → {vn(tong_moi,1)} triệu). Nhấn **Lưu** lần nữa để xác nhận."
                            )
                            if "m07_confirmed" not in st.session_state:
                                st.session_state["m07_confirmed"] = False
                                st.stop()

                lan_moi = len(lich_su) + 1
                entry = {
                    "lan":      lan_moi,
                    "loai":     loai_van_ban,
                    "ngay":     datetime.now().isoformat(),
                    "username": username,
                    "data":     data_nhap,
                }
                ok1 = _luu_kv(kv_key_ht, data_nhap, username)
                ok2 = _luu_kv(kv_key_ls, lich_su + [entry], username)
                _sync_khtd_xa(xa_chon_raw, data_nhap, username)
                db.ghi_audit(
                    username,
                    "luu_khtd_mau07",
                    f"{loai_van_ban} lần {lan_moi} — {xa_chon_raw} ({pgd_chon})",
                )
                if ok1 and ok2:
                    st.success(f"✅ Đã lưu lần {lan_moi} thành công!")
                    st.cache_data.clear()
                    st.rerun()
                else:
                    st.error("❌ Lỗi khi lưu dữ liệu. Vui lòng thử lại.")

        # ── Xử lý Xuất Word ───────────────────────────────────────────────────
        if btn_word:
            xuat_data_src = data_nhap if lan_xuat == "Đang nhập" else (
                lich_su[int(lan_xuat.split()[1]) - 1]["data"] if lich_su else data_nhap
            )
            if not xuat_data_src:
                st.error("❌ Không có dữ liệu để xuất.")
            else:
                lan_chon_val = "dang_nhap" if lan_xuat == "Đang nhập" else int(lan_xuat.split()[1])
                try:
                    doc_bytes = xuat_mau07_word(
                        xa=xa_chon_raw,
                        nam=nam_kh,
                        loai_van_ban=loai_van_ban,
                        data_dict=xuat_data_src,
                        du_no_baseline=du_no_baseline,
                        lich_su=lich_su,
                        lan_chon=lan_chon_val,
                    )
                    ten_file = (
                        f"Mau07_{_slug(xa_chon_raw)}_{loai_van_ban}_"
                        f"{datetime.now().strftime('%d%m%Y')}.docx"
                    )
                    st.download_button(
                        label="⬇️ Tải về Word (.docx)",
                        data=doc_bytes,
                        file_name=ten_file,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="m07_dl_word",
                    )
                    st.success("✅ Đã tạo file Word — nhấn nút trên để tải về!")
                except Exception as e:
                    st.error(f"❌ Lỗi xuất Word: {e}")
                    st.exception(e)

        # ── ⑦ Lịch sử ────────────────────────────────────────────────────────
        st.divider()
        with st.expander("📜 Lịch sử giao/điều chỉnh", expanded=False):
            _render_lich_su(lich_su)
