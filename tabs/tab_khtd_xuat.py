"""Xuất báo cáo / export cho tab Kế hoạch Tín dụng."""


from __future__ import annotations

from logger import get_logger
logger = get_logger(__name__)

import os
from datetime import datetime
from io import BytesIO

import pandas as pd
import streamlit as st
from openpyxl.styles import Font, PatternFill

from state_manager import SCMStateManager
from config import TEN_CHINH_THUC_CT, CHUONG_TRINH_KHTD, DS_PGD, COT_TEN_PGD, CACHE_HSTD, CACHE_GQVL, XA_TO_PGD, PGD_XA_MAP, DON_VI_CHI_NHANH, COT_TONG_DU_NO, COT_DU_NO_TH, COT_MA_CHUONG_TRINH, COT_NGUON_VON
from data.core import ts_file
from utils import lay_ngay_so_lieu


@st.cache_resource(show_spinner=False)
def _doc_hstd_cached(ts: float = 0) -> pd.DataFrame:
    _ = ts
    try:
        return pd.read_parquet(CACHE_HSTD)
    except Exception as e:
        logger.error("Lỗi trong khối except: %s", e, exc_info=True)
        return pd.DataFrame()


@st.cache_resource(show_spinner=False)
def _doc_gqvl_cached(ts: float = 0) -> pd.DataFrame:
    _ = ts
    try:
        return pd.read_parquet(CACHE_GQVL)
    except Exception as e:
        logger.error("Lỗi trong khối except: %s", e, exc_info=True)
        return pd.DataFrame()
from pdf_service import xuat_pdf
from utils import hien_thi_dataframe_phan_trang, xuat_excel, ten_file_xuat

from tabs.tab_khtd import (

    KV_KEY_CN,
    KV_KEY_XA,
    _doc_kv,
    _dong_bo_gqvl_tong_keys,
    _fvn,
    _fmt_vn,
    _fmt_vn_signed,
    _iter_khtd_cn_group_rows,
    _nv_int_tu_ma_key,
    _quet_ct_co_du_no,
    _tinh_thuc_hien_khtd_cn,
    _tinh_thuc_hien_theo_ct,
)


def _hien_thi_bang_cn_readonly(
    kh_cn: dict,
    th_cn: dict[str, float] | None = None,
    ds_ct_loc: list[str] | None = None,
    df_loc: "pd.DataFrame | None" = None,
    th_gqvl: dict[str, float] | None = None,
    username: str = "",
) -> None:
    """Bảng tóm tắt KHTD Chi nhánh — HTML thuần: STT | Chỉ tiêu | KH | Thực hiện | Còn phải thực hiện | TL%."""
    kh_d = _dong_bo_gqvl_tong_keys(dict(kh_cn or {}))
    th_d = _dong_bo_gqvl_tong_keys(dict(th_cn or {}))

    if df_loc is not None and (not th_d or th_gqvl is None):
        df_gqvl = _doc_gqvl_cached(ts_file(CACHE_GQVL))
        th_d_moi, th_gqvl = _tinh_thuc_hien_khtd_cn(df_loc, df_gqvl)
        if not th_d:
            th_d = th_d_moi

    if th_gqvl:
        th_d.update(_dong_bo_gqvl_tong_keys(dict(th_gqvl)))

    if not kh_d and not th_d:
        st.info("Chưa có dữ liệu.")
        return

    BD = "#d1d5db"
    H_BG = "#003D7A"
    NHOM_BG = "#e8f0f8"
    TONG_BG = "#E8F4FD"
    RED = "#DC2626"
    AMBER = "#D97706"
    GREEN = "#16A34A"

    tong_kh = 0.0
    tong_th = 0.0
    tong_kh_i = 0.0
    tong_th_i = 0.0
    tong_kh_ii = 0.0
    tong_th_ii = 0.0
    so_ct_co_kh = 0
    tat_ca_rows = _iter_khtd_cn_group_rows()
    tong_ct = len({int(row.get("ma_ct")) for _g, rows in tat_ca_rows for row in rows if row.get("ma_ct") is not None})
    stt_no = 0

    html_rows: list[str] = []
    nhom_hien = ""

    def _tl_color(tl: float | None) -> str:
        if tl is None:
            return "#9ca3af"
        if tl >= 100:
            return GREEN
        if tl >= 95:
            return AMBER
        return RED

    def _td(v: str, align: str = "right", color: str = "", bg: str = "", fw: str = "") -> str:
        s = f'text-align:{align};padding:5px 8px;border:1px solid {BD};font-size:0.82rem;white-space:nowrap'
        if color:
            s += f";color:{color}"
        if bg:
            s += f";background:{bg}"
        if fw:
            s += f";font-weight:{fw}"
        return f"<td style='{s}'>{v}</td>"

    def _add_group_hdr(label: str) -> None:
        nonlocal nhom_hien
        nhom_hien = label
        tds = (
            _td("", "center", "", NHOM_BG, "bold")
            + _td(label, "left", "#fff", H_BG, "bold")
            + _td("", "right", "", NHOM_BG)
            + _td("", "right", "", NHOM_BG)
            + _td("", "right", "", NHOM_BG)
            + _td("", "right", "", NHOM_BG)
        )
        html_rows.append(f"<tr>{tds}</tr>")

    def _add_row(
        ten: str, kh_vnd: float, th_vnd: float, thu_hoi_nq11_vnd: float = 0.0
    ) -> None:
        nonlocal tong_kh, tong_th, tong_kh_i, tong_th_i, tong_kh_ii, tong_th_ii, stt_no
        kh_v = kh_vnd / 1e6
        th_v = th_vnd / 1e6
        con_phai_th_v = (kh_vnd - th_vnd - thu_hoi_nq11_vnd) / 1e6
        tl = th_v / kh_v * 100 if kh_v > 0 else None

        if kh_v > 0 or th_v > 0:
            tong_kh += kh_vnd
            tong_th += th_vnd
            if nhom_hien.startswith("I."):
                tong_kh_i += kh_vnd
                tong_th_i += th_vnd
            elif nhom_hien.startswith("II."):
                tong_kh_ii += kh_vnd
                tong_th_ii += th_vnd

        stt_no += 1

        kh_str = _fvn(kh_v, 0) if kh_v > 0 else "—"
        th_str = _fvn(th_v, 0) if th_v > 0 else "—"
        con_phai_th_str = _fvn(con_phai_th_v, 0)
        tl_str = f"{_fvn(tl, 1)}%" if tl is not None else "—"
        tl_c = _tl_color(tl)

        tds = (
            _td(str(stt_no), "center", "", "", "") +
            _td(ten, "left", "", "", "") +
            _td(kh_str, "right", "", "", "") +
            _td(th_str, "right", "", "", "") +
            _td(con_phai_th_str, "right", "", "", "") +
            _td(tl_str, "right", tl_c, "", "")
        )
        html_rows.append(f"<tr>{tds}</tr>")

    def _add_section_total(label: str, kh_vnd: float, th_vnd: float) -> None:
        kh_v = kh_vnd / 1e6
        th_v = th_vnd / 1e6
        con_phai_th_v = (kh_vnd - th_vnd) / 1e6
        tl = th_v / kh_v * 100 if kh_v > 0 else None
        tds = (
            _td("", "center", "#1f2937", TONG_BG, "bold") +
            _td(label, "left", "#1f2937", TONG_BG, "bold") +
            _td(_fvn(kh_v, 0) if kh_v > 0 else "—", "right", "#1f2937", TONG_BG, "bold") +
            _td(_fvn(th_v, 0) if th_v > 0 else "—", "right", "#1f2937", TONG_BG, "bold") +
            _td(_fvn(con_phai_th_v, 0), "right", "#1f2937", TONG_BG, "bold") +
            _td(f"{_fvn(tl, 1)}%" if tl is not None else "—", "right", _tl_color(tl), TONG_BG, "bold")
        )
        html_rows.append(f"<tr>{tds}</tr>")

    loc_set = set(ds_ct_loc or [])

    def _duoc_loc(key_name: str | None) -> bool:
        if not loc_set:
            return True
        if not key_name:
            return False
        return key_name in loc_set

    # Pre-calculate NQ11 dư nợ theo (ma_ct, nv_int) từ df_loc
    nq11_by_mact_nv: dict[tuple[int, int], float] = {}
    if df_loc is not None and not df_loc.empty and "__is_nq11" in df_loc.columns:
        _col_dn = (
            COT_TONG_DU_NO if COT_TONG_DU_NO in df_loc.columns
            else COT_DU_NO_TH if COT_DU_NO_TH in df_loc.columns
            else None
        )
        if _col_dn and COT_MA_CHUONG_TRINH in df_loc.columns and COT_NGUON_VON in df_loc.columns:
            _mask = df_loc["__is_nq11"].fillna(False).astype(bool)
            _df_nq = df_loc[_mask]
            if not _df_nq.empty:
                _ma_s = pd.to_numeric(_df_nq[COT_MA_CHUONG_TRINH], errors="coerce").fillna(0).astype(int)
                _nv_s = pd.to_numeric(_df_nq[COT_NGUON_VON], errors="coerce").fillna(0).astype(int)
                _dn_s = pd.to_numeric(_df_nq[_col_dn], errors="coerce").fillna(0)
                _tmp = pd.DataFrame({"ma_ct": _ma_s, "nv": _nv_s, "dn": _dn_s})
                _tmp = _tmp[(_tmp["ma_ct"] > 0) & _tmp["nv"].isin([1, 2])]
                for (_mc, _nv), _v in _tmp.groupby(["ma_ct", "nv"])["dn"].sum().items():
                    nq11_by_mact_nv[(_mc, _nv)] = float(_v)

    # Pre-calculate NQ11 thu hồi trong năm theo (ma_ct, nv_int) từ df_loc
    nq11_thuhoi_nam_by_mact_nv: dict[tuple[int, int], float] = {}
    if df_loc is not None and not df_loc.empty and "__is_nq11" in df_loc.columns:
        _cols_th_nam = ["Thu nợ TH Năm", "Thu nợ QH Năm", "Thu nợ Khoanh Năm"]
        _cols_ok = [c for c in _cols_th_nam if c in df_loc.columns]
        if _cols_ok and COT_MA_CHUONG_TRINH in df_loc.columns and COT_NGUON_VON in df_loc.columns:
            _mask2 = df_loc["__is_nq11"].fillna(False).astype(bool)
            _df_nq2 = df_loc[_mask2]
            if not _df_nq2.empty:
                _ma_s2 = pd.to_numeric(_df_nq2[COT_MA_CHUONG_TRINH], errors="coerce").fillna(0).astype(int)
                _nv_s2 = pd.to_numeric(_df_nq2[COT_NGUON_VON], errors="coerce").fillna(0).astype(int)
                _th_s2 = sum(
                    pd.to_numeric(_df_nq2[c], errors="coerce").fillna(0) for c in _cols_ok
                )
                _tmp2 = pd.DataFrame({"ma_ct": _ma_s2, "nv": _nv_s2, "th": _th_s2})
                _tmp2 = _tmp2[(_tmp2["ma_ct"] > 0) & _tmp2["nv"].isin([1, 2]) & (_tmp2["th"] > 0)]
                for (_mc2, _nv2), _v2 in _tmp2.groupby(["ma_ct", "nv"])["th"].sum().items():
                    nq11_thuhoi_nam_by_mact_nv[(_mc2, _nv2)] = float(_v2)

    NQ11_BG = "#EFF6FF"

    def _add_nq11_subrow(
        dn_vnd: float,
        thu_hoi_nam_vnd: float = 0.0,
        da_tru_o_dong_tren: bool = False,
    ) -> None:
        dn_trieu = dn_vnd / 1e6
        if da_tru_o_dong_tren:
            # Với hàng "normal": thu hồi đã trừ vào Còn phải TH của dòng trên
            label = "&nbsp;&nbsp;&nbsp;↳ Trong đó: Dư nợ ch.trình NQ11 (thu hồi NQ11 đã trừ vào Còn phải TH)"
            con_str = "—"
            con_color = "#9ca3af"
        elif thu_hoi_nam_vnd > 0:
            # Với GQVL sub-rows: hiển thị điều chỉnh Còn phải TH tại đây
            thu_hoi_trieu = thu_hoi_nam_vnd / 1e6
            label = (
                f"&nbsp;&nbsp;&nbsp;↳ Trong đó: Dư nợ NQ11 = {_fvn(dn_trieu, 0)} | "
                f"Thu hồi NQ11 trong năm (điều chỉnh Còn phải TH): −{_fvn(thu_hoi_trieu, 0)}"
            )
            con_str = f"−{_fvn(thu_hoi_trieu, 0)}"
            con_color = "#dc2626"
        else:
            label = "&nbsp;&nbsp;&nbsp;↳ Trong đó: Dư nợ ch.trình các món vay Nghị quyết 11"
            con_str = "—"
            con_color = "#9ca3af"

        dn_td = "—" if thu_hoi_nam_vnd > 0 else _fvn(dn_trieu, 0)
        tds = (
            _td("", "center", "", NQ11_BG, "") +
            _td(label, "left", "#1D4ED8", NQ11_BG, "400") +
            _td("—", "right", "#9ca3af", NQ11_BG, "") +
            _td(dn_td, "right", "#1D4ED8", NQ11_BG, "") +
            _td(con_str, "right", con_color, NQ11_BG, "bold" if con_color != "#9ca3af" else "") +
            _td("—", "right", "#9ca3af", NQ11_BG, "")
        )
        html_rows.append(f"<tr>{tds}</tr>")

    for tieu_de, side_key, tong_label in [
        ("I. Nguồn vốn Trung ương", "key_tw", "TỔNG CỘNG PHẦN I"),
        ("II. Nguồn vốn Địa phương", "key_dp", "TỔNG CỘNG PHẦN II"),
    ]:
        nv_int_cur = 1 if side_key == "key_tw" else 2
        da_co_dong = False

        # Pre-collect visible rows để có look-ahead (biết hàng nào là hàng cuối của mỗi ma_ct)
        vis: list[dict] = []
        for _ten_nhom, ds_rows in _iter_khtd_cn_group_rows():
            for row in ds_rows:
                key_name = row.get(side_key)
                if not _duoc_loc(str(key_name) if key_name else None):
                    continue
                kh_vnd = float(kh_d.get(str(key_name), 0.0) or 0.0) if key_name else 0.0
                th_vnd = float(th_d.get(str(key_name), 0.0) or 0.0) if key_name else 0.0
                if kh_vnd > 0 or th_vnd > 0:
                    vis.append(row)

        for idx, row in enumerate(vis):
            key_name = row.get(side_key)
            kh_vnd = float(kh_d.get(str(key_name), 0.0) or 0.0) if key_name else 0.0
            th_vnd = float(th_d.get(str(key_name), 0.0) or 0.0) if key_name else 0.0
            ma_ct = int(row.get("ma_ct") or 0)
            row_type = row.get("type", "normal")

            next_ma_ct = int(vis[idx + 1].get("ma_ct") or 0) if idx + 1 < len(vis) else None
            is_last_mact = next_ma_ct != ma_ct

            # Với hàng "normal" (không phải gqvl_sub): trừ thu hồi NQ11 vào Còn phải TH
            thu_hoi_row = 0.0
            if is_last_mact and row_type != "gqvl_sub":
                thu_hoi_row = nq11_thuhoi_nam_by_mact_nv.get((ma_ct, nv_int_cur), 0.0)

            if not da_co_dong:
                _add_group_hdr(tieu_de)
                da_co_dong = True
            _add_row(str(row.get("label", "") or ""), kh_vnd, th_vnd, thu_hoi_row)

            # Chèn hàng NQ11 sau hàng cuối của mỗi nhóm ma_ct
            if is_last_mact and nq11_by_mact_nv:
                nq11_vnd = nq11_by_mact_nv.get((ma_ct, nv_int_cur), 0.0)
                if nq11_vnd > 0:
                    thu_hoi_nq11 = nq11_thuhoi_nam_by_mact_nv.get((ma_ct, nv_int_cur), 0.0)
                    _add_nq11_subrow(
                        nq11_vnd,
                        thu_hoi_nam_vnd=thu_hoi_nq11 if row_type == "gqvl_sub" else 0.0,
                        da_tru_o_dong_tren=(row_type != "gqvl_sub" and thu_hoi_nq11 > 0),
                    )

        if da_co_dong:
            if side_key == "key_tw":
                _add_section_total(tong_label, tong_kh_i, tong_th_i)
            else:
                _add_section_total(tong_label, tong_kh_ii, tong_th_ii)

    dem_ma_ct: set[int] = set()
    for _ten_nhom, ds_rows in tat_ca_rows:
        for row in ds_rows:
            ma_ct = row.get("ma_ct")
            key_tw = str(row.get("key_tw") or "")
            key_dp = str(row.get("key_dp") or "")
            kh_val = float(kh_d.get(key_tw, 0.0) or 0.0) + float(kh_d.get(key_dp, 0.0) or 0.0)
            if ma_ct is not None and kh_val > 0:
                dem_ma_ct.add(int(ma_ct))
    so_ct_co_kh = len(dem_ma_ct)
    tong_ct = len({int(row.get("ma_ct")) for _g, rows in tat_ca_rows for row in rows if row.get("ma_ct") is not None})

    tong_tl = tong_th / tong_kh * 100 if tong_kh > 0 else None
    tong_con_phai_th = (tong_kh - tong_th) / 1e6

    tds_tong = (
        _td("", "center", "#1f2937", TONG_BG, "bold") +
        _td("TỔNG CỘNG", "left", "#1f2937", TONG_BG, "bold") +
        _td(_fvn(tong_kh / 1e6, 0), "right", "#1f2937", TONG_BG, "bold") +
        _td(_fvn(tong_th / 1e6, 0), "right", "#1f2937", TONG_BG, "bold") +
        _td(_fvn(tong_con_phai_th, 0), "right", "#1f2937", TONG_BG, "bold") +
        _td(f"{_fvn(tong_tl, 1)}%" if tong_tl is not None else "—", "right", _tl_color(tong_tl), TONG_BG, "bold")
    )
    # Chèn TỔNG CỘNG vào cuối bảng (sau TỔNG CỘNG PHẦN II)
    insert_idx = None
    for _i, _row in enumerate(html_rows):
        if "TỔNG CỘNG PHẦN II" in _row:
            insert_idx = _i + 1
            break
    if insert_idx is not None:
        html_rows.insert(insert_idx, f"<tr>{tds_tong}</tr>")
    else:
        html_rows.append(f"<tr>{tds_tong}</tr>")

    k1, k2, k3, k4 = st.columns(4)
    k1.metric("Tổng KH (triệu đồng)", f"{_fvn(tong_kh / 1e6, 0)}")
    k2.metric("Tổng TH (triệu đồng)", f"{_fvn(tong_th / 1e6, 0)}")
    k3.metric(
        "Tỷ lệ đạt KH",
        f"{_fvn(tong_tl, 1)}%" if tong_tl is not None else "—",
    )
    k4.metric("Số CT có KH / tổng", f"{so_ct_co_kh}/{tong_ct}")

    headers = ["STT", "Chỉ tiêu", "KH (triệu đồng)", "Thực hiện (triệu đồng)", "Còn phải thực hiện (triệu đồng)", "TL%"]
    thead = "".join(
        f'<th style="background:{H_BG};color:#fff;text-align:{"center" if i==0 else "left" if i==1 else "right"};'
        f'padding:6px 8px;border:1px solid {BD};font-size:0.82rem;white-space:nowrap">{h}</th>'
        for i, h in enumerate(headers)
    )

    html = f"""
<div style="overflow-x:auto;margin:8px 0">
<table style="border-collapse:collapse;width:100%;font-family:'Inter','Segoe UI',sans-serif;font-size:0.82rem">
  <thead><tr>{thead}</tr></thead>
  <tbody>{"".join(html_rows)}</tbody>
</table>
<p style="font-size:0.78rem;color:var(--text-color, #6B7280);margin:4px 0 0 0">
  * Đơn vị: triệu đồng · KH từ nhập liệu, Thực hiện từ Tổng dư nợ HSTD + GQVL phân tầng · Còn phải thực hiện = KH - Thực hiện<br>
  TL% tô màu: <span style="color:{GREEN}">xanh</span> ≥ 100% &nbsp;
  <span style="color:{AMBER}">vàng</span> ≥ 95% &nbsp;
  <span style="color:{RED}">đỏ</span> &lt; 95%
</p>
</div>
"""
    render_html = getattr(st, "html", None)
    if callable(render_html):
        render_html(html)
    else:
        st.markdown(html, unsafe_allow_html=True)


def _tab_canh_bao_chenh_lech() -> None:
    st.subheader("⚠️ Cảnh báo Chênh lệch Phân bổ KHTD")
    st.caption(
        "So sánh kế hoạch chi nhánh (tổng) với tổng kế hoạch đã phân bổ xuống xã "
        "theo từng chương trình tín dụng."
    )

    kh_cn = _doc_kv(KV_KEY_CN)
    kh_xa = _doc_kv(KV_KEY_XA)

    if not kh_cn:
        st.info("Chưa có dữ liệu kế hoạch Chi nhánh. Vui lòng nhập ở tab **KHTD Chi nhánh**.")
        return

    # ── Tổng hợp tổng xã theo từng mã CT ─────────────────────────────────
    tong_xa_theo_ct: dict[str, float] = {}
    for khoa, gia_tri in kh_xa.items():
        phan = khoa.split("|")
        if len(phan) != 2:
            continue
        _, ma_ct = phan
        tong_xa_theo_ct[ma_ct] = tong_xa_theo_ct.get(ma_ct, 0.0) + float(gia_tri)

    # ── Xây dựng bảng so sánh ─────────────────────────────────────────────
    rows = []
    co_canh_bao_do = False
    co_canh_bao_vang = False

    keys_all = sorted(
        set(kh_cn.keys()) | set(tong_xa_theo_ct.keys()),
        key=lambda k: (
            _nv_int_tu_ma_key(k) or 9,
            int(k.split("_", 1)[0]) if "_" in k and k.split("_", 1)[0].isdigit()
            else (int(k.split("|", 1)[0]) if "|" in k and k.split("|", 1)[0].isdigit() else 9_999),
            k,
        ),
    )
    for ma_key in keys_all:
        ten_ct = TEN_CHINH_THUC_CT.get(ma_key, ma_key)
        nv_int = _nv_int_tu_ma_key(ma_key)
        nguon_von = "Trung ương" if nv_int == 1 else ("Địa phương" if nv_int == 2 else "—")
        gia_tri_cn = float(kh_cn.get(ma_key, 0.0))
        gia_tri_xa = tong_xa_theo_ct.get(ma_key, 0.0)

        if gia_tri_cn == 0 and gia_tri_xa == 0:
            continue

        chenh_lech = gia_tri_xa - gia_tri_cn
        ty_le_phan_bo = (gia_tri_xa / gia_tri_cn * 100) if gia_tri_cn > 0 else None

        if gia_tri_xa > gia_tri_cn:
            co_canh_bao_do = True
            trang_thai = "🔴 Vượt CN"
        elif ty_le_phan_bo is not None and ty_le_phan_bo < 95:
            co_canh_bao_vang = True
            trang_thai = "🟡 Chưa đủ 95%"
        else:
            trang_thai = "🟢 Đạt"

        rows.append({
            "Chương trình": ten_ct,
            "Nguồn vốn": nguon_von,
            "Chi nhánh (triệu)": round(gia_tri_cn / 1_000_000, 1),
            "Tổng xã (triệu)": round(gia_tri_xa / 1_000_000, 1),
            "Chênh lệch (triệu)": round(chenh_lech / 1_000_000, 1),
            "Tỷ lệ phân bổ %": round(ty_le_phan_bo, 1) if ty_le_phan_bo is not None else None,
            "Trạng thái": trang_thai,
        })

    if not rows:
        st.info("Không có dữ liệu để so sánh.")
        return

    # ── Hiển thị cảnh báo tổng quan ──────────────────────────────────────
    if co_canh_bao_do:
        st.error(
            "🔴 **Cảnh báo nghiêm trọng:** Một số chương trình có tổng kế hoạch xã "
            "**vượt quá** kế hoạch Chi nhánh đã duyệt!"
        )
    if co_canh_bao_vang:
        st.warning(
            "🟡 **Lưu ý:** Một số chương trình chưa phân bổ đủ 95% kế hoạch Chi nhánh "
            "xuống cấp xã."
        )
    if not co_canh_bao_do and not co_canh_bao_vang:
        st.success("🟢 Tất cả chương trình đã phân bổ đạt yêu cầu (≥ 95% và không vượt CN).")

    st.divider()

    # ── Metrics tổng ─────────────────────────────────────────────────────
    df = pd.DataFrame(rows)
    tong_cn_all = df["Chi nhánh (triệu)"].sum()
    tong_xa_all = df["Tổng xã (triệu)"].sum()
    ty_le_all = (tong_xa_all / tong_cn_all * 100) if tong_cn_all > 0 else 0.0

    m1, m2, m3, m4 = st.columns(4)
    m1.metric("Tổng KH Chi nhánh (triệu)", _fmt_vn(tong_cn_all, 1))
    m2.metric("Tổng KH Xã (triệu)", _fmt_vn(tong_xa_all, 1))
    m3.metric("Chênh lệch (triệu)", _fmt_vn_signed(tong_xa_all - tong_cn_all, 1))
    m4.metric("Tỷ lệ phân bổ tổng", f"{_fmt_vn(ty_le_all, 1)}%")

    st.divider()

    # ── Tô màu bảng theo trạng thái ──────────────────────────────────────
    def _to_mau_hang(row: pd.Series) -> list[str]:
        if "🔴" in str(row.get("Trạng thái", "")):
            return ["background-color: #ffd6d6; color:#1f2937"] * len(row)
        if "🟡" in str(row.get("Trạng thái", "")):
            return ["background-color: #fff9d6; color:#1f2937"] * len(row)
        return [""] * len(row)

    # Column config cho bảng cảnh báo
    column_config_cb: dict[str, st.column_config.Column] = {
        "Chi nhánh (triệu)": st.column_config.NumberColumn(
            "Chi nhánh (triệu)",
            format=",.1f",
            help="Kế hoạch Chi nhánh (triệu đồng)"
        ),
        "Tổng xã (triệu)": st.column_config.NumberColumn(
            "Tổng xã (triệu)",
            format=",.1f",
            help="Tổng kế hoạch xã (triệu đồng)"
        ),
        "Chênh lệch (triệu)": st.column_config.NumberColumn(
            "Chênh lệch (triệu)",
            format=",.1f",
            help="Chênh lệch (triệu đồng)"
        ),
        "Tỷ lệ phân bổ %": st.column_config.NumberColumn(
            "Tỷ lệ phân bổ %",
            format=".1f",
            help="Tỷ lệ phân bổ %"
        ),
    }

    hien_thi_dataframe_phan_trang(
        df.style.apply(_to_mau_hang, axis=1),
        key="khtd_cbtd_view",
        column_config=column_config_cb,
        height=450,
    )

    # ── Xuất Excel ────────────────────────────────────────────────────────
    st.divider()
    if st.button("📥 Xuất báo cáo chênh lệch", key="btn_xuat_chenh_lech"):
        state = SCMStateManager()
        buf = BytesIO()
        with pd.ExcelWriter(buf, engine="openpyxl") as writer:
            df.to_excel(writer, index=False, sheet_name="Chênh lệch KHTD")
        state.downloads.set(
            "khtd_chenh_lech_excel",
            buf.getvalue(),
            f"CanhBao_KHTD_{datetime.today().strftime('%d%m%Y')}.xlsx",
        )

    state = SCMStateManager()
    if state.downloads.has("khtd_chenh_lech_excel"):
        if st.download_button(
            label="⬇ Tải file Excel",
            data=state.downloads.get_bytes("khtd_chenh_lech_excel"),
            file_name=state.downloads.get_filename("khtd_chenh_lech_excel") or f"CanhBao_KHTD_{datetime.today().strftime('%d%m%Y')}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            key="dl_chenh_lech_excel",
        ):
            state.downloads.clear("khtd_chenh_lech_excel")


def _tab_tien_do_kh_th() -> None:
    """Dashboard cảnh báo tiến độ KH vs TH thực hiện theo PGD."""
    st.subheader("🎯 Tiến độ Kế hoạch vs Thực hiện")
    st.caption(
        "So sánh KH đã nhập với TH thực tế từ HSTD + GQVL. "
        "Cảnh báo 🔴 khi PGD đạt < 95% KH."
    )

    kh_cn = _dong_bo_gqvl_tong_keys(_doc_kv(KV_KEY_CN) or {})
    if not kh_cn:
        st.warning("⚠️ Chưa có KH Chi nhánh. Vào tab **🏛️ KHTD Chi nhánh** nhập trước.")
        return

    df_hstd = _doc_hstd_cached(ts_file(CACHE_HSTD))
    if df_hstd.empty:
        st.warning("⚠️ Chưa có dữ liệu HSTD. Upload file trước.")
        return

    df_gqvl = _doc_gqvl_cached(ts_file(CACHE_GQVL))
    th_cn, _th_gqvl = _tinh_thuc_hien_khtd_cn(df_hstd, df_gqvl)

    chi_tieu_theo_ma_ct: dict[int, dict[str, float]] = {}
    for _ten_nhom, ds_rows in _iter_khtd_cn_group_rows():
        for row in ds_rows:
            ma_ct = row.get("ma_ct")
            if ma_ct is None:
                continue
            key_tw = str(row.get("key_tw") or "")
            key_dp = str(row.get("key_dp") or "")
            entry = chi_tieu_theo_ma_ct.setdefault(int(ma_ct), {"kh": 0.0, "th": 0.0})
            entry["kh"] += float(kh_cn.get(key_tw, 0.0) or 0.0) + float(kh_cn.get(key_dp, 0.0) or 0.0)
            entry["th"] += float(th_cn.get(key_tw, 0.0) or 0.0) + float(th_cn.get(key_dp, 0.0) or 0.0)

    tong_kh = sum(v["kh"] for v in chi_tieu_theo_ma_ct.values())
    tong_th = sum(v["th"] for v in chi_tieu_theo_ma_ct.values())
    tl_cn = tong_th / tong_kh * 100 if tong_kh > 0 else 0

    so_chua = sum(
        1
        for vals in chi_tieu_theo_ma_ct.values()
        if vals["kh"] > 0 and (vals["th"] / vals["kh"] * 100) < 95
    )

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Tổng KH (triệu đồng)", f"{_fvn(tong_kh/1e6, 0)}")
    c2.metric("Tổng TH (triệu đồng)", f"{_fvn(tong_th/1e6, 0)}")
    c3.metric("Tỷ lệ CN", f"{_fvn(tl_cn, 1)}%",
              delta=f"{'✅' if tl_cn >= 95 else '⚠️'}")
    c4.metric("CT chưa đạt 95%", str(so_chua),
              delta_color="inverse" if so_chua > 0 else "off")

    st.divider()

    st.markdown("#### 📋 Chi tiết theo Chương trình")

    rows_ct: list[dict[str, object]] = []
    for tieu_de, side_key in [("I. Trung ương", "key_tw"), ("II. Địa phương", "key_dp")]:
        da_co_dong = False
        for _ten_nhom, ds_rows in _iter_khtd_cn_group_rows():
            for row in ds_rows:
                key_name = str(row.get(side_key) or "")
                if not key_name:
                    continue
                kh_val = float(kh_cn.get(key_name, 0.0) or 0.0)
                th_val = float(th_cn.get(key_name, 0.0) or 0.0)
                if kh_val == 0 and th_val == 0:
                    continue
                if not da_co_dong:
                    da_co_dong = True
                    rows_ct.append({
                        "STT": tieu_de,
                        "Chỉ tiêu": tieu_de,
                        "KH (triệu đồng)": None,
                        "TH (triệu đồng)": None,
                        "TL%": None,
                        "Trạng thái": "",
                        "_nhom": True,
                    })
                tl_val = th_val / kh_val * 100 if kh_val > 0 else None
                if tl_val is None:
                    trang_thai = "—"
                elif tl_val >= 100:
                    trang_thai = "🟢 Đạt"
                elif tl_val >= 95:
                    trang_thai = "🟡 Đang thực hiện"
                else:
                    trang_thai = "🔴 Chậm"
                rows_ct.append({
                    "STT": "",
                    "Chỉ tiêu": str(row.get("label", "") or ""),
                    "KH (triệu đồng)": round(kh_val / 1e6, 0) if kh_val else None,
                    "TH (triệu đồng)": round(th_val / 1e6, 0) if th_val else None,
                    "TL%": round(tl_val, 1) if tl_val is not None else None,
                    "Trạng thái": trang_thai,
                    "_nhom": False,
                })

    rows_ct.append({
        "STT": "", "Chỉ tiêu": "TỔNG CỘNG",
        "KH (triệu đồng)": round(tong_kh/1e6, 0),
        "TH (triệu đồng)": round(tong_th/1e6, 0),
        "TL%": round(tl_cn, 1),
        "Trạng thái": "🟢 Đạt" if tl_cn >= 95 else "🔴 Chậm",
        "_nhom": False,
    })

    df_ct = pd.DataFrame(rows_ct)

    def _to_mau_ct(row):
        if row.get("_nhom"):
            return ["background-color: #D9E1F2; color:#1f2937; font-weight: bold"] * len(row)
        if "🔴" in str(row.get("Trạng thái", "")):
            return ["background-color: #ffd6d6; color:#1f2937"] * len(row)
        if "🟡" in str(row.get("Trạng thái", "")):
            return ["background-color: #fff9d6; color:#1f2937"] * len(row)
        return [""] * len(row)

    cols_show = ["Chỉ tiêu", "KH (triệu đồng)", "TH (triệu đồng)", "TL%", "Trạng thái"]
    hien_thi_dataframe_phan_trang(
        df_ct[cols_show].style.apply(_to_mau_ct, axis=1),
        key="khtd_tien_do_ct",
        height=480,
        column_config={
            "KH (triệu đồng)": st.column_config.NumberColumn(format=",.0f"),
            "TH (triệu đồng)": st.column_config.NumberColumn(format=",.0f"),
            "TL%":     st.column_config.ProgressColumn(
                           min_value=0, max_value=100, format=".1f"),
        },
    )

    st.divider()
    st.markdown("#### 🔴 Cảnh báo PGD chậm tiến độ (< 95% KH)")

    if COT_TEN_PGD not in df_hstd.columns:
        st.info("Không có cột PGD trong HSTD.")
        return

    rows_pgd = []
    for ten_pgd in DS_PGD:
        df_pgd = df_hstd[df_hstd[COT_TEN_PGD] == ten_pgd]
        try:
            df_gqvl_pgd = (
                df_gqvl[df_gqvl[COT_TEN_PGD] == ten_pgd]
                if not df_gqvl.empty and COT_TEN_PGD in df_gqvl.columns
                else pd.DataFrame()
            )
        except Exception as e:
            logger.error("Lỗi trong khối except: %s", e, exc_info=True)
            df_gqvl_pgd = pd.DataFrame()
        th_pgd, _ = _tinh_thuc_hien_khtd_cn(df_pgd, df_gqvl_pgd)

        tong_kh_pgd = tong_kh
        tong_th_pgd = sum(
            float(th_pgd.get(str(row.get("key_tw") or ""), 0.0) or 0.0)
            + float(th_pgd.get(str(row.get("key_dp") or ""), 0.0) or 0.0)
            for _ten_nhom, ds_rows in _iter_khtd_cn_group_rows()
            for row in ds_rows
        )
        tl_pgd = tong_th_pgd / tong_kh_pgd * 100 if tong_kh_pgd > 0 else 0

        if tl_pgd < 95:
            rows_pgd.append({
                "PGD": ten_pgd,
                "TH (triệu đồng)": round(tong_th_pgd/1e6, 0),
                "KH CN (triệu đồng)": round(tong_kh_pgd/1e6, 0),
                "TL%": round(tl_pgd, 1),
            })

    if rows_pgd:
        df_pgd_cb = pd.DataFrame(rows_pgd).sort_values("TL%")
        st.dataframe(
            df_pgd_cb,
            use_container_width=True,
            hide_index=True,
            column_config={
                "TL%": st.column_config.ProgressColumn(
                    min_value=0, max_value=100, format=".1f"),
            },
        )
    else:
        st.success("🟢 Tất cả PGD đang đạt ≥ 95% KH Chi nhánh.")

    st.divider()
    col_ex, col_pdf = st.columns(2)

    with col_ex:
        if st.button("📥 Xuất Excel", key="btn_xuat_tien_do_excel"):
            try:
                buf = xuat_excel({
                    "KH vs TH": df_ct[cols_show],
                    "PGD chậm": pd.DataFrame(rows_pgd) if rows_pgd else pd.DataFrame(),
                })
                SCMStateManager().downloads.set("tien_do_kh_th_excel", buf, ten_file_xuat("TienDo_KH_TH"))
            except Exception as e:
                logger.error("Lỗi trong khối except: %s", e, exc_info=True)
                st.error(f"❌ Lỗi: {e}")
        state = SCMStateManager()
        if state.downloads.has("tien_do_kh_th_excel"):
            if st.download_button(
                "⬇ Tải Excel",
                data=state.downloads.get_bytes("tien_do_kh_th_excel"),
                file_name=state.downloads.get_filename("tien_do_kh_th_excel") or ten_file_xuat("TienDo_KH_TH"),
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                key="dl_tien_do_excel",
            ):
                state.downloads.clear("tien_do_kh_th_excel")

    with col_pdf:
        if st.button("📄 Xuất PDF", key="btn_xuat_tien_do_pdf", type="primary"):
            try:
                with st.spinner("⏳ Đang tạo PDF..."):
                    pdf_bytes = xuat_pdf(
                        df_ct[cols_show].dropna(subset=["KH (triệu đồng)"]),
                        "Báo cáo Tiến độ Kế hoạch vs Thực hiện",
                        username="VBSP-SCM",
                        cols_tien=["KH (triệu đồng)", "TH (triệu đồng)"],
                    )
                SCMStateManager().downloads.set(
                    "tien_do_kh_th_pdf",
                    pdf_bytes,
                    ten_file_xuat("TienDo_KH_TH", ext=".pdf"),
                )
            except Exception as e:
                logger.error("Lỗi trong khối except: %s", e, exc_info=True)
                st.error(f"❌ Lỗi PDF: {e}")
        state = SCMStateManager()
        if state.downloads.has("tien_do_kh_th_pdf"):
            if st.download_button(
                "⬇ Tải PDF",
                data=state.downloads.get_bytes("tien_do_kh_th_pdf"),
                file_name=state.downloads.get_filename("tien_do_kh_th_pdf") or "TienDo_KH_TH.pdf",
                mime="application/pdf",
                key="dl_tien_do_pdf",
            ):
                state.downloads.clear("tien_do_kh_th_pdf")


def xuat_khtd_theo_xa(role: str, username: str, df_full: "pd.DataFrame | None" = None) -> bytes:
    """Xuất Excel Kế hoạch Tín dụng theo Xã — ma trận Xã × Chương trình."""
    kh_xa = _doc_kv(KV_KEY_XA)
    kh_cn = _doc_kv(KV_KEY_CN)

    ds_xa = sorted(XA_TO_PGD.keys())

    ct_tw = [(mk, ten) for mk, _, ten, nv, _ in CHUONG_TRINH_KHTD if nv == "TW"]
    ct_dp = [(mk, ten) for mk, _, ten, nv, _ in CHUONG_TRINH_KHTD if nv == "DP"]

    def _lookup(ten_xa: str, ma_key: str) -> float:
        val = kh_xa.get(f"{ten_xa}|{ma_key}", None)
        if val is not None:
            return float(val)
        return 0.0

    def _cn_val(ma_key: str) -> float:
        return float(kh_cn.get(ma_key, 0.0))

    rows_data = []
    so_xa_co_kh = 0
    for ten_xa in ds_xa:
        pgd = XA_TO_PGD.get(ten_xa, "")
        row = {"Xã/Phường": ten_xa, "PGD": pgd}
        tong_xa = 0.0
        for mk, ten_ct in ct_tw + ct_dp:
            trieu = round(_lookup(ten_xa, mk) / 1_000_000, 1)
            row[ten_ct] = trieu
            tong_xa += trieu
        row["Tổng"] = round(tong_xa, 1)
        if tong_xa > 0:
            so_xa_co_kh += 1
        rows_data.append(row)

    row_tong_cn = {"Xã/Phường": "TỔNG CHI NHÁNH", "PGD": ""}
    tong_cn_all = 0.0
    for mk, ten_ct in ct_tw + ct_dp:
        trieu = round(_cn_val(mk) / 1_000_000, 1)
        row_tong_cn[ten_ct] = trieu
        tong_cn_all += trieu
    row_tong_cn["Tổng"] = round(tong_cn_all, 1)
    rows_data.append(row_tong_cn)

    col_order = ["Xã/Phường", "PGD"] + [ten for _, ten in ct_tw + ct_dp] + ["Tổng"]
    df_xa = pd.DataFrame(rows_data, columns=col_order)

    # ── Dùng xuat_excel() từ utils.py để build multi-sheet (23 sheet) ──────
    DS_TT_22 = [DON_VI_CHI_NHANH] + DS_PGD

    sheets: dict[str, pd.DataFrame] = {}

    # Sheet 1: Tổng hợp toàn CN (giữ cả dòng TỔNG CHI NHÁNH)
    sheets["Tổng hợp CN"] = df_xa

    # Sheet 2..23: từng đơn vị trong DS_TT_22 (22 đơn vị)
    df_xa_chi = df_xa[df_xa["Xã/Phường"] != "TỔNG CHI NHÁNH"].copy()

    for ten_dv in DS_TT_22:
        ds_xa_dv = PGD_XA_MAP.get(ten_dv, [])
        if not ds_xa_dv:
            continue
        df_dv = df_xa_chi[df_xa_chi["Xã/Phường"].isin(ds_xa_dv)].copy()
        if df_dv.empty:
            continue

        # Thêm dòng tổng đơn vị
        row_tong: dict = {"Xã/Phường": f"TỔNG {ten_dv}", "PGD": ""}
        for col in col_order[2:]:  # bỏ 2 cột đầu Xã/Phường và PGD
            row_tong[col] = round(float(df_dv[col].sum()), 1)
        df_dv = pd.concat(
            [df_dv, pd.DataFrame([row_tong])],
            ignore_index=True
        )

        # Tên sheet: tối đa 31 ký tự (giới hạn Excel)
        ten_sheet = (
            ten_dv.replace("Hội sở Chi nhánh tỉnh", "Hội sở CN tỉnh")
                  .replace("PGD ", "")
                  .strip()[:31]
        )
        sheets[ten_sheet] = df_dv

    return xuat_excel(sheets)


def xuat_to_trinh_bgd_word(username: str = "unknown") -> bytes:
    """Xuất Tờ trình BGĐ Word (.docx) tổng hợp KHTD vs thực hiện thực tế."""
    import db as _db
    from docx import Document
    from docx.shared import Pt, Cm, Mm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from datetime import date as _date
    from config import TEN_CHI_NHANH_HIEN_THI, COT_TONG_DU_NO

    kh_cn = _dong_bo_gqvl_tong_keys(_doc_kv(KV_KEY_CN) or {})
    if not kh_cn:
        raise ValueError("Chưa có dữ liệu KHTD Chi nhánh. Hãy giao KHTD trước.")

    th_cn: dict[str, float] = {}
    ten_map_word: dict[str, str] = {}
    ngay_sl = ""
    du_no_pgd: dict[str, float] = {}

    if os.path.exists(CACHE_HSTD):
        try:
            df_h = pd.read_parquet(CACHE_HSTD)
            if not df_h.empty:
                df_gqvl_word = _doc_gqvl_cached(ts_file(CACHE_GQVL))
                th_cn, _ = _tinh_thuc_hien_khtd_cn(df_h, df_gqvl_word)
                _, ten_map_word = _quet_ct_co_du_no(df_h)
                if COT_TEN_PGD in df_h.columns and COT_TONG_DU_NO in df_h.columns:
                    du_no_pgd = (
                        pd.to_numeric(df_h[COT_TONG_DU_NO], errors="coerce")
                        .fillna(0)
                        .groupby(df_h[COT_TEN_PGD])
                        .sum()
                        .to_dict()
                    )
                meta = _db.doc_kv("merge_meta_hstd") or {}
                if meta.get("ngay_sl"):
                    ngay_sl = str(meta["ngay_sl"])
                if not ngay_sl:
                    _dt_ns = lay_ngay_so_lieu(df_h)
                    ngay_sl = _dt_ns.strftime("%d/%m/%Y") if _dt_ns else ""
        except Exception as e:
            logger.error("xuat_to_trinh_bgd_word: đọc HSTD lỗi — %s", e, exc_info=True)

    today = _date.today()
    GREEN = RGBColor(0x1B, 0x5E, 0x20)
    DARK = RGBColor(0x21, 0x21, 0x21)
    GRAY = RGBColor(0x75, 0x75, 0x75)

    doc = Document()

    # ── Trang A4, lề chuẩn hành chính ────────────────────────────────────────
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.left_margin = Cm(3)
    sec.right_margin = Cm(2)
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)

    def _p(text="", align=WD_ALIGN_PARAGRAPH.LEFT, bold=False,
           size=11, color=None, space_after=2):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        if text:
            run = p.add_run(text)
            run.bold = bold
            run.font.size = Pt(size)
            if color:
                run.font.color.rgb = color
        return p

    def _shade(cell, fill_hex: str):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill_hex)
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)

    def _fmt_ty(vnd: float) -> str:
        if vnd == 0:
            return "—"
        return f"{vnd / 1e9:,.3f}".replace(",", ".")

    def _fmt_pct(kh: float, th: float) -> str:
        if kh <= 0:
            return "—"
        return f"{th / kh * 100:.1f}%".replace(".", ",")

    def _cell_fmt(cell, text: str, align=WD_ALIGN_PARAGRAPH.LEFT, bold=False,
                  size=9, color: RGBColor | None = None):
        cell.text = text
        for p in cell.paragraphs:
            p.alignment = align
            for run in p.runs:
                run.bold = bold
                run.font.size = Pt(size)
                if color:
                    run.font.color.rgb = color

    # ── Header ───────────────────────────────────────────────────────────────
    _p("NGÂN HÀNG CHÍNH SÁCH XÃ HỘI", WD_ALIGN_PARAGRAPH.CENTER, True, 12, GREEN)
    _p(TEN_CHI_NHANH_HIEN_THI.upper(), WD_ALIGN_PARAGRAPH.CENTER, True, 11, GREEN)
    _p("────────────────────────────────────────────",
       WD_ALIGN_PARAGRAPH.CENTER, False, 9, GRAY, 6)

    _p("TỜ TRÌNH", WD_ALIGN_PARAGRAPH.CENTER, True, 16, DARK, 2)
    _p("V/v: Báo cáo tiến độ thực hiện Kế hoạch Tín dụng",
       WD_ALIGN_PARAGRAPH.CENTER, False, 11)
    _p(f"(Tính đến ngày {ngay_sl})", WD_ALIGN_PARAGRAPH.CENTER, False, 10, GRAY, 8)

    _p("Kính gửi: Ban Giám đốc Chi nhánh NHCSXH tỉnh Đồng Nai",
       WD_ALIGN_PARAGRAPH.LEFT, False, 11, space_after=6)

    _p(
        f"Căn cứ số liệu Kế hoạch Tín dụng được giao và kết quả thực hiện tính đến "
        f"ngày {ngay_sl}, Phòng Kế hoạch - Nghiệp vụ báo cáo Ban Giám đốc như sau:",
        WD_ALIGN_PARAGRAPH.JUSTIFY, False, 11, space_after=8,
    )

    # ── Phần I: Tổng hợp theo chương trình ───────────────────────────────────
    _p("I. KẾT QUẢ THỰC HIỆN KẾ HOẠCH TÍN DỤNG TOÀN CHI NHÁNH",
       WD_ALIGN_PARAGRAPH.LEFT, True, 12, GREEN, 2)
    _p("Bảng 1: Tổng hợp thực hiện kế hoạch tín dụng theo chương trình (đơn vị: tỷ đồng)",
       WD_ALIGN_PARAGRAPH.LEFT, False, 10, GRAY, 4)

    HDR1 = ["STT", "Chương trình tín dụng", "NV", "KH giao", "Thực hiện", "Tỷ lệ %"]
    WIDTHS1 = [1.0, 8.5, 1.2, 2.2, 2.2, 1.8]
    t1 = doc.add_table(rows=1, cols=len(HDR1))
    t1.style = "Table Grid"
    for i, h in enumerate(HDR1):
        c = t1.rows[0].cells[i]
        _cell_fmt(c, h, WD_ALIGN_PARAGRAPH.CENTER, True, 9, RGBColor(0xFF, 0xFF, 0xFF))
        _shade(c, "1B5E20")
    for row in t1.rows:
        for i, w in enumerate(WIDTHS1):
            if i < len(row.cells):
                row.cells[i].width = Cm(w)

    stt = 0
    tong_kh = tong_th = 0.0

    row_groups = _iter_khtd_cn_group_rows(ten_map_word)
    for nv_label, side_key in [("TW", "key_tw"), ("ĐP", "key_dp")]:
        # Group header
        gr = t1.add_row()
        merged = gr.cells[0].merge(gr.cells[-1])
        label = "NGUỒN VỐN TRUNG ƯƠNG" if nv_label == "TW" else "NGUỒN VỐN ĐỊA PHƯƠNG"
        _cell_fmt(merged, label, WD_ALIGN_PARAGRAPH.LEFT, True, 9)
        _shade(merged, "E8F4FD")

        sub_kh = sub_th = 0.0
        for _ten_nhom, ds_rows in row_groups:
            for row_model in ds_rows:
                mk = str(row_model.get(side_key) or "")
                if not mk:
                    continue
                ten = str(row_model.get("label", "") or "")
                kh_v = float(kh_cn.get(mk, 0) or 0)
                th_v = float(th_cn.get(mk, 0) or 0)
                if kh_v == 0 and th_v == 0:
                    continue
                stt += 1
                sub_kh += kh_v
                sub_th += th_v
                tong_kh += kh_v
                tong_th += th_v
                dr = t1.add_row()
                vals = [str(stt), ten, nv_label, _fmt_ty(kh_v), _fmt_ty(th_v), _fmt_pct(kh_v, th_v)]
                for i, v in enumerate(vals):
                    alg = WD_ALIGN_PARAGRAPH.RIGHT if i >= 3 else WD_ALIGN_PARAGRAPH.LEFT
                    _cell_fmt(dr.cells[i], v, alg)
                for i, w in enumerate(WIDTHS1):
                    if i < len(dr.cells):
                        dr.cells[i].width = Cm(w)

        # Subtotal
        sr = t1.add_row()
        sv = ["", f"Cộng {nv_label}", "", _fmt_ty(sub_kh), _fmt_ty(sub_th), _fmt_pct(sub_kh, sub_th)]
        for i, v in enumerate(sv):
            alg = WD_ALIGN_PARAGRAPH.RIGHT if i >= 3 else WD_ALIGN_PARAGRAPH.LEFT
            _cell_fmt(sr.cells[i], v, alg, True)
            _shade(sr.cells[i], "F1F8E9")
        for i, w in enumerate(WIDTHS1):
            if i < len(sr.cells):
                sr.cells[i].width = Cm(w)

    # Grand total
    tr = t1.add_row()
    tv = ["", "TỔNG CỘNG", "", _fmt_ty(tong_kh), _fmt_ty(tong_th), _fmt_pct(tong_kh, tong_th)]
    for i, v in enumerate(tv):
        alg = WD_ALIGN_PARAGRAPH.RIGHT if i >= 3 else WD_ALIGN_PARAGRAPH.LEFT
        _cell_fmt(tr.cells[i], v, alg, True, 9, RGBColor(0xFF, 0xFF, 0xFF))
        _shade(tr.cells[i], "1B5E20")
    for i, w in enumerate(WIDTHS1):
        if i < len(tr.cells):
            tr.cells[i].width = Cm(w)

    doc.add_paragraph()

    # ── Phần II: Theo PGD ────────────────────────────────────────────────────
    if du_no_pgd:
        _p("II. TIẾN ĐỘ THỰC HIỆN THEO ĐƠN VỊ",
           WD_ALIGN_PARAGRAPH.LEFT, True, 12, GREEN, 2)
        _p("Bảng 2: Tổng hợp dư nợ theo đơn vị (đơn vị: tỷ đồng)",
           WD_ALIGN_PARAGRAPH.LEFT, False, 10, GRAY, 4)

        tong_dn = sum(du_no_pgd.values())
        HDR2 = ["STT", "Đơn vị", "Dư nợ (tỷ đồng)", "Tỷ trọng"]
        WIDTHS2 = [1.0, 8.5, 3.5, 2.5]
        t2 = doc.add_table(rows=1, cols=4)
        t2.style = "Table Grid"
        for i, h in enumerate(HDR2):
            c = t2.rows[0].cells[i]
            _cell_fmt(c, h, WD_ALIGN_PARAGRAPH.CENTER, True, 9, RGBColor(0xFF, 0xFF, 0xFF))
            _shade(c, "1B5E20")
        for row in t2.rows:
            for i, w in enumerate(WIDTHS2):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)

        stt2 = 0
        for pgd in [DON_VI_CHI_NHANH] + DS_PGD:
            dn = float(du_no_pgd.get(pgd, 0))
            if dn == 0:
                continue
            stt2 += 1
            pct_s = f"{dn / tong_dn * 100:.1f}%".replace(".", ",") if tong_dn else "—"
            dr2 = t2.add_row()
            v2 = [str(stt2), pgd, _fmt_ty(dn), pct_s]
            for i, v in enumerate(v2):
                alg = WD_ALIGN_PARAGRAPH.RIGHT if i >= 2 else WD_ALIGN_PARAGRAPH.LEFT
                _cell_fmt(dr2.cells[i], v, alg)
            for i, w in enumerate(WIDTHS2):
                if i < len(dr2.cells):
                    dr2.cells[i].width = Cm(w)

        tot2 = t2.add_row()
        tv2 = ["", "TỔNG CỘNG", _fmt_ty(tong_dn), "100%"]
        for i, v in enumerate(tv2):
            alg = WD_ALIGN_PARAGRAPH.RIGHT if i >= 2 else WD_ALIGN_PARAGRAPH.LEFT
            _cell_fmt(tot2.cells[i], v, alg, True)
            _shade(tot2.cells[i], "C8E6C9")
        for i, w in enumerate(WIDTHS2):
            if i < len(tot2.cells):
                tot2.cells[i].width = Cm(w)

        doc.add_paragraph()

    # ── Phần III: Kết luận ───────────────────────────────────────────────────
    _p("III. NHẬN XÉT VÀ KIẾN NGHỊ",
       WD_ALIGN_PARAGRAPH.LEFT, True, 12, GREEN, 4)
    pct_chung = tong_th / tong_kh * 100 if tong_kh > 0 else 0
    _p(
        f"Tính đến ngày {ngay_sl}, toàn Chi nhánh đã thực hiện được "
        f"{_fmt_ty(tong_th)} tỷ đồng dư nợ, đạt "
        f"{_fmt_pct(tong_kh, tong_th)} so với kế hoạch được giao "
        f"({_fmt_ty(tong_kh)} tỷ đồng).",
        WD_ALIGN_PARAGRAPH.JUSTIFY, False, 11, space_after=6,
    )
    _p(
        "Phòng Kế hoạch - Nghiệp vụ kính trình Ban Giám đốc xem xét, chỉ đạo.",
        WD_ALIGN_PARAGRAPH.JUSTIFY, False, 11, space_after=16,
    )

    # ── Chữ ký ──────────────────────────────────────────────────────────────
    sig_t = doc.add_table(rows=1, cols=2)
    sig_t.cell(0, 0).text = (
        "Nơi nhận:\n"
        "- Ban Giám đốc Chi nhánh (để b/c);\n"
        "- Lưu: VT, KHNV."
    )
    for p in sig_t.cell(0, 0).paragraphs:
        for run in p.runs:
            run.font.size = Pt(10)

    sig_r = sig_t.cell(0, 1)
    sig_r.text = (
        f"Đồng Nai, ngày {today.day:02d} tháng {today.month:02d} năm {today.year}\n"
        "TM. PHÒNG KẾ HOẠCH - NGHIỆP VỤ\n"
        "TRƯỞNG PHÒNG\n\n\n"
        "(Ký, ghi rõ họ tên)"
    )
    for p in sig_r.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(11)
    # In đậm dòng chức danh
    sig_r.paragraphs[1].runs[0].bold = True

    buf = BytesIO()
    doc.save(buf)
    _db.ghi_audit(username, "xuat_to_trinh_bgd_word",
                  f"Tờ trình BGĐ ngày {today.isoformat()}")
    return buf.getvalue()


def _tab_xuat_khtd_xa(role: str, username: str, df_full: "pd.DataFrame | None") -> None:
    """Tab xuất Excel KHTD phân bổ theo Xã — ma trận 95 xã × chương trình."""
    st.subheader("📍 KHTD phân bổ theo Xã/Phường")
    st.caption("Ma trận kế hoạch tín dụng phân bổ đến từng xã/phường — đơn vị: triệu đồng · 23 sheet (Tổng hợp CN + 22 đơn vị)")

    with st.expander("👁️ Xem trước dữ liệu", expanded=False):
        try:
            kh_xa_prev = _doc_kv(KV_KEY_XA)
            if not kh_xa_prev:
                st.info("Chưa có dữ liệu kế hoạch xã.")
            else:
                ct_all = [(mk, ten) for mk, _, ten, nv, _ in CHUONG_TRINH_KHTD]
                rows_prev = []
                for ten_xa, ten_dv in sorted(XA_TO_PGD.items()):
                    row = {"Xã/Phường": ten_xa, "Đơn vị": ten_dv}
                    tong = 0.0
                    for mk, ten_ct in ct_all:
                        v = round(float(kh_xa_prev.get(f"{ten_xa}|{mk}", 0)) / 1e6, 1)
                        row[ten_ct] = v
                        tong += v
                    row["Tổng"] = round(tong, 1)
                    if tong > 0:
                        rows_prev.append(row)
                if rows_prev:
                    col_order_prev = ["Xã/Phường", "Đơn vị"] + [t for _, t in ct_all] + ["Tổng"]
                    df_prev = pd.DataFrame(rows_prev, columns=col_order_prev)
                    so_xa = len(df_prev)
                    tong_all = df_prev["Tổng"].sum()
                    st.caption(
                        f"**{so_xa}** xã/phường có kế hoạch · "
                        f"Tổng: **{tong_all:,.1f}** triệu đồng"
                    )
                    st.dataframe(df_prev, use_container_width=True,
                                 hide_index=True, height=300)
                else:
                    st.info("Chưa có xã/phường nào được giao kế hoạch.")
        except Exception as e:
            logger.error("Lỗi trong khối except: %s", e, exc_info=True)
            st.warning(f"Không thể xem trước: {e}")

    col1, _col2 = st.columns([1, 3])
    with col1:
        if st.button("📥 Xuất Excel KHTD/Xã", key="btn_xuat_khtd_xa"):
            with st.spinner("Đang tạo file..."):
                try:
                    excel_bytes = xuat_khtd_theo_xa(role, username, df_full)
                    ten_file = ten_file_xuat("KHTD_theo_Xa")
                    state = SCMStateManager()
                    state.downloads.set("khtd_xa_excel", excel_bytes, ten_file)
                    st.success(f"✅ Đã tạo: {ten_file}")
                except Exception as e:
                    logger.error("Lỗi trong khối except: %s", e, exc_info=True)
                    st.error(f"❌ Lỗi khi tạo file: {e}")

    state = SCMStateManager()
    if state.downloads.has("khtd_xa_excel"):
        if st.download_button(
            label="⬇️ Tải file Excel",
            data=state.downloads.get_bytes("khtd_xa_excel"),
            file_name=state.downloads.get_filename("khtd_xa_excel") or "KHTD_theo_Xa.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            key="dl_khtd_xa",
        ):
            state.downloads.clear("khtd_xa_excel")


def _tab_xuat_to_trinh_bgd(username: str) -> None:
    """Tab xuất Tờ trình BGĐ dạng Word."""
    st.subheader("📄 Tờ trình Ban Giám đốc")
    st.caption(
        "Tờ trình tổng hợp Kế hoạch Tín dụng vs Thực hiện — "
        "định dạng .docx chuẩn hành chính, ký số trực tiếp."
    )

    c1, _ = st.columns([1, 3])
    with c1:
        if st.button("📄 Tạo Tờ trình Word", key="btn_xuat_to_trinh_bgd", type="primary"):
            with st.spinner("Đang tạo tài liệu..."):
                try:
                    word_bytes = xuat_to_trinh_bgd_word(username=username or "unknown")
                    ten_file = ten_file_xuat("ToTrinh_BGD", ext="docx")
                    state2 = SCMStateManager()
                    state2.downloads.set("to_trinh_bgd_word", word_bytes, ten_file)
                    st.success(f"✅ Đã tạo: {ten_file}")
                except ValueError as e:
                    st.warning(f"⚠️ {e}")
                except Exception as e:
                    logger.error("xuat_to_trinh_bgd_word UI: %s", e, exc_info=True)
                    st.error(f"❌ Lỗi: {e}")

    state2 = SCMStateManager()
    if state2.downloads.has("to_trinh_bgd_word"):
        if st.download_button(
            label="⬇️ Tải Tờ trình (.docx)",
            data=state2.downloads.get_bytes("to_trinh_bgd_word"),
            file_name=state2.downloads.get_filename("to_trinh_bgd_word") or "ToTrinh_BGD.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="dl_to_trinh_bgd",
        ):
            state2.downloads.clear("to_trinh_bgd_word")


def render_xuat_baocao(role: str = "", username: str = "", df_full: "pd.DataFrame | None" = None) -> None:
    sub1, sub2, sub3, sub4 = st.tabs([
        "📋 Tiến độ KH vs TH",
        "⚠️ Chênh lệch phân bổ",
        "📍 KHTD theo Xã",
        "📄 Tờ trình BGĐ",
    ])
    with sub1:
        _tab_tien_do_kh_th()
    with sub2:
        _tab_canh_bao_chenh_lech()
    with sub3:
        _tab_xuat_khtd_xa(role, username, df_full)
    with sub4:
        _tab_xuat_to_trinh_bgd(username)

