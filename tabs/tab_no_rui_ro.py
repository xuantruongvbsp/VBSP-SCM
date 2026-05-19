"""Xử lý nợ rủi ro theo QĐ 62/2015/QĐ-TTg — 5 bước: lọc, chọn, nhập, xuất, xem lại."""
from __future__ import annotations

import io
from datetime import date, datetime

import pandas as pd
import streamlit as st
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from streamlit.delta_generator import DeltaGenerator

import db
from config import (
    COT_DIA_CHI,
    COT_DU_NO_QH,
    COT_DU_NO_TH,
    COT_LAI_TON,
    COT_MA_KH,
    COT_MUC_VAY,
    COT_NGAY_DH,
    COT_NGAY_VAY,
    COT_NGUON_VON,
    COT_SDT,
    COT_TEN_XA,
    COT_TEN_TO,
    COT_TEN_KH,
    COT_SO_KU,
    COT_TEN_CT,
    COT_TONG_DU_NO,
    COT_DU_NO_QH,
    COT_TEN_PGD,
    DS_PGD,
    NGUYEN_NHAN_RR,
)
from data.pgd import pgd_slug
from auth import la_phan_he_cn, la_phan_he_pgd, normalize_role
from utils import fmt, fmt_bang_ty, fmt_ngay, hien_thi_dataframe_phan_trang, xuat_excel
from services.template_service import (
    docx_bytes_to_pdf,
    nut_tai_word_va_pdf,
    hien_thi_nut_tai,
)

NGUON_TW = 1
NGUON_DP = 2
LABEL_TW = "Trung ương"
LABEL_DP = "Địa phương"


def _pgd_plain(ten: str) -> str:
    s = str(ten or "").strip()
    if s.lower().startswith("pgd "):
        return s[4:].strip()
    return s


def _pgd_line(ten: str) -> str:
    s = str(ten or "").strip()
    if not s:
        return "PGD"
    if s.lower().startswith("pgd "):
        return s
    return f"PGD {s}"


def _style_doc_xln(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)


# ── Helpers dùng chung cho Word XLN ─────────────────────────────
def _bo_border_cell(cell) -> None:
    """Xóa border khỏi 4 cạnh của một cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for bn in ["top", "left", "bottom", "right"]:
        b = OxmlElement(f"w:{bn}")
        b.set(qn("w:val"), "none")
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _set_cell(
    cell, text: str, *,
    bold: bool = False, italic: bool = False,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    v_align=WD_ALIGN_VERTICAL.CENTER,
    font_size: int = 10,
) -> None:
    """Gán text + format cho một cell bảng."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(font_size)
    cell.vertical_alignment = v_align


def _set_row_font(row, font_size: int = 10) -> None:
    """Đồng bộ font cho toàn bộ row."""
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(font_size)


def _num(v) -> float:
    """Chuyển đổi an toàn string → float."""
    if v is None:
        return 0.0
    if isinstance(v, (int, float)):
        return float(v)
    s = str(v).strip()
    if not s:
        return 0.0
    s = s.replace(" ", "").replace(".", "").replace(",", ".")
    try:
        return float(s)
    except Exception:
        return 0.0


def _add_header_xln(doc: Document, dia_danh: str, ngay_ky: date) -> None:
    """Header Quốc hiệu + ngày tháng, không có số VB (mẫu đơn cá nhân)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    for cell in t.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for bn in ["top", "left", "bottom", "right"]:
            b = OxmlElement(f"w:{bn}")
            b.set(qn("w:val"), "none")
            tcBorders.append(b)
        tcPr.append(tcBorders)

    cell_l = t.rows[0].cells[0]
    cell_l.paragraphs[0].add_run("Mẫu số 01/XLN")

    cell_r = t.rows[0].cells[1]
    p3 = cell_r.paragraphs[0]
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run("CỘNG HOÀ XÃ HỘI CHỦ NGHĨA VIỆT NAM").bold = True
    p4 = cell_r.add_paragraph("Độc lập - Tự do - Hạnh phúc")
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.runs[0].bold = True
    p5 = cell_r.add_paragraph(
        f"{dia_danh}, ngày {ngay_ky.day} tháng {ngay_ky.month} năm {ngay_ky.year}"
    )
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def _tao_word_01xln(du_lieu: dict) -> bytes:
    """Mẫu 01/XLN — Đơn đề nghị xử lý nợ (KH tự viết)."""
    doc = Document()
    _style_doc_xln(doc)
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    _add_header_xln(
        doc,
        dia_danh=du_lieu.get("dia_danh", ""),
        ngay_ky=du_lieu.get("ngay_ky", date.today()),
    )

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("ĐƠN ĐỀ NGHỊ XỬ LÝ NỢ")
    r.bold = True
    r.font.size = Pt(14)

    doc.add_paragraph(
        f"Kính gửi: Ngân hàng Chính sách xã hội {du_lieu.get('ten_nhcsxh','')}"
    )
    doc.add_paragraph()

    doc.add_paragraph(
        f"Tên tôi là: {du_lieu.get('ten_kh','')}\n"
        f"Hiện cư trú tại: {du_lieu.get('dia_chi','')}\n"
        f"Là thành viên của Tổ TK&VV: {du_lieu.get('ten_to','')} "
        f"do ông (bà): {du_lieu.get('to_truong','')} làm Tổ trưởng"
    )
    doc.add_paragraph(
        f"1. Theo HĐTD (sổ Vay vốn) số {du_lieu.get('so_ku','')}, "
        f"ngày {du_lieu.get('ngay_vay','')}, tôi có đứng tên vay vốn "
        f"chương trình {du_lieu.get('ten_ct','')} tại NHCSXH "
        f"{du_lieu.get('ten_nhcsxh','')}.\n"
        f"    Số tiền vay: {du_lieu.get('muc_vay','')} đồng; "
        f"Hạn trả nợ: {du_lieu.get('ngay_dh','')}; "
        f"Mục đích vay vốn: {du_lieu.get('muc_dich_vay','')}\n"
        f"    Hiện nay, tôi còn nợ Ngân hàng số tiền: "
        f"{du_lieu.get('tong_du_no','')} đồng\n"
        f"    (Trong đó: Nợ gốc: {du_lieu.get('du_no_goc','')} đồng; "
        f"Nợ lãi: {du_lieu.get('lai_ton','')} đồng)"
    )
    doc.add_paragraph(f"2. Trong thời gian vừa qua do:\n{du_lieu.get('nguyen_nhan','')}")
    doc.add_paragraph(
        "3. Số vốn, tài sản của dự án bị thiệt hại:\n"
        f"    - Số vốn và tài sản bị thiệt hại: {du_lieu.get('so_tien_thiet_hai','')} đồng\n"
        f"    - Tổng số vốn thực hiện dự án: {du_lieu.get('muc_vay','')} đồng\n"
        f"    - Mức độ thiệt hại: {du_lieu.get('muc_do_thiet_hai','')}%"
    )
    doc.add_paragraph(
        "4. Tình hình kinh tế, khả năng trả nợ sau khi gặp rủi ro:\n"
        f"{du_lieu.get('kha_nang_tra_no','')}"
    )
    doc.add_paragraph(
        f"Vậy tôi làm đơn này đề nghị NHCSXH {du_lieu.get('ten_nhcsxh','')} "
        f"xem xét {du_lieu.get('bien_phap','')} số nợ bị rủi ro, cụ thể:\n"
        f"    - Số tiền đề nghị: {du_lieu.get('so_tien_de_nghi','')} đồng\n"
        f"      (Nợ gốc: {du_lieu.get('du_no_goc','')} đồng; "
        f"Nợ lãi: {du_lieu.get('lai_ton','')} đồng)\n"
        f"    - Thời gian đề nghị: {du_lieu.get('so_thang','')} tháng\n"
        f"    - Kế hoạch trả nợ: {du_lieu.get('ke_hoach_tra_no','')}"
    )
    doc.add_paragraph(
        "Tôi xin cam đoan và chịu trách nhiệm trước pháp luật về nội dung "
        "kê khai trên đơn và các hồ sơ giấy tờ chứng minh là đúng."
    )

    doc.add_paragraph()
    p_ky = doc.add_paragraph()
    p_ky.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    d_ky = du_lieu.get("ngay_ky", date.today())
    p_ky.add_run(
        f"Ngày {d_ky.day} tháng {d_ky.month} năm {d_ky.year}\n"
        "Người làm đơn\n(Ký, ghi rõ họ tên)\n\n\n\n"
        f"{du_lieu.get('ten_kh','')}"
    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _tao_word_02xln(du_lieu: dict) -> bytes:
    """Mẫu 02/XLN — Biên bản đề nghị xử lý nợ bị rủi ro (nhiều bên ký)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document()
    _style_doc_xln(doc)
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    _add_header_xln(
        doc, dia_danh=du_lieu.get("dia_danh", ""), ngay_ky=du_lieu.get("ngay_lap", date.today())
    )
    if doc.tables:
        cell = doc.tables[0].rows[0].cells[0]
        if cell.paragraphs and cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].text = "Mẫu số 02/XLN"
            for rr in cell.paragraphs[0].runs[1:]:
                rr.text = ""
        else:
            cell.text = "Mẫu số 02/XLN"

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.add_run("BIÊN BẢN\nĐề nghị xử lý nợ bị rủi ro").bold = True
    p_ct = doc.add_paragraph(f"(Chương trình {du_lieu.get('ten_ct','')})")
    p_ct.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ngay_lap = du_lieu.get("ngay_lap", date.today())
    doc.add_paragraph(
        f"Hôm nay, ngày {ngay_lap.day} tháng {ngay_lap.month} "
        f"năm {ngay_lap.year}, tại {du_lieu.get('dia_diem','')}, "
        "chúng tôi gồm có:"
    )

    thanh_phan = du_lieu.get("thanh_phan", [])
    for i, tp in enumerate(thanh_phan, 1):
        doc.add_paragraph(
            f"{i}. Ông (bà) {tp.get('ho_ten','')}   "
            f"Chức vụ: {tp.get('chuc_vu','')}   "
            f"Đại diện: {tp.get('dai_dien','')}"
        )
    for i in range(len(thanh_phan) + 1, 8):
        doc.add_paragraph(
            f"{i}. Ông (bà) ....................................   "
            "Chức vụ: ....................   "
            "Đại diện: ......................"
        )

    doc.add_paragraph(
        "Đã tiến hành thẩm tra và lập biên bản đề nghị xử lý nợ bị rủi ro "
        f"của ông (bà): {du_lieu.get('ten_kh','')}  "
        f"địa chỉ: {du_lieu.get('dia_chi','')}\n"
        "Là đại diện hộ gia đình vay vốn NHCSXH theo HĐTD số "
        f"{du_lieu.get('so_ku','')} ngày {du_lieu.get('ngay_vay','')}, "
        f"có mã món vay: {du_lieu.get('so_ku','')}. Cụ thể như sau:"
    )

    p = doc.add_paragraph("I. Nguyên nhân khách hàng bị rủi ro:")
    p.runs[0].bold = True
    doc.add_paragraph(du_lieu.get("nguyen_nhan", ""))

    p = doc.add_paragraph("II. Xác định mức độ thiệt hại về vốn và tài sản:")
    p.runs[0].bold = True
    doc.add_paragraph(
        "1. Số vốn và tài sản bị thiệt hại: "
        f"{du_lieu.get('so_tien_thiet_hai','')} đồng\n"
        f"2. Tổng số vốn thực hiện dự án: {du_lieu.get('muc_vay','')} đồng\n"
        f"3. Đánh giá mức độ thiệt hại: {du_lieu.get('muc_do_thiet_hai','')}%"
    )

    p = doc.add_paragraph("III. Dư nợ tại NHCSXH đến ngày lập biên bản:")
    p.runs[0].bold = True
    doc.add_paragraph(
        f"Tổng số nợ còn phải trả: {du_lieu.get('tong_du_no','')} đồng\n"
        f"    Trong đó:  + Nợ gốc: {du_lieu.get('du_no_goc','')} đồng\n"
        f"               + Nợ lãi: {du_lieu.get('lai_ton','')} đồng"
    )

    p = doc.add_paragraph("IV. Đánh giá thực trạng dự án, tài sản và khả năng trả nợ:")
    p.runs[0].bold = True
    doc.add_paragraph(
        "1. Đánh giá thực trạng dự án / phương án khôi phục:\n"
        f"{du_lieu.get('thuc_trang_du_an','')}\n\n"
        "2. Tài sản hiện tại của khách hàng:\n"
        f"{du_lieu.get('tai_san_hien_tai','')}\n\n"
        "3. Đánh giá khả năng trả nợ:\n"
        f"{du_lieu.get('kha_nang_tra_no','')}\n\n"
        "4. Về việc áp dụng biện pháp thu hồi nợ:\n"
        f"{du_lieu.get('bien_phap_thu_hoi','')}"
    )

    p = doc.add_paragraph("V. Đề xuất biện pháp xử lý:")
    p.runs[0].bold = True
    doc.add_paragraph(
        f"Chúng tôi nhất trí đề nghị NHCSXH xem xét {du_lieu.get('bien_phap','')} "
        f"cho ông (bà) {du_lieu.get('ten_kh','')} với thời gian "
        f"{du_lieu.get('so_thang','')} tháng, số tiền "
        f"{du_lieu.get('so_tien_de_nghi','')} đồng.\n"
        f"    Trong đó:  + Nợ gốc: {du_lieu.get('du_no_goc','')} đồng\n"
        f"               + Nợ lãi: {du_lieu.get('lai_ton','')} đồng\n"
        "Biên bản này lập thành 02 bản có giá trị pháp lý như nhau."
    )

    doc.add_paragraph()
    ky = doc.add_table(rows=1, cols=3)
    ky.style = "Table Grid"
    for cell in ky.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for bn in ["top", "left", "bottom", "right"]:
            b = OxmlElement(f"w:{bn}")
            b.set(qn("w:val"), "none")
            tcBorders.append(b)
        tcPr.append(tcBorders)

    def _ky(cell, nhan: str, ten: str = ""):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"{nhan}\n(Ký, ghi rõ họ tên)\n\n\n\n{ten}")

    _ky(
        ky.rows[0].cells[0],
        "ĐẠI DIỆN KHÁCH HÀNG\nĐẠI DIỆN UBND CẤP XÃ",
        du_lieu.get("ten_kh", ""),
    )
    _ky(ky.rows[0].cells[1], "TỔ TRƯỞNG TỔ TK&VV\nĐẠI DIỆN HỘI ĐOÀN THỂ")
    _ky(
        ky.rows[0].cells[2],
        "CÁN BỘ TÍN DỤNG\nĐẠI DIỆN NHCSXH",
        du_lieu.get("can_bo_td", ""),
    )

    ky2 = doc.add_table(rows=1, cols=2)
    ky2.style = "Table Grid"
    for cell in ky2.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for bn in ["top", "left", "bottom", "right"]:
            b = OxmlElement(f"w:{bn}")
            b.set(qn("w:val"), "none")
            tcBorders.append(b)
        tcPr.append(tcBorders)
    _ky(ky2.rows[0].cells[0], "ĐẠI DIỆN CƠ QUAN CÔNG AN CẤP XÃ\n(Xác nhận, ký tên, đóng dấu)")
    _ky(ky2.rows[0].cells[1], "ĐẠI DIỆN TỔ CHỨC, CÁ NHÂN LIÊN QUAN (nếu có)")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _set_margins(
    doc: Document,
    left_cm: float = 3.0,
    right_cm: float = 2.0,
    top_cm: float = 2.5,
    bottom_cm: float = 2.5,
) -> None:
    for section in doc.sections:
        section.left_margin = Cm(left_cm)
        section.right_margin = Cm(right_cm)
        section.top_margin = Cm(top_cm)
        section.bottom_margin = Cm(bottom_cm)


def _loc_theo_nguon(df_rows: list[dict], nguon: int) -> list[dict]:
    """Lọc danh sách hồ sơ theo nguồn vốn (1=TW, 2=ĐP)."""
    return [r for r in df_rows if int(r.get("nguon_von", 0)) == nguon]


def _tong_hop_no(ds: list[dict]) -> dict:
    """Tổng hợp số liệu theo nhóm chương trình."""
    from collections import defaultdict

    nhom = defaultdict(lambda: {"so_ho": 0, "goc": 0.0, "lai": 0.0, "ds": []})
    for r in ds:
        ct = r.get("ten_ct", "Khác") or "Khác"
        nhom[ct]["so_ho"] += 1
        nhom[ct]["goc"] += float(r.get("du_no_goc", 0) or 0)
        nhom[ct]["lai"] += float(r.get("lai_ton", 0) or 0)
        nhom[ct]["ds"].append(r)
    tong_goc = sum(v["goc"] for v in nhom.values())
    tong_lai = sum(v["lai"] for v in nhom.values())
    return {
        "nhom_ct": dict(nhom),
        "tong_ho": sum(v["so_ho"] for v in nhom.values()),
        "tong_goc": tong_goc,
        "tong_lai": tong_lai,
        "tong_tien": tong_goc + tong_lai,
    }


def _tao_word_xln_bao_cao(
    tong_hop: dict,
    ten_pgd: str,
    nguon_label: str,
    so_qd: str,
    ngay_qd: date,
    ngay_bat_dau: date,
    ngay_ket_thuc: date,
    mau_so: str,
    tieu_de: str,
) -> bytes:
    doc = Document()
    _style_doc_xln(doc)
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width = Cm(29.7)
    sec.page_height = Cm(21.0)
    _set_margins(doc, left_cm=2.0, right_cm=2.0, top_cm=2.0, bottom_cm=2.0)

    p_ms = doc.add_paragraph()
    p_ms.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = p_ms.add_run(f"Mẫu số: {mau_so}\n")
    r1.italic = True
    r2 = p_ms.add_run("PGD gửi tỉnh, tỉnh tổng hợp gửi TW")
    r2.italic = True

    p = doc.add_paragraph("NGÂN HÀNG CHÍNH SÁCH XÃ HỘI")
    if p.runs:
        p.runs[0].bold = True
    if ten_pgd:
        p2 = doc.add_paragraph(_pgd_line(ten_pgd))
        if p2.runs:
            p2.runs[0].bold = True
    else:
        doc.add_paragraph("PGD")

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(tieu_de)
    r.bold = True

    ngay_qd_txt = ngay_qd.strftime("%d/%m/%Y")
    t2 = doc.add_paragraph(
        f"(theo Quyết định số {so_qd} ngày {ngay_qd_txt} của Chủ tịch Hội đồng quản trị NHCSXH)"
    )
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    tu_ngay = ngay_bat_dau.strftime("%d/%m/%Y")
    den_ngay = ngay_ket_thuc.strftime("%d/%m/%Y")
    t3 = doc.add_paragraph(f"Từ ngày {tu_ngay} đến ngày {den_ngay}")
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Nguồn vốn: {nguon_label}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Đơn vị tính: hộ, đồng").alignment = WD_ALIGN_PARAGRAPH.RIGHT

    tbl = doc.add_table(rows=3, cols=15)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr0 = tbl.rows[0].cells
    hdr1 = tbl.rows[1].cells
    hdr2 = tbl.rows[2].cells

    hdr0[0].merge(hdr2[0])
    hdr0[1].merge(hdr2[1])
    hdr0[14].merge(hdr2[14])

    hdr0[2].merge(hdr0[5])
    hdr0[6].merge(hdr0[9])
    hdr0[10].merge(hdr0[13])

    hdr1[2].merge(hdr2[2])
    hdr1[6].merge(hdr2[6])
    hdr1[10].merge(hdr2[10])

    hdr1[3].merge(hdr1[5])
    hdr1[7].merge(hdr1[9])
    hdr1[11].merge(hdr1[13])

    _set_cell(hdr0[0], "STT\n(1)", bold=True)
    _set_cell(hdr0[1], "Chi nhánh tỉnh, thành phố\n(2)", bold=True)
    _set_cell(hdr0[2], "Số được thông báo", bold=True)
    _set_cell(hdr0[6], "Số hạch toán thực tế", bold=True)
    _set_cell(hdr0[10], "Chênh lệch", bold=True)
    _set_cell(hdr0[14], "Nguyên nhân\nchênh lệch\n(15)", bold=True)

    _set_cell(hdr1[2], "Số hộ\n(3)", bold=True)
    _set_cell(hdr1[3], "Số tiền", bold=True)
    _set_cell(hdr1[6], "Số hộ\n(7)", bold=True)
    _set_cell(hdr1[7], "Số tiền", bold=True)
    _set_cell(hdr1[10], "Số hộ\n(11)", bold=True)
    _set_cell(hdr1[11], "Số tiền", bold=True)

    _set_cell(hdr2[3], "Tổng\n(4)", bold=True)
    _set_cell(hdr2[4], "Gốc\n(5)", bold=True)
    _set_cell(hdr2[5], "Lãi\n(6)", bold=True)
    _set_cell(hdr2[7], "Tổng\n(8)", bold=True)
    _set_cell(hdr2[8], "Gốc\n(9)", bold=True)
    _set_cell(hdr2[9], "Lãi\n(10)", bold=True)
    _set_cell(hdr2[11], "Tổng\n(12)", bold=True)
    _set_cell(hdr2[12], "Gốc\n(13)", bold=True)
    _set_cell(hdr2[13], "Lãi\n(14)", bold=True)

    _set_row_font(tbl.rows[0], 10)
    _set_row_font(tbl.rows[1], 10)
    _set_row_font(tbl.rows[2], 10)

    stt = 0
    nhom_ct = tong_hop.get("nhom_ct", {}) or {}
    for ten_ct, v in nhom_ct.items():
        stt += 1

        row_nhom = tbl.add_row()
        _set_row_font(row_nhom, 10)
        _set_cell(row_nhom.cells[0], str(stt), bold=True)
        cell_nhom = row_nhom.cells[1].merge(row_nhom.cells[-1])
        _set_cell(cell_nhom, str(ten_ct), bold=True, italic=True, align=WD_ALIGN_PARAGRAPH.LEFT)
        so_ho = int(v.get("so_ho", 0) or 0)
        goc = float(v.get("goc", 0) or 0)
        lai = float(v.get("lai", 0) or 0)
        tong = goc + lai


        row = tbl.add_row()
        _set_row_font(row, 10)
        _set_cell(row.cells[0], str(stt), align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell(row.cells[1], ten_pgd or "", align=WD_ALIGN_PARAGRAPH.LEFT)

        _set_cell(row.cells[2], str(so_ho), align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell(row.cells[3], fmt(tong), align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell(row.cells[4], fmt(goc), align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell(row.cells[5], fmt(lai), align=WD_ALIGN_PARAGRAPH.RIGHT)

        _set_cell(row.cells[6], str(so_ho), align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell(row.cells[7], fmt(tong), align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell(row.cells[8], fmt(goc), align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell(row.cells[9], fmt(lai), align=WD_ALIGN_PARAGRAPH.RIGHT)

        _set_cell(row.cells[10], "0", align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell(row.cells[11], "0", align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell(row.cells[12], "0", align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell(row.cells[13], "0", align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell(row.cells[14], "", align=WD_ALIGN_PARAGRAPH.LEFT)

    row_tong = tbl.add_row()
    _set_row_font(row_tong, 10)
    _set_cell(row_tong.cells[0], "", bold=True)
    _set_cell(row_tong.cells[1], "Tổng cộng:", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)

    tong_ho = int(tong_hop.get("tong_ho", 0) or 0)
    tong_goc = float(tong_hop.get("tong_goc", 0) or 0)
    tong_lai = float(tong_hop.get("tong_lai", 0) or 0)
    tong_tien = tong_goc + tong_lai

    _set_cell(row_tong.cells[2], str(tong_ho), bold=True)
    _set_cell(row_tong.cells[3], fmt(tong_tien), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell(row_tong.cells[4], fmt(tong_goc), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell(row_tong.cells[5], fmt(tong_lai), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)

    _set_cell(row_tong.cells[6], str(tong_ho), bold=True)
    _set_cell(row_tong.cells[7], fmt(tong_tien), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell(row_tong.cells[8], fmt(tong_goc), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell(row_tong.cells[9], fmt(tong_lai), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)

    _set_cell(row_tong.cells[10], "0", bold=True)
    _set_cell(row_tong.cells[11], "0", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell(row_tong.cells[12], "0", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell(row_tong.cells[13], "0", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell(row_tong.cells[14], "", bold=True)

    doc.add_paragraph()
    ky = doc.add_table(rows=1, cols=3)
    ky.style = "Table Grid"
    for cell in ky.rows[0].cells:
        _bo_border_cell(cell)
        _set_cell(cell, "", font_size=11)

    _set_cell(ky.rows[0].cells[0], "LẬP BIỂU\n\n\n\n", bold=True, font_size=11)
    _set_cell(ky.rows[0].cells[1], "KIỂM SOÁT\n\n\n\n", bold=True, font_size=11)
    _set_cell(ky.rows[0].cells[2], "GIÁM ĐỐC\n\n\n\n", bold=True, font_size=11)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _tao_word_13xln(
    tong_hop: dict,
    ten_pgd: str,
    nguon_label: str,
    so_qd: str,
    ngay_qd: date,
    ngay_bat_dau: date,
    ngay_ket_thuc: date,
) -> bytes:
    return _tao_word_xln_bao_cao(
        tong_hop=tong_hop,
        ten_pgd=ten_pgd,
        nguon_label=nguon_label,
        so_qd=so_qd,
        ngay_qd=ngay_qd,
        ngay_bat_dau=ngay_bat_dau,
        ngay_ket_thuc=ngay_ket_thuc,
        mau_so="13/XLN",
        tieu_de="BÁO CÁO CÁC KHOẢN NỢ SAU KHI THỰC HIỆN HẠCH TOÁN KHOANH NỢ",
    )


def _tao_word_14xln(
    tong_hop: dict,
    ten_pgd: str,
    nguon_label: str,
    so_qd: str,
    ngay_qd: date,
    ngay_bat_dau: date,
    ngay_ket_thuc: date,
) -> bytes:
    return _tao_word_xln_bao_cao(
        tong_hop=tong_hop,
        ten_pgd=ten_pgd,
        nguon_label=nguon_label,
        so_qd=so_qd,
        ngay_qd=ngay_qd,
        ngay_bat_dau=ngay_bat_dau,
        ngay_ket_thuc=ngay_ket_thuc,
        mau_so="14/XLN",
        tieu_de="BÁO CÁO CÁC KHOẢN NỢ SAU KHI THỰC HIỆN HẠCH TOÁN XÓA NỢ",
    )


def _tao_word_04xln(
    tong_hop: dict,
    ten_pgd: str,
    nguon_label: str,
    dot: int,
    nam: int,
) -> bytes:
    doc = Document()
    _style_doc_xln(doc)
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    _set_margins(doc, left_cm=3.0, right_cm=1.5, top_cm=2.0, bottom_cm=2.0)

    hdr = doc.add_table(rows=2, cols=2)
    hdr.style = "Table Grid"
    for row in hdr.rows:
        for cell in row.cells:
            _bo_border_cell(cell)

    p0l = hdr.rows[0].cells[0].paragraphs[0]
    p0l.add_run("NGÂN HÀNG CHÍNH SÁCH XÃ HỘI")
    p0r = hdr.rows[0].cells[1].paragraphs[0]
    p0r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r0r = p0r.add_run("Mẫu số: 04/XLN")
    r0r.italic = True

    hdr.rows[1].cells[0].paragraphs[0].add_run("Chi nhánh NHCSXH tỉnh Đồng Nai")
    p1r = hdr.rows[1].cells[1].paragraphs[0]
    p1r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1r = p1r.add_run("PGD gửi tỉnh, tỉnh tổng hợp gửi TW")
    r1r.italic = True

    doc.add_paragraph(_pgd_line(ten_pgd))

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("BIỂU TỔNG HỢP ĐỀ NGHỊ KHOANH NỢ ĐỐI VỚI KHÁCH HÀNG")
    r.bold = True
    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("VAY VỐN BỊ RỦI RO DO NGUYÊN NHÂN KHÁCH QUAN")
    r2.bold = True
    doc.add_paragraph(f"Đợt {dot} năm {nam} — Nguồn vốn: {nguon_label}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Đơn vị tính: đồng").alignment = WD_ALIGN_PARAGRAPH.RIGHT

    tbl = doc.add_table(rows=2, cols=13)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    h0 = tbl.rows[0].cells
    h1 = tbl.rows[1].cells
    h0[0].merge(h1[0])
    h0[1].merge(h1[1])
    h0[2].merge(h1[2])
    h0[3].merge(h1[3])
    h0[4].merge(h1[4])
    h0[5].merge(h1[5])
    h0[12].merge(h1[12])

    h0[6].merge(h0[8])
    h0[9].merge(h0[11])

    _set_cell(h0[0], "STT", bold=True)
    _set_cell(h0[1], "Chương trình; Huyện, thị xã;\nHọ và tên", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell(h0[2], "Địa chỉ (Xã, phường)", bold=True)
    _set_cell(h0[3], "Mã món vay", bold=True)
    _set_cell(h0[4], "Ngày vay", bold=True)
    _set_cell(h0[5], "Mức độ thiệt hại (%)", bold=True)
    _set_cell(h0[6], "Số dư nợ tại NHCS", bold=True)
    _set_cell(h0[9], "Số nợ đề nghị xử lý", bold=True)
    _set_cell(h0[12], "Ghi chú", bold=True)

    _set_cell(h1[6], "Số tiền", bold=True)
    _set_cell(h1[7], "Gốc", bold=True)
    _set_cell(h1[8], "Lãi", bold=True)
    _set_cell(h1[9], "Số tiền", bold=True)
    _set_cell(h1[10], "Gốc", bold=True)
    _set_cell(h1[11], "Lãi", bold=True)

    ds_all: list[dict] = []
    for v in (tong_hop.get("nhom_ct", {}) or {}).values():
        ds_all.extend(list(v.get("ds", []) or []))

    nhom_3nam = [r for r in ds_all if int(r.get("so_thang", 0) or 0) <= 36]
    nhom_5nam = [r for r in ds_all if int(r.get("so_thang", 0) or 0) > 36]
    th_3 = _tong_hop_no(nhom_3nam)
    th_5 = _tong_hop_no(nhom_5nam)

    def _roman(n: int) -> str:
        pairs = [
            (1000, "M"),
            (900, "CM"),
            (500, "D"),
            (400, "CD"),
            (100, "C"),
            (90, "XC"),
            (50, "L"),
            (40, "XL"),
            (10, "X"),
            (9, "IX"),
            (5, "V"),
            (4, "IV"),
            (1, "I"),
        ]
        out = ""
        x = n
        for v, s in pairs:
            while x >= v:
                out += s
                x -= v
        return out

    def _muc_do_txt(v: str) -> str:
        s = (v or "").strip()
        if not s or "Không" in s:
            return ""
        if "40%" in s and "80%" in s:
            return "40-<80"
        if "80%" in s and "100%" in s:
            return "80-100"
        return s

    def _render_nhom(tong_hop_nhom: dict, label: str) -> None:
        row_t = tbl.add_row()
        cell = row_t.cells[0].merge(row_t.cells[-1])
        _set_cell(
            cell,
            f"{label}  —  Số tiền {fmt(float(tong_hop_nhom.get('tong_tien', 0) or 0))}  "
            f"(Gốc {fmt(float(tong_hop_nhom.get('tong_goc', 0) or 0))}, "
            f"Lãi {fmt(float(tong_hop_nhom.get('tong_lai', 0) or 0))})",
            bold=True,
            align=WD_ALIGN_PARAGRAPH.LEFT,
        )

        row_a = tbl.add_row()
        cell_a = row_a.cells[0].merge(row_a.cells[-1])
        _set_cell(cell_a, "A  Khoanh nợ lần đầu", italic=True, align=WD_ALIGN_PARAGRAPH.LEFT)

        for i_ct, (ten_ct, v) in enumerate((tong_hop_nhom.get("nhom_ct", {}) or {}).items(), 1):
            goc_ct = float(v.get("goc", 0) or 0)
            lai_ct = float(v.get("lai", 0) or 0)
            tong_ct = goc_ct + lai_ct

            row_ct = tbl.add_row()
            _set_cell(row_ct.cells[0], _roman(i_ct), italic=True)
            _set_cell(row_ct.cells[1], str(ten_ct), italic=True, align=WD_ALIGN_PARAGRAPH.LEFT)
            for j in range(2, 6):
                _set_cell(row_ct.cells[j], "", align=WD_ALIGN_PARAGRAPH.LEFT)
            _set_cell(row_ct.cells[6], fmt(tong_ct), align=WD_ALIGN_PARAGRAPH.RIGHT)
            _set_cell(row_ct.cells[7], fmt(goc_ct), align=WD_ALIGN_PARAGRAPH.RIGHT)
            _set_cell(row_ct.cells[8], fmt(lai_ct), align=WD_ALIGN_PARAGRAPH.RIGHT)
            _set_cell(row_ct.cells[9], fmt(tong_ct), align=WD_ALIGN_PARAGRAPH.RIGHT)
            _set_cell(row_ct.cells[10], fmt(goc_ct), align=WD_ALIGN_PARAGRAPH.RIGHT)
            _set_cell(row_ct.cells[11], fmt(lai_ct), align=WD_ALIGN_PARAGRAPH.RIGHT)
            _set_cell(row_ct.cells[12], "", align=WD_ALIGN_PARAGRAPH.LEFT)

            ds_ct = list(v.get("ds", []) or [])
            for idx, r0 in enumerate(ds_ct, 1):
                goc = float(r0.get("du_no_goc", 0) or 0)
                lai = float(r0.get("lai_ton", 0) or 0)
                tong = goc + lai
                row = tbl.add_row()
                _set_cell(row.cells[0], str(idx))
                _set_cell(row.cells[1], str(r0.get("ten_kh", "")), align=WD_ALIGN_PARAGRAPH.LEFT)
                _set_cell(row.cells[2], str(r0.get("dia_chi", "")), align=WD_ALIGN_PARAGRAPH.LEFT)
                _set_cell(row.cells[3], str(r0.get("so_ku", "")))
                _set_cell(row.cells[4], str(r0.get("ngay_vay", "")))
                _set_cell(row.cells[5], _muc_do_txt(str(r0.get("muc_do", ""))))
                _set_cell(row.cells[6], fmt(tong), align=WD_ALIGN_PARAGRAPH.RIGHT)
                _set_cell(row.cells[7], fmt(goc), align=WD_ALIGN_PARAGRAPH.RIGHT)
                _set_cell(row.cells[8], fmt(lai), align=WD_ALIGN_PARAGRAPH.RIGHT)
                _set_cell(row.cells[9], fmt(tong), align=WD_ALIGN_PARAGRAPH.RIGHT)
                _set_cell(row.cells[10], fmt(goc), align=WD_ALIGN_PARAGRAPH.RIGHT)
                _set_cell(row.cells[11], fmt(lai), align=WD_ALIGN_PARAGRAPH.RIGHT)
                _set_cell(row.cells[12], str(r0.get("ghi_chu", "")), align=WD_ALIGN_PARAGRAPH.LEFT)

    _render_nhom(th_3, "Khoanh nợ tối đa 3 năm")
    _render_nhom(th_5, "Khoanh nợ tối đa 5 năm")

    row_tong = tbl.add_row()
    cell_t = row_tong.cells[0].merge(row_tong.cells[5])
    _set_cell(cell_t, "TỔNG CỘNG", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    tong_goc = float(tong_hop.get("tong_goc", 0) or 0)
    tong_lai = float(tong_hop.get("tong_lai", 0) or 0)
    tong_tien = tong_goc + tong_lai
    _set_cell(row_tong.cells[6], fmt(tong_tien), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell(row_tong.cells[7], fmt(tong_goc), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell(row_tong.cells[8], fmt(tong_lai), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell(row_tong.cells[9], fmt(tong_tien), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell(row_tong.cells[10], fmt(tong_goc), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell(row_tong.cells[11], fmt(tong_lai), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell(row_tong.cells[12], "", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)

    doc.add_paragraph()
    ky = doc.add_table(rows=2, cols=3)
    ky.style = "Table Grid"
    for row in ky.rows:
        for cell in row.cells:
            _bo_border_cell(cell)

    cell_ngay = ky.rows[0].cells[0].merge(ky.rows[0].cells[2])
    _set_cell(cell_ngay, f"{_pgd_plain(ten_pgd)}, ngày ... tháng ... năm ...", align=WD_ALIGN_PARAGRAPH.RIGHT, font_size=11)

    _set_cell(ky.rows[1].cells[0], "LẬP BIỂU\n(Ký, ghi rõ họ tên)", bold=True, font_size=11)
    _set_cell(ky.rows[1].cells[1], "KIỂM SOÁT\n(Ký, ghi rõ họ tên)", bold=True, font_size=11)
    _set_cell(ky.rows[1].cells[2], "GIÁM ĐỐC\n(Ký tên, đóng dấu)", bold=True, font_size=11)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _tao_word_05xln(
    tong_hop: dict,
    ten_pgd: str,
    nguon_label: str,
    dot: int,
    nam: int,
) -> bytes:
    doc = Document()
    _style_doc_xln(doc)
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    _set_margins(doc, left_cm=3.0, right_cm=1.5, top_cm=2.0, bottom_cm=2.0)

    hdr = doc.add_table(rows=2, cols=2)
    hdr.style = "Table Grid"
    for row in hdr.rows:
        for cell in row.cells:
            _bo_border_cell(cell)

    hdr.rows[0].cells[0].paragraphs[0].add_run("NGÂN HÀNG CHÍNH SÁCH XÃ HỘI")
    p0r = hdr.rows[0].cells[1].paragraphs[0]
    p0r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r0r = p0r.add_run("Mẫu số: 05/XLN")
    r0r.italic = True

    hdr.rows[1].cells[0].paragraphs[0].add_run("Chi nhánh NHCSXH tỉnh Đồng Nai")
    p1r = hdr.rows[1].cells[1].paragraphs[0]
    p1r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1r = p1r.add_run("PGD gửi tỉnh, tỉnh tổng hợp gửi TW")
    r1r.italic = True

    doc.add_paragraph(_pgd_line(ten_pgd))

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("BIỂU TỔNG HỢP ĐỀ NGHỊ XÓA NỢ ĐỐI VỚI KHÁCH HÀNG")
    r.bold = True
    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("VAY VỐN BỊ RỦI RO DO NGUYÊN NHÂN KHÁCH QUAN")
    r2.bold = True
    doc.add_paragraph(f"Đợt {dot} năm {nam} — Nguồn vốn: {nguon_label}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Đơn vị tính: đồng").alignment = WD_ALIGN_PARAGRAPH.RIGHT

    tbl = doc.add_table(rows=2, cols=13)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    h0 = tbl.rows[0].cells
    h1 = tbl.rows[1].cells
    h0[0].merge(h1[0])
    h0[1].merge(h1[1])
    h0[2].merge(h1[2])
    h0[3].merge(h1[3])
    h0[4].merge(h1[4])
    h0[5].merge(h1[5])
    h0[12].merge(h1[12])

    h0[6].merge(h0[8])
    h0[9].merge(h0[11])

    _set_cell(h0[0], "STT", bold=True)
    _set_cell(h0[1], "Chương trình; Huyện, thị xã;\nHọ và tên", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell(h0[2], "Địa chỉ (Xã, phường)", bold=True)
    _set_cell(h0[3], "Mã món vay", bold=True)
    _set_cell(h0[4], "Ngày vay", bold=True)
    _set_cell(h0[5], "Mức độ thiệt hại (%)", bold=True)
    _set_cell(h0[6], "Số dư nợ tại NHCS", bold=True)
    _set_cell(h0[9], "Số nợ đề nghị xử lý", bold=True)
    _set_cell(h0[12], "Ghi chú", bold=True)

    _set_cell(h1[6], "Số tiền", bold=True)
    _set_cell(h1[7], "Gốc", bold=True)
    _set_cell(h1[8], "Lãi", bold=True)
    _set_cell(h1[9], "Số tiền", bold=True)
    _set_cell(h1[10], "Gốc", bold=True)
    _set_cell(h1[11], "Lãi", bold=True)

    def _muc_do_txt(v: str) -> str:
        s = (v or "").strip()
        if not s or "Không" in s:
            return ""
        if "40%" in s and "80%" in s:
            return "40-<80"
        if "80%" in s and "100%" in s:
            return "80-100"
        return s

    for i_ct, (ten_ct, v) in enumerate((tong_hop.get("nhom_ct", {}) or {}).items(), 1):
        goc_ct = float(v.get("goc", 0) or 0)
        lai_ct = float(v.get("lai", 0) or 0)
        tong_ct = goc_ct + lai_ct
        row_ct = tbl.add_row()
        _set_cell(row_ct.cells[0], str(i_ct))
        _set_cell(row_ct.cells[1], str(ten_ct), italic=True, align=WD_ALIGN_PARAGRAPH.LEFT)
        for j in range(2, 6):
            _set_cell(row_ct.cells[j], "", align=WD_ALIGN_PARAGRAPH.LEFT)
        _set_cell(row_ct.cells[6], fmt(tong_ct), align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell(row_ct.cells[7], fmt(goc_ct), align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell(row_ct.cells[8], fmt(lai_ct), align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell(row_ct.cells[9], fmt(tong_ct), align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell(row_ct.cells[10], fmt(goc_ct), align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell(row_ct.cells[11], fmt(lai_ct), align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell(row_ct.cells[12], "", align=WD_ALIGN_PARAGRAPH.LEFT)

        ds_ct = list(v.get("ds", []) or [])
        for idx, r0 in enumerate(ds_ct, 1):
            goc = float(r0.get("du_no_goc", 0) or 0)
            lai = float(r0.get("lai_ton", 0) or 0)
            tong = goc + lai
            row = tbl.add_row()
            _set_cell(row.cells[0], str(idx))
            _set_cell(row.cells[1], str(r0.get("ten_kh", "")), align=WD_ALIGN_PARAGRAPH.LEFT)
            _set_cell(row.cells[2], str(r0.get("dia_chi", "")), align=WD_ALIGN_PARAGRAPH.LEFT)
            _set_cell(row.cells[3], str(r0.get("so_ku", "")))
            _set_cell(row.cells[4], str(r0.get("ngay_vay", "")))
            _set_cell(row.cells[5], _muc_do_txt(str(r0.get("muc_do", ""))))
            _set_cell(row.cells[6], fmt(tong), align=WD_ALIGN_PARAGRAPH.RIGHT)
            _set_cell(row.cells[7], fmt(goc), align=WD_ALIGN_PARAGRAPH.RIGHT)
            _set_cell(row.cells[8], fmt(lai), align=WD_ALIGN_PARAGRAPH.RIGHT)
            _set_cell(row.cells[9], fmt(tong), align=WD_ALIGN_PARAGRAPH.RIGHT)
            _set_cell(row.cells[10], fmt(goc), align=WD_ALIGN_PARAGRAPH.RIGHT)
            _set_cell(row.cells[11], fmt(lai), align=WD_ALIGN_PARAGRAPH.RIGHT)
            _set_cell(row.cells[12], str(r0.get("ghi_chu", "")), align=WD_ALIGN_PARAGRAPH.LEFT)

    row_tong = tbl.add_row()
    cell_t = row_tong.cells[0].merge(row_tong.cells[5])
    _set_cell(cell_t, "TỔNG CỘNG", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    tong_goc = float(tong_hop.get("tong_goc", 0) or 0)
    tong_lai = float(tong_hop.get("tong_lai", 0) or 0)
    tong_tien = tong_goc + tong_lai
    _set_cell(row_tong.cells[6], fmt(tong_tien), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell(row_tong.cells[7], fmt(tong_goc), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell(row_tong.cells[8], fmt(tong_lai), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell(row_tong.cells[9], fmt(tong_tien), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell(row_tong.cells[10], fmt(tong_goc), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell(row_tong.cells[11], fmt(tong_lai), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell(row_tong.cells[12], "", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)

    doc.add_paragraph()
    ky = doc.add_table(rows=2, cols=3)
    ky.style = "Table Grid"
    for row in ky.rows:
        for cell in row.cells:
            _bo_border_cell(cell)

    cell_ngay = ky.rows[0].cells[0].merge(ky.rows[0].cells[2])
    _set_cell(cell_ngay, f"{_pgd_plain(ten_pgd)}, ngày ... tháng ... năm ...", align=WD_ALIGN_PARAGRAPH.RIGHT, font_size=11)

    _set_cell(ky.rows[1].cells[0], "LẬP BIỂU\n(Ký, ghi rõ họ tên)", bold=True, font_size=11)
    _set_cell(ky.rows[1].cells[1], "KIỂM SOÁT\n(Ký, ghi rõ họ tên)", bold=True, font_size=11)
    _set_cell(ky.rows[1].cells[2], "GIÁM ĐỐC\n(Ký tên, đóng dấu)", bold=True, font_size=11)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _tao_word_to_trinh_pgd(
    tong_hop_khoanh: dict,
    tong_hop_xoa: dict,
    ds_khoanh: list[dict],
    ten_pgd: str,
    nguon_label: str,
    dot: int,
    nam: int,
) -> bytes:
    def _set_run(p, text: str, *, bold: bool = False, italic: bool = False) -> None:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    def _p(doc_or_cell, text: str = "", *, bold: bool = False, align=None) -> None:
        if hasattr(doc_or_cell, "add_paragraph"):
            p = doc_or_cell.add_paragraph()
        else:
            p = doc_or_cell.paragraphs[0]
            p.text = ""
        if align is not None:
            p.alignment = align
        _set_run(p, text, bold=bold)

    ten_pgd_plain = _pgd_plain(ten_pgd)
    ten_pgd_line = _pgd_line(ten_pgd)

    nhom_3nam = [r for r in ds_khoanh if int(r.get("so_thang", 0) or 0) <= 36]
    nhom_5nam = [r for r in ds_khoanh if int(r.get("so_thang", 0) or 0) > 36]
    th_3 = _tong_hop_no(nhom_3nam)
    th_5 = _tong_hop_no(nhom_5nam)

    doc = Document()
    _style_doc_xln(doc)
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    _set_margins(doc, left_cm=3.0, right_cm=2.0, top_cm=2.5, bottom_cm=2.5)

    ngay_hn = date.today()
    hdr = doc.add_table(rows=1, cols=2)
    hdr.style = "Table Grid"
    for cell in hdr.rows[0].cells:
        _bo_border_cell(cell)

    cell_l = hdr.rows[0].cells[0]
    cell_r = hdr.rows[0].cells[1]

    p_l = cell_l.paragraphs[0]
    p_l.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_run(p_l, "CHI NHÁNH NHCSXH TỈNH ĐỒNG NAI", bold=True)
    p_l2 = cell_l.add_paragraph()
    _set_run(p_l2, ten_pgd_line, bold=True)
    p_l3 = cell_l.add_paragraph()
    _set_run(p_l3, "Số: .../TTr-NHCS", bold=False)

    p_r = cell_r.paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(p_r, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", bold=True)
    p_r2 = cell_r.add_paragraph()
    p_r2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(p_r2, "Độc lập - Tự do - Hạnh phúc", bold=True)
    p_r3 = cell_r.add_paragraph()
    p_r3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(p_r3, f"{ten_pgd_plain}, ngày ... tháng ... năm ...", bold=False)

    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(t, "TỜ TRÌNH", bold=True)
    t.runs[0].font.size = Pt(14)
    doc.add_paragraph(f"Về việc đề nghị xử lý nợ bị rủi ro đợt {dot} năm {nam}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Nguồn vốn: {nguon_label}").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph("Kính gửi: Giám đốc Chi nhánh Ngân hàng Chính sách xã hội tỉnh Đồng Nai")

    for line in [
        "Thực hiện Quyết định số 62/QĐ-HĐQT ngày 27/9/2021 của Chủ tịch Hội đồng",
        "quản trị NHCSXH về việc Ban hành Quy định xử lý nợ bị rủi ro trong hệ thống",
        "NHCSXH và các Quyết định sửa đổi, bổ sung Quyết định số 62/QĐ-HĐQT của Hội",
        "đồng quản trị NHCSXH.",
        "",
        "Căn cứ tình hình nợ bị rủi ro tại Phòng giao dịch; Sau khi nhận được hồ sơ",
        "đề nghị xử lý rủi ro của khách hàng, Phòng giao dịch NHCSXH "
        f"{ten_pgd_plain} đã",
        "phối hợp với khách hàng, cá nhân, tổ chức có thẩm quyền theo quy định tại",
        "Điều 7 Quyết định số 62/QĐ-HĐQT tổ chức kiểm tra thực tế, kiểm tra điều",
        "kiện, tính đầy đủ, chính xác, hợp pháp, hợp lệ của toàn bộ khoản nợ đề nghị",
        f"xử lý rủi ro đợt {dot} năm {nam}. Phòng giao dịch NHCSXH {ten_pgd_plain} báo cáo",
        "và trình Giám đốc Chi nhánh NHCSXH tỉnh Đồng Nai các nội dung như sau:",
    ]:
        if line:
            doc.add_paragraph(line)
        else:
            doc.add_paragraph()

    so_mon_khoanh = int(tong_hop_khoanh.get("tong_ho", 0) or 0)
    so_mon_xoa = int(tong_hop_xoa.get("tong_ho", 0) or 0)
    so_mon_tong = so_mon_khoanh + so_mon_xoa

    goc_khoanh = float(tong_hop_khoanh.get("tong_goc", 0) or 0)
    lai_khoanh = float(tong_hop_khoanh.get("tong_lai", 0) or 0)
    tien_khoanh = goc_khoanh + lai_khoanh

    goc_xoa = float(tong_hop_xoa.get("tong_goc", 0) or 0)
    lai_xoa = float(tong_hop_xoa.get("tong_lai", 0) or 0)
    tien_xoa = goc_xoa + lai_xoa

    goc_tong = goc_khoanh + goc_xoa
    lai_tong = lai_khoanh + lai_xoa
    tien_tong = goc_tong + lai_tong

    doc.add_paragraph(
        f"1. Tổng số đề nghị xử lý rủi ro đợt {dot} năm {nam}: {so_mon_tong} món vay, số"
        f" tiền {fmt(tien_tong)} đồng (gốc {fmt(goc_tong)} đồng, lãi {fmt(lai_tong)} đồng). Cụ thể:"
    )

    if so_mon_khoanh:
        so_mon_3 = int(th_3.get("tong_ho", 0) or 0)
        so_mon_5 = int(th_5.get("tong_ho", 0) or 0)
        goc_3 = float(th_3.get("tong_goc", 0) or 0)
        lai_3 = float(th_3.get("tong_lai", 0) or 0)
        goc_5 = float(th_5.get("tong_goc", 0) or 0)
        lai_5 = float(th_5.get("tong_lai", 0) or 0)
        doc.add_paragraph(
            f"a) Khoanh nợ: {so_mon_khoanh} món vay, số tiền {fmt(tien_khoanh)} đồng, gốc"
            f" {fmt(goc_khoanh)} đồng, lãi {fmt(lai_khoanh)} đồng. Trong đó:"
        )
        doc.add_paragraph(
            f"+ Khoanh nợ tối đa 3 năm: {so_mon_3} món vay, số tiền {fmt(goc_3 + lai_3)}"
            f" đồng, gốc {fmt(goc_3)} đồng, lãi {fmt(lai_3)} đồng."
        )
        doc.add_paragraph(
            f"+ Khoanh nợ tối đa 5 năm: {so_mon_5} món vay, số tiền {fmt(goc_5 + lai_5)}"
            f" đồng, gốc {fmt(goc_5)} đồng, lãi {fmt(lai_5)} đồng."
        )

    if so_mon_xoa:
        doc.add_paragraph(
            f"b) Xóa nợ: {so_mon_xoa} món vay, số tiền {fmt(tien_xoa)} đồng, gốc"
            f" {fmt(goc_xoa)} đồng, lãi {fmt(lai_xoa)} đồng."
        )

    doc.add_paragraph("(Biểu tổng hợp đề nghị xử lý nợ mẫu số 03, 04, 05/XLN đính kèm)")

    doc.add_paragraph(
        "2. Các khoản nợ bị rủi ro đề nghị xử lý đợt "
        f"{dot} năm {nam} của Phòng giao"
        " dịch đảm bảo đủ điều kiện, đúng thực tế, hồ sơ được thiết lập đúng và đầy đủ"
        " theo quy định."
    )
    doc.add_paragraph(
        "3. Phòng giao dịch chịu trách nhiệm trước pháp luật, trước Tổng Giám đốc về"
        " tính đầy đủ, chính xác, hợp pháp, hợp lệ, đúng thực tế của toàn bộ hồ sơ và"
        f" số liệu đề nghị xử lý rủi ro đợt {dot} năm {nam}."
    )

    doc.add_paragraph()
    doc.add_paragraph(
        f"Phòng giao dịch NHCSXH {ten_pgd_plain} kính trình Giám đốc chi nhánh NHCSXH tỉnh"
        " Đồng Nai trình cấp có thẩm quyền xem xét, quyết định./."
    )

    doc.add_paragraph()
    ky = doc.add_table(rows=1, cols=2)
    ky.style = "Table Grid"
    for cell in ky.rows[0].cells:
        _bo_border_cell(cell)

    ky.rows[0].cells[0].paragraphs[0].add_run("Nơi nhận:\n- Như kính gửi;\n- Lưu VT")
    pky = ky.rows[0].cells[1].paragraphs[0]
    pky.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(pky, "GIÁM ĐỐC", bold=True)
    pky2 = ky.rows[0].cells[1].add_paragraph()
    pky2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(pky2, "(Ký tên, đóng dấu)", bold=False)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _tao_word_to_trinh_cn(
    tong_hop_khoanh: dict,
    tong_hop_xoa: dict,
    ds_khoanh: list[dict],
    ten_tinh: str,
    nguon_label: str,
    dot: int,
    nam: int,
) -> bytes:
    def _set_run(p, text: str, *, bold: bool = False, italic: bool = False) -> None:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    nhom_3nam = [r for r in ds_khoanh if int(r.get("so_thang", 0) or 0) <= 36]
    nhom_5nam = [r for r in ds_khoanh if int(r.get("so_thang", 0) or 0) > 36]
    th_3 = _tong_hop_no(nhom_3nam)
    th_5 = _tong_hop_no(nhom_5nam)

    doc = Document()
    _style_doc_xln(doc)
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    _set_margins(doc, left_cm=3.0, right_cm=2.0, top_cm=2.5, bottom_cm=2.5)

    hdr = doc.add_table(rows=1, cols=2)
    hdr.style = "Table Grid"
    for cell in hdr.rows[0].cells:
        _bo_border_cell(cell)

    cell_l = hdr.rows[0].cells[0]
    cell_r = hdr.rows[0].cells[1]

    p_l = cell_l.paragraphs[0]
    p_l.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_run(p_l, "NGÂN HÀNG CHÍNH SÁCH XÃ HỘI", bold=True)
    p_l2 = cell_l.add_paragraph()
    _set_run(p_l2, f"CHI NHÁNH TỈNH {ten_tinh.upper()}", bold=True)
    p_l3 = cell_l.add_paragraph()
    _set_run(p_l3, "Số: .../TTr-NHCS", bold=False)

    p_r = cell_r.paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(p_r, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", bold=True)
    p_r2 = cell_r.add_paragraph()
    p_r2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(p_r2, "Độc lập - Tự do - Hạnh phúc", bold=True)
    p_r3 = cell_r.add_paragraph()
    p_r3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(p_r3, f"{ten_tinh}, ngày ... tháng ... năm ...", bold=False)

    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(t, "TỜ TRÌNH", bold=True)
    t.runs[0].font.size = Pt(14)
    doc.add_paragraph(f"Về việc đề nghị xử lý nợ bị rủi ro đợt {dot} năm {nam}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Nguồn vốn: {nguon_label}").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph("Kính gửi: Tổng Giám đốc Ngân hàng Chính sách xã hội")

    for line in [
        "Thực hiện Quyết định số 62/QĐ-HĐQT ngày 27/9/2021 của Chủ tịch Hội đồng",
        "quản trị NHCSXH về việc Ban hành Quy định xử lý nợ bị rủi ro trong hệ thống",
        "NHCSXH và các Quyết định sửa đổi, bổ sung Quyết định số 62/QĐ-HĐQT của Hội",
        "đồng quản trị NHCSXH.",
        "",
        "Căn cứ tình hình nợ bị rủi ro tại chi nhánh; Sau khi nhận được hồ sơ",
        "đề nghị xử lý rủi ro của khách hàng, chi nhánh NHCSXH tỉnh đã tổ chức",
        "kiểm tra thực tế khách hàng, kiểm tra hồ sơ của toàn bộ khoản nợ... "
        f"Chi nhánh NHCSXH tỉnh {ten_tinh} kính trình Tổng Giám đốc các nội dung như sau:",
    ]:
        if line:
            doc.add_paragraph(line)
        else:
            doc.add_paragraph()

    so_mon_khoanh = int(tong_hop_khoanh.get("tong_ho", 0) or 0)
    so_mon_xoa = int(tong_hop_xoa.get("tong_ho", 0) or 0)
    so_mon_tong = so_mon_khoanh + so_mon_xoa

    goc_khoanh = float(tong_hop_khoanh.get("tong_goc", 0) or 0)
    lai_khoanh = float(tong_hop_khoanh.get("tong_lai", 0) or 0)
    tien_khoanh = goc_khoanh + lai_khoanh

    goc_xoa = float(tong_hop_xoa.get("tong_goc", 0) or 0)
    lai_xoa = float(tong_hop_xoa.get("tong_lai", 0) or 0)
    tien_xoa = goc_xoa + lai_xoa

    goc_tong = goc_khoanh + goc_xoa
    lai_tong = lai_khoanh + lai_xoa
    tien_tong = goc_tong + lai_tong

    doc.add_paragraph(
        f"1. Tổng số đề nghị xử lý rủi ro đợt {dot} năm {nam}: {so_mon_tong} món vay, số"
        f" tiền {fmt(tien_tong)} đồng (gốc {fmt(goc_tong)} đồng, lãi {fmt(lai_tong)} đồng). Cụ thể:"
    )

    if so_mon_khoanh:
        so_mon_3 = int(th_3.get("tong_ho", 0) or 0)
        so_mon_5 = int(th_5.get("tong_ho", 0) or 0)
        goc_3 = float(th_3.get("tong_goc", 0) or 0)
        lai_3 = float(th_3.get("tong_lai", 0) or 0)
        goc_5 = float(th_5.get("tong_goc", 0) or 0)
        lai_5 = float(th_5.get("tong_lai", 0) or 0)
        doc.add_paragraph(
            f"a) Khoanh nợ: {so_mon_khoanh} món vay, số tiền {fmt(tien_khoanh)} đồng, gốc"
            f" {fmt(goc_khoanh)} đồng, lãi {fmt(lai_khoanh)} đồng. Trong đó:"
        )
        doc.add_paragraph(
            f"+ Khoanh nợ tối đa 3 năm: {so_mon_3} món vay, số tiền {fmt(goc_3 + lai_3)}"
            f" đồng, gốc {fmt(goc_3)} đồng, lãi {fmt(lai_3)} đồng."
        )
        doc.add_paragraph(
            f"+ Khoanh nợ tối đa 5 năm: {so_mon_5} món vay, số tiền {fmt(goc_5 + lai_5)}"
            f" đồng, gốc {fmt(goc_5)} đồng, lãi {fmt(lai_5)} đồng."
        )

    if so_mon_xoa:
        doc.add_paragraph(
            f"b) Xóa nợ: {so_mon_xoa} món vay, số tiền {fmt(tien_xoa)} đồng, gốc"
            f" {fmt(goc_xoa)} đồng, lãi {fmt(lai_xoa)} đồng."
        )

    doc.add_paragraph("(Biểu tổng hợp đề nghị xử lý nợ mẫu số 03, 04, 05/XLN đính kèm)")

    doc.add_paragraph(
        "2. Các khoản nợ bị rủi ro đề nghị xử lý đợt "
        f"{dot} năm {nam} của chi nhánh đảm bảo đủ điều kiện, đúng thực tế, hồ sơ được"
        " thiết lập đúng và đầy đủ theo quy định."
    )
    doc.add_paragraph(
        "3. Chi nhánh chịu trách nhiệm trước pháp luật, trước Tổng Giám đốc về tính"
        " đầy đủ, chính xác, hợp pháp, hợp lệ, đúng thực tế của toàn bộ hồ sơ và số"
        f" liệu đề nghị xử lý rủi ro đợt {dot} năm {nam}."
    )

    doc.add_paragraph()
    doc.add_paragraph(
        f"Chi nhánh NHCSXH tỉnh {ten_tinh} kính trình Tổng Giám đốc trình cấp có thẩm quyền"
        " xem xét, quyết định./."
    )

    doc.add_paragraph()
    ky = doc.add_table(rows=1, cols=2)
    ky.style = "Table Grid"
    for cell in ky.rows[0].cells:
        _bo_border_cell(cell)

    ky.rows[0].cells[0].paragraphs[0].add_run("Nơi nhận:\n- Như kính gửi;\n- Lưu VT")
    pky = ky.rows[0].cells[1].paragraphs[0]
    pky.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(pky, "GIÁM ĐỐC CHI NHÁNH", bold=True)
    pky2 = ky.rows[0].cells[1].add_paragraph()
    pky2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(pky2, "(Ký tên, đóng dấu)", bold=False)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _lay_pgd_tu_user(role: str, pgd_user: str | None, df: pd.DataFrame) -> str | None:
    if pgd_user:
        return pgd_user
    if la_phan_he_cn(role) and df is not None and COT_TEN_PGD in df.columns:
        ds = df[COT_TEN_PGD].dropna().unique().tolist()
        if len(ds) == 1:
            return str(ds[0])
    return None


def _loc_df_theo_pgd(df: pd.DataFrame, role: str, pgd_user: str | None) -> pd.DataFrame:
    if df is None or df.empty:
        return df
    if la_phan_he_pgd(role) and pgd_user and COT_TEN_PGD in df.columns:
        return df[df[COT_TEN_PGD] == pgd_user].copy()
    return df


def _tao_kv_key(ten_pgd: str) -> str:
    now = datetime.now()
    return f"no_rui_ro_{pgd_slug(ten_pgd)}_{now.year}_{now.month:02d}"


def _hien_thi_chi_tiet(ds: list[dict]) -> None:
    if not ds:
        st.info("ℹ️ Chưa có hồ sơ nào.")
        return
    df_xem = pd.DataFrame(ds)
    cols_xem = [c for c in [
        "ten_kh", "so_ku", "ten_ct", "du_no", "bien_phap",
        "nguyen_nhan", "muc_do", "so_thang", "ngay_rr", "ghi_chu",
    ] if c in df_xem.columns]
    if "du_no" in df_xem.columns:
        df_xem["du_no"] = df_xem["du_no"].apply(lambda x: fmt(x) if pd.notna(x) else "")
    st.dataframe(df_xem[cols_xem], use_container_width=True, hide_index=True)


# ── Hàm tách từ render() — Bước 4: xuất 04/05 XLN + Tờ trình ─────────
def _render_04_05_tt(
    ds_khoanh, ds_xoa, ten_don_vi, nguon_label,
    key_prefix, dot_xuat, nam_xuat, la_cn=False,
) -> None:
    ngay_hom_nay = date.today()
    st.markdown(f"**📄 04/XLN — Tổng hợp đề nghị khoanh nợ ({nguon_label})**")
    if st.button(f"📥 Xuất 04/XLN — {nguon_label}", use_container_width=True,
                 key=f"{key_prefix}_04xln"):
        if not ds_khoanh:
            st.warning("⚠️ Không có hồ sơ khoanh nợ.")
        else:
            docx_b = _tao_word_04xln(_tong_hop_no(ds_khoanh), ten_don_vi,
                                      nguon_label, dot_xuat, nam_xuat)
            nut_tai_word_va_pdf(docx_b,
                f"Mau04XLN_{nguon_label[:2]}_{ten_don_vi}_{ngay_hom_nay:%d%m%Y}",
                f"{key_prefix}_04xln")
    hien_thi_nut_tai(f"{key_prefix}_04xln")

    st.markdown(f"**📄 05/XLN — Tổng hợp đề nghị xóa nợ ({nguon_label})**")
    if st.button(f"📥 Xuất 05/XLN — {nguon_label}", use_container_width=True,
                 key=f"{key_prefix}_05xln"):
        if not ds_xoa:
            st.warning("⚠️ Không có hồ sơ xóa nợ.")
        else:
            docx_b = _tao_word_05xln(_tong_hop_no(ds_xoa), ten_don_vi,
                                      nguon_label, dot_xuat, nam_xuat)
            nut_tai_word_va_pdf(docx_b,
                f"Mau05XLN_{nguon_label[:2]}_{ten_don_vi}_{ngay_hom_nay:%d%m%Y}",
                f"{key_prefix}_05xln")
    hien_thi_nut_tai(f"{key_prefix}_05xln")

    ten_tt = "02/TT" if la_cn else "01/TT"
    st.markdown(f"**📄 Tờ trình {ten_tt} ({nguon_label})**")
    if st.button(f"📥 Xuất Tờ trình — {nguon_label}", use_container_width=True,
                 key=f"{key_prefix}_tt"):
        if not ds_khoanh and not ds_xoa:
            st.warning("⚠️ Không có hồ sơ nào.")
        else:
            if la_cn:
                docx_b = _tao_word_to_trinh_cn(
                    _tong_hop_no(ds_khoanh), _tong_hop_no(ds_xoa),
                    ds_khoanh, "Đồng Nai", nguon_label, dot_xuat, nam_xuat,
                )
            else:
                docx_b = _tao_word_to_trinh_pgd(
                    _tong_hop_no(ds_khoanh), _tong_hop_no(ds_xoa),
                    ds_khoanh, ten_don_vi, nguon_label, dot_xuat, nam_xuat,
                )
            nut_tai_word_va_pdf(docx_b,
                f"ToTrinh{ten_tt.replace('/','')}_{nguon_label[:2]}_{ten_don_vi}_{ngay_hom_nay:%d%m%Y}",
                f"{key_prefix}_tt")
    hien_thi_nut_tai(f"{key_prefix}_tt")


# ── Hàm tách từ render() — Bước 1→5 (luồng nhập hồ sơ) ──────────────
def _render_luong_nhap_ho_so(
    df_pgd: pd.DataFrame,
    ten_pgd: str,
    kv_key: str,
    username: str,
    la_cn: bool = False,
    key_prefix: str = "",
) -> None:
    """Luồng nhập hồ sơ rủi ro 5 bước — dùng chung cho PGD và CN.
    key_prefix: thêm vào đầu mọi st.* key để tránh conflict widget.
    la_cn: True → Tờ trình 02/TT + kv key riêng của CN.
    """
    df = df_pgd

    # ── Bước 1: Lọc hộ vay ──────────────────────────────────────────
    with st.expander("🔎 Lọc hộ vay", expanded=True):
        c1, c2, c3 = st.columns(3)
        with c1:
            ds_xa = sorted(df[COT_TEN_XA].dropna().unique().tolist()) if COT_TEN_XA in df.columns else []
            chon_xa = st.selectbox("Xã/Phường", [""] + ds_xa, key=f"{key_prefix}nrr_xa")
        with c2:
            df_loc = df[df[COT_TEN_XA] == chon_xa] if chon_xa and COT_TEN_XA in df.columns else df
            ds_to = sorted(df_loc[COT_TEN_TO].dropna().unique().tolist()) if COT_TEN_TO in df_loc.columns else []
            chon_to = st.selectbox("Tổ TK&VV", [""] + ds_to, key=f"{key_prefix}nrr_to")
        with c3:
            tim_kh = st.text_input("Tìm tên KH", placeholder="Nhập tên...", key=f"{key_prefix}nrr_tim")

    df_hien = df.copy()
    if chon_xa and COT_TEN_XA in df_hien.columns:
        df_hien = df_hien[df_hien[COT_TEN_XA] == chon_xa]
    if chon_to and COT_TEN_TO in df_hien.columns:
        df_hien = df_hien[df_hien[COT_TEN_TO] == chon_to]
    if tim_kh and COT_TEN_KH in df_hien.columns:
        df_hien = df_hien[df_hien[COT_TEN_KH].str.contains(tim_kh, case=False, na=False)]

    if df_hien.empty:
        st.info("ℹ️ Không tìm thấy hộ vay nào phù hợp.")
        return

    # ── Bước 2: Bảng chọn hộ vay ────────────────────────────────────
    st.markdown("#### 📋 Danh sách hộ vay")
    cot_hien = [c for c in [COT_TEN_KH, COT_SO_KU, COT_TEN_CT, COT_TONG_DU_NO, COT_DU_NO_QH] if c in df_hien.columns]
    df_editor = df_hien[cot_hien].copy()
    for c in [COT_TONG_DU_NO, COT_DU_NO_QH]:
        if c in df_editor.columns:
            df_editor[c] = df_editor[c].apply(lambda x: fmt(x) if pd.notna(x) else "")
    df_editor.insert(0, "Chọn", False)
    edited = st.data_editor(
        df_editor,
        use_container_width=True,
        hide_index=True,
        height=300,
        column_config={"Chọn": st.column_config.CheckboxColumn("Chọn")},
        key=f"{key_prefix}nrr_editor",
    )
    ds_chon = edited[edited["Chọn"] == True]
    if ds_chon.empty:
        st.info("👆 Tích chọn ít nhất 1 hộ vay để nhập thông tin rủi ro.")
        return
    st.success(f"✅ Đã chọn **{len(ds_chon)}** hộ vay.")

    # ── Bước 3: Form nhập thông tin rủi ro ──────────────────────────
    st.markdown("#### 📝 Thông tin rủi ro")
    with st.form(f"{key_prefix}form_no_rui_ro"):
        col1, col2 = st.columns(2)
        with col1:
            bien_phap = st.selectbox(
                "Biện pháp xử lý",
                ["Khoanh nợ (QĐ62)", "Xóa nợ (QĐ62)"],
                key=f"{key_prefix}nrr_bien_phap",
            )
            nguyen_nhan = st.selectbox(
                "Nguyên nhân rủi ro",
                NGUYEN_NHAN_RR,
                key=f"{key_prefix}nrr_nguyen_nhan",
            )
        with col2:
            ngay_rr = st.date_input(
                "Ngày xảy ra rủi ro",
                value=date.today(),
                key=f"{key_prefix}nrr_ngay_rr",
            )
        muc_do = ""
        so_thang = 0
        if "Khoanh nợ" in bien_phap:
            st.markdown("**Mức độ thiệt hại (khoanh nợ)**")
            mc1, mc2 = st.columns(2)
            with mc1:
                muc_do = st.radio(
                    "Mức độ thiệt hại",
                    ["Từ 40% đến <80%", "Từ 80% đến 100%", "Không áp dụng"],
                    key=f"{key_prefix}nrr_muc_do",
                )
            with mc2:
                goi_y = 60 if "80%" in muc_do else 36
                so_thang = st.number_input(
                    "Số tháng đề nghị khoanh",
                    min_value=0, max_value=120, value=goi_y, step=6,
                    key=f"{key_prefix}nrr_so_thang",
                    help=f"Gợi ý: {goi_y} tháng theo mức độ đã chọn",
                )
        ghi_chu = st.text_area(
            "Ghi chú / Tóm tắt nguyên nhân",
            placeholder="Nhập tối thiểu 20 ký tự...",
            height=100,
            key=f"{key_prefix}nrr_ghi_chu",
        )
        submitted = st.form_submit_button("💾 Lưu hồ sơ", type="primary")

    if submitted:
        if len(ghi_chu.strip()) < 20:
            st.error("⚠️ Ghi chú phải có ít nhất 20 ký tự.")
            st.stop()

        ds_luu = []
        for _, row in ds_chon.iterrows():
            so_ku_r = str(row.get(COT_SO_KU, ""))
            row_full = row
            if so_ku_r and df is not None and not df.empty and COT_SO_KU in df.columns:
                df_tmp = df[df[COT_SO_KU].astype(str) == so_ku_r]
                if not df_tmp.empty:
                    row_full = df_tmp.iloc[0]
            ds_luu.append({
                "ma_kh":   so_ku_r,
                "ten_kh":  str(row_full.get(COT_TEN_KH, "")),
                "so_ku":   so_ku_r,
                "ten_ct":  str(row_full.get(COT_TEN_CT, "")),
                "du_no":   _num(row_full.get(COT_TONG_DU_NO, 0) or 0),
                "dia_chi": str(row_full.get(COT_DIA_CHI, "")),
                "ngay_vay": fmt_ngay(row_full.get(COT_NGAY_VAY, "")),
                "du_no_goc": _num(row_full.get(COT_TONG_DU_NO, 0) or 0),
                "lai_ton": _num(row_full.get(COT_LAI_TON, 0) or 0),
                "nguon_von": int(row_full.get(COT_NGUON_VON, 0) or 0),
                "bien_phap":   bien_phap,
                "nguyen_nhan": nguyen_nhan,
                "muc_do":      muc_do,
                "so_thang":    int(so_thang),
                "ngay_rr":     ngay_rr.isoformat(),
                "ghi_chu":     ghi_chu.strip(),
            })
        db.ghi_kv(kv_key, {"danh_sach": ds_luu, "ngay_tao": datetime.now().isoformat()}, username)
        db.ghi_audit(username, "luu_no_rui_ro", f"{len(ds_luu)} hồ sơ — {ten_pgd or 'unknown'}")
        st.cache_data.clear()
        st.success(f"✅ Đã lưu **{len(ds_luu)}** hồ sơ xử lý nợ rủi ro.")
        st.balloons()

    # ── Bước 4: Xuất biểu mẫu ───────────────────────────────────────
    if ds_chon is not None and not ds_chon.empty:
        st.markdown("#### 📄 Xuất biểu mẫu")
        st.info("Soạn mẫu 01/02 XLN đã được chuyển sang phần '🧾 Soạn mẫu 01/XLN và 02/XLN' ở phía trên.")

        ds_xuat_full = []
        so_hs_khong_nguon = 0
        for _, row in ds_chon.iterrows():
            so_ku_r = str(row.get(COT_SO_KU, ""))
            row_full = row
            if so_ku_r and df is not None and COT_SO_KU in df.columns:
                df_tmp = df[df[COT_SO_KU].astype(str) == so_ku_r]
                if not df_tmp.empty:
                    row_full = df_tmp.iloc[0]
            try:
                nguon_von_int = int(row_full.get(COT_NGUON_VON, 0) or 0)
            except (ValueError, TypeError):
                nguon_von_int = 0
            if nguon_von_int not in (NGUON_TW, NGUON_DP):
                so_hs_khong_nguon += 1
                nguon_von_int = NGUON_TW
            ds_xuat_full.append({
                "ten_ct":    str(row_full.get(COT_TEN_CT, "")),
                "ten_kh":    str(row_full.get(COT_TEN_KH, "")),
                "dia_chi":   str(row_full.get(COT_DIA_CHI, "")),
                "so_ku":     so_ku_r,
                "ngay_vay":  fmt_ngay(row_full.get(COT_NGAY_VAY, "")),
                "du_no_goc": float(row_full.get(COT_TONG_DU_NO, 0) or 0),
                "lai_ton":   float(row_full.get(COT_LAI_TON, 0) or 0),
                "bien_phap": bien_phap,
                "muc_do":    muc_do,
                "so_thang":  int(so_thang),
                "ghi_chu":   ghi_chu,
                "nguon_von": nguon_von_int,
            })
        if so_hs_khong_nguon:
            st.warning(
                f"Có {so_hs_khong_nguon} hồ sơ không xác định được nguồn vốn "
                f"(cột 'Nguồn vốn' trong HSTD gốc bị trống hoặc sai). "
                f"Mặc định gán về Trung ương. Kiểm tra lại file HSTD gốc nếu cần."
            )

        ds_tw = _loc_theo_nguon(ds_xuat_full, NGUON_TW)
        ds_dp = _loc_theo_nguon(ds_xuat_full, NGUON_DP)
        ds_khoanh_tw = [r for r in ds_tw if "Khoanh" in r.get("bien_phap", "")]
        ds_xoa_tw    = [r for r in ds_tw if "Xóa"    in r.get("bien_phap", "")]
        ds_khoanh_dp = [r for r in ds_dp if "Khoanh" in r.get("bien_phap", "")]
        ds_xoa_dp    = [r for r in ds_dp if "Xóa"    in r.get("bien_phap", "")]

        dot_xuat = 1
        nam_xuat = date.today().year

        st.markdown("#### 📋 Biểu đề nghị xử lý nợ + Tờ trình")

        col_tw, col_dp = st.columns(2)
        with col_tw:
            st.markdown("##### 🔵 Trung ương")
            _render_04_05_tt(ds_khoanh_tw, ds_xoa_tw, ten_pgd, LABEL_TW,
                             f"{key_prefix}tw", dot_xuat, nam_xuat, la_cn)
        with col_dp:
            st.markdown("##### 🟢 Địa phương")
            _render_04_05_tt(ds_khoanh_dp, ds_xoa_dp, ten_pgd, LABEL_DP,
                             f"{key_prefix}dp", dot_xuat, nam_xuat, la_cn)

        st.markdown("---")
        st.markdown("#### 📊 Báo cáo sau hạch toán (13/XLN · 14/XLN)")
        st.caption("Xuất sau khi có Quyết định của Hội đồng quản trị NHCSXH.")

        with st.expander("⚙️ Thông tin Quyết định HĐQT", expanded=False):
            col_qd1, col_qd2, col_qd3, col_qd4 = st.columns(4)
            with col_qd1:
                so_qd = st.text_input("Số QĐ HĐQT", placeholder="vd: 123/QĐ-HĐQT", key=f"{key_prefix}nrr_so_qd")
            with col_qd2:
                ngay_qd = st.date_input("Ngày ký QĐ", value=date.today(), key=f"{key_prefix}nrr_ngay_qd")
            with col_qd3:
                ngay_bd = st.date_input("Từ ngày", value=date.today(), key=f"{key_prefix}nrr_ngay_bd")
            with col_qd4:
                ngay_kt = st.date_input("Đến ngày", value=date.today(), key=f"{key_prefix}nrr_ngay_kt")

        col13_tw, col13_dp, col14_tw, col14_dp = st.columns(4)
        with col13_tw:
            if st.button("📥 13/XLN\nTrung ương", use_container_width=True, key=f"{key_prefix}nrr_13xln_tw"):
                if not ds_khoanh_tw:
                    st.warning("⚠️ Không có hồ sơ khoanh nợ TW.")
                else:
                    docx_b = _tao_word_13xln(
                        _tong_hop_no(ds_khoanh_tw), ten_pgd, LABEL_TW,
                        so_qd, ngay_qd, ngay_bd, ngay_kt,
                    )
                    nut_tai_word_va_pdf(docx_b, f"Mau13XLN_TW_{ten_pgd}_{date.today():%d%m%Y}", f"{key_prefix}nrr_13xln_tw")
            hien_thi_nut_tai(f"{key_prefix}nrr_13xln_tw")
        with col13_dp:
            if st.button("📥 13/XLN\nĐịa phương", use_container_width=True, key=f"{key_prefix}nrr_13xln_dp"):
                if not ds_khoanh_dp:
                    st.warning("⚠️ Không có hồ sơ khoanh nợ ĐP.")
                else:
                    docx_b = _tao_word_13xln(
                        _tong_hop_no(ds_khoanh_dp), ten_pgd, LABEL_DP,
                        so_qd, ngay_qd, ngay_bd, ngay_kt,
                    )
                    nut_tai_word_va_pdf(docx_b, f"Mau13XLN_DP_{ten_pgd}_{date.today():%d%m%Y}", f"{key_prefix}nrr_13xln_dp")
            hien_thi_nut_tai(f"{key_prefix}nrr_13xln_dp")
        with col14_tw:
            if st.button("📥 14/XLN\nTrung ương", use_container_width=True, key=f"{key_prefix}nrr_14xln_tw"):
                if not ds_xoa_tw:
                    st.warning("⚠️ Không có hồ sơ xóa nợ TW.")
                else:
                    docx_b = _tao_word_14xln(
                        _tong_hop_no(ds_xoa_tw), ten_pgd, LABEL_TW,
                        so_qd, ngay_qd, ngay_bd, ngay_kt,
                    )
                    nut_tai_word_va_pdf(docx_b, f"Mau14XLN_TW_{ten_pgd}_{date.today():%d%m%Y}", f"{key_prefix}nrr_14xln_tw")
            hien_thi_nut_tai(f"{key_prefix}nrr_14xln_tw")
        with col14_dp:
            if st.button("📥 14/XLN\nĐịa phương", use_container_width=True, key=f"{key_prefix}nrr_14xln_dp"):
                if not ds_xoa_dp:
                    st.warning("⚠️ Không có hồ sơ xóa nợ ĐP.")
                else:
                    docx_b = _tao_word_14xln(
                        _tong_hop_no(ds_xoa_dp), ten_pgd, LABEL_DP,
                        so_qd, ngay_qd, ngay_bd, ngay_kt,
                    )
                    nut_tai_word_va_pdf(docx_b, f"Mau14XLN_DP_{ten_pgd}_{date.today():%d%m%Y}", f"{key_prefix}nrr_14xln_dp")
            hien_thi_nut_tai(f"{key_prefix}nrr_14xln_dp")

    # ── Bước 5: Xem lại hồ sơ đã lưu ─────────────────────────────────
    st.markdown("---")
    with st.expander("📋 Hồ sơ đã lập kỳ này", expanded=False):
        du_lieu_cu = db.doc_kv(kv_key)
        if du_lieu_cu and "danh_sach" in du_lieu_cu:
            ds_cu = du_lieu_cu["danh_sach"]
            st.caption(f"🕐 {du_lieu_cu.get('ngay_tao', '')} — {len(ds_cu)} hồ sơ")
            _hien_thi_chi_tiet(ds_cu)
            if st.button("🗑️ Xóa bản ghi", key=f"{key_prefix}nrr_btn_xoa", type="secondary"):
                st.session_state[f"{key_prefix}nrr_xac_nhan_xoa"] = True
            if st.session_state.get(f"{key_prefix}nrr_xac_nhan_xoa"):
                st.warning("⚠️ Bạn có chắc chắn muốn xóa toàn bộ hồ sơ kỳ này?")
                c_xc1, c_xc2 = st.columns(2)
                with c_xc1:
                    if st.button("✅ Xác nhận xóa", key=f"{key_prefix}nrr_btn_xac_nhan"):
                        db.ghi_kv(kv_key, {}, username)
                        db.ghi_audit(username, "xoa_no_rui_ro",
                                     f"Xóa {len(ds_cu)} hồ sơ — {ten_pgd or 'unknown'}")
                        st.session_state.pop(f"{key_prefix}nrr_xac_nhan_xoa", None)
                        st.cache_data.clear()
                        st.success("✅ Đã xóa hồ sơ.")
                        st.rerun()
                with c_xc2:
                    if st.button("❌ Hủy", key=f"{key_prefix}nrr_btn_huy"):
                        st.session_state.pop(f"{key_prefix}nrr_xac_nhan_xoa", None)
                        st.rerun()
        else:
            st.info("ℹ️ Chưa có hồ sơ nào trong kỳ này.")


# ── Workspace cho Phòng KH-NV (CN) ─────────────────────────────────
def _render_workspace_cn(tab, **kwargs) -> None:
    df       = kwargs.get("df")
    username = kwargs.get("username", "unknown")

    _tab_ctx = tab if tab is not None else st.container()
    with _tab_ctx:
        t0, t1, t2, t3 = st.tabs([
            "📝 Nhập hồ sơ theo PGD",
            "📊 Tổng quan toàn tỉnh",
            "📋 Biểu đề nghị + Tờ trình CN",
            "📊 13/XLN · 14/XLN",
        ])

        # ── T0: Nhập hồ sơ theo PGD ──────────────────────────────────
        with t0:
            st.caption("Phòng KH-NV nhập thay PGD hoặc nhập hồ sơ Hội sở tỉnh.")

            pgd_chon = st.selectbox(
                "📍 Chọn PGD",
                options=DS_PGD,
                key="cn_nrr_chon_pgd",
            )
            df_pgd = df[df[COT_TEN_PGD] == pgd_chon].copy() \
                     if COT_TEN_PGD in df.columns else pd.DataFrame()

            if df_pgd.empty:
                st.warning(f"⚠️ Không có dữ liệu HSTD cho {pgd_chon}.")
            else:
                now = datetime.now()
                kv_key_cn = f"no_rui_ro_{pgd_slug(pgd_chon)}_{now.year}_{now.month:02d}"
                _render_luong_nhap_ho_so(
                    df_pgd=df_pgd,
                    ten_pgd=pgd_chon,
                    kv_key=kv_key_cn,
                    username=username,
                    la_cn=True,
                    key_prefix="cn_",
                )

        # ── T1: Tổng quan toàn tỉnh ──────────────────────────────────
        with t1:
            col_thang, col_nam = st.columns(2)
            with col_thang:
                thang_xem = st.selectbox("Tháng", list(range(1, 13)),
                    index=datetime.now().month - 1, key="cn_nrr_thang")
            with col_nam:
                nam_xem = st.number_input("Năm", min_value=2020,
                    max_value=2030, value=datetime.now().year, key="cn_nrr_nam")

            ds_all: list[dict] = []
            pgd_co_du_lieu: list[str] = []
            for pgd in DS_PGD:
                key = f"no_rui_ro_{pgd_slug(pgd)}_{nam_xem}_{thang_xem:02d}"
                data = db.doc_kv(key)
                if data and "danh_sach" in data:
                    for item in data["danh_sach"]:
                        item["_pgd"] = pgd
                    ds_all.extend(data["danh_sach"])
                    pgd_co_du_lieu.append(pgd)

            tong_hs    = len(ds_all)
            so_khoanh  = sum(1 for r in ds_all if "Khoanh" in r.get("bien_phap",""))
            so_xoa     = sum(1 for r in ds_all if "Xóa"    in r.get("bien_phap",""))
            tien_khoanh = sum(r.get("du_no", 0) for r in ds_all
                              if "Khoanh" in r.get("bien_phap",""))
            tien_xoa    = sum(r.get("du_no", 0) for r in ds_all
                              if "Xóa"    in r.get("bien_phap",""))

            c1, c2, c3, c4 = st.columns(4)
            c1.metric("Tổng hồ sơ", tong_hs)
            c2.metric("Khoanh nợ", f"{so_khoanh} món",
                      fmt(tien_khoanh))
            c3.metric("Xóa nợ", f"{so_xoa} món",
                      fmt(tien_xoa))
            c4.metric("PGD có hồ sơ",
                      f"{len(pgd_co_du_lieu)}/{len(DS_PGD)}")

            if not ds_all:
                st.info("ℹ️ Chưa có PGD nào nhập hồ sơ trong kỳ này.")
            else:
                rows_pgd = []
                for pgd in DS_PGD:
                    ds_pgd = [r for r in ds_all if r.get("_pgd") == pgd]
                    if not ds_pgd:
                        continue
                    kh  = [r for r in ds_pgd if "Khoanh" in r.get("bien_phap","")]
                    xoa = [r for r in ds_pgd if "Xóa"    in r.get("bien_phap","")]
                    tw  = [r for r in ds_pgd if r.get("nguon_von") == NGUON_TW]
                    dp  = [r for r in ds_pgd if r.get("nguon_von") == NGUON_DP]
                    rows_pgd.append({
                        "PGD":             pgd,
                        "Khoanh (món)":    len(kh),
                        "Khoanh (triệu)":  sum(r.get("du_no",0) for r in kh) / 1e6,
                        "Xóa (món)":       len(xoa),
                        "Xóa (triệu)":     sum(r.get("du_no",0) for r in xoa) / 1e6,
                        "TW (triệu)":      sum(r.get("du_no",0) for r in tw) / 1e6,
                        "ĐP (triệu)":      sum(r.get("du_no",0) for r in dp) / 1e6,
                        "Tổng (triệu)":    sum(r.get("du_no",0) for r in ds_pgd) / 1e6,
                    })
                df_th = pd.DataFrame(rows_pgd)
                hien_thi_dataframe_phan_trang(df_th, key="cn_nrr_th_pgd")

                if st.button("📥 Xuất Excel tổng hợp", key="cn_nrr_xuat_xl"):
                    df_ct = pd.DataFrame([
                        {k: v for k, v in r.items() if k != "_pgd"}
                        for r in ds_all
                    ])
                    buf = xuat_excel({
                        "Tổng hợp PGD": df_th,
                        "Chi tiết":     df_ct,
                    })
                    st.session_state["_cn_nrr_xl"] = buf
                    db.ghi_audit(username, "xuat_bieu_cn",
                                 f"Excel tổng hợp NRR {thang_xem}/{nam_xem}")
                if st.session_state.get("_cn_nrr_xl"):
                    st.download_button(
                        "⬇ Tải Excel",
                        data=st.session_state["_cn_nrr_xl"],
                        file_name=f"TongHop_NRR_{thang_xem:02d}{nam_xem}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        key="cn_nrr_dl_xl",
                    )

        # ── T2: Biểu đề nghị + Tờ trình CN ──────────────────────────
        with t2:
            col_f1, col_f2, col_f3, col_f4 = st.columns(4)
            with col_f1:
                pgd_t2 = st.selectbox("PGD", ["Tất cả"] + DS_PGD,
                                       key="cn_nrr_t2_pgd")
            with col_f2:
                nv_t2 = st.selectbox("Nguồn vốn",
                    ["Tất cả", "Trung ương", "Địa phương"],
                    key="cn_nrr_t2_nv")
            with col_f3:
                dot_t2 = st.number_input("Đợt", min_value=1, max_value=4,
                                          value=1, key="cn_nrr_t2_dot")
            with col_f4:
                nam_t2 = st.number_input("Năm", min_value=2020, max_value=2030,
                    value=datetime.now().year, key="cn_nrr_t2_nam")

            ds_t2: list[dict] = []
            thang_ht = datetime.now().month
            for pgd in (DS_PGD if pgd_t2 == "Tất cả" else [pgd_t2]):
                key = f"no_rui_ro_{pgd_slug(pgd)}_{nam_t2}_{thang_ht:02d}"
                data = db.doc_kv(key)
                if data and "danh_sach" in data:
                    ds_t2.extend(data["danh_sach"])

            if nv_t2 == "Trung ương":
                ds_t2 = [r for r in ds_t2 if r.get("nguon_von") == NGUON_TW]
            elif nv_t2 == "Địa phương":
                ds_t2 = [r for r in ds_t2 if r.get("nguon_von") == NGUON_DP]

            ds_khoanh_tw = [r for r in ds_t2
                if "Khoanh" in r.get("bien_phap","")
                and r.get("nguon_von") == NGUON_TW]
            ds_xoa_tw    = [r for r in ds_t2
                if "Xóa"    in r.get("bien_phap","")
                and r.get("nguon_von") == NGUON_TW]
            ds_khoanh_dp = [r for r in ds_t2
                if "Khoanh" in r.get("bien_phap","")
                and r.get("nguon_von") == NGUON_DP]
            ds_xoa_dp    = [r for r in ds_t2
                if "Xóa"    in r.get("bien_phap","")
                and r.get("nguon_von") == NGUON_DP]

            ten_don_vi = "Chi nhánh Đồng Nai"
            col_tw, col_dp = st.columns(2)
            with col_tw:
                st.markdown("##### 🔵 Trung ương")
                _render_04_05_tt(
                    ds_khoanh_tw, ds_xoa_tw, ten_don_vi, LABEL_TW,
                    key_prefix="cn_t2_tw",
                    dot_xuat=int(dot_t2), nam_xuat=int(nam_t2),
                    la_cn=True,
                )
            with col_dp:
                st.markdown("##### 🟢 Địa phương")
                _render_04_05_tt(
                    ds_khoanh_dp, ds_xoa_dp, ten_don_vi, LABEL_DP,
                    key_prefix="cn_t2_dp",
                    dot_xuat=int(dot_t2), nam_xuat=int(nam_t2),
                    la_cn=True,
                )

        # ── T3: 13/XLN · 14/XLN ────────────────────────────────────
        with t3:
            st.caption("Xuất sau khi có Quyết định của Hội đồng quản trị NHCSXH.")

            pgd_t3 = st.selectbox("PGD", ["Tất cả"] + DS_PGD,
                                    key="cn_nrr_t3_pgd")

            with st.expander("⚙️ Thông tin Quyết định HĐQT", expanded=True):
                c1, c2, c3, c4 = st.columns(4)
                with c1:
                    so_qd  = st.text_input("Số QĐ", placeholder="123/QĐ-HĐQT",
                                            key="cn_nrr_so_qd")
                with c2:
                    ngay_qd = st.date_input("Ngày ký", value=date.today(),
                                             format="DD/MM/YYYY", key="cn_nrr_ngay_qd")
                with c3:
                    ngay_bd = st.date_input("Từ ngày", value=date.today(),
                                             format="DD/MM/YYYY", key="cn_nrr_ngay_bd")
                with c4:
                    ngay_kt = st.date_input("Đến ngày", value=date.today(),
                                             format="DD/MM/YYYY", key="cn_nrr_ngay_kt")

            thang_ht = datetime.now().month
            nam_ht   = datetime.now().year
            ds_t3: list[dict] = []
            for pgd in (DS_PGD if pgd_t3 == "Tất cả" else [pgd_t3]):
                key = f"no_rui_ro_{pgd_slug(pgd)}_{nam_ht}_{thang_ht:02d}"
                data = db.doc_kv(key)
                if data and "danh_sach" in data:
                    ds_t3.extend(data["danh_sach"])

            ds_kh_tw = [r for r in ds_t3
                if "Khoanh" in r.get("bien_phap","")
                and r.get("nguon_von") == NGUON_TW]
            ds_kh_dp = [r for r in ds_t3
                if "Khoanh" in r.get("bien_phap","")
                and r.get("nguon_von") == NGUON_DP]
            ds_xo_tw = [r for r in ds_t3
                if "Xóa" in r.get("bien_phap","")
                and r.get("nguon_von") == NGUON_TW]
            ds_xo_dp = [r for r in ds_t3
                if "Xóa" in r.get("bien_phap","")
                and r.get("nguon_von") == NGUON_DP]

            ten_don_vi = "Chi nhánh Đồng Nai"
            ngay_hom_nay = date.today()

            col13_tw, col13_dp, col14_tw, col14_dp = st.columns(4)
            with col13_tw:
                if st.button("📥 13/XLN TW", use_container_width=True,
                              key="cn_nrr_13tw"):
                    if not ds_kh_tw:
                        st.warning("⚠️ Không có hồ sơ khoanh nợ TW.")
                    else:
                        docx_b = _tao_word_13xln(
                            _tong_hop_no(ds_kh_tw), ten_don_vi, LABEL_TW,
                            so_qd, ngay_qd, ngay_bd, ngay_kt,
                        )
                        nut_tai_word_va_pdf(docx_b,
                            f"Mau13XLN_TW_CN_{ngay_hom_nay:%d%m%Y}",
                            "cn_nrr_13tw")
                        db.ghi_audit(username, "xuat_bieu_cn", "13XLN TW")
                hien_thi_nut_tai("cn_nrr_13tw")

            with col13_dp:
                if st.button("📥 13/XLN ĐP", use_container_width=True,
                              key="cn_nrr_13dp"):
                    if not ds_kh_dp:
                        st.warning("⚠️ Không có hồ sơ khoanh nợ ĐP.")
                    else:
                        docx_b = _tao_word_13xln(
                            _tong_hop_no(ds_kh_dp), ten_don_vi, LABEL_DP,
                            so_qd, ngay_qd, ngay_bd, ngay_kt,
                        )
                        nut_tai_word_va_pdf(docx_b,
                            f"Mau13XLN_DP_CN_{ngay_hom_nay:%d%m%Y}",
                            "cn_nrr_13dp")
                        db.ghi_audit(username, "xuat_bieu_cn", "13XLN ĐP")
                hien_thi_nut_tai("cn_nrr_13dp")

            with col14_tw:
                if st.button("📥 14/XLN TW", use_container_width=True,
                              key="cn_nrr_14tw"):
                    if not ds_xo_tw:
                        st.warning("⚠️ Không có hồ sơ xóa nợ TW.")
                    else:
                        docx_b = _tao_word_14xln(
                            _tong_hop_no(ds_xo_tw), ten_don_vi, LABEL_TW,
                            so_qd, ngay_qd, ngay_bd, ngay_kt,
                        )
                        nut_tai_word_va_pdf(docx_b,
                            f"Mau14XLN_TW_CN_{ngay_hom_nay:%d%m%Y}",
                            "cn_nrr_14tw")
                        db.ghi_audit(username, "xuat_bieu_cn", "14XLN TW")
                hien_thi_nut_tai("cn_nrr_14tw")

            with col14_dp:
                if st.button("📥 14/XLN ĐP", use_container_width=True,
                              key="cn_nrr_14dp"):
                    if not ds_xo_dp:
                        st.warning("⚠️ Không có hồ sơ xóa nợ ĐP.")
                    else:
                        docx_b = _tao_word_14xln(
                            _tong_hop_no(ds_xo_dp), ten_don_vi, LABEL_DP,
                            so_qd, ngay_qd, ngay_bd, ngay_kt,
                        )
                        nut_tai_word_va_pdf(docx_b,
                            f"Mau14XLN_DP_CN_{ngay_hom_nay:%d%m%Y}",
                            "cn_nrr_14dp")
                        db.ghi_audit(username, "xuat_bieu_cn", "14XLN ĐP")
                hien_thi_nut_tai("cn_nrr_14dp")


def render(tab: DeltaGenerator, **kwargs) -> None:
    df = kwargs.get("df")
    role_raw = str(kwargs.get("role", "user") or "user")
    role = normalize_role(role_raw)
    username = kwargs.get("username", "unknown")
    pgd_user = kwargs.get("pgd_user")

    _tab_ctx = tab if tab is not None else __import__('streamlit').container()
    with _tab_ctx:
        st.subheader("💳 Xử lý nợ rủi ro — QĐ 62/2015/QĐ-TTg")
        st.caption(
            "Khoanh nợ / Xóa nợ cho hộ vay gặp rủi ro theo Quyết định 62. "
            "Dữ liệu được lưu theo kỳ (tháng hiện tại)."
        )

        if df is None or df.empty:
            st.warning("⚠️ Chưa có dữ liệu HSTD.")
            return
        if la_phan_he_cn(role):
            _render_workspace_cn(tab, df=df, role=role, username=username, **kwargs)
            return


        df = _loc_df_theo_pgd(df, role, pgd_user)
        if df.empty:
            st.warning("⚠️ Không có dữ liệu HSTD cho đơn vị hiện tại.")
            return

        ten_pgd = _lay_pgd_tu_user(role, pgd_user, df)
        kv_key = _tao_kv_key(ten_pgd or "unknown")

        st.markdown("#### 🧾 Soạn mẫu 01/XLN và 02/XLN")
        so_ku_list = []
        if COT_SO_KU in df.columns:
            try:
                so_ku_list = sorted(
                    [str(x).strip() for x in df[COT_SO_KU].dropna().astype(str).unique().tolist()]
                )
            except Exception:
                so_ku_list = []

        if not so_ku_list:
            st.warning("⚠️ Không có món vay nào để soạn mẫu (thiếu Số khế ước hoặc dữ liệu rỗng).")
        else:
            so_ku_chon = st.selectbox(
                "BƯỚC 1 — Chọn món vay (Số khế ước)",
                options=so_ku_list,
                key="xln_so_ku_chon",
            )

            row_hstd = None
            try:
                df_row = df[df[COT_SO_KU].astype(str) == str(so_ku_chon)]
                if not df_row.empty:
                    row_hstd = df_row.iloc[0]
            except Exception:
                row_hstd = None

            if row_hstd is None:
                st.warning("⚠️ Không tìm thấy dữ liệu HSTD tương ứng Số khế ước đã chọn.")
            else:
                ten_kh = str(row_hstd.get(COT_TEN_KH, "") or "")
                dia_chi = str(row_hstd.get(COT_DIA_CHI, "") or "")
                sdt = str(row_hstd.get(COT_SDT, "") or "")

                ngay_vay_str = ""
                try:
                    nv = pd.to_datetime(row_hstd.get(COT_NGAY_VAY, ""), errors="coerce", dayfirst=True)
                    if pd.notna(nv):
                        ngay_vay_str = nv.strftime("%d/%m/%Y")
                    else:
                        ngay_vay_str = str(row_hstd.get(COT_NGAY_VAY, "") or "")
                except Exception:
                    ngay_vay_str = str(row_hstd.get(COT_NGAY_VAY, "") or "")

                ngay_dh_str = ""
                try:
                    ndh = pd.to_datetime(row_hstd.get(COT_NGAY_DH, ""), errors="coerce", dayfirst=True)
                    if pd.notna(ndh):
                        ngay_dh_str = ndh.strftime("%d/%m/%Y")
                    else:
                        ngay_dh_str = str(row_hstd.get(COT_NGAY_DH, "") or "")
                except Exception:
                    ngay_dh_str = str(row_hstd.get(COT_NGAY_DH, "") or "")

                ten_ct = str(row_hstd.get(COT_TEN_CT, "") or "")
                muc_vay_vnd = _num(row_hstd.get(COT_MUC_VAY, 0) or 0)
                du_no_th_vnd = _num(row_hstd.get(COT_DU_NO_TH, 0) or 0)
                du_no_qh_vnd = _num(row_hstd.get(COT_DU_NO_QH, 0) or 0)
                du_no_goc_vnd = du_no_th_vnd + du_no_qh_vnd
                tong_du_no_vnd = _num(row_hstd.get(COT_TONG_DU_NO, 0) or 0)

                co_cot_lai_ton = COT_LAI_TON in df.columns
                lai_ton_vnd = _num(row_hstd.get(COT_LAI_TON, 0) or 0) if co_cot_lai_ton else 0.0

                ten_pgd_row = str(row_hstd.get(COT_TEN_PGD, "") or "")
                ten_nhcsxh = ten_pgd_row or (ten_pgd or "")

                st.info("BƯỚC 2 — Thông tin tự động điền (từ HSTD)")
                a1, a2 = st.columns(2)
                with a1:
                    st.markdown(f"**Tên KH:** {ten_kh or '—'}")
                    st.markdown(f"**Địa chỉ:** {dia_chi or '—'}")
                    st.markdown(f"**Số điện thoại:** {sdt or '—'}")
                    st.markdown(f"**Số khế ước:** {str(so_ku_chon) or '—'}")
                    st.markdown(f"**Ngày vay:** {ngay_vay_str or '—'}")
                    st.markdown(f"**Ngày đến hạn:** {ngay_dh_str or '—'}")
                with a2:
                    st.markdown(f"**Chương trình tín dụng:** {ten_ct or '—'}")
                    st.markdown(f"**Mức vay:** {fmt(muc_vay_vnd) if muc_vay_vnd else '—'}")
                    st.markdown(f"**Dư nợ gốc:** {fmt(du_no_goc_vnd) if du_no_goc_vnd else '—'}")
                    if co_cot_lai_ton:
                        st.markdown(f"**Lãi tồn:** {fmt(lai_ton_vnd) if lai_ton_vnd else '—'}")
                    else:
                        st.markdown("**Lãi tồn:** 0 (không có cột Lãi tồn)")
                    st.markdown(f"**Tổng dư nợ:** {fmt(tong_du_no_vnd) if tong_du_no_vnd else '—'}")
                    st.markdown(f"**NHCSXH/PGD:** {ten_nhcsxh or '—'}")

                with st.form("form_xln_soan", clear_on_submit=False):
                    st.markdown("BƯỚC 3 — Nhập phần tự thuật")
                    c_trai, c_phai = st.columns(2)
                    with c_trai:
                        nguyen_nhan = st.text_area(
                            "Nguyên nhân rủi ro",
                            height=110,
                            key="xln_nguyen_nhan",
                        )
                        thuc_trang = st.text_area(
                            "Thực trạng dự án / tài sản (chỉ dùng cho 02/XLN)",
                            height=90,
                            key="xln_thuc_trang",
                        )
                        kha_nang = st.text_area(
                            "Khả năng trả nợ",
                            height=70,
                            key="xln_kha_nang",
                        )
                    with c_phai:
                        muc_do_pct = st.number_input(
                            "Mức độ thiệt hại %",
                            min_value=0,
                            max_value=100,
                            value=0,
                            step=1,
                            key="xln_muc_do_pct",
                        )
                        so_tien_thiet_hai_trieu = st.number_input(
                            "Số tiền thiệt hại (triệu đồng)",
                            min_value=0.0,
                            value=0.0,
                            step=1.0,
                            key="xln_thiet_hai_trieu",
                        )
                        bien_phap = st.selectbox(
                            "Biện pháp đề nghị",
                            options=["Khoanh nợ", "Xóa nợ"],
                            key="xln_bien_phap",
                        )
                        so_thang = st.number_input(
                            "Số tháng đề nghị",
                            min_value=0,
                            max_value=120,
                            value=36,
                            step=1,
                            key="xln_so_thang",
                        )
                        ke_hoach = st.text_input(
                            "Kế hoạch trả nợ",
                            key="xln_ke_hoach",
                        )
                        ngay_lap = st.date_input(
                            "Ngày lập",
                            value=date.today(),
                            key="xln_ngay_lap",
                        )
                        dia_danh_default = _pgd_plain(ten_pgd_row or ten_pgd or "")
                        dia_danh = st.text_input(
                            "Địa danh",
                            value=dia_danh_default,
                            key="xln_dia_danh",
                        )

                    st.markdown("BƯỚC 4 — Xuất")
                    b1, b2, b3 = st.columns(3)
                    xuat_01 = b1.form_submit_button("📄 Xuất Word 01/XLN", type="primary")
                    xuat_02 = b2.form_submit_button("📋 Xuất Word 02/XLN")
                    xuat_pdf = b3.form_submit_button("📕 Xuất PDF")

                so_tien_thiet_hai_vnd = float(so_tien_thiet_hai_trieu) * 1_000_000.0

                du_lieu_xln = {
                    "ma_kh": str(row_hstd.get(COT_MA_KH, "") or ""),
                    "ten_kh": ten_kh,
                    "dia_chi": dia_chi,
                    "sdt": sdt,
                    "so_ku": str(so_ku_chon),
                    "ten_ct": ten_ct,
                    "muc_vay": fmt(muc_vay_vnd),
                    "tong_du_no": fmt(tong_du_no_vnd),
                    "du_no_goc": fmt(du_no_goc_vnd),
                    "lai_ton": fmt(lai_ton_vnd),
                    "nqh": fmt(du_no_qh_vnd),
                    "ngay_vay": ngay_vay_str,
                    "ngay_dh": ngay_dh_str,
                    "ten_to": str(row_hstd.get(COT_TEN_TO, "") or ""),
                    "dia_danh": str(dia_danh or ""),
                    "ten_nhcsxh": ten_nhcsxh,
                    "nguyen_nhan": str(nguyen_nhan or ""),
                    "bien_phap": str(bien_phap or ""),
                    "so_thang": str(int(so_thang)),
                    "so_tien_de_nghi": fmt(tong_du_no_vnd),
                    "so_tien_thiet_hai": fmt(so_tien_thiet_hai_vnd),
                    "muc_do_thiet_hai": str(int(muc_do_pct)),
                    "kha_nang_tra_no": str(kha_nang or ""),
                    "thuc_trang_du_an": str(thuc_trang or ""),
                    "ke_hoach_tra_no": str(ke_hoach or ""),
                    "dia_diem": ten_nhcsxh,
                    "can_bo_td": st.session_state.get("username", ""),
                    "ngay_ky": ngay_lap,
                    "ngay_lap": ngay_lap,
                    "thanh_phan": [
                        {
                            "stt": 1,
                            "ho_ten": st.session_state.get("username", ""),
                            "chuc_vu": "Cán bộ tín dụng",
                            "dai_dien": ten_nhcsxh,
                        },
                        {
                            "stt": 7,
                            "ho_ten": ten_kh,
                            "chuc_vu": "",
                            "dai_dien": "Khách hàng vay vốn",
                        },
                    ],
                }

                if xuat_01:
                    with st.spinner("Đang tạo 01/XLN..."):
                        docx_b = _tao_word_01xln(du_lieu_xln)
                    ten_file = f"Mau01XLN_{str(so_ku_chon) or 'KH'}_{date.today():%d%m%Y}"
                    nut_tai_word_va_pdf(docx_b, ten_file, "xln_01")
                    st.session_state["_xln_last_docx"] = docx_b
                    st.session_state["_xln_last_name"] = ten_file
                    db.ghi_audit(
                        username,
                        "xuat_01xln",
                        f"so_ku={str(so_ku_chon)} · pgd={ten_nhcsxh} · bien_phap={bien_phap}",
                    )

                if xuat_02:
                    with st.spinner("Đang tạo 02/XLN..."):
                        docx_b = _tao_word_02xln(du_lieu_xln)
                    ten_file = f"Mau02XLN_{str(so_ku_chon) or 'KH'}_{date.today():%d%m%Y}"
                    nut_tai_word_va_pdf(docx_b, ten_file, "xln_02")
                    st.session_state["_xln_last_docx"] = docx_b
                    st.session_state["_xln_last_name"] = ten_file
                    db.ghi_audit(
                        username,
                        "xuat_02xln",
                        f"so_ku={str(so_ku_chon)} · pgd={ten_nhcsxh} · bien_phap={bien_phap}",
                    )

                if xuat_pdf:
                    src_docx = st.session_state.get("_xln_last_docx")
                    ten_file = st.session_state.get("_xln_last_name")
                    if not src_docx:
                        with st.spinner("Đang tạo 02/XLN để xuất PDF..."):
                            src_docx = _tao_word_02xln(du_lieu_xln)
                        ten_file = f"Mau02XLN_{str(so_ku_chon) or 'KH'}_{date.today():%d%m%Y}"
                    pdf_b = docx_bytes_to_pdf(src_docx)
                    st.session_state["_xln_pdf_docx"] = src_docx
                    st.session_state["_xln_pdf_pdf"] = pdf_b
                    st.session_state["_xln_pdf_name"] = ten_file

                pdf_docx = st.session_state.get("_xln_pdf_docx")
                if pdf_docx:
                    ten_pdf = st.session_state.get("_xln_pdf_name", "MauXLN")
                    pdf_bytes = st.session_state.get("_xln_pdf_pdf")
                    d1, d2 = st.columns(2)
                    with d1:
                        st.download_button(
                            "⬇️ Tải Word (.docx)",
                            data=pdf_docx,
                            file_name=f"{ten_pdf}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True,
                            key="xln_pdf_dl_docx",
                        )
                    with d2:
                        if pdf_bytes:
                            st.download_button(
                                "⬇️ Tải PDF",
                                data=pdf_bytes,
                                file_name=f"{ten_pdf}.pdf",
                                mime="application/pdf",
                                use_container_width=True,
                                key="xln_pdf_dl_pdf",
                            )
                        else:
                            st.caption("⚠️ PDF: cần MS Word trên server. Vẫn có thể tải Word.")

                col_dl_01, col_dl_02 = st.columns(2)
                with col_dl_01:
                    hien_thi_nut_tai("xln_01")
                with col_dl_02:
                    hien_thi_nut_tai("xln_02")

        st.divider()

        _render_luong_nhap_ho_so(
            df_pgd=df,
            ten_pgd=ten_pgd or "",
            kv_key=kv_key,
            username=username,
            la_cn=False,
            key_prefix="",
        )
