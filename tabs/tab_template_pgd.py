"""
Tab Quản lý Template — Phân hệ PGD.
─────────────────────────────────────
Cho phép CBTD địa bàn:
  • Xem danh sách template Word có sẵn
  • Test template với dữ liệu mẫu
  • Tải template về máy

Không cho phép upload template mới (chỉ admin CN mới có quyền này).
"""

from __future__ import annotations

from logger import get_logger
logger = get_logger(__name__)

import os
import re
from io import BytesIO
from pathlib import Path

import streamlit as st
import pandas as pd

from config import TEMPLATES_DIR, COT_TEN_PGD, TAG_MAP
from utils import quet_templates, auto_fill_document
from auth import la_phan_he_pgd


def _lay_danh_sach_templates() -> list:
    """Quét thư mục templates và trả về danh sách [(tên_hiển_thị, Path)]."""
    try:
        return quet_templates(TEMPLATES_DIR)
    except Exception as e:
        logger.error("Lỗi quét templates: %s", e, exc_info=True)
        return []


def _doc_template_mau() -> BytesIO | None:
    """Tạo file template mẫu để test."""
    try:
        from docx import Document
        doc = Document()
        doc.add_heading("Mẫu test template", 0)
        doc.add_paragraph("Tên PGD: {{Tên PGD}}")
        doc.add_paragraph("Ngày báo cáo: {{ngay_bao_cao}}")
        doc.add_paragraph("Tên KH: {{Tên KH}}")
        doc.add_paragraph("Mã KH: {{Mã KH}}")
        doc.add_paragraph("Số KU: {{Số khế ước}}")
        
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf
    except Exception as e:
        logger.error("Lỗi tạo template mẫu: %s", e, exc_info=True)
        return None


def render(tab=None, **kwargs) -> None:
    """
    Render tab Quản lý Template cho PGD.
    
    Args:
        tab: Streamlit container (nếu None dùng st trực tiếp)
        **kwargs: Bao gồm df, pgd_user, role, username
    """
    ctx = tab if tab is not None else st
    
    df = kwargs.get("df")
    pgd_user = kwargs.get("pgd_user", "")
    role = kwargs.get("role", "")
    username = kwargs.get("username", "")
    
    ctx.header("📄 Quản lý Template")
    ctx.caption(f"PGD: **{pgd_user or 'Chưa xác định'}**")
    
    # ── Phần 1: Danh sách template ─────────────────────────────────────
    ctx.subheader("📋 Danh sách Template có sẵn")
    
    templates = _lay_danh_sach_templates()
    
    if not templates:
        ctx.info("ℹ️ Chưa có template nào trong hệ thống. Liên hệ Phòng KH-NV để được hỗ trợ.")
    else:
        # quet_templates trả về list[(tên_hiển_thị, Path)] — không phải list[dict]
        df_templates = pd.DataFrame(
            [(str(t[0]), Path(t[1]).name) for t in templates],
            columns=["Tên template", "Tên file"]
        )
        ctx.dataframe(df_templates, use_container_width=True, hide_index=True)
        
        ctx.divider()
        
        # ── Phần 2: Test template ─────────────────────────────────────────
        ctx.subheader("🧪 Test Template")
        
        col1, col2 = ctx.columns([2, 1])
        
        with col1:
            selected_template = ctx.selectbox(
                "Chọn template để test",
                options=[str(t[0]) for t in templates],
                key="template_pgd_test_select",
                help="Chọn template Word để xem kết quả render với dữ liệu mẫu"
            )
        
        with col2:
            ctx.markdown("<br>", unsafe_allow_html=True)
            test_clicked = ctx.button("▶️ Test template", type="primary", key="template_pgd_test_btn")
        
        if test_clicked and selected_template:
            # Tìm path của template được chọn — tuples (tên_hiển_thị, Path)
            template_path = None
            for t in templates:
                if str(t[0]) == selected_template:
                    template_path = str(t[1])
                    break
            
            if template_path and os.path.exists(template_path):
                try:
                    # Tạo dữ liệu mẫu
                    sample_data = {
                        "Tên PGD": pgd_user or "PGD Mẫu",
                        "ngay_bao_cao": pd.Timestamp.now().strftime("%d/%m/%Y"),
                        "Tên KH": "Nguyễn Văn A",
                        "Mã KH": "00012345",
                        "Số khế ước": "KU.2024.001",
                        "Tên chương trình": "Hộ nghèo",
                        "Tổng dư nợ": 50000000,
                        "Dư nợ trong hạn": 45000000,
                        "Dư nợ quá hạn": 5000000,
                    }
                    
                    with ctx.spinner("Đang render template..."):
                        # auto_fill_document(data_row, template_path, tag_map) -> bytes
                        output = auto_fill_document(sample_data, template_path, TAG_MAP)

                    if output:
                        ctx.success("✅ Template render thành công với dữ liệu mẫu!")
                        ctx.download_button(
                            label="📥 Tải xuống file test",
                            data=output,
                            file_name=f"test_{selected_template}",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="template_pgd_download_test"
                        )
                        
                        # Hiển thị preview các trường được thay thế
                        with ctx.expander("🔍 Xem các trường được thay thế"):
                            st.code("\n".join([f"{{{k}}} → {v}" for k, v in sample_data.items()]))
                    else:
                        ctx.error("❌ Không thể render template. Kiểm tra file template có đúng định dạng không.")
                        
                except Exception as e:
                    logger.error("Lỗi test template: %s", e, exc_info=True)
                    ctx.error(f"❌ Lỗi khi test template: {e}")
            else:
                ctx.warning("⚠️ Không tìm thấy file template.")
    
    ctx.divider()
    
    # ── Phần 3: Hướng dẫn sử dụng ───────────────────────────────────────
    with ctx.expander("📖 Hướng dẫn sử dụng Template"):
        st.markdown("""
        **Cách sử dụng template trong Trung tâm mẫu biểu:**
        
        1. **Template Word** (.docx) chứa các placeholder dạng `{{Tên_Cột}}`
        2. Hệ thống tự động thay thế bằng dữ liệu từ HSTD
        3. Các trường phổ biến:
           - `{{Tên PGD}}` — Tên Phòng Giao dịch
           - `{{Tên KH}}` — Tên khách hàng
           - `{{Mã KH}}` — Mã khách hàng
           - `{{Số khế ước}}` — Số khế ước
           - `{{Tổng dư nợ}}` — Tổng dư nợ
        
        **Lưu ý:**
        - Tên placeholder phải khớp chính xác với tên cột trong HSTD
        - Nếu cột không tồn tại, placeholder sẽ được giữ nguyên
        - Liên hệ Phòng KH-NV để thêm template mới
        """)
        
        # Tải template mẫu
        sample_buf = _doc_template_mau()
        if sample_buf:
            st.download_button(
                label="📥 Tải template mẫu (để tham khảo)",
                data=sample_buf.getvalue(),
                file_name="template_mau_tham_khao.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="template_pgd_download_sample"
            )


def render_tab(tab=None, **kwargs) -> None:
    """Alias cho render() để tương thích với lazy_tabs."""
    render(tab, **kwargs)
