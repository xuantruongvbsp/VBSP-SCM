"""Tab Ủy thác — Theo dõi Hội đoàn thể và các mẫu biểu kiểm tra."""
from __future__ import annotations
import io, pickle, uuid
from datetime import date, datetime, timedelta
import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from streamlit.delta_generator import DeltaGenerator

import db
from auth import la_phan_he_cn, normalize_role
from data.pgd import pgd_slug
from config import (
    COT_TEN_PGD, COT_TEN_KH, COT_SO_KU, COT_TEN_CT,
    COT_TONG_DU_NO, COT_DU_NO_QH, COT_LAI_TON, COT_LAI_TON_QH,
    COT_SO_DU_TG, COT_NGAY_VAY, COT_TEN_TO, COT_DVUT,
    COT_TEN_XA, COT_TEN_THON, COT_MUC_VAY,
    TEN_CHI_NHANH_HIEN_THI, DS_PGD,
)
from utils import fmt, fmt_bang_ty, fmt_so, xuat_excel
from services.template_service import docx_bytes_to_pdf

# ── Hằng số ──────────────────────────────────────────────────────────────────
DVUT_ORDER = [
    "Hội nông dân",
    "Hội liên hiệp phụ nữ",
    "Hội cựu chiến binh",
    "Đoàn thanh niên",
]


# ══════════════════════════════════════════════════════════════════════════════
# CACHE FUNCTIONS
# ══════════════════════════════════════════════════════════════════════════════

@st.cache_data(show_spinner=False, ttl=300)
def _tinh_theo_dvut(_df_bytes: bytes) -> bytes:
    df = pickle.loads(_df_bytes)
    if COT_DVUT not in df.columns:
        return pickle.dumps(pd.DataFrame())
    agg = {}
    if COT_TEN_TO      in df.columns: agg["so_to"]   = (COT_TEN_TO,     "nunique")
    if COT_SO_KU       in df.columns: agg["so_kh"]   = (COT_SO_KU,      "nunique")
    if COT_TONG_DU_NO  in df.columns: agg["tong_dn"] = (COT_TONG_DU_NO, "sum")
    if COT_DU_NO_QH    in df.columns: agg["nqh"]     = (COT_DU_NO_QH,   "sum")
    if COT_LAI_TON     in df.columns: agg["lai_ton"] = (COT_LAI_TON,    "sum")
    if not agg:
        return pickle.dumps(pd.DataFrame())
    t = df.groupby(COT_DVUT).agg(**agg).reset_index()
    t["_ord"] = t[COT_DVUT].apply(
        lambda x: DVUT_ORDER.index(x) if x in DVUT_ORDER else 99)
    return pickle.dumps(t.sort_values("_ord").drop(columns="_ord"))


@st.cache_data(show_spinner=False, ttl=300)
def _loc_mau06(_df_bytes: bytes, ngay_tu: str, ngay_den: str) -> bytes:
    df = pickle.loads(_df_bytes)
    if COT_NGAY_VAY not in df.columns:
        return pickle.dumps(pd.DataFrame())
    ngay_vay = pd.to_datetime(df[COT_NGAY_VAY], errors="coerce")
    mask = (ngay_vay >= pd.Timestamp(ngay_tu)) & \
           (ngay_vay <= pd.Timestamp(ngay_den))
    cols = [c for c in [
        COT_TEN_TO, COT_DVUT, COT_TEN_XA, COT_TEN_KH,
        COT_SO_KU, COT_TEN_CT, COT_NGAY_VAY, COT_MUC_VAY,
        COT_TONG_DU_NO, COT_DU_NO_QH, COT_LAI_TON,
    ] if c in df.columns]
    result = df.loc[mask, cols].copy()
    return pickle.dumps(result.sort_values(COT_NGAY_VAY, ascending=False))


@st.cache_data(show_spinner=False, ttl=300)
def _loc_mau15(_df_bytes: bytes, ten_to: str) -> bytes:
    df = pickle.loads(_df_bytes)
    if COT_TEN_TO not in df.columns:
        return pickle.dumps(pd.DataFrame())
    df_to = df[df[COT_TEN_TO] == ten_to].copy()
    # Tính nợ lãi = Lãi tồn TH + Lãi tồn QH
    if COT_LAI_TON in df_to.columns and COT_LAI_TON_QH in df_to.columns:
        df_to["Nợ lãi"] = df_to[COT_LAI_TON].fillna(0) + \
                           df_to[COT_LAI_TON_QH].fillna(0)
    elif COT_LAI_TON in df_to.columns:
        df_to["Nợ lãi"] = df_to[COT_LAI_TON].fillna(0)
    else:
        df_to["Nợ lãi"] = 0
    cols = [c for c in [
        COT_TEN_KH, COT_TEN_CT, COT_SO_KU,
        COT_TONG_DU_NO, "Nợ lãi", COT_SO_DU_TG,
    ] if c in df_to.columns or c == "Nợ lãi"]
    return pickle.dumps(df_to[cols].reset_index(drop=True))


# ══════════════════════════════════════════════════════════════════════════════
# WORD EXPORT FUNCTIONS
# ══════════════════════════════════════════════════════════════════════════════

def _style_doc(doc: Document) -> None:
    """Thiết lập font mặc định cho toàn bộ document."""
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)


def _add_header_quoc_hieu(doc: Document, don_vi: str, so_vb: str,
                           dia_danh: str, ngay_ky: date) -> None:
    """Thêm phần Quốc hiệu, Tiêu ngữ, số VB chuẩn hành chính."""
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    # Bỏ border
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    for cell in t.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for border_name in ["top", "left", "bottom", "right"]:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "none")
            tcBorders.append(border)
        tcPr.append(tcBorders)

    # Cột trái: Đơn vị + Số VB
    cell_l = t.rows[0].cells[0]
    p = cell_l.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run((don_vi or "").upper())
    run.bold = True
    run.font.size = Pt(12)
    p2 = cell_l.add_paragraph(f"Số: {so_vb}")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Cột phải: Quốc hiệu
    cell_r = t.rows[0].cells[1]
    p3 = cell_r.paragraphs[0]
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    run3.bold = True
    run3.font.size = Pt(12)
    p4 = cell_r.add_paragraph("Độc lập - Tự do - Hạnh phúc")
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.runs[0] if p4.runs else p4.add_run("Độc lập - Tự do - Hạnh phúc")
    run4.bold = True

    p5 = cell_r.add_paragraph(
        f"{dia_danh}, ngày {ngay_ky.day} tháng {ngay_ky.month} "
        f"năm {ngay_ky.year}"
    )
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def _tao_word_ke_hoach(du_lieu: dict, ds_to: list) -> bytes:
    """Tạo file Word Kế hoạch kiểm tra giám sát ủy thác."""
    doc = Document()
    _style_doc(doc)
    # Lề trang
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2)

    _add_header_quoc_hieu(
        doc,
        don_vi   = du_lieu.get("don_vi_kt", ""),
        so_vb    = du_lieu.get("so_vb", ""),
        dia_danh = du_lieu.get("dia_danh", ""),
        ngay_ky  = du_lieu.get("ngay_ky", date.today()),
    )

    # Tiêu đề
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = title.add_run(
        f"KẾ HOẠCH\nKiểm tra giám sát hoạt động ủy thác "
        f"năm {du_lieu.get('nam_kh', date.today().year)}"
    )
    run_t.bold = True
    run_t.font.size = Pt(13)
    doc.add_paragraph(
        f"Căn cứ văn bản số 727/HD-NHCS ngày 11/02/2026 của Tổng Giám đốc NHCSXH;"
    )

    # I. Mục đích yêu cầu
    doc.add_paragraph("I. MỤC ĐÍCH, YÊU CẦU").runs[0].bold = True
    doc.add_paragraph(f"1. Mục đích\n{du_lieu.get('muc_dich','')}")
    doc.add_paragraph(f"2. Yêu cầu\n{du_lieu.get('yeu_cau','')}")

    # II. Kế hoạch kiểm tra
    doc.add_paragraph("II. KẾ HOẠCH KIỂM TRA").runs[0].bold = True
    doc.add_paragraph(
        f"1. Nội dung, thời hiệu kiểm tra\n{du_lieu.get('noi_dung_kt','')}"
    )

    # Bảng đối tượng kiểm tra từ hệ thống
    doc.add_paragraph("2. Đối tượng được kiểm tra")
    if ds_to:
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Table Grid"
        hdrs = ["STT", "Hội đoàn thể", "Xã/Phường", "Tên Tổ TK&VV"]
        for i, h in enumerate(hdrs):
            cell = tbl.rows[0].cells[i]
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for idx, to in enumerate(ds_to, 1):
            row = tbl.add_row()
            row.cells[0].text = str(idx)
            row.cells[1].text = str(to.get(COT_DVUT, ""))
            row.cells[2].text = str(to.get(COT_TEN_XA, ""))
            row.cells[3].text = str(to.get(COT_TEN_TO, ""))
        doc.add_paragraph()

    doc.add_paragraph(
        f"3. Thành phần Đoàn kiểm tra\n{du_lieu.get('thanh_phan','')}"
    )

    # III. Kế hoạch giám sát
    doc.add_paragraph("III. KẾ HOẠCH GIÁM SÁT").runs[0].bold = True
    doc.add_paragraph(
        f"1. Nội dung, thời hiệu giám sát\n{du_lieu.get('noi_dung_gs','')}"
    )
    doc.add_paragraph(
        f"2. Phân công cán bộ giám sát\n{du_lieu.get('phan_cong_gs','')}"
    )

    # IV. Tổ chức thực hiện
    doc.add_paragraph("IV. TỔ CHỨC THỰC HIỆN").runs[0].bold = True
    doc.add_paragraph(du_lieu.get("to_chuc", ""))

    # Ký tên
    doc.add_paragraph()
    ky = doc.add_table(rows=1, cols=2)
    ky.style = "Table Grid"
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    for cell in ky.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for border_name in ["top","left","bottom","right"]:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "none")
            tcBorders.append(border)
        tcPr.append(tcBorders)
    ky.rows[0].cells[0].text = "Nơi nhận:\n- NHCSXH;\n- Lưu: VT."
    p_ky = ky.rows[0].cells[1].paragraphs[0]
    p_ky.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ky.add_run("CHỦ TỊCH\n(Ký, ghi rõ họ tên)\n\n\n\n").bold = True
    p_ky.add_run(du_lieu.get("chu_tich", ""))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _tao_word_mau15(du_lieu: dict, df_to: pd.DataFrame) -> bytes:
    """Tạo file Word Mẫu 15/TD — Danh sách đối chiếu số dư."""
    doc = Document()
    _style_doc(doc)
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2)
        section.right_margin  = Cm(1.5)
        section.page_width    = Cm(29.7)  # A4 ngang
        section.page_height   = Cm(21)

    # Header
    _add_header_quoc_hieu(
        doc,
        don_vi   = du_lieu.get("pgd", ""),
        so_vb    = "15/TD",
        dia_danh = du_lieu.get("ten_xa", ""),
        ngay_ky  = du_lieu.get("ngay_chot", date.today()),
    )

    # Tiêu đề
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = title.add_run("DANH SÁCH ĐỐI CHIẾU SỐ DƯ TIỀN VAY VÀ TIỀN GỬI")
    run_t.bold = True
    run_t.font.size = Pt(13)
    doc.add_paragraph(
        f"Đến ngày {du_lieu.get('ngay_chot', date.today()).strftime('%d/%m/%Y')}"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        f"Tổ trưởng: {du_lieu.get('to_truong','')}    "
        f"Mã Tổ: {du_lieu.get('ma_to','')}    "
        f"Địa chỉ: {du_lieu.get('dia_chi','')}"
    )
    doc.add_paragraph("Đơn vị tính: đồng")

    # Bảng dữ liệu
    headers = [
        "STT", "Họ và tên KH / Chương trình", "Mã khoản vay",
        "Nợ gốc (NH)", "Nợ lãi (NH)", "Số dư TG (NH)",
        "Nợ gốc (KH)", "Nợ lãi (KH)", "Số dư TG (KH)",
        "CL Nợ gốc", "CL Nợ lãi", "CL Số dư TG",
        "Nguyên nhân CL", "Chữ ký KH",
    ]
    tbl = doc.add_table(rows=2, cols=len(headers))
    tbl.style = "Table Grid"

    # Dòng header số thứ tự cột
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].font.size = Pt(9)

    stt_row = tbl.rows[1]
    for i in range(len(headers)):
        stt_row.cells[i].text = str(i + 1)
        stt_row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        stt_row.cells[i].paragraphs[0].runs[0].font.size = Pt(9)

    # Dữ liệu
    tong_goc = tong_lai = tong_tg = 0
    for idx, row in df_to.iterrows():
        r = tbl.add_row()
        ten_kh = str(row.get(COT_TEN_KH, ""))
        ten_ct = str(row.get(COT_TEN_CT, ""))
        ma_ku  = str(row.get(COT_SO_KU, ""))
        goc    = float(row.get(COT_TONG_DU_NO, 0) or 0)
        lai    = float(row.get("Nợ lãi", 0) or 0)
        tg     = float(row.get(COT_SO_DU_TG, 0) or 0)
        tong_goc += goc
        tong_lai += lai
        tong_tg  += tg

        vals = [
            str(idx + 1),
            f"{ten_kh}\n{ten_ct}",
            ma_ku,
            f"{goc:,.0f}", f"{lai:,.0f}", f"{tg:,.0f}",
            "", "", "",   # cột KH điền tay
            "", "", "",   # cột chênh lệch
            "",           # nguyên nhân
            "",           # chữ ký
        ]
        for i, v in enumerate(vals):
            r.cells[i].text = v
            r.cells[i].paragraphs[0].runs[0].font.size = Pt(9)

    # Dòng tổng cộng
    r_tong = tbl.add_row()
    r_tong.cells[0].merge(r_tong.cells[2])
    r_tong.cells[0].text = "Tổng cộng"
    r_tong.cells[0].paragraphs[0].runs[0].bold = True
    r_tong.cells[3].text = f"{tong_goc:,.0f}"
    r_tong.cells[4].text = f"{tong_lai:,.0f}"
    r_tong.cells[5].text = f"{tong_tg:,.0f}"

    doc.add_paragraph()
    # Ký tên
    p_kt = doc.add_paragraph()
    p_kt.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_kt.add_run(
        "CÁN BỘ ĐỐI CHIẾU\n(Ký, ghi rõ họ tên)\n\n\n\n"
        f"{du_lieu.get('can_bo_kt','')}"
    ).bold = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _tao_word_mau06(du_lieu: dict, df_m06: pd.DataFrame,
                    loai: str = "06") -> bytes:
    """Tạo file Word Mẫu 06/TD hoặc 06A/TD."""
    doc = Document()
    _style_doc(doc)
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2)
        section.right_margin  = Cm(1.5)
        section.page_width    = Cm(29.7)
        section.page_height   = Cm(21)

    # Header
    tbl_h = doc.add_table(rows=1, cols=2)
    tbl_h.style = "Table Grid"
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    for cell in tbl_h.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for bn in ["top","left","bottom","right"]:
            b = OxmlElement(f"w:{bn}")
            b.set(qn("w:val"), "none")
            tcBorders.append(b)
        tcPr.append(tcBorders)

    tbl_h.rows[0].cells[0].text = (
        f"Đơn vị kiểm tra: {du_lieu.get('don_vi_kt','')}\n"
        f"{du_lieu.get('ten_xa','')}"
    )
    p_r = tbl_h.rows[0].cells[1].paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_r.add_run(f"Mẫu số: {loai}/TD\nLập 02 liên:\n"
                "- 01 liên chính lưu NH;\n"
                "- 01 liên phô tô lưu Đ.v k.tra")

    # Tiêu đề
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("PHIẾU KIỂM TRA SỬ DỤNG VỐN VAY").bold = True

    # Thông tin đoàn kiểm tra
    ngay_kt = du_lieu.get("ngay_kt", date.today())
    doc.add_paragraph(
        f"Người kiểm tra: 1. {du_lieu.get('can_bo_1','')}   "
        f"Chức vụ: {du_lieu.get('chuc_vu_1','')}\n"
        f"                        2. {du_lieu.get('can_bo_2','')}   "
        f"Chức vụ: {du_lieu.get('chuc_vu_2','')}"
    )
    doc.add_paragraph(
        f"Thời điểm kiểm tra: ............   "
        f"Địa bàn kiểm tra: {du_lieu.get('dia_ban','')}   "
        f"Tổ TK&VV: {du_lieu.get('ten_to','')}"
    )
    doc.add_paragraph("Đơn vị tính: triệu đồng")

    if loai == "06":
        # Mẫu 06 — bảng nhiều KH
        headers = [
            "STT", "Họ và tên người vay",
            "Mã khoản vay", "Chương trình cho vay",
            "Số tiền giải ngân", "Dư nợ đến ngày KT",
            "Mục đích sử dụng vốn",
            "Tổng tiền thực nhận", "Dư nợ thực tế",
            "Vào việc", "Số tiền đúng MĐ", "Số tiền sai MĐ",
            "Hiệu quả ĐT", "Nợ lãi", "Chữ ký KH",
        ]
        tbl = doc.add_table(rows=1, cols=len(headers))
        tbl.style = "Table Grid"
        for i, h in enumerate(headers):
            c = tbl.rows[0].cells[i]
            c.text = h
            c.paragraphs[0].runs[0].bold = True
            c.paragraphs[0].runs[0].font.size = Pt(8)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        tong_gn = tong_dn = 0
        for idx, row in df_m06.iterrows():
            r = tbl.add_row()
            gn = float(row.get(COT_MUC_VAY, 0) or 0)
            dn = float(row.get(COT_TONG_DU_NO, 0) or 0)
            tong_gn += gn
            tong_dn += dn
            vals = [
                str(idx + 1),
                str(row.get(COT_TEN_KH, "")),
                str(row.get(COT_SO_KU, "")),
                str(row.get(COT_TEN_CT, "")),
                f"{gn:,.0f}", f"{dn:,.0f}",
                str(row.get("Mục đích sử dụng vốn vay", "")),
                "", "", "", "", "", "", "", "",
            ]
            for i, v in enumerate(vals):
                r.cells[i].text = v
                r.cells[i].paragraphs[0].runs[0].font.size = Pt(8)

        # Dòng cộng
        r_c = tbl.add_row()
        r_c.cells[0].merge(r_c.cells[3])
        r_c.cells[0].text = "Cộng"
        r_c.cells[0].paragraphs[0].runs[0].bold = True
        r_c.cells[4].text = f"{tong_gn:,.0f}"
        r_c.cells[5].text = f"{tong_dn:,.0f}"

    else:
        # Mẫu 06A — từng KH riêng lẻ, lấy KH đầu tiên
        if not df_m06.empty:
            row = df_m06.iloc[0]
            doc.add_paragraph(
                f"Họ và tên: {row.get(COT_TEN_KH,'')}\n"
                f"Mã khoản vay: {row.get(COT_SO_KU,'')}\n"
                f"Chương trình: {row.get(COT_TEN_CT,'')}\n"
                f"Số tiền giải ngân: {row.get(COT_MUC_VAY,0):,.0f} đồng\n"
                f"Dư nợ gốc: {row.get(COT_TONG_DU_NO,0):,.0f} đồng\n"
                f"Nợ lãi: {row.get('Nợ lãi',0):,.0f} đồng\n"
                f"Số dư tiền gửi TK: {row.get(COT_SO_DU_TG,0):,.0f} đồng"
            )
            doc.add_paragraph("Mục đích vay vốn: "
                               + str(row.get("Mục đích sử dụng vốn vay","")))
            doc.add_paragraph("Thực tế sử dụng vốn:\n"
                               "- Sử dụng đúng mục đích: ............... đồng\n"
                               "- Sử dụng sai mục đích: ................. đồng")
            doc.add_paragraph("Hiệu quả đầu tư: ................................")
            doc.add_paragraph("Khả năng trả nợ: ................................")

    # Nhận xét
    doc.add_paragraph(
        "\nNhận xét:\n"
        "1. Tình hình thực hiện phương án: ........................................\n"
        "2. Kiểm tra, đối chiếu thực tế được ........ KH, số tiền ........ đồng.\n"
        "Biện pháp xử lý: ............................................................"
    )

    # Ký tên
    doc.add_paragraph()
    ky = doc.add_table(rows=1, cols=2)
    ky.style = "Table Grid"
    for cell in ky.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for bn in ["top","left","bottom","right"]:
            b = OxmlElement(f"w:{bn}")
            b.set(qn("w:val"), "none")
            tcBorders.append(b)
        tcPr.append(tcBorders)
    ky.rows[0].cells[0].text = (
        "CÁN BỘ CHỨNG KIẾN (nếu có)\n(Ký, ghi rõ họ tên)\n\n\n"
    )
    p_r2 = ky.rows[0].cells[1].paragraphs[0]
    p_r2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_r2.add_run(
        f"Ngày {ngay_kt.day} tháng {ngay_kt.month} năm {ngay_kt.year}\n"
        "CÁN BỘ KIỂM TRA\n(Ký, ghi rõ họ tên)\n\n\n"
        f"{du_lieu.get('can_bo_1','')}"
    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _tao_word_mau16(du_lieu: dict, df_to: pd.DataFrame) -> bytes:
    doc = Document()
    _style_doc(doc)
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)

    ngay_kt = du_lieu.get("ngay_kt", date.today())
    _add_header_quoc_hieu(
        doc,
        don_vi=du_lieu.get("don_vi_kt", ""),
        so_vb="16/TD",
        dia_danh=du_lieu.get("ten_xa", ""),
        ngay_ky=ngay_kt,
    )

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = title.add_run("BIÊN BẢN KIỂM TRA\nHoạt động tín dụng chính sách tại Tổ TK&VV")
    run_t.bold = True
    run_t.font.size = Pt(13)

    ten_to = du_lieu.get("ten_to", "")
    ten_xa = du_lieu.get("ten_xa", "")
    doc.add_paragraph(
        f"Hôm nay, ngày {ngay_kt.day} tháng {ngay_kt.month} năm {ngay_kt.year}\n"
        f"Đoàn kiểm tra: {du_lieu.get('don_vi_kt', '')}\n"
        f"Cán bộ kiểm tra: {du_lieu.get('can_bo_kt', '')} — Chức vụ: {du_lieu.get('chuc_vu', '')}\n"
        f"Địa điểm kiểm tra: Tổ TK&VV {ten_to}, xã/phường {ten_xa}\n"
        f"Đơn vị được kiểm tra: Tổ trưởng Tổ TK&VV {ten_to}"
    )

    doc.add_paragraph("I. NỘI DUNG KIỂM TRA").runs[0].bold = True
    doc.add_paragraph("1. Kết quả hoạt động tín dụng").runs[0].bold = True

    headers = [
        "STT",
        "Họ và tên KH",
        "Mã khoản vay",
        "Chương trình",
        "Dư nợ (đồng)",
        "Nợ quá hạn (đồng)",
        "Lãi tồn (đồng)",
        "Ghi chú",
    ]
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if p.runs:
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(9)

    tong_dn = tong_nqh = tong_lai_ton = 0.0
    for stt, (_, row) in enumerate(df_to.iterrows(), 1):
        r = tbl.add_row()
        dn = float(row.get(COT_TONG_DU_NO, 0) or 0)
        nqh = float(row.get(COT_DU_NO_QH, 0) or 0)
        lai_ton = float(row.get(COT_LAI_TON, 0) or 0)
        tong_dn += dn
        tong_nqh += nqh
        tong_lai_ton += lai_ton

        vals = [
            str(stt),
            str(row.get(COT_TEN_KH, "")),
            str(row.get(COT_SO_KU, "")),
            str(row.get(COT_TEN_CT, "")),
            f"{dn:,.0f}",
            f"{nqh:,.0f}",
            f"{lai_ton:,.0f}",
            "",
        ]
        for i, v in enumerate(vals):
            r.cells[i].text = v
            r.cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = r.cells[i].paragraphs[0]
            if i in (0, 4, 5, 6):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i in (4, 5, 6) else WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if p.runs:
                p.runs[0].font.size = Pt(9)

    r_tong = tbl.add_row()
    r_tong.cells[0].merge(r_tong.cells[3])
    r_tong.cells[0].text = "Tổng cộng"
    p_t = r_tong.cells[0].paragraphs[0]
    p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if p_t.runs:
        p_t.runs[0].bold = True
        p_t.runs[0].font.size = Pt(9)
    r_tong.cells[4].text = f"{tong_dn:,.0f}"
    r_tong.cells[5].text = f"{tong_nqh:,.0f}"
    r_tong.cells[6].text = f"{tong_lai_ton:,.0f}"
    for i in (4, 5, 6):
        p = r_tong.cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if p.runs:
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(9)

    doc.add_paragraph()
    doc.add_paragraph("2. Nhận xét").runs[0].bold = True
    doc.add_paragraph(du_lieu.get("nhan_xet", "................") or "................")
    doc.add_paragraph("II. KẾT LUẬN VÀ KIẾN NGHỊ").runs[0].bold = True
    doc.add_paragraph(du_lieu.get("ket_luan", "................") or "................")

    doc.add_paragraph()
    ky = doc.add_table(rows=1, cols=2)
    ky.style = "Table Grid"
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    for cell in ky.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for border_name in ["top", "left", "bottom", "right"]:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "none")
            tcBorders.append(border)
        tcPr.append(tcBorders)

    p_l = ky.rows[0].cells[0].paragraphs[0]
    p_l.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_l.add_run("TỔ TRƯỞNG TỔ TK&VV\n(Ký, ghi rõ họ tên)\n\n\n\n").bold = True

    p_r = ky.rows[0].cells[1].paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_r.add_run("CÁN BỘ KIỂM TRA\n(Ký, ghi rõ họ tên)\n\n\n\n").bold = True
    p_r.add_run(str(du_lieu.get("can_bo_kt", "")))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _tao_word_bien_ban_xac_minh(du_lieu: dict) -> bytes:
    doc = Document()
    _style_doc(doc)
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)

    ngay_lap = du_lieu.get("ngay_lap", date.today())
    _add_header_quoc_hieu(
        doc,
        don_vi=du_lieu.get("pgd_user", ""),
        so_vb="",
        dia_danh="",
        ngay_ky=ngay_lap,
    )

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = title.add_run("BIÊN BẢN\nXác minh nợ chiếm dụng")
    run_t.bold = True
    run_t.font.size = Pt(13)

    doc.add_paragraph(
        f"Hôm nay, ngày {ngay_lap.day} tháng {ngay_lap.month} năm {ngay_lap.year}\n"
        f"Cán bộ lập biên bản: {du_lieu.get('can_bo_lap', '')}\n"
        f"Làm việc với: Khách hàng {du_lieu.get('ten_kh', '')}\n"
        f"Số khế ước: {du_lieu.get('so_ku', '')}\n"
        "NỘI DUNG XÁC MINH:\n\n"
        f"Số tiền chiếm dụng: {du_lieu.get('so_tien', '')} triệu đồng\n"
        "Lý do / Hoàn cảnh:\n"
        f"{du_lieu.get('ly_do', '')}\n"
        "Biện pháp xử lý đã thống nhất:\n"
        f"{du_lieu.get('bien_phap', '')}\n\n"
        "Biên bản được lập thành 02 liên, mỗi bên giữ 01 liên."
    )

    doc.add_paragraph()
    ky = doc.add_table(rows=1, cols=2)
    ky.style = "Table Grid"
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    for cell in ky.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for border_name in ["top", "left", "bottom", "right"]:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "none")
            tcBorders.append(border)
        tcPr.append(tcBorders)

    p_l = ky.rows[0].cells[0].paragraphs[0]
    p_l.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_l.add_run("KHÁCH HÀNG\n(Ký, ghi rõ họ tên)\n\n\n\n").bold = True
    p_l.add_run(str(du_lieu.get("ten_kh", "")))

    p_r = ky.rows[0].cells[1].paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_r.add_run("CÁN BỘ LẬP BIÊN BẢN\n(Ký, ghi rõ họ tên)\n\n\n\n").bold = True
    p_r.add_run(str(du_lieu.get("can_bo_lap", "")))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Helpers ───────────────────────────────────────────────────────────────────

def _parse_date(s: str | None) -> date:
    if not s:
        return date.today()
    try:
        return datetime.strptime(s, "%Y-%m-%d").date()
    except Exception:
        return date.today()


def _xoa_border_table(tbl) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    for cell in tbl.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for bn in ["top", "left", "bottom", "right"]:
            b = OxmlElement(f"w:{bn}")
            b.set(qn("w:val"), "none")
            tcBorders.append(b)
        tcPr.append(tcBorders)


# ── Mẫu 02/BB-CT & 03/BB-CX ──────────────────────────────────────────────────

def _tao_word_bb_ct_cx(du_lieu: dict, cap: str = "tinh") -> bytes:
    """Mẫu 02/BB-CT (cap='tinh') hoặc 03/BB-CX (cap='xa')."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    _style_doc(doc)
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2)
        section.page_width    = Cm(21)
        section.page_height   = Cm(29.7)

    # Header 2 cột (không border)
    tbl_h = doc.add_table(rows=1, cols=2)
    tbl_h.style = "Table Grid"
    _xoa_border_table(tbl_h)
    cell_l = tbl_h.rows[0].cells[0]
    p_l = cell_l.paragraphs[0]
    p_l.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_l = p_l.add_run("ĐƠN VỊ KIỂM TRA\n")
    r_l.bold = True
    r_l2 = p_l.add_run((du_lieu.get("don_vi_kt") or "").upper())
    r_l2.bold = True

    cell_r = tbl_h.rows[0].cells[1]
    p_r = cell_r.paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_r.add_run("CỘNG HOÀ XÃ HỘI CHỦ NGHĨA VIỆT NAM").bold = True
    p_r2 = cell_r.add_paragraph("Độc lập - Tự do - Hạnh phúc")
    p_r2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if p_r2.runs:
        p_r2.runs[0].bold = True
    ngay_kt = du_lieu.get("ngay_kt", date.today())
    if isinstance(ngay_kt, str):
        ngay_kt = _parse_date(ngay_kt)
    p_r3 = cell_r.add_paragraph(
        f"{du_lieu.get('dia_danh', '')}, ngày {ngay_kt.day} tháng {ngay_kt.month} năm {ngay_kt.year}"
    )
    p_r3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Tiêu đề
    ten_cap = "cấp tỉnh" if cap == "tinh" else "cấp xã"
    so_hieu = "02/BB-CT" if cap == "tinh" else "03/BB-CX"
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = title.add_run(
        f"BIÊN BẢN KIỂM TRA\n"
        f"Hoạt động nhận ủy thác với NHCSXH của tổ chức CT-XH {ten_cap}"
    )
    run_t.bold = True
    run_t.font.size = Pt(13)

    doc.add_paragraph(
        f"Hôm nay, ngày {ngay_kt.day} tháng {ngay_kt.month} năm {ngay_kt.year}, chúng tôi gồm:"
    )
    doc.add_paragraph(
        f"ĐOÀN KIỂM TRA: {du_lieu.get('don_vi_kt', '')}\n"
        f"- Ông (bà): {du_lieu.get('truong_doan', '')}   Chức vụ: Trưởng đoàn\n"
        f"- Ông (bà): {du_lieu.get('can_bo_2', '')}   Chức vụ: {du_lieu.get('chuc_vu_2', '')}"
    )
    doc.add_paragraph(
        f"ĐƠN VỊ ĐƯỢC KIỂM TRA: {du_lieu.get('ten_don_vi', '')}\n"
        f"- Ông (bà): {du_lieu.get('dai_dien_dc', '')}   Chức vụ: {du_lieu.get('chuc_vu_dc', '')}"
    )
    doc.add_paragraph(
        f"Cùng tiến hành kiểm tra việc thực hiện các nội dung công việc được ủy thác của "
        f"{du_lieu.get('ten_don_vi', '')}, thống nhất kết quả kiểm tra như sau:"
    )

    # I. Kết quả hoạt động ủy thác
    p_i = doc.add_paragraph("I. KẾT QUẢ HOẠT ĐỘNG ỦY THÁC")
    p_i.runs[0].bold = True
    doc.add_paragraph(
        f"(Đến thời điểm {ngay_kt.strftime('%d/%m/%Y')})\n"
        "- Tổng số Tổ TK&VV do Hội quản lý: ........... tổ\n"
        "- Tổng số khách hàng vay vốn: ........... người\n"
        "- Tổng dư nợ nhận ủy thác: ........... triệu đồng\n"
        "  Trong đó: Nợ quá hạn ........... (tỷ lệ .......%), Nợ khoanh ..........."
    )

    # II. Kết quả thực hiện nội dung nhận ủy thác
    ten_khoan = "khoản 1" if cap == "tinh" else "khoản 2"
    p_ii = doc.add_paragraph(
        f"II. KẾT QUẢ THỰC HIỆN CÁC NỘI DUNG NHẬN ỦY THÁC CỦA TỔ CHỨC CT-XH {ten_cap.upper()}"
    )
    p_ii.runs[0].bold = True
    doc.add_paragraph(
        f"Đoàn kiểm tra thực hiện kiểm tra theo nội dung kiểm tra, giám sát tại {ten_khoan} "
        f"Phụ lục I văn bản số 727/HD-NHCS ngày 11/02/2026. Cụ thể:"
    )

    muc_noi_dung = [
        ("tuyen_truyen",    "1. Công tác tuyên truyền, vận động"),
        ("kiem_tra_giam_sat", "2. Công tác kiểm tra, giám sát hoạt động ủy thác"),
        ("tap_huan",        "3. Công tác tập huấn"),
        ("phoi_hop_nhcs",   "4. Hoạt động phối hợp thực hiện cùng NHCSXH"),
    ]
    if cap == "tinh":
        muc_noi_dung.append(("trach_nhiem", "5. Trách nhiệm của tổ chức CT-XH cấp tỉnh"))

    for field_key, ten_muc in muc_noi_dung:
        p_muc = doc.add_paragraph(ten_muc)
        p_muc.runs[0].bold = True
        data = du_lieu.get(field_key) or {}
        doc.add_paragraph(f"a) Kết quả đạt được\n{data.get('ket_qua', '.....')}")
        doc.add_paragraph(f"b) Tồn tại\n{data.get('ton_tai', '.....')}")

    # III. Đánh giá, Nhận xét, Kiến nghị
    p_iii = doc.add_paragraph("III. ĐÁNH GIÁ, NHẬN XÉT CỦA ĐOÀN KIỂM TRA")
    p_iii.runs[0].bold = True
    han_str = du_lieu.get("han_hoan_thanh", "....../....../20......")
    if han_str and len(han_str) == 10 and han_str[4] == "-":
        try:
            han_str = datetime.strptime(han_str, "%Y-%m-%d").strftime("%d/%m/%Y")
        except Exception:
            pass
    doc.add_paragraph(
        f"1. Ưu điểm\n{du_lieu.get('uu_diem', '.....')}\n\n"
        f"2. Tồn tại\n{du_lieu.get('ton_tai_chung', '.....')}\n\n"
        f"3. Kiến nghị của Đoàn kiểm tra\n{du_lieu.get('kien_nghi', '.....')}\n\n"
        f"Đơn vị được kiểm tra hoàn thành các kiến nghị và báo cáo kết quả "
        f"trước ngày {han_str}."
    )

    # IV. Ý kiến đơn vị được kiểm tra
    p_iv = doc.add_paragraph("IV. Ý KIẾN CỦA ĐƠN VỊ ĐƯỢC KIỂM TRA")
    p_iv.runs[0].bold = True
    doc.add_paragraph(du_lieu.get("y_kien_don_vi_dc", ".....") or ".....")

    doc.add_paragraph(
        "Biên bản được lập thành 02 bản "
        "(01 bản lưu Đoàn kiểm tra, 01 bản lưu Đơn vị được kiểm tra)."
    )

    # Ký tên
    doc.add_paragraph()
    ky = doc.add_table(rows=1, cols=2)
    ky.style = "Table Grid"
    _xoa_border_table(ky)
    p_kl = ky.rows[0].cells[0].paragraphs[0]
    p_kl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_kl.add_run("TRƯỞNG ĐOÀN KIỂM TRA\n(Ký, ghi rõ họ tên)\n\n\n\n").bold = True
    p_kl.add_run(du_lieu.get("truong_doan", ""))
    p_kr = ky.rows[0].cells[1].paragraphs[0]
    p_kr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_kr.add_run("ĐƠN VỊ ĐƯỢC KIỂM TRA\n(Ký tên, đóng dấu)\n\n\n\n").bold = True
    p_kr.add_run(du_lieu.get("dai_dien_dc", ""))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Mẫu 04/BC-TH ─────────────────────────────────────────────────────────────

def _tao_word_bc_th(du_lieu: dict, ds_bien_ban: list) -> bytes:
    """Mẫu 04/BC-TH — Báo cáo tổng hợp kết quả kiểm tra."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    _style_doc(doc)
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2)
        section.page_width    = Cm(21)
        section.page_height   = Cm(29.7)

    ngay_bc = du_lieu.get("ngay_bc", date.today())
    if isinstance(ngay_bc, str):
        ngay_bc = _parse_date(ngay_bc)

    tbl_h = doc.add_table(rows=1, cols=2)
    tbl_h.style = "Table Grid"
    _xoa_border_table(tbl_h)
    cell_l = tbl_h.rows[0].cells[0]
    cell_l.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_l.paragraphs[0].add_run(
        f"ĐƠN VỊ KIỂM TRA\n{(du_lieu.get('don_vi_kt') or '').upper()}"
    ).bold = True
    cell_r = tbl_h.rows[0].cells[1]
    p_r = cell_r.paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_r.add_run("CỘNG HOÀ XÃ HỘI CHỦ NGHĨA VIỆT NAM").bold = True
    p_r2 = cell_r.add_paragraph("Độc lập - Tự do - Hạnh phúc")
    p_r2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if p_r2.runs:
        p_r2.runs[0].bold = True
    p_r3 = cell_r.add_paragraph(
        f"{du_lieu.get('dia_danh', '')}, ngày {ngay_bc.day} "
        f"tháng {ngay_bc.month} năm {ngay_bc.year}"
    )
    p_r3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = title.add_run("BÁO CÁO TỔNG HỢP\nKết quả kiểm tra hoạt động nhận ủy thác cho vay")
    run_t.bold = True
    run_t.font.size = Pt(13)
    doc.add_paragraph()

    # I. Thành phần
    p_i = doc.add_paragraph("I. THÀNH PHẦN")
    p_i.runs[0].bold = True
    doc.add_paragraph(
        f"1. Đoàn kiểm tra: {du_lieu.get('don_vi_kt', '')}\n"
        f"   Trưởng đoàn: {du_lieu.get('truong_doan', '')}\n"
        f"2. Cấp ủy, chính quyền địa phương (nếu có): {du_lieu.get('cap_uy', '')}"
    )

    # II. Thời gian, địa điểm, đơn vị được kiểm tra
    p_ii = doc.add_paragraph("II. THỜI GIAN, ĐỊA ĐIỂM, ĐƠN VỊ ĐƯỢC KIỂM TRA")
    p_ii.runs[0].bold = True
    tbl_dv = doc.add_table(rows=1, cols=4)
    tbl_dv.style = "Table Grid"
    for i, h in enumerate(["STT", "Thời gian", "Đơn vị được kiểm tra", "Địa điểm"]):
        c = tbl_dv.rows[0].cells[i]
        c.text = h
        if c.paragraphs[0].runs:
            c.paragraphs[0].runs[0].bold = True
            c.paragraphs[0].runs[0].font.size = Pt(10)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for idx, bb in enumerate(ds_bien_ban, 1):
        ngay_str = bb.get("ngay_kt", "")
        try:
            ngay_str = datetime.strptime(ngay_str, "%Y-%m-%d").strftime("%d/%m/%Y")
        except Exception:
            pass
        r = tbl_dv.add_row()
        r.cells[0].text = str(idx)
        r.cells[1].text = ngay_str
        r.cells[2].text = bb.get("ten_don_vi", "")
        r.cells[3].text = bb.get("dia_danh", "")
        for cell in r.cells:
            if cell.paragraphs and cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_paragraph()

    # III. Nội dung kiểm tra
    p_iii = doc.add_paragraph("III. NỘI DUNG KIỂM TRA")
    p_iii.runs[0].bold = True
    doc.add_paragraph(
        du_lieu.get("noi_dung_kt",
                    "Theo Phụ lục I văn bản số 727/HD-NHCS ngày 11/02/2026.")
    )

    # IV. Kết quả kiểm tra
    p_iv = doc.add_paragraph("IV. KẾT QUẢ KIỂM TRA")
    p_iv.runs[0].bold = True
    doc.add_paragraph("1. Đánh giá, nhận xét của Đoàn kiểm tra").runs[0].bold = True
    doc.add_paragraph(
        f"a) Đối với tổ chức CT-XH được kiểm tra\n{du_lieu.get('nx_ctxh', '.....')}\n\n"
        f"b) Đối với Tổ TK&VV\n{du_lieu.get('nx_to', '.....')}\n\n"
        f"c) Đối với tổ viên Tổ TK&VV\n{du_lieu.get('nx_to_vien', '.....')}"
    )
    doc.add_paragraph("2. Kiến nghị của Đoàn kiểm tra").runs[0].bold = True
    doc.add_paragraph(
        f"a) Đối với tổ chức CT-XH được kiểm tra\n{du_lieu.get('kn_ctxh', '.....')}\n\n"
        f"b) Đối với Tổ TK&VV\n{du_lieu.get('kn_to', '.....')}\n\n"
        f"c) Đối với tổ viên Tổ TK&VV\n{du_lieu.get('kn_to_vien', '.....')}\n\n"
        f"d) Đối với NHCSXH\n{du_lieu.get('kn_nhcs', '.....')}\n\n"
        f"đ) Đối với tổ chức CT-XH cấp trên\n{du_lieu.get('kn_cap_tren', '.....')}"
    )
    doc.add_paragraph("3. Kiến nghị của Đơn vị được kiểm tra").runs[0].bold = True
    doc.add_paragraph(
        f"a) Đối với NHCSXH\n.....\n\nb) Đối với tổ chức CT-XH cấp trên\n....."
    )

    # VI. Tài liệu kèm theo
    p_vi = doc.add_paragraph("VI. TÀI LIỆU KÈM THEO (nếu có)")
    p_vi.runs[0].bold = True
    doc.add_paragraph(
        f"1. Phiếu kiểm tra sử dụng vốn vay (mẫu 06/TD, 06A/TD): ............... phiếu.\n"
        f"2. Danh sách đối chiếu (mẫu 15/TD): ...................... danh sách.\n"
        f"3. Biên bản kiểm tra tổ chức CT-XH: {len(ds_bien_ban)} biên bản."
    )

    doc.add_paragraph()
    ky = doc.add_table(rows=1, cols=1)
    ky.style = "Table Grid"
    _xoa_border_table(ky)
    p_ky = ky.rows[0].cells[0].paragraphs[0]
    p_ky.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_ky.add_run("TRƯỞNG ĐOÀN KIỂM TRA\n(Ký, ghi rõ họ tên)\n\n\n\n").bold = True
    p_ky.add_run(du_lieu.get("truong_doan", ""))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ══════════════════════════════════════════════════════════════════════════════
# SUB-TAB RENDERERS
# ══════════════════════════════════════════════════════════════════════════════

def _render_theo_dvut(df: pd.DataFrame) -> None:
    st.markdown("#### 📊 Thống kê theo Hội đoàn thể")
    if df is None or df.empty:
        st.warning("Chưa có dữ liệu."); return
    try:
        t = pickle.loads(_tinh_theo_dvut(pickle.dumps(df)))
    except Exception as e:
        st.error(f"Lỗi: {e}"); return
    if t.empty:
        st.info("Không có dữ liệu."); return

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Hội đoàn thể", len(t))
    c2.metric("Tổng Tổ TK&VV", fmt_so(int(t.get("so_to", pd.Series([0])).sum())))
    c3.metric("Tổng KH", fmt_so(int(t.get("so_kh", pd.Series([0])).sum())))
    c4.metric("Tổng dư nợ (triệu đồng)", fmt(t.get("tong_dn", pd.Series([0])).sum()))
    st.divider()
    hien = t.rename(columns={
        COT_DVUT: "Hội đoàn thể", "so_to": "Số Tổ",
        "so_kh": "Số KH", "tong_dn": "Dư nợ (tỷ)",
        "nqh": "NQH (tỷ)", "lai_ton": "Lãi tồn (tỷ)",
    })
    for col in ["Dư nợ (tỷ)", "NQH (tỷ)", "Lãi tồn (tỷ)"]:
        if col in hien.columns:
            hien[col] = hien[col].apply(fmt_bang_ty)
    st.dataframe(hien, use_container_width=True, hide_index=True)


def _render_ke_hoach(df: pd.DataFrame, pgd_user: str) -> None:
    st.markdown("#### 📋 Kế hoạch kiểm tra giám sát ủy thác")
    st.caption("Hội đoàn thể cấp xã lập (PGD) hoặc cấp tỉnh lập (CN). "
               "Danh sách Tổ TK&VV tự động lấy từ hệ thống.")

    # Lấy danh sách Tổ từ df
    ds_to = []
    if df is not None and not df.empty:
        grp = [c for c in [COT_DVUT, COT_TEN_XA, COT_TEN_TO] if c in df.columns]
        if grp:
            ds_to = (df[grp].drop_duplicates()
                     .sort_values(grp).to_dict("records"))

    with st.form("form_ke_hoach_kt"):
        c1, c2 = st.columns(2)
        don_vi_kt = c1.selectbox(
            "Hội đoàn thể kiểm tra",
            options=DVUT_ORDER,
            key="kh_don_vi_kt"
        )
        so_vb      = c1.text_input("Số văn bản", placeholder="VD: 12/KH-HND")
        ds_xa_kh = sorted(df[COT_TEN_XA].dropna().unique().tolist()) \
                   if df is not None and COT_TEN_XA in df.columns else []
        dia_danh = c1.selectbox(
            "Địa danh (xã/phường)",
            options=ds_xa_kh,
            key="kh_dia_danh",
            help="Xã/phường nơi Hội đóng trụ sở — dùng làm địa danh ký văn bản"
        )
        nam_kh     = c2.number_input("Năm kế hoạch",
                                      value=date.today().year,
                                      min_value=2020, max_value=2035, step=1)
        ngay_ky    = c2.date_input("Ngày ký", value=date.today())
        chu_tich   = c2.text_input("Chủ tịch ký",
                                    placeholder="Họ và tên Chủ tịch Hội")

        st.markdown("**I. Mục đích, yêu cầu**")
        muc_dich   = st.text_area("Mục đích", height=70)
        yeu_cau    = st.text_area("Yêu cầu", height=70)

        st.markdown("**II. Kế hoạch kiểm tra**")
        noi_dung_kt = st.text_area("Nội dung, thời hiệu kiểm tra", height=70)
        thanh_phan  = st.text_area("Thành phần Đoàn kiểm tra", height=60)
        st.info(f"📋 Hệ thống tìm thấy **{len(ds_to)}** Tổ TK&VV "
                f"— sẽ tự động điền vào bảng Đối tượng kiểm tra.")

        st.markdown("**III. Kế hoạch giám sát**")
        noi_dung_gs  = st.text_area("Nội dung, thời hiệu giám sát", height=70)
        phan_cong_gs = st.text_area("Phân công cán bộ giám sát", height=60)

        st.markdown("**IV. Tổ chức thực hiện**")
        to_chuc = st.text_area("Tổ chức thực hiện", height=60)

        submitted = st.form_submit_button("📄 Tạo Word")

    if submitted:
        # Build context cho template Kế hoạch kiểm tra
        context = {
            "don_vi_kt":   don_vi_kt,
            "so_vb":       so_vb,
            "dia_danh":    dia_danh,
            "nam_kh":      int(nam_kh),
            "ngay_ky":     ngay_ky.strftime("%d/%m/%Y"),
            "ngay":        ngay_ky.day,
            "thang":       ngay_ky.month,
            "nam":         ngay_ky.year,
            "muc_dich":    muc_dich,
            "yeu_cau":     yeu_cau,
            "noi_dung_kt": noi_dung_kt,
            "thanh_phan":  thanh_phan,
            "noi_dung_gs": noi_dung_gs,
            "phan_cong_gs": phan_cong_gs,
            "to_chuc":     to_chuc,
            "chu_tich":    chu_tich,
            "so_to":       len(ds_to),
            "ds_to":       ds_to,
        }

        with st.spinner("Đang tạo file..."):
            du_lieu_word = {**context, "ngay_ky": ngay_ky}
            docx_bytes = _tao_word_ke_hoach(du_lieu_word, ds_to)
        ten_file = f"KH_KiemTra_UyThac_{int(nam_kh)}_{don_vi_kt[:20].replace(' ','_')}"
        col1, col2 = st.columns(2)
        with col1:
            st.download_button(
                "⬇️ Tải Word (.docx)",
                data=docx_bytes,
                file_name=ten_file + ".docx",
                mime="application/vnd.openxmlformats-officedocument"
                     ".wordprocessingml.document",
                key="kh_docx",
            )
        with col2:
            with st.spinner("Đang tạo PDF..."):
                pdf_bytes = docx_bytes_to_pdf(docx_bytes)
            if pdf_bytes:
                st.download_button(
                    "⬇️ Tải PDF",
                    data=pdf_bytes,
                    file_name=ten_file + ".pdf",
                    mime="application/pdf",
                    key="kh_pdf",
                )
            else:
                st.caption("⚠️ PDF không khả dụng — cần MS Word trên server")


def _render_mau06(df: pd.DataFrame, pgd_user: str) -> None:
    st.markdown("#### 📋 Mẫu 06/TD & 06A/TD — Phiếu kiểm tra sử dụng vốn")
    st.caption("Quy định: kiểm tra 100% món vay trong 30 ngày sau giải ngân. "
               "Thời điểm kiểm tra cụ thể do CBTD nhập khi đi thực địa.")
    if df is None or df.empty:
        st.warning("Chưa có dữ liệu HSTD."); return
    if COT_NGAY_VAY not in df.columns:
        st.warning(f"Không tìm thấy cột '{COT_NGAY_VAY}'."); return

    c1, c2 = st.columns(2)
    loai_mau = c1.radio("Loại mẫu", ["06/TD (bảng nhiều KH)",
                                       "06A/TD (từng KH riêng)"],
                         key="m06_loai")
    so_ngay  = c2.slider("Giải ngân trong N ngày qua", 7, 30, 30,
                          key="m06_ngay")
    st.caption("Ngày kiểm tra thực tế do Cán bộ hội đi kiểm tra ghi vào mẫu.")

    ngay_den = date.today()
    ngay_tu  = date.today() - timedelta(days=so_ngay)
    st.caption(f"📅 {ngay_tu.strftime('%d/%m/%Y')} → {ngay_den.strftime('%d/%m/%Y')}")

    try:
        raw    = _loc_mau06(pickle.dumps(df), str(ngay_tu), str(ngay_den))
        df_m06 = pickle.loads(raw)
    except Exception as e:
        st.error(f"Lỗi: {e}"); return

    if df_m06.empty:
        st.success("✅ Không có món vay nào cần kiểm tra."); return

    tong_dn = df_m06[COT_TONG_DU_NO].sum() \
              if COT_TONG_DU_NO in df_m06.columns else 0
    ca, cb  = st.columns(2)
    ca.metric("Số món cần KT", fmt_so(len(df_m06)))
    cb.metric("Tổng dư nợ (triệu đồng)", fmt(tong_dn))
    st.dataframe(df_m06, use_container_width=True,
                 hide_index=True, height=300)

    # Form thông tin người kiểm tra
    with st.form("form_xuat_m06"):
        st.markdown("**Thông tin xuất mẫu:**")
        f1, f2 = st.columns(2)
        don_vi_kt = f1.selectbox(
            "Hội đoàn thể kiểm tra",
            options=DVUT_ORDER,
            key="m06_don_vi_kt"
        )
        ds_xa_m06 = [""] + sorted(df_m06[COT_TEN_XA].dropna().unique().tolist()) \
                    if COT_TEN_XA in df_m06.columns else [""]
        ten_xa = f1.selectbox(
            "Xã/Phường",
            options=ds_xa_m06,
            key="m06_ten_xa"
        )
        # Lọc Tổ theo Xã đã chọn (dùng session_state vì trong form)
        ten_xa_filter = st.session_state.get("m06_ten_xa", "")
        df_to_filter = df_m06[df_m06[COT_TEN_XA] == ten_xa_filter] \
                       if ten_xa_filter and COT_TEN_XA in df_m06.columns else df_m06
        ds_to_m06 = [""] + sorted(df_to_filter[COT_TEN_TO].dropna().unique().tolist()) \
                    if COT_TEN_TO in df_to_filter.columns else [""]
        ten_to = f1.selectbox(
            "Tổ TK&VV",
            options=ds_to_m06,
            key="m06_chon_to"
        )
        can_bo_1  = f2.text_input("Cán bộ kiểm tra 1", key="m06_can_bo_1")
        chuc_vu_1 = f2.text_input("Chức vụ 1", key="m06_chuc_vu_1")
        can_bo_2  = f2.text_input("Cán bộ kiểm tra 2 (nếu có)", key="m06_can_bo_2")
        dia_ban   = f2.text_input("Địa bàn kiểm tra",
                                   placeholder="Ấp..., xã...",
                                   key="m06_dia_ban")
        ngay_kt   = f2.date_input("Ngày kiểm tra",
                                   value=date.today(), key="m06_ngay_kt")
        submitted = st.form_submit_button("📄 Tạo Word")

    if submitted:
        df_xuat = df_m06.copy()
        if ten_to:
            df_xuat = df_xuat[df_xuat[COT_TEN_TO] == ten_to] \
                      if COT_TEN_TO in df_xuat.columns else df_xuat
        # Tính Nợ lãi
        if COT_LAI_TON in df_xuat.columns and COT_LAI_TON_QH in df_xuat.columns:
            df_xuat["Nợ lãi"] = (df_xuat[COT_LAI_TON].fillna(0) +
                                  df_xuat[COT_LAI_TON_QH].fillna(0))
        # Build context cho template
        ds_kh = []
        for i, (_, row) in enumerate(df_xuat.iterrows(), 1):
            ds_kh.append({
                "stt":      i,
                "ten_kh":  str(row.get(COT_TEN_KH, "")),
                "so_ku":   str(row.get(COT_SO_KU, "")),
                "ten_ct":  str(row.get(COT_TEN_CT, "")),
                "muc_vay": fmt(row.get(COT_MUC_VAY, 0)),
                "du_no":   fmt(row.get(COT_TONG_DU_NO, 0)),
                "muc_dich": str(row.get("Mục đích sử dụng vốn vay", "")),
                "no_lai":  fmt(row.get("Nợ lãi", 0)),
            })
        context = {
            "don_vi_kt":   don_vi_kt,
            "ten_xa":      ten_xa,
            "ten_to":      ten_to,
            "can_bo_1":    can_bo_1,
            "chuc_vu_1":   chuc_vu_1,
            "can_bo_2":    can_bo_2,
            "dia_ban":     dia_ban,
            "ngay_kt":     ngay_kt.strftime("%d/%m/%Y"),
            "ngay":        ngay_kt.day,
            "thang":       ngay_kt.month,
            "nam":         ngay_kt.year,
            "so_kh_kt":    len(df_xuat),
            "ds_kh":       ds_kh,
            "tong_muc_vay": fmt(df_xuat[COT_MUC_VAY].sum()
                               if COT_MUC_VAY in df_xuat.columns else 0),
            "tong_du_no":  fmt(df_xuat[COT_TONG_DU_NO].sum()
                               if COT_TONG_DU_NO in df_xuat.columns else 0),
        }
        loai_word = "06" if "06/TD" in loai_mau else "06A"
        with st.spinner("Đang tạo file..."):
            du_lieu_word = {
                "don_vi_kt": don_vi_kt,
                "ten_xa": ten_xa,
                "ten_to": ten_to,
                "can_bo_1": can_bo_1,
                "chuc_vu_1": chuc_vu_1,
                "can_bo_2": can_bo_2,
                "dia_ban": dia_ban,
                "ngay_kt": ngay_kt,
            }
            docx_bytes = _tao_word_mau06(du_lieu_word, df_xuat, loai=loai_word)
        ten_file = f"Mau06TD_{pgd_user}_{ngay_kt.strftime('%d%m%Y')}"
        col1, col2 = st.columns(2)
        with col1:
            st.download_button(
                "⬇️ Tải Word (.docx)",
                data=docx_bytes,
                file_name=ten_file + ".docx",
                mime="application/vnd.openxmlformats-officedocument"
                     ".wordprocessingml.document",
                key="m06_docx",
            )
        with col2:
            with st.spinner("Đang tạo PDF..."):
                pdf_bytes = docx_bytes_to_pdf(docx_bytes)
            if pdf_bytes:
                st.download_button(
                    "⬇️ Tải PDF",
                    data=pdf_bytes,
                    file_name=ten_file + ".pdf",
                    mime="application/pdf",
                    key="m06_pdf",
                )
            else:
                st.caption("⚠️ PDF không khả dụng — cần MS Word trên server")


def _render_mau15(df: pd.DataFrame, pgd_user: str) -> None:
    st.markdown("#### 📋 Mẫu 15/TD — Danh sách đối chiếu số dư")
    st.caption("Đối chiếu nợ gốc, nợ lãi, số dư tiền gửi TK từng tổ viên.")
    if df is None or df.empty:
        st.warning("Chưa có dữ liệu HSTD."); return

    # Chọn Tổ TK&VV
    ds_to = sorted(df[COT_TEN_TO].dropna().unique().tolist()) \
            if COT_TEN_TO in df.columns else []
    if not ds_to:
        st.warning("Không có dữ liệu Tổ TK&VV."); return

    c1, c2 = st.columns(2)
    chon_dvut = c1.selectbox("Hội đoàn thể", ["Tất cả"] + DVUT_ORDER,
                               key="m15_dvut")
    # Lọc Tổ theo DVUT
    df_filter = df.copy()
    if chon_dvut != "Tất cả" and COT_DVUT in df_filter.columns:
        df_filter = df_filter[df_filter[COT_DVUT] == chon_dvut]
    ds_to_filter = sorted(df_filter[COT_TEN_TO].dropna().unique().tolist()) \
                   if COT_TEN_TO in df_filter.columns else []
    chon_to = c2.selectbox("Tổ TK&VV", ds_to_filter, key="m15_to")

    if not chon_to:
        st.info("Chọn Tổ TK&VV để xem dữ liệu."); return

    try:
        df_to = pickle.loads(_loc_mau15(pickle.dumps(df), chon_to))
    except Exception as e:
        st.error(f"Lỗi: {e}"); return

    if df_to.empty:
        st.info(f"Không có dữ liệu cho Tổ **{chon_to}**."); return

    # KPI
    tong_goc = df_to[COT_TONG_DU_NO].sum() if COT_TONG_DU_NO in df_to.columns else 0
    tong_lai = df_to["Nợ lãi"].sum() if "Nợ lãi" in df_to.columns else 0
    tong_tg  = df_to[COT_SO_DU_TG].sum() if COT_SO_DU_TG in df_to.columns else 0
    k1, k2, k3, k4 = st.columns(4)
    k1.metric("Số KH", fmt_so(len(df_to)))
    k2.metric("Tổng nợ gốc (triệu đồng)", fmt(tong_goc))
    k3.metric("Tổng nợ lãi (triệu đồng)", fmt(tong_lai))
    k4.metric("Tổng TG TK (triệu đồng)", fmt(tong_tg))
    st.dataframe(df_to, use_container_width=True, hide_index=True, height=350)

    # Tự động lấy xã và tổ trưởng từ Tổ đang chọn
    xa_cua_to = ""
    ten_to_truong = ""
    if chon_to:
        if COT_TEN_XA in df.columns and COT_TEN_TO in df.columns:
            s_xa = df[df[COT_TEN_TO] == chon_to][COT_TEN_XA].dropna()
            xa_cua_to = s_xa.iloc[0] if not s_xa.empty else ""
        for cot in ["Tên Tổ trưởng", "Tổ trưởng", "Họ tên Tổ trưởng"]:
            if cot in df.columns:
                s_tt = df[df[COT_TEN_TO] == chon_to][cot].dropna()
                ten_to_truong = str(s_tt.iloc[0]) if not s_tt.empty else ""
                break

    # Form xuất Word
    with st.form("form_xuat_m15"):
        st.markdown("**Thông tin xuất mẫu:**")
        f1, f2 = st.columns(2)
        pgd = f1.text_input(
            "PGD",
            value=pgd_user,
            disabled=True,
            key="m15_pgd"
        )
        ten_xa = f1.text_input(
            "Xã/Phường",
            value=xa_cua_to,
            disabled=True,
            help="Tự động lấy theo Tổ TK&VV đã chọn",
            key="m15_ten_xa"
        )
        to_truong = f1.text_input(
            "Tổ trưởng",
            value=ten_to_truong,
            help="Tự động lấy từ HSTD, có thể sửa lại nếu cần",
            key="m15_to_truong"
        )
        ma_to     = f1.text_input("Mã Tổ")
        dia_chi   = f2.text_input("Địa chỉ Tổ")
        can_bo_kt = f2.text_input("Cán bộ đối chiếu")
        ngay_chot = f2.date_input("Ngày chốt số liệu", value=date.today(),
                                   key="m15_ngay_chot")
        submitted = st.form_submit_button("📄 Tạo Word")

    if submitted:
        # Build context cho template Mẫu 15/TD
        context = {
            "pgd":         pgd,
            "ten_xa":      ten_xa,
            "ten_to":      chon_to,
            "to_truong":   to_truong,
            "ma_to":       ma_to,
            "dia_chi":     dia_chi,
            "can_bo_kt":   can_bo_kt,
            "ngay_chot":   ngay_chot.strftime("%d/%m/%Y"),
            "ngay":        ngay_chot.day,
            "thang":       ngay_chot.month,
            "nam":         ngay_chot.year,
            "tong_du_no":  fmt(tong_goc),
            "tong_lai":    fmt(tong_lai),
            "tong_tg":     fmt(tong_tg),
            "so_kh":       len(df_to),
            "ds_kh": [
                {
                    "stt":    i,
                    "ten_kh": str(row.get(COT_TEN_KH, "")),
                    "ten_ct": str(row.get(COT_TEN_CT, "")),
                    "so_ku":  str(row.get(COT_SO_KU, "")),
                    "du_no":  fmt(row.get(COT_TONG_DU_NO, 0)),
                    "lai":    fmt(row.get("Nợ lãi", 0)),
                    "tg":     fmt(row.get(COT_SO_DU_TG, 0)),
                }
                for i, (_, row) in enumerate(df_to.iterrows(), 1)
            ],
        }

        with st.spinner("Đang tạo file..."):
            du_lieu_word = {
                "pgd": pgd,
                "ten_xa": ten_xa,
                "ten_to": chon_to,
                "to_truong": to_truong,
                "ma_to": ma_to,
                "dia_chi": dia_chi,
                "can_bo_kt": can_bo_kt,
                "ngay_chot": ngay_chot,
            }
            docx_bytes = _tao_word_mau15(du_lieu_word, df_to)
        ten_file = f"Mau15TD_{chon_to.replace(' ','_')}_{ngay_chot.strftime('%d%m%Y')}"
        col1, col2 = st.columns(2)
        with col1:
            st.download_button(
                "⬇️ Tải Word (.docx)",
                data=docx_bytes,
                file_name=ten_file + ".docx",
                mime="application/vnd.openxmlformats-officedocument"
                     ".wordprocessingml.document",
                key="m15_docx",
            )
        with col2:
            with st.spinner("Đang tạo PDF..."):
                pdf_bytes = docx_bytes_to_pdf(docx_bytes)
            if pdf_bytes:
                st.download_button(
                    "⬇️ Tải PDF",
                    data=pdf_bytes,
                    file_name=ten_file + ".pdf",
                    mime="application/pdf",
                    key="m15_pdf",
                )
            else:
                st.caption("⚠️ PDF không khả dụng — cần MS Word trên server")


def _render_bien_ban(df: pd.DataFrame, pgd_user: str) -> None:
    st.markdown("#### 📋 Biên bản & Báo cáo tổng hợp")
    loai = st.radio(
        "Loại biên bản",
        [
            "📋 Mẫu 16/TD — Kiểm tra CT-XH Tổ TK&VV",
            "📄 Biên bản xác minh nợ chiếm dụng",
            "📊 Báo cáo tổng hợp ủy thác (Excel)",
        ],
        horizontal=True,
        key="bb_loai",
    )

    if loai.startswith("📋"):
        if df is None or df.empty:
            st.warning("Chưa có dữ liệu HSTD.")
            return

        with st.form("form_bb_m16"):
            c1, c2 = st.columns(2)
            don_vi_kt = c1.selectbox("Hội đoàn thể kiểm tra", DVUT_ORDER, key="bb_dvkt")
            ds_xa_bb = (
                sorted(df[COT_TEN_XA].dropna().unique().tolist())
                if COT_TEN_XA in df.columns
                else []
            )
            ten_xa = c1.selectbox("Xã/Phường", [""] + ds_xa_bb, key="bb_xa")
            ten_xa_filter = st.session_state.get("bb_xa", "")
            df_to_filter = (
                df[df[COT_TEN_XA] == ten_xa_filter]
                if ten_xa_filter and COT_TEN_XA in df.columns
                else df
            )
            ds_to_bb = (
                sorted(df_to_filter[COT_TEN_TO].dropna().unique().tolist())
                if COT_TEN_TO in df_to_filter.columns
                else []
            )
            ten_to = c1.selectbox("Tổ TK&VV", ["Tất cả"] + ds_to_bb, key="bb_to")

            can_bo_kt = c2.text_input("Cán bộ kiểm tra", key="bb_cb")
            chuc_vu = c2.text_input("Chức vụ", key="bb_cv")
            ngay_kt = c2.date_input("Ngày kiểm tra", value=date.today(), key="bb_ngay")

            nhan_xet = st.text_area("Nhận xét", height=80, key="bb_nhanxet")
            ket_luan = st.text_area("Kết luận & Kiến nghị", height=80, key="bb_ketluan")
            submitted_m16 = st.form_submit_button("📄 Tạo Word", type="primary")

        if submitted_m16:
            df_xuat = df.copy()
            if ten_xa and COT_TEN_XA in df_xuat.columns:
                df_xuat = df_xuat[df_xuat[COT_TEN_XA] == ten_xa]
            if ten_to != "Tất cả" and COT_TEN_TO in df_xuat.columns:
                df_xuat = df_xuat[df_xuat[COT_TEN_TO] == ten_to]

            du_lieu = {
                "don_vi_kt": don_vi_kt,
                "ten_xa": ten_xa,
                "ten_to": ten_to if ten_to != "Tất cả" else "",
                "can_bo_kt": can_bo_kt,
                "chuc_vu": chuc_vu,
                "ngay_kt": ngay_kt,
                "ngay": ngay_kt.day,
                "thang": ngay_kt.month,
                "nam": ngay_kt.year,
                "nhan_xet": nhan_xet,
                "ket_luan": ket_luan,
            }

            with st.spinner("Đang tạo file..."):
                docx_bytes = _tao_word_mau16(du_lieu, df_xuat)
            ten_to_file = (ten_to if ten_to != "Tất cả" else "TatCa").replace(" ", "_")
            ten_file = f"Mau16TD_{ten_to_file}_{ngay_kt.strftime('%d%m%Y')}"

            col1, col2 = st.columns(2)
            with col1:
                st.download_button(
                    "⬇️ Tải Word (.docx)",
                    data=docx_bytes,
                    file_name=ten_file + ".docx",
                    mime="application/vnd.openxmlformats-officedocument"
                    ".wordprocessingml.document",
                    key="bb_m16_docx",
                )
            with col2:
                with st.spinner("Đang tạo PDF..."):
                    pdf_bytes = docx_bytes_to_pdf(docx_bytes)
                if pdf_bytes:
                    st.download_button(
                        "⬇️ Tải PDF",
                        data=pdf_bytes,
                        file_name=ten_file + ".pdf",
                        mime="application/pdf",
                        key="bb_m16_pdf",
                    )
                else:
                    st.caption("⚠️ PDF không khả dụng — cần MS Word trên server")

        return

    if loai.startswith("📄"):
        with st.form("form_bb_xm"):
            c1, c2 = st.columns(2)
            ten_kh = c1.text_input("Họ tên khách hàng", key="xm_kh")
            so_ku = c1.text_input("Số khế ước", key="xm_sku")
            so_tien = c1.number_input(
                "Số tiền chiếm dụng (triệu đồng)",
                min_value=0.0,
                step=0.1,
                key="xm_sotien",
            )
            can_bo_lap = c2.text_input("Cán bộ lập biên bản", key="xm_cb")
            ngay_lap = c2.date_input("Ngày lập", value=date.today(), key="xm_ngay")
            ly_do = st.text_area("Lý do / Hoàn cảnh", height=80, key="xm_lydo")
            bien_phap = st.text_area("Biện pháp xử lý", height=80, key="xm_bien_phap")
            submitted_xm = st.form_submit_button("📄 Tạo Word", type="primary")

        if submitted_xm:
            du_lieu = {
                "ten_kh": ten_kh,
                "so_ku": so_ku,
                "so_tien": f"{so_tien:,.1f}",
                "ly_do": ly_do,
                "bien_phap": bien_phap,
                "can_bo_lap": can_bo_lap,
                "pgd_user": pgd_user,
                "ngay_lap": ngay_lap,
                "ngay": ngay_lap.day,
                "thang": ngay_lap.month,
                "nam": ngay_lap.year,
            }
            with st.spinner("Đang tạo file..."):
                docx_bytes = _tao_word_bien_ban_xac_minh(du_lieu)
            ten_file = f"BBXacMinh_{so_ku}_{ngay_lap.strftime('%d%m%Y')}"

            col1, col2 = st.columns(2)
            with col1:
                st.download_button(
                    "⬇️ Tải Word (.docx)",
                    data=docx_bytes,
                    file_name=ten_file + ".docx",
                    mime="application/vnd.openxmlformats-officedocument"
                    ".wordprocessingml.document",
                    key="xm_docx",
                )
            with col2:
                with st.spinner("Đang tạo PDF..."):
                    pdf_bytes = docx_bytes_to_pdf(docx_bytes)
                if pdf_bytes:
                    st.download_button(
                        "⬇️ Tải PDF",
                        data=pdf_bytes,
                        file_name=ten_file + ".pdf",
                        mime="application/pdf",
                        key="xm_pdf",
                    )
                else:
                    st.caption("⚠️ PDF không khả dụng — cần MS Word trên server")
        return

    if df is None or df.empty:
        st.warning("Chưa có dữ liệu.")
        return

    try:
        df_th = pickle.loads(_tinh_theo_dvut(pickle.dumps(df)))
    except Exception as e:
        st.error(f"Lỗi: {e}")
        df_th = pd.DataFrame()

    if df_th.empty:
        st.info("Không có dữ liệu.")
        return

    hien = df_th.rename(
        columns={
            COT_DVUT: "Hội đoàn thể",
            "so_to": "Số Tổ TK&VV",
            "so_kh": "Số hộ vay",
            "tong_dn": "Dư nợ (triệu đồng)",
            "nqh": "NQH (triệu đồng)",
            "lai_ton": "Lãi tồn (triệu đồng)",
        }
    )
    for col in ["Dư nợ (triệu đồng)", "NQH (triệu đồng)", "Lãi tồn (triệu đồng)"]:
        if col in hien.columns:
            hien[col] = hien[col].apply(fmt)
    st.dataframe(hien, use_container_width=True, hide_index=True)

    buf = xuat_excel({"Tổng hợp ủy thác": df_th})
    ten_file = f"TongHopUyThac_{pgd_user}_{date.today().strftime('%d%m%Y')}.xlsx"
    st.download_button(
        "📥 Tải Excel",
        data=buf,
        file_name=ten_file,
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        key="th_dl",
    )


def _render_bb_ct_cx(df: pd.DataFrame, pgd_user: str | None,
                     username: str, role: str) -> None:
    """Sub-tab 3 — Nhập + lưu Mẫu 02/BB-CT và 03/BB-CX với theo dõi tiến độ."""
    st.markdown("#### 📝 Biên bản kiểm tra tổ chức CT-XH (Mẫu 02/BB-CT & 03/BB-CX)")
    st.caption(
        "Nhập kết quả kiểm tra và lưu vào hệ thống để theo dõi tiến độ xử lý kiến nghị. "
        "Xuất Word/PDF trực tiếp từ dữ liệu đã lưu."
    )

    loai_sel = st.radio(
        "Loại biên bản",
        ["02/BB-CT — Tổ chức CT-XH cấp tỉnh", "03/BB-CX — Tổ chức CT-XH cấp xã"],
        horizontal=True, key="bbct_loai",
    )
    cap = "tinh" if "CT" in loai_sel else "xa"
    kv_prefix = "ut_bbct" if cap == "tinh" else "ut_bbcx"

    c_nam, c_pgd = st.columns(2)
    nam = int(c_nam.number_input(
        "Năm", value=date.today().year,
        min_value=2020, max_value=2035, step=1, key="bbct_nam",
    ))

    if pgd_user:
        scope = pgd_user
        c_pgd.info(f"Đơn vị: **{pgd_user}**")
    else:
        scope = c_pgd.selectbox(
            "PGD / Đơn vị quản lý hồ sơ",
            options=DS_PGD, key="bbct_pgd_sel",
        )

    slug = pgd_slug(scope) if scope else "cn"
    kv_key = f"{kv_prefix}_{slug}_{nam}"

    ds_luu: list = db.doc_kv(kv_key) or []

    # ── Danh sách biên bản đã lưu ──────────────────────────────────────────
    if ds_luu:
        so_hieu_mau = "02/BB-CT" if cap == "tinh" else "03/BB-CX"
        st.markdown(f"##### Biên bản {so_hieu_mau} đã lưu — {scope} ({nam})")
        for bb in reversed(ds_luu):
            tt = bb.get("trang_thai", "cho_xu_ly")
            tt_label = {
                "cho_xu_ly": "🔴 Chờ xử lý",
                "da_xu_ly": "✅ Đã xử lý",
                "khong_ton_tai": "🟢 Không tồn tại",
            }.get(tt, tt)
            ngay_str = bb.get("ngay_kt", "")
            ten_dv = bb.get("ten_don_vi", "")
            with st.expander(
                f"[{so_hieu_mau}] {ngay_str} — {ten_dv} | {tt_label}"
            ):
                col_i1, col_i2 = st.columns(2)
                col_i1.markdown(f"**Đơn vị KT:** {bb.get('don_vi_kt', '')}")
                col_i1.markdown(f"**Trưởng đoàn:** {bb.get('truong_doan', '')}")
                col_i2.markdown(f"**Đại diện được KT:** {bb.get('dai_dien_dc', '')}")
                col_i2.markdown(f"**Hạn hoàn thành:** {bb.get('han_hoan_thanh', '')}")
                if bb.get("kien_nghi"):
                    st.markdown(f"**Kiến nghị:** {bb['kien_nghi']}")
                if bb.get("ket_qua_xu_ly"):
                    st.markdown(f"**Kết quả xử lý:** {bb['ket_qua_xu_ly']}")

                rec_id = bb.get("id", "")
                ten_file = (
                    f"{'BB_CT' if cap == 'tinh' else 'BB_CX'}"
                    f"_{ten_dv[:20].replace(' ', '_')}"
                    f"_{ngay_str.replace('-', '')}"
                )
                ss_key = f"bbct_bytes_{rec_id}"
                if st.button("📄 Tạo Word / PDF", key=f"bbct_gen_{rec_id}"):
                    st.session_state[ss_key] = _tao_word_bb_ct_cx(
                        bb, cap=bb.get("loai_cap", cap)
                    )
                if ss_key in st.session_state:
                    docx_b = st.session_state[ss_key]
                    col_e1, col_e2 = st.columns(2)
                    with col_e1:
                        st.download_button(
                            "⬇️ Tải Word",
                            data=docx_b,
                            file_name=ten_file + ".docx",
                            mime="application/vnd.openxmlformats-officedocument"
                                 ".wordprocessingml.document",
                            key=f"bbct_dl_{rec_id}",
                        )
                    with col_e2:
                        pdf_b = docx_bytes_to_pdf(docx_b)
                        if pdf_b:
                            st.download_button(
                                "⬇️ Tải PDF",
                                data=pdf_b,
                                file_name=ten_file + ".pdf",
                                mime="application/pdf",
                                key=f"bbct_pdf_{rec_id}",
                            )
        st.divider()

    # ── Form nhập biên bản mới ─────────────────────────────────────────────
    so_hieu_label = "02/BB-CT" if cap == "tinh" else "03/BB-CX"
    st.markdown(f"##### Nhập biên bản {so_hieu_label} mới")

    with st.form(f"form_bb_{cap}_moi", clear_on_submit=True):
        st.markdown("**Thông tin chung**")
        fc1, fc2 = st.columns(2)
        dvut      = fc1.selectbox("Hội đoàn thể kiểm tra", DVUT_ORDER,  key=f"bbct_dvut_{cap}")
        don_vi_kt = fc1.text_input("Tên đơn vị kiểm tra (đầy đủ)",      key=f"bbct_dvkt_{cap}")
        ten_don_vi = fc1.text_input("Đơn vị được kiểm tra",              key=f"bbct_dvdc_{cap}",
                                    placeholder="Hội ... xã/tỉnh ...")
        ngay_kt    = fc2.date_input("Ngày kiểm tra", value=date.today(), key=f"bbct_ngay_{cap}")
        truong_doan = fc2.text_input("Trưởng đoàn kiểm tra",             key=f"bbct_td_{cap}")
        can_bo_2   = fc2.text_input("Cán bộ kiểm tra 2 (nếu có)",       key=f"bbct_cb2_{cap}")
        dai_dien_dc = fc2.text_input("Đại diện đơn vị được kiểm tra",   key=f"bbct_dddc_{cap}")
        chuc_vu_dc = fc2.text_input("Chức vụ đại diện",                  key=f"bbct_cvdc_{cap}")

        st.markdown("**II. Kết quả thực hiện (theo Phụ lục I VB 727)**")
        muc_list = [
            ("tuyen_truyen",    "1. Công tác tuyên truyền, vận động"),
            ("kiem_tra_giam_sat", "2. Công tác kiểm tra, giám sát"),
            ("tap_huan",        "3. Công tác tập huấn"),
            ("phoi_hop_nhcs",   "4. Hoạt động phối hợp với NHCSXH"),
        ]
        if cap == "tinh":
            muc_list.append(("trach_nhiem", "5. Trách nhiệm của tổ chức CT-XH cấp tỉnh"))

        nd_results: dict[str, dict] = {}
        for field_key, ten_muc in muc_list:
            st.markdown(f"*{ten_muc}*")
            mc1, mc2 = st.columns(2)
            kq = mc1.text_area("a) Kết quả", height=60, key=f"bbct_{field_key}_kq_{cap}")
            tt_nd = mc2.text_area("b) Tồn tại", height=60, key=f"bbct_{field_key}_tt_{cap}")
            nd_results[field_key] = {"ket_qua": kq, "ton_tai": tt_nd}

        st.markdown("**III. Đánh giá, Nhận xét & Kiến nghị**")
        ek1, ek2 = st.columns(2)
        uu_diem     = ek1.text_area("Ưu điểm",        height=80, key=f"bbct_uu_{cap}")
        ton_tai_ch  = ek2.text_area("Tồn tại chung",  height=80, key=f"bbct_tt_{cap}")
        kien_nghi   = st.text_area("Kiến nghị",       height=80, key=f"bbct_kn_{cap}")

        hk1, hk2 = st.columns(2)
        han_ht = hk1.date_input("Hạn hoàn thành kiến nghị",
                                  value=date.today(), key=f"bbct_han_{cap}")
        tt_sel = hk2.selectbox(
            "Trạng thái tồn tại",
            options=["cho_xu_ly", "khong_ton_tai"],
            format_func=lambda x: {
                "cho_xu_ly": "🔴 Có tồn tại — chờ xử lý",
                "khong_ton_tai": "🟢 Không có tồn tại",
            }.get(x, x),
            key=f"bbct_tt_select_{cap}",
        )
        y_kien = st.text_area("IV. Ý kiến đơn vị được kiểm tra", height=60,
                               key=f"bbct_ykien_{cap}")
        submitted = st.form_submit_button("💾 Lưu biên bản", type="primary")

    if submitted:
        new_id = uuid.uuid4().hex[:8]
        record = {
            "id":           new_id,
            "kv_key":       kv_key,
            "loai":         "CT" if cap == "tinh" else "CX",
            "loai_cap":     cap,
            "ngay_kt":      ngay_kt.strftime("%Y-%m-%d"),
            "dvut":         dvut,
            "don_vi_kt":    don_vi_kt,
            "ten_don_vi":   ten_don_vi,
            "truong_doan":  truong_doan,
            "can_bo_2":     can_bo_2,
            "dai_dien_dc":  dai_dien_dc,
            "chuc_vu_dc":   chuc_vu_dc,
            "dia_danh":     scope,
            **nd_results,
            "uu_diem":       uu_diem,
            "ton_tai_chung": ton_tai_ch,
            "kien_nghi":     kien_nghi,
            "han_hoan_thanh": han_ht.strftime("%Y-%m-%d"),
            "trang_thai":    tt_sel,
            "y_kien_don_vi_dc": y_kien,
            "ket_qua_xu_ly": "",
            "ngay_cap_nhat": date.today().strftime("%Y-%m-%d"),
            "nguoi_cap_nhat": username,
        }
        db.ghi_kv(kv_key, ds_luu + [record], username)
        loai_str = "bb_ct" if cap == "tinh" else "bb_cx"
        db.ghi_audit(username, f"luu_{loai_str}",
                      f"Mẫu {'02/BB-CT' if cap == 'tinh' else '03/BB-CX'} — "
                      f"{ten_don_vi} ngày {ngay_kt.strftime('%d/%m/%Y')}")
        st.success(
            f"✅ Đã lưu biên bản {'02/BB-CT' if cap == 'tinh' else '03/BB-CX'} — {ten_don_vi}"
        )
        st.rerun()


def _render_theo_doi_bc_th(pgd_user: str | None,
                            username: str, role: str) -> None:
    """Sub-tab 7 — Theo dõi tiến độ xử lý kiến nghị + xuất Mẫu 04/BC-TH."""
    st.markdown("#### 📊 Theo dõi tiến độ & Báo cáo tổng hợp (Mẫu 04/BC-TH)")

    # ── Section 1: Theo dõi ────────────────────────────────────────────────
    st.markdown("##### I. Theo dõi tiến độ xử lý kiến nghị")

    tc1, tc2, tc3 = st.columns(3)
    nam_td = int(tc1.number_input(
        "Năm", value=date.today().year,
        min_value=2020, max_value=2035, step=1, key="td_nam",
    ))
    loai_td = tc2.selectbox(
        "Loại", ["Tất cả", "BB-CT (cấp tỉnh)", "BB-CX (cấp xã)"], key="td_loai"
    )
    tt_td = tc3.selectbox(
        "Trạng thái",
        ["Tất cả", "Chờ xử lý", "Đã xử lý", "Không tồn tại"],
        key="td_tt",
    )

    # Load records
    all_records: list[dict] = []
    if pgd_user:
        slug = pgd_slug(pgd_user)
        for pref in ["ut_bbct", "ut_bbcx"]:
            recs = db.doc_kv(f"{pref}_{slug}_{nam_td}") or []
            all_records.extend(recs)
    else:
        for pref in ["ut_bbct", "ut_bbcx"]:
            all_kv: dict = db.doc_kv_prefix(f"{pref}_") or {}
            for key, recs in all_kv.items():
                if key.endswith(f"_{nam_td}") and isinstance(recs, list):
                    all_records.extend(recs)

    # Filter theo loại
    if "BB-CT" in loai_td:
        all_records = [r for r in all_records if r.get("loai") == "CT"]
    elif "BB-CX" in loai_td:
        all_records = [r for r in all_records if r.get("loai") == "CX"]

    # Filter theo trạng thái
    tt_map = {"Chờ xử lý": "cho_xu_ly", "Đã xử lý": "da_xu_ly",
              "Không tồn tại": "khong_ton_tai"}
    if tt_td != "Tất cả":
        all_records = [r for r in all_records if r.get("trang_thai") == tt_map.get(tt_td)]

    if not all_records:
        st.info("Chưa có biên bản nào trong kỳ này.")
    else:
        rows = []
        for r in sorted(all_records, key=lambda x: x.get("ngay_kt", ""), reverse=True):
            tt = r.get("trang_thai", "cho_xu_ly")
            tt_label = {
                "cho_xu_ly": "🔴 Chờ xử lý",
                "da_xu_ly": "✅ Đã xử lý",
                "khong_ton_tai": "🟢 Không tồn tại",
            }.get(tt, tt)
            mau_so = "02/BB-CT" if r.get("loai") == "CT" else "03/BB-CX"
            kn_text = r.get("kien_nghi") or ""
            rows.append({
                "ID": r.get("id", ""),
                "Mẫu số": mau_so,
                "Ngày KT": r.get("ngay_kt", ""),
                "Đơn vị được KT": r.get("ten_don_vi", ""),
                "Kiến nghị": kn_text[:80] + "..." if len(kn_text) > 80 else kn_text,
                "Hạn hoàn thành": r.get("han_hoan_thanh", ""),
                "Trạng thái": tt_label,
                "Kết quả xử lý": r.get("ket_qua_xu_ly", ""),
            })
        df_td = pd.DataFrame(rows)
        st.dataframe(
            df_td.drop(columns=["ID"]),
            use_container_width=True, hide_index=True,
        )

        # Cập nhật trạng thái (chỉ CN role)
        if la_phan_he_cn(role):
            st.markdown("**Cập nhật trạng thái xử lý:**")
            cho_xu_ly = [r for r in all_records if r.get("trang_thai") == "cho_xu_ly"]
            if not cho_xu_ly:
                st.success("✅ Tất cả kiến nghị đã được xử lý.")
            else:
                opt_map = {
                    f"{r.get('ngay_kt','')} — {r.get('ten_don_vi','')} [{r.get('id','')}]": r
                    for r in cho_xu_ly
                }
                chon_label = st.selectbox(
                    "Chọn biên bản cần cập nhật",
                    options=[""] + list(opt_map.keys()),
                    key="td_chon_label",
                )
                if chon_label:
                    target = opt_map[chon_label]
                    ket_qua_xl = st.text_area(
                        "Kết quả xử lý", height=60, key="td_kq_xl"
                    )
                    if st.button("✅ Đánh dấu đã xử lý", key="td_btn_xl"):
                        kv_key_t = target.get("kv_key", "")
                        if kv_key_t:
                            ds_cur = db.doc_kv(kv_key_t) or []
                            updated = [
                                {**rec,
                                 "trang_thai": "da_xu_ly",
                                 "ket_qua_xu_ly": ket_qua_xl,
                                 "ngay_cap_nhat": date.today().strftime("%Y-%m-%d"),
                                 "nguoi_cap_nhat": username}
                                if rec.get("id") == target.get("id")
                                else rec
                                for rec in ds_cur
                            ]
                            db.ghi_kv(kv_key_t, updated, username)
                            db.ghi_audit(
                                username, "cap_nhat_trang_thai_bb",
                                f"ID {target.get('id')} — {target.get('ten_don_vi','')} → Đã xử lý"
                            )
                            st.success("✅ Đã cập nhật trạng thái!")
                            st.rerun()

    st.divider()

    # ── Section 2: Xuất BC-TH ──────────────────────────────────────────────
    st.markdown("##### II. Xuất Báo cáo tổng hợp (Mẫu 04/BC-TH)")

    all_for_bc: list[dict] = []
    if pgd_user:
        slug = pgd_slug(pgd_user)
        for pref in ["ut_bbct", "ut_bbcx"]:
            recs = db.doc_kv(f"{pref}_{slug}_{nam_td}") or []
            all_for_bc.extend(recs)
    else:
        for pref in ["ut_bbct", "ut_bbcx"]:
            all_kv2: dict = db.doc_kv_prefix(f"{pref}_") or {}
            for key, recs in all_kv2.items():
                if key.endswith(f"_{nam_td}") and isinstance(recs, list):
                    all_for_bc.extend(recs)

    if not all_for_bc:
        st.info("Không có biên bản nào để lập báo cáo tổng hợp.")
        return

    opt_bc = {
        f"[{'02/BB-CT' if r.get('loai')=='CT' else '03/BB-CX'}] "
        f"{r.get('ngay_kt','')} — {r.get('ten_don_vi','')}": r
        for r in all_for_bc
    }
    chon_bc = st.multiselect(
        "Chọn biên bản đưa vào báo cáo tổng hợp",
        options=list(opt_bc.keys()), key="bcth_chon",
    )
    if not chon_bc:
        st.info("Chọn ít nhất 1 biên bản để tạo báo cáo.")
        return

    ds_chon = [opt_bc[k] for k in chon_bc]

    with st.form("form_bc_th"):
        st.markdown("**Thông tin báo cáo:**")
        bc1, bc2 = st.columns(2)
        don_vi_kt_bc  = bc1.text_input("Đơn vị kiểm tra", key="bcth_dvkt")
        truong_doan_bc = bc1.text_input("Trưởng đoàn kiểm tra", key="bcth_td")
        cap_uy        = bc1.text_input("Cấp ủy, chính quyền tham dự (nếu có)", key="bcth_capuy")
        dia_danh_bc   = bc2.text_input("Địa danh ký", placeholder="Biên Hòa", key="bcth_dd")
        ngay_bc       = bc2.date_input("Ngày báo cáo", value=date.today(), key="bcth_ngay")
        noi_dung_kt   = st.text_area(
            "III. Nội dung kiểm tra",
            value="Theo Phụ lục I văn bản số 727/HD-NHCS ngày 11/02/2026.",
            height=60, key="bcth_ndkt",
        )
        st.markdown("**IV. Đánh giá & Kiến nghị:**")
        r1, r2 = st.columns(2)
        nx_ctxh    = r1.text_area("Nhận xét đối với CT-XH",       height=60, key="bcth_nx_ctxh")
        nx_to      = r2.text_area("Nhận xét đối với Tổ TK&VV",    height=60, key="bcth_nx_to")
        nx_to_vien = r1.text_area("Nhận xét đối với tổ viên",     height=60, key="bcth_nx_tov")
        kn_ctxh    = r2.text_area("Kiến nghị với CT-XH",          height=60, key="bcth_kn_ctxh")
        kn_nhcs    = r1.text_area("Kiến nghị với NHCSXH",         height=60, key="bcth_kn_nhcs")
        kn_cap_tren = r2.text_area("Kiến nghị với CT-XH cấp trên", height=60, key="bcth_kn_ct")
        submitted_bc = st.form_submit_button("📄 Tạo Báo cáo tổng hợp Word", type="primary")

    if submitted_bc:
        du_lieu_bc = {
            "don_vi_kt": don_vi_kt_bc, "truong_doan": truong_doan_bc,
            "dia_danh": dia_danh_bc, "ngay_bc": ngay_bc, "cap_uy": cap_uy,
            "noi_dung_kt": noi_dung_kt,
            "nx_ctxh": nx_ctxh, "nx_to": nx_to, "nx_to_vien": nx_to_vien,
            "kn_ctxh": kn_ctxh, "kn_nhcs": kn_nhcs, "kn_cap_tren": kn_cap_tren,
        }
        with st.spinner("Đang tạo file..."):
            docx_bytes = _tao_word_bc_th(du_lieu_bc, ds_chon)
        ten_file = f"BaoCaoTH_UyThac_{nam_td}"
        bc_c1, bc_c2 = st.columns(2)
        with bc_c1:
            st.download_button(
                "⬇️ Tải Word (.docx)",
                data=docx_bytes,
                file_name=ten_file + ".docx",
                mime="application/vnd.openxmlformats-officedocument"
                     ".wordprocessingml.document",
                key="bcth_dl_docx",
            )
        with bc_c2:
            pdf_bc = docx_bytes_to_pdf(docx_bytes)
            if pdf_bc:
                st.download_button(
                    "⬇️ Tải PDF",
                    data=pdf_bc,
                    file_name=ten_file + ".pdf",
                    mime="application/pdf",
                    key="bcth_dl_pdf",
                )
            else:
                st.caption("⚠️ PDF không khả dụng — cần MS Word trên server")
        db.ghi_audit(username, "xuat_bc_th",
                      f"Báo cáo tổng hợp năm {nam_td} — {len(ds_chon)} biên bản")


# ══════════════════════════════════════════════════════════════════════════════
# ENTRY POINT
# ══════════════════════════════════════════════════════════════════════════════

from tabs.base_tab import TabContext

def render(tab: DeltaGenerator, **kwargs) -> None:
    """Entry point — dùng chung cho ws_operation và ws_management."""
    ctx = TabContext(tab, **kwargs)
    df       = kwargs.get("df")
    pgd_user = ctx.pgd_user
    username = kwargs.get("username", "unknown")
    role     = normalize_role(str(kwargs.get("role", "user") or "user"))

    with ctx:
        st.subheader("🤝 Ủy thác — Hội đoàn thể")
        st.caption(
            "Theo dõi hoạt động ủy thác và các mẫu biểu kiểm tra "
            "theo văn bản 727/HD-NHCS."
        )
        sub1, sub2, sub3, sub4, sub5, sub6, sub7 = st.tabs([
            "📊 Theo Hội đoàn thể",
            "📋 Kế hoạch (01/KH)",
            "📝 Biên bản CT-XH",
            "📋 Mẫu 06/TD & 06A/TD",
            "📋 Mẫu 15/TD",
            "📋 Biên bản & Báo cáo",
            "📊 Theo dõi & BC-TH",
        ])
        with sub1: _render_theo_dvut(df)
        with sub2: _render_ke_hoach(df, pgd_user)
        with sub3: _render_bb_ct_cx(df, pgd_user, username, role)
        with sub4: _render_mau06(df, pgd_user)
        with sub5: _render_mau15(df, pgd_user)
        with sub6: _render_bien_ban(df, pgd_user)
        with sub7: _render_theo_doi_bc_th(pgd_user, username, role)
