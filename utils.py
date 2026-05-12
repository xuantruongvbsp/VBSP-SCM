"""Hàm tiện ích dùng chung toàn app."""
from io import BytesIO
import functools
import inspect
import pandas as pd
from datetime import datetime

import streamlit as st


def norm_col_header(s: str) -> str:
    """Chuẩn hóa tên cột: NBSP/thin-space → space, strip."""
    return str(s).replace("\u00a0", " ").replace("\u202f", " ").strip()


# ══════════════════════════════════════════════════════════════════════════════
# DECORATOR — Tự động ghi audit log + clear cache sau mỗi thao tác ghi
# ══════════════════════════════════════════════════════════════════════════════


def auto_audit(action: str = "", clear_cache: bool = True):
    """
    Decorator cho các hàm ghi dữ liệu trong VBSP-SCM.

    Tự động thực hiện sau khi hàm được bọc chạy thành công:
      1. Gọi db.ghi_audit(username, action, detail)
      2. Gọi st.cache_data.clear() nếu clear_cache=True

    Quy ước tham số hàm được bọc:
      - Tham số `username` (str): bắt buộc, dùng để ghi audit.
        Có thể là positional hoặc keyword argument.
      - Giá trị trả về của hàm được bọc sẽ được truyền nguyên vào `detail`.

    Cách dùng:
        @auto_audit(action="luu_khtd_pgd")
        def luu_khtd_pgd(pgd: str, data: dict, username: str) -> bool:
            db.ghi_kv(f"khtd_{pgd}", data, username)
            return True

    Nếu hàm raise exception, decorator KHÔNG ghi audit (để tránh log sai).
    """

    def decorator(func):
        @functools.wraps(func)
        def wrapper(*args, **kwargs):
            import db as _db
            import streamlit as _st

            # Xác định action label
            _action = action or func.__name__

            # Lấy username từ kwargs hoặc args theo tên tham số
            sig = inspect.signature(func)
            params = list(sig.parameters.keys())
            _username = kwargs.get("username", "system")
            if _username == "system" and "username" in params:
                idx = params.index("username")
                if idx < len(args):
                    _username = args[idx]

            # Chạy hàm gốc
            result = func(*args, **kwargs)

            # Ghi audit — detail là repr ngắn của result
            try:
                _detail = str(result)[:200] if result is not None else ""
                _db.ghi_audit(_username, _action, _detail)
            except Exception:
                pass  # Không để lỗi audit phá vỡ luồng chính

            # Xóa cache Streamlit
            if clear_cache:
                try:
                    _st.cache_data.clear()
                except Exception:
                    pass

            return result

        return wrapper

    return decorator


# ── Helper đọc cấu hình động ─────────────────────────────────────────────────
def lay_config(key: str, fallback):
    """
    Đọc cấu hình động theo key:
      1. Ưu tiên kv_store (Admin có thể chỉnh sửa qua giao diện).
      2. Fallback về giá trị mặc định từ config.py nếu kv_store chưa có.

    Params
    ------
    key      : key trong bảng kv_store (vd: "ds_pgd", "pgd_xa_map")
    fallback : giá trị mặc định lấy từ config.py

    Returns
    -------
    Dữ liệu đã parse từ JSON (list / dict) hoặc fallback.
    """
    try:
        import db
        val = db.doc_kv(key)
        return val if val is not None else fallback
    except Exception:
        return fallback


def format_df_vn(df: pd.DataFrame) -> pd.DataFrame:
    """
    Tự động format tất cả cột số trong DataFrame sang chuẩn Việt Nam.
    - Số nguyên lớn: dấu chấm phân cách hàng nghìn (1.234.567)
    - Số thập phân: dấu phẩy cho thập phân (1.234,56)
    - Cột tên chứa % / "tỷ lệ" / pct: thêm ký hiệu %
    """
    df = df.copy()
    for col in df.columns:
        if not pd.api.types.is_numeric_dtype(df[col]):
            continue
        if pd.api.types.is_bool_dtype(df[col]):
            continue
        col_str = str(col)
        col_lower = col_str.lower()
        if "%" in col_str or "tỷ lệ" in col_lower or "pct" in col_lower:
            df[col] = df[col].map(
                lambda x: (
                    f"{float(x):,.2f}%".replace(",", "X").replace(".", ",").replace("X", ".")
                    if pd.notna(x)
                    else ""
                )
            )
        elif "tỷ" in col_lower:
            df[col] = df[col].map(
                lambda x: (
                    f"{float(x):,.3f}".replace(",", "X").replace(".", ",").replace("X", ".")
                    if pd.notna(x)
                    else ""
                )
            )
        elif "triệu" in col_lower:
            df[col] = df[col].map(
                lambda x: (
                    f"{float(x):,.1f}".replace(",", "X").replace(".", ",").replace("X", ".")
                    if pd.notna(x)
                    else ""
                )
            )
        else:
            df[col] = df[col].map(
                lambda x: (
                    f"{float(x):,.0f}".replace(",", ".")
                    if pd.notna(x)
                    else ""
                )
            )
    return df


def hien_thi_dataframe_phan_trang(df, so_dong_moi_trang=500, key="df", **kwargs):
    """Hiển thị dataframe có phân trang, mặc định 500 dòng/trang.

    ``kwargs`` truyền thêm cho ``st.dataframe`` (vd: ``height``, ``column_config``,
    ``hide_index``). Với ``Styler``, không phân trang (hiển thị một lần).
    """
    try:
        from pandas.io.formats.style import Styler
    except ImportError:
        Styler = None  # type: ignore[misc, assignment]

    opts = {"use_container_width": True, "hide_index": True}
    opts.update(kwargs)

    if Styler is not None and isinstance(df, Styler):
        st.dataframe(df, **opts)
        return

    def _disp(chunk: pd.DataFrame) -> None:
        st.dataframe(format_df_vn(chunk), **opts)

    tong = len(df)
    if tong <= so_dong_moi_trang:
        _disp(df)
        return

    tong_trang = max(1, (tong + so_dong_moi_trang - 1) // so_dong_moi_trang)
    trang = st.number_input(
        f"Trang (1-{tong_trang}) · Tổng {tong:,} dòng",
        min_value=1,
        max_value=tong_trang,
        value=1,
        key=f"{key}_trang",
    )
    start = (trang - 1) * so_dong_moi_trang
    end = start + so_dong_moi_trang
    _disp(df.iloc[start:end])


# ── Định dạng số kiểu Việt Nam ───────────────────────────────────────────────
def vn(x, d=1, show_sign: bool = False):
    """Số thực → chuỗi VN (dấu . nghìn, dấu , thập phân). show_sign=True → +/- rõ ràng."""
    try:
        x = float(x)
        s = f"{x:,.{d}f}".replace(",", "X").replace(".", ",").replace("X", ".")
        s = s.rstrip("0").rstrip(",") if "," in s else s
        if show_sign and x > 0:
            return f"+{s}"
        return s
    except:
        return "—"


def fmt_tien(x):
    """Đồng → tỷ/triệu, định dạng VN."""
    try:
        x = float(x)
        if abs(x) >= 1e9:
            s = f"{x/1e9:,.3f}".replace(",","X").replace(".",",").replace("X",".")
            return f"{s.rstrip('0').rstrip(',') if ',' in s else s} tỷ đồng"
        if abs(x) >= 1e6:
            s = f"{x/1e6:,.1f}".replace(",","X").replace(".",",").replace("X",".")
            return f"{s} triệu đồng"
        if abs(x) > 0:
            return f"{x:,.0f}".replace(",",".") + " đồng"
        return "—"
    except:
        return "—"


def fmt_ty(x):
    """Đồng → tỷ ngắn gọn (dùng trong bảng so sánh Điện báo)."""
    try:
        x = float(x)
        ty = x / 1e9
        if abs(ty) >= 1:
            return f"{vn(ty, 3)} tỷ"
        if abs(x) >= 1e6:
            return f"{vn(x/1e6, 1)} triệu"
        if abs(x) > 0:
            return f"{vn(x/1e6, 3)} triệu"
        return "—"
    except:
        return "—"


def fmt_bang_ty(x, so_le: int = 3) -> str:
    """
    Format số tiền (đồng) → tỷ đồng, CỐ ĐỊNH đơn vị tỷ.
    Dùng trong cột bảng để đảm bảo đồng nhất đơn vị.
    Ví dụ:
        28_100_000_000 → "28,100"
        190_000_000    → "0,190"
        14_000_000     → "0,014"
        0              → "—"
    Tham số so_le: số chữ số thập phân (mặc định 3).
    Header cột phải ghi rõ "(tỷ đồng)" hoặc "(tỷ)".
    """
    try:
        x = float(x)
        if x == 0:
            return "—"
        ty = x / 1e9
        s = f"{ty:,.{so_le}f}".replace(",", "X").replace(".", ",").replace("X", ".")
        return s
    except Exception:
        return "—"


def get_tab_context(tab):
    """
    Trả về context manager cho tab.
    Nếu tab is None, trả về st.container() thay thế.
    """
    import streamlit as st
    return tab if tab is not None else st.container()


def fmt_cl(x):
    """Chênh lệch tỷ có dấu + / -."""
    try:
        x = float(x)
        s = fmt_ty(x)
        return ("+" + s if x > 0 else s) if s != "—" else "—"
    except:
        return "—"


def fmt_pct(x):
    """Tỷ lệ % có dấu +/-."""
    try:
        x = float(x)
        return (f"+{vn(x,1)}%" if x > 0 else f"{vn(x,1)}%") if x != 0 else "0%"
    except:
        return "—"


def fmt_so(x):
    """Số nguyên dấu . nghìn."""
    try:
        return f"{int(x):,}".replace(",",".")
    except:
        return "—"


def fmt(x):
    """
    Format số tiền (đồng) → tỷ/triệu ngắn gọn.
    Dùng thay cho các _fmt() định nghĩa rải rác ở các tab.
      >= 1 tỷ  → "1,234 tỷ"
      >= 1 triệu → "123,4 triệu"
      > 0      → "123.456"
      0 / lỗi  → "—"
    Hỗ trợ số âm (dùng trong cột chênh lệch).
    """
    try:
        x = float(x)
        if abs(x) >= 1e9:
            s = f"{x/1e9:,.3f}".replace(",","X").replace(".",",").replace("X",".")
            return f"{s.rstrip('0').rstrip(',')} tỷ"
        if abs(x) >= 1e6:
            s = f"{x/1e6:,.1f}".replace(",","X").replace(".",",").replace("X",".")
            return f"{s} triệu"
        if abs(x) > 0:
            return f"{x:,.0f}".replace(",",".")
        return "—"
    except:
        return "—"


def fmt_tl(th, kh):
    """Tỷ lệ thực hiện/kế hoạch → chuỗi '87,5%' hoặc '—'."""
    try:
        if float(kh) <= 0:
            return "—"
        return f"{float(th)/float(kh)*100:,.1f}%".replace(",","X").replace(".",",").replace("X",".")
    except:
        return "—"


# ── Xuất Excel BytesIO ────────────────────────────────────────────────────────
def xuat_excel(sheets: dict) -> bytes:
    """
    sheets = {"Tên sheet": dataframe, ...}
    Trả về bytes để dùng với st.download_button.
    """
    buf = BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as w:
        for sheet_name, df in sheets.items():
            df.to_excel(w, index=False, sheet_name=sheet_name[:31])
    return buf.getvalue()


def ten_file_xuat(prefix: str, ext="xlsx") -> str:
    return f"{prefix}_{datetime.today().strftime('%d%m%Y')}.{ext}"


# ═══════════════════════════════════════════════════════════════════════════════
# AUTO FILL DOCUMENT — điền dữ liệu vào file Word template
# ═══════════════════════════════════════════════════════════════════════════════
from pathlib import Path
from datetime import datetime as _dt
import copy


def _replace_in_para(para, replacements: dict):
    """Thay thế tag trong 1 paragraph, giữ nguyên định dạng."""
    full = "".join(r.text for r in para.runs)
    if not any(tag in full for tag in replacements):
        return
    for tag, val in replacements.items():
        full = full.replace(tag, str(val))
    # Ghi lại vào run đầu tiên, xóa runs còn lại
    if para.runs:
        para.runs[0].text = full
        for r in para.runs[1:]:
            r.text = ""


def _build_replacements(data_row, tag_map: dict, extra: dict = None) -> dict:
    """
    Tạo dict {tag → giá trị thực} từ 1 dòng dữ liệu (pd.Series hoặc dict).
    extra: các tag tự do không có trong tag_map (vd: ngày in, tên cán bộ...)
    """
    repl = {}
    for tag, col in tag_map.items():
        if hasattr(data_row, "get"):
            val = data_row.get(col, "")
        else:
            val = getattr(data_row, col, "")
        if val is None or (hasattr(val, "__class__") and val.__class__.__name__ == "float"
                           and str(val) == "nan"):
            val = ""
        repl[tag] = val
    # Luôn bổ sung ngày in tự động
    repl["{{ngay_in}}"]     = _dt.today().strftime("%d/%m/%Y")
    repl["{{ngay_in_day}}"] = _dt.today().strftime("%d")
    repl["{{ngay_in_mon}}"] = _dt.today().strftime("%m")
    repl["{{ngay_in_yr}}"]  = _dt.today().strftime("%Y")
    if extra:
        repl.update(extra)
    return repl


def auto_fill_document(data_row, template_path: str, tag_map: dict,
                       extra: dict = None) -> bytes:
    """
    Điền dữ liệu vào file Word template (.docx).

    Params
    ------
    data_row      : pd.Series hoặc dict — 1 hàng dữ liệu khách hàng
    template_path : đường dẫn tới file .docx mẫu
    tag_map       : dict {tag → tên cột}, lấy từ config.TAG_MAP
    extra         : dict tag bổ sung tự do

    Returns
    -------
    bytes — nội dung file .docx đã điền, dùng trực tiếp với st.download_button
    """
    try:
        from docx import Document
        from io import BytesIO
    except ImportError:
        raise ImportError("Cần cài: pip install python-docx")

    doc  = Document(str(template_path))
    repl = _build_replacements(data_row, tag_map, extra)

    # Xử lý tất cả paragraphs (kể cả trong header/footer)
    for para in doc.paragraphs:
        _replace_in_para(para, repl)

    # Xử lý bảng trong document
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_para(para, repl)

    # Header / Footer
    for section in doc.sections:
        for hdr_ftr in [section.header, section.footer,
                        section.even_page_header, section.even_page_footer,
                        section.first_page_header, section.first_page_footer]:
            if hdr_ftr:
                for para in hdr_ftr.paragraphs:
                    _replace_in_para(para, repl)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def auto_fill_batch(df_rows, template_path: str, tag_map: dict,
                    extra: dict = None) -> bytes:
    """
    Điền nhiều hồ sơ vào 1 file Word — mỗi hồ sơ 1 trang (page break giữa các trang).
    Trả về bytes file .docx tổng hợp.
    """
    try:
        from docx import Document
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from io import BytesIO
        import lxml.etree as etree
    except ImportError:
        raise ImportError("Cần cài: pip install python-docx lxml")

    def add_page_break(doc):
        para = doc.add_paragraph()
        run  = para.add_run()
        br   = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        run._r.append(br)

    master = None
    for i, (_, row) in enumerate(df_rows.iterrows()):
        doc  = Document(str(template_path))
        repl = _build_replacements(row, tag_map, extra)
        for para in doc.paragraphs:
            _replace_in_para(para, repl)
        for table in doc.tables:
            for tr in table.rows:
                for cell in tr.cells:
                    for para in cell.paragraphs:
                        _replace_in_para(para, repl)
        if master is None:
            master = doc
        else:
            add_page_break(master)
            for elem in doc.element.body:
                master.element.body.append(copy.deepcopy(elem))

    if master is None:
        return b""
    buf = BytesIO()
    master.save(buf)
    return buf.getvalue()


def quet_templates(templates_dir) -> list:
    """Quét thư mục templates/, trả về list (tên_hien_thi, path)."""
    p = Path(str(templates_dir))
    if not p.exists():
        return []
    files = sorted(p.glob("*.docx"))
    return [(f.stem.replace("_", " ").title(), f) for f in files]


# ══════════════════════════════════════════════════════════════════════════════
# AUTO-FILL MẪU KL GIAO BAN — Bảng II cột 3 tháng không hoạt động
# ══════════════════════════════════════════════════════════════════════════════

def _tinh_so_lieu_klgb(df: "pd.DataFrame") -> dict:
    """
    Tính toán các chỉ số cần điền vào mẫu KL giao ban.
    Trả về dict {tag_key: giá_trị_chuỗi} để merge vào replacements.

    Tên ĐVUT thực tế trong HSTD (đã xác nhận với dữ liệu):
        "Hội nông dân", "Hội liên hiệp phụ nữ",
        "Hội cựu chiến binh", "Đoàn thanh niên"
    """
    from datetime import date
    from data import danh_dau_khong_hd, tong_hop_khong_hd, canh_bao_migration
    from config import (COT_TONG_DU_NO, COT_DU_NO_TH, COT_DU_NO_QH,
                        DVUT_TAG_KEY)

    today = date.today()
    so_lieu = {
        "ngay_in":       today.strftime("%d/%m/%Y"),
        "thang_bao_cao": today.strftime("%m/%Y"),
    }

    if df is None or df.empty:
        return so_lieu

    # Tổng dư nợ toàn PGD
    tdn = df[COT_TONG_DU_NO].sum() if COT_TONG_DU_NO in df.columns else 0
    dqh = df[COT_DU_NO_QH].sum()   if COT_DU_NO_QH   in df.columns else 0
    so_lieu["ty_le_nqh"] = f"{dqh/tdn*100:.3f}%" if tdn > 0 else "0.000%"

    # 3 tháng không hoạt động — tổng hợp theo ĐVUT
    df_kh = danh_dau_khong_hd(df)
    khd   = tong_hop_khong_hd(df_kh, nhom_theo="Tên ĐVUT")

    # Điền từng ĐVUT theo đúng tên thực tế trong file HSTD
    tong_3m = 0
    for ten_dvut_thuc, tag_key in DVUT_TAG_KEY.items():
        so_mon = 0
        if not khd.empty and "Tên ĐVUT" in khd.columns:
            r = khd[khd["Tên ĐVUT"].str.lower() == ten_dvut_thuc.lower()]
            so_mon = int(r["Món_3m_KHĐ"].sum()) if not r.empty else 0
        so_lieu[tag_key] = str(so_mon)
        tong_3m += so_mon
    so_lieu["mon_3m_tong"] = str(tong_3m)

    # Cảnh báo amber
    try:
        df_amber = canh_bao_migration(df_kh)
        so_lieu["mon_amber_tong"] = str(len(df_amber))
    except Exception:
        so_lieu["mon_amber_tong"] = "0"

    return so_lieu


def auto_fill_klgb(df: "pd.DataFrame", template_path: str,
                   ten_pgd: str = "") -> bytes:
    """
    Tạo file Word KL giao ban đã điền đầy đủ số liệu Bảng II.

    Params
    ------
    df            : DataFrame HSTD của PGD (đã lọc theo PGD nếu cần)
    template_path : đường dẫn file .docx mẫu KL giao ban
    ten_pgd       : tên PGD để điền vào {{ten_pgd}}

    Returns
    -------
    bytes — nội dung file .docx đã điền
    """
    try:
        from docx import Document
        from io import BytesIO
    except ImportError:
        raise ImportError("Cần cài: pip install python-docx")

    from config import TAG_MAP_KLGB

    # Tính số liệu
    so_lieu = _tinh_so_lieu_klgb(df)
    if ten_pgd:
        so_lieu["ten_pgd_override"] = ten_pgd

    # Build replacements từ TAG_MAP_KLGB
    repl = {}
    for tag, key in TAG_MAP_KLGB.items():
        # Nếu key trỏ vào cột DataFrame
        if df is not None and key in df.columns:
            val = df[key].sum() if pd.api.types.is_numeric_dtype(df[key]) \
                  else df[key].iloc[0] if len(df) > 0 else ""
            repl[tag] = fmt(val) if isinstance(val, (int, float)) else str(val)
        # Nếu key là computed value
        elif key in so_lieu:
            repl[tag] = str(so_lieu[key])
        else:
            repl[tag] = "—"

    # Ghi đè ten_pgd nếu truyền thủ công
    if ten_pgd:
        repl["{{ten_pgd}}"] = ten_pgd

    # Điền vào Word
    doc = Document(str(template_path))
    for para in doc.paragraphs:
        _replace_in_para(para, repl)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_para(para, repl)
    for section in doc.sections:
        for hdr_ftr in [section.header, section.footer]:
            if hdr_ftr:
                for para in hdr_ftr.paragraphs:
                    _replace_in_para(para, repl)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def pick_hstd_column(df: pd.DataFrame, *candidates: str) -> str | None:
    """
    Tìm tên cột đầu tiên khớp (case-insensitive, strip) trong candidates.
    Dùng để chịu lệch "Tên Thôn" vs "Tên thôn", "Tên xã" vs "Tên Xã"…
    """
    cols_stripped = {str(c).strip(): c for c in df.columns}
    for cand in candidates:
        key = cand.strip()
        if key in cols_stripped:
            return cols_stripped[key]
        for col_s, col_orig in cols_stripped.items():
            if col_s.lower() == key.lower():
                return col_orig
    return None
